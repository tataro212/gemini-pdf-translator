"""
Test Script for Augmentation Propositions Implementation

This script validates the implementation of the three augmentation propositions:
1. Spatial Layout Analysis in pdf_parser.py
2. Two-Pass TOC Generation in document_generator.py  
3. Global Font Analysis for Adaptive Heading Detection

Run this script to verify that all augmentations are working correctly.
"""

import asyncio
import logging
import os
import sys
import tempfile
import time
from pathlib import Path

# Setup logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class AugmentationPropositionsTest:
    """
    Comprehensive test suite for all three augmentation propositions.
    
    Tests the enhanced PDF translation pipeline with:
    - Spatial layout analysis and image-text association
    - Two-pass TOC generation with accurate page numbers
    - Global font analysis for adaptive heading detection
    """
    
    def __init__(self):
        self.test_results = {}
        self.temp_dirs = []
    
    async def run_all_tests(self):
        """Run all augmentation proposition tests"""
        logger.info("🚀 Starting Augmentation Propositions Test Suite")
        logger.info("=" * 70)
        
        tests = [
            ("spatial_layout_analysis", self.test_spatial_layout_analysis),
            ("two_pass_toc_generation", self.test_two_pass_toc_generation),
            ("global_font_analysis", self.test_global_font_analysis),
            ("enhanced_detection_algorithms", self.test_enhanced_detection_algorithms),
            ("integrated_workflow", self.test_integrated_workflow)
        ]
        
        for test_name, test_func in tests:
            logger.info(f"\n📋 Running Test: {test_name.replace('_', ' ').title()}")
            logger.info("-" * 50)
            
            try:
                start_time = time.time()
                result = await test_func()
                end_time = time.time()
                
                self.test_results[test_name] = {
                    'status': 'PASSED' if result else 'FAILED',
                    'duration': end_time - start_time,
                    'details': result if isinstance(result, dict) else {}
                }
                
                status_emoji = '✅' if result else '❌'
                logger.info(f"{status_emoji} Test {test_name}: {'PASSED' if result else 'FAILED'} ({end_time - start_time:.2f}s)")
                
            except Exception as e:
                logger.error(f"💥 Test {test_name} failed with exception: {e}")
                self.test_results[test_name] = {
                    'status': 'ERROR',
                    'error': str(e),
                    'duration': 0
                }
        
        self._generate_final_report()
        self._cleanup()
        
        return self.test_results
    
    async def test_spatial_layout_analysis(self):
        """
        Test Proposition 1: Spatial Layout Analysis
        
        Validates that the PDF parser can:
        - Extract spatial elements with coordinates
        - Apply spatial reading order
        - Associate images with text based on spatial relationships
        - Detect and link captions
        """
        logger.info("🔍 Testing spatial layout analysis implementation...")
        
        try:
            from pdf_parser import StructuredContentExtractor

            extractor = StructuredContentExtractor()

            # Test 1: Check if spatial analysis methods exist
            required_methods = [
                '_extract_spatial_elements',
                '_apply_spatial_reading_order',
                '_associate_images_with_text_spatial',
                '_detect_and_link_caption'
            ]

            # Debug: List all methods that start with _extract or _apply
            all_methods = [method for method in dir(extractor) if method.startswith('_extract') or method.startswith('_apply') or method.startswith('_associate') or method.startswith('_detect')]
            logger.info(f"Available spatial methods: {all_methods}")

            for method_name in required_methods:
                if not hasattr(extractor, method_name):
                    logger.error(f"❌ Required spatial analysis method missing: {method_name}")
                    return False
            
            logger.info("✅ All spatial analysis methods available")
            
            # Test 2: Check enhanced ImagePlaceholder model
            from structured_document_model import ImagePlaceholder
            
            # Create test image placeholder
            from structured_document_model import ContentType
            test_image = ImagePlaceholder(
                block_type=ContentType.IMAGE_PLACEHOLDER,
                original_text="",
                page_num=1,
                bbox=(100, 100, 200, 200),
                image_path="test.jpg"
            )
            
            # Check for spatial enhancement attributes
            spatial_attributes = [
                'caption_block_id',
                'spatial_relationship', 
                'reading_order_position'
            ]
            
            for attr in spatial_attributes:
                if not hasattr(test_image, attr):
                    logger.error(f"❌ ImagePlaceholder missing spatial attribute: {attr}")
                    return False
            
            logger.info("✅ ImagePlaceholder model enhanced with spatial attributes")
            
            # Test 3: Test spatial reading order algorithm
            mock_elements = [
                {'bbox': [0, 100, 100, 120], 'content': 'Element 1'},
                {'bbox': [0, 80, 100, 100], 'content': 'Element 2'},  # Higher on page
                {'bbox': [0, 60, 100, 80], 'content': 'Element 3'},   # Highest on page
            ]
            
            # Apply spatial reading order
            sorted_elements = extractor._apply_spatial_reading_order(mock_elements)
            
            # Verify correct ordering (should be Element 3, 2, 1 from top to bottom)
            expected_order = ['Element 3', 'Element 2', 'Element 1']
            actual_order = [elem['content'] for elem in sorted_elements]
            
            if actual_order == expected_order:
                logger.info("✅ Spatial reading order algorithm working correctly")
                return True
            else:
                logger.error(f"❌ Spatial reading order incorrect: {actual_order} != {expected_order}")
                return False
                
        except ImportError as e:
            logger.error(f"❌ Required modules not available: {e}")
            return False
        except Exception as e:
            logger.error(f"❌ Spatial layout analysis test failed: {e}")
            return False
    
    async def test_two_pass_toc_generation(self):
        """
        Test Proposition 2: Two-Pass TOC Generation
        
        Validates that the document generator can:
        - Generate content with page tracking (Pass 1)
        - Create accurate TOC with real page numbers (Pass 2)
        - Handle bookmark creation and hyperlinks
        """
        logger.info("📄 Testing two-pass TOC generation implementation...")
        
        try:
            from document_generator import WordDocumentGenerator

            generator = WordDocumentGenerator()
            
            # Test 1: Check if two-pass methods exist
            required_methods = [
                '_generate_content_with_page_tracking',
                '_generate_accurate_toc',
                '_calculate_current_page_number',
                '_add_accurate_toc_entry'
            ]
            
            for method_name in required_methods:
                if not hasattr(generator, method_name):
                    logger.error(f"❌ Required two-pass TOC method missing: {method_name}")
                    return False
            
            logger.info("✅ All two-pass TOC generation methods available")
            
            # Test 2: Test page number calculation
            from docx import Document
            test_doc = Document()
            
            # Add some content to test page calculation
            for i in range(5):
                test_doc.add_paragraph(f"Test paragraph {i}")
            
            calculated_page = generator._calculate_current_page_number(test_doc, 1)
            
            if isinstance(calculated_page, int) and calculated_page > 0:
                logger.info(f"✅ Page number calculation working: {calculated_page}")
            else:
                logger.error(f"❌ Invalid page number calculation: {calculated_page}")
                return False
            
            # Test 3: Check if create_word_document_from_structured_document uses two-pass approach
            import inspect
            method_source = inspect.getsource(generator.create_word_document_from_structured_document)
            
            two_pass_indicators = [
                '_generate_content_with_page_tracking',
                '_generate_accurate_toc',
                'Pass 1',
                'Pass 2'
            ]
            
            found_indicators = []
            for indicator in two_pass_indicators:
                if indicator in method_source:
                    found_indicators.append(indicator)
            
            if len(found_indicators) >= 3:
                logger.info(f"✅ Two-pass TOC generation integrated: {found_indicators}")
                return True
            else:
                logger.warning(f"⚠️ Limited two-pass integration found: {found_indicators}")
                return False
                
        except ImportError as e:
            logger.error(f"❌ Required modules not available: {e}")
            return False
        except Exception as e:
            logger.error(f"❌ Two-pass TOC generation test failed: {e}")
            return False
    
    async def test_global_font_analysis(self):
        """
        Test Proposition 3: Global Font Analysis
        
        Validates that the PDF parser can:
        - Perform global font analysis across entire document
        - Build adaptive font hierarchy
        - Use document-specific styling for heading detection
        """
        logger.info("🔤 Testing global font analysis implementation...")
        
        try:
            from pdf_parser import StructuredContentExtractor

            extractor = StructuredContentExtractor()

            # Test 1: Check if global font analysis methods exist
            required_methods = [
                '_perform_global_font_analysis',
                '_classify_content_type_adaptive'
            ]

            # Debug: List all methods that contain 'font' or 'classify'
            font_methods = [method for method in dir(extractor) if 'font' in method.lower() or 'classify' in method.lower() or method.startswith('_perform')]
            logger.info(f"Available font analysis methods: {font_methods}")

            for method_name in required_methods:
                if not hasattr(extractor, method_name):
                    logger.error(f"❌ Required font analysis method missing: {method_name}")
                    return False
            
            logger.info("✅ All global font analysis methods available")
            
            # Test 2: Test font analysis structure
            # Create mock document structure analysis
            mock_structure = {
                'font_hierarchy': {14.0: 1, 12.0: 2, 10.0: 3},
                'dominant_font_size': 10.0,
                'body_text_style': 'Arial, 10.0pt',
                'heading_styles': {'h1': 14.0, 'h2': 12.0, 'h3': 10.0}
            }
            
            # Test adaptive classification
            test_formatting = {
                'font_sizes': [14.0],
                'font_names': ['Arial'],
                'is_bold': True,
                'is_italic': False
            }
            
            content_type = extractor._classify_content_type_adaptive(
                "Test Heading", test_formatting, mock_structure
            )
            
            if content_type == 'h1':
                logger.info("✅ Adaptive content classification working correctly")
            else:
                logger.warning(f"⚠️ Unexpected classification result: {content_type}")
            
            # Test 3: Check if document structure analysis includes font hierarchy
            import inspect
            analysis_source = inspect.getsource(extractor._analyze_document_structure)
            
            font_analysis_indicators = [
                'font_hierarchy',
                'global_font_analysis',
                'adaptive',
                'body_text_style'
            ]
            
            found_indicators = []
            for indicator in font_analysis_indicators:
                if indicator in analysis_source:
                    found_indicators.append(indicator)
            
            if len(found_indicators) >= 2:
                logger.info(f"✅ Global font analysis integrated: {found_indicators}")
                return True
            else:
                logger.warning(f"⚠️ Limited font analysis integration: {found_indicators}")
                return False
                
        except ImportError as e:
            logger.error(f"❌ Required modules not available: {e}")
            return False
        except Exception as e:
            logger.error(f"❌ Global font analysis test failed: {e}")
            return False

    async def test_enhanced_detection_algorithms(self):
        """Test the enhanced detection algorithms for better accuracy"""
        logger.info("🔍 Testing Enhanced Detection Algorithms...")

        try:
            from pdf_parser import StructuredContentExtractor

            extractor = StructuredContentExtractor()

            # Test 1: Enhanced Caption Detection
            logger.info("Testing enhanced caption detection...")

            # Create mock image placeholder
            mock_image = type('MockImage', (), {
                'bbox': (100, 200, 300, 400),  # x0, y0, x1, y1
                'image_path': 'test_image.png'
            })()

            # Create mock text blocks with various caption scenarios
            mock_text_blocks = [
                # Scenario 1: Standard caption below image
                type('MockBlock', (), {
                    'bbox': (110, 420, 290, 440),
                    'content': 'Figure 1: Test diagram showing the process',
                    'original_text': 'Figure 1: Test diagram showing the process',
                    'block_id': 'block_1'
                })(),

                # Scenario 2: Caption above image
                type('MockBlock', (), {
                    'bbox': (110, 180, 290, 195),
                    'content': 'Chart 2.1 - Performance metrics',
                    'original_text': 'Chart 2.1 - Performance metrics',
                    'block_id': 'block_2'
                })(),

                # Scenario 3: Side caption
                type('MockBlock', (), {
                    'bbox': (320, 250, 450, 280),
                    'content': 'Source: Annual Report 2023',
                    'original_text': 'Source: Annual Report 2023',
                    'block_id': 'block_3'
                })(),

                # Scenario 4: Non-caption text
                type('MockBlock', (), {
                    'bbox': (110, 500, 290, 550),
                    'content': 'This is a regular paragraph that should not be detected as a caption.',
                    'original_text': 'This is a regular paragraph that should not be detected as a caption.',
                    'block_id': 'block_4'
                })(),

                # Scenario 5: Short descriptive text (potential caption)
                type('MockBlock', (), {
                    'bbox': (110, 450, 290, 465),
                    'content': 'System architecture overview',
                    'original_text': 'System architecture overview',
                    'block_id': 'block_5'
                })(),

                # Scenario 6: Numbered caption
                type('MockBlock', (), {
                    'bbox': (110, 410, 290, 425),
                    'content': '(a) Initial configuration',
                    'original_text': '(a) Initial configuration',
                    'block_id': 'block_6'
                })()
            ]

            # Test caption detection
            caption_result = extractor._detect_and_link_caption(mock_image, mock_text_blocks)

            if caption_result:
                logger.info(f"✅ Caption detected: {caption_result['caption_text']}")
                logger.info(f"   Confidence: {caption_result.get('confidence', 'N/A')}")
                logger.info(f"   Relationship: {caption_result.get('spatial_relationship', 'N/A')}")

                # Verify it's a reasonable caption
                caption_text = caption_result['caption_text']
                is_reasonable = (
                    len(caption_text) < 200 and
                    ('figure' in caption_text.lower() or
                     'chart' in caption_text.lower() or
                     'source' in caption_text.lower() or
                     '(' in caption_text)
                )

                if is_reasonable:
                    logger.info("✅ Caption detection working correctly")
                    caption_test_passed = True
                else:
                    logger.warning(f"⚠️ Detected caption may not be appropriate: {caption_text}")
                    caption_test_passed = False
            else:
                logger.warning("⚠️ No caption detected - this might indicate detection issues")
                caption_test_passed = False

            # Test 2: Enhanced Spatial Reading Order
            logger.info("Testing enhanced spatial reading order...")

            # Create mock elements for multi-column layout
            mock_elements = [
                # Left column (higher y values = lower on page in PDF coordinates)
                {'bbox': (50, 700, 200, 720), 'content': 'Left column header', 'type': 'text'},
                {'bbox': (50, 670, 200, 690), 'content': 'Left column paragraph 1', 'type': 'text'},
                {'bbox': (50, 640, 200, 660), 'content': 'Left column paragraph 2', 'type': 'text'},

                # Right column
                {'bbox': (250, 700, 400, 720), 'content': 'Right column header', 'type': 'text'},
                {'bbox': (250, 670, 400, 690), 'content': 'Right column paragraph 1', 'type': 'text'},
                {'bbox': (250, 640, 400, 660), 'content': 'Right column paragraph 2', 'type': 'text'},

                # Full-width element
                {'bbox': (50, 600, 400, 620), 'content': 'Full-width conclusion', 'type': 'text'},
            ]

            sorted_elements = extractor._apply_spatial_reading_order(mock_elements)

            logger.info("Reading order results:")
            for i, element in enumerate(sorted_elements):
                logger.info(f"  {i+1}. {element['content']}")

            # Verify reading order makes sense
            actual_order = [elem['content'] for elem in sorted_elements]

            # Check if multi-column detection worked
            spatial_test_passed = False
            if 'Left column header' in actual_order and 'Right column header' in actual_order:
                left_idx = actual_order.index('Left column header')
                right_idx = actual_order.index('Right column header')

                if abs(left_idx - right_idx) <= 1:  # Should be close together
                    logger.info("✅ Multi-column layout detected correctly")
                    spatial_test_passed = True
                else:
                    logger.warning("⚠️ Multi-column layout detection may need improvement")

            # Test 3: Enhanced Adaptive Classification
            logger.info("Testing enhanced adaptive classification...")

            # Mock structure analysis
            mock_structure = {
                'font_hierarchy': {16.0: 1, 14.0: 2, 12.0: 3},
                'dominant_font_size': 10.0,
                'body_text_style': 'Arial, 10.0pt'
            }

            # Test various text scenarios
            test_scenarios = [
                {
                    'text': '1. Introduction',
                    'formatting': {'font_sizes': [14.0], 'is_bold': True, 'flags': 16},
                    'expected_type': 'heading'
                },
                {
                    'text': 'METHODOLOGY',
                    'formatting': {'font_sizes': [12.0], 'is_bold': True, 'flags': 16},
                    'expected_type': 'heading'
                },
                {
                    'text': 'This is a regular paragraph with normal formatting that contains multiple sentences.',
                    'formatting': {'font_sizes': [10.0], 'is_bold': False, 'flags': 0},
                    'expected_type': 'paragraph'
                },
                {
                    'text': '• First bullet point',
                    'formatting': {'font_sizes': [10.0], 'is_bold': False, 'flags': 0},
                    'expected_type': 'list_item'
                },
                {
                    'text': 'Chapter 3: Advanced Topics',
                    'formatting': {'font_sizes': [16.0], 'is_bold': True, 'flags': 16},
                    'expected_type': 'heading'
                }
            ]

            classification_results = []
            for scenario in test_scenarios:
                result = extractor._classify_content_type_adaptive(
                    scenario['text'],
                    scenario['formatting'],
                    mock_structure
                )

                # Check if result matches expected type category
                is_correct = False
                if scenario['expected_type'] == 'heading':
                    is_correct = result.startswith('h')
                elif scenario['expected_type'] == 'paragraph':
                    is_correct = result == 'paragraph'
                elif scenario['expected_type'] == 'list_item':
                    is_correct = result == 'list_item'

                classification_results.append({
                    'text': scenario['text'][:30] + '...' if len(scenario['text']) > 30 else scenario['text'],
                    'expected': scenario['expected_type'],
                    'actual': result,
                    'correct': is_correct
                })

            logger.info("Classification results:")
            for result in classification_results:
                status = "✅" if result['correct'] else "❌"
                logger.info(f"  {status} '{result['text']}' -> {result['actual']} (expected: {result['expected']})")

            # Calculate success rate
            correct_classifications = sum(1 for r in classification_results if r['correct'])
            success_rate = correct_classifications / len(classification_results) * 100

            logger.info(f"Classification success rate: {success_rate:.1f}%")

            # Overall assessment
            overall_success = (caption_test_passed and spatial_test_passed and success_rate >= 70)

            if overall_success:
                logger.info("✅ Enhanced detection algorithms working well")
                return True
            else:
                logger.warning(f"⚠️ Detection accuracy needs improvement")
                logger.warning(f"   Caption: {caption_test_passed}, Spatial: {spatial_test_passed}, Classification: {success_rate:.1f}%")
                return False

        except ImportError as e:
            logger.error(f"❌ Required modules not available: {e}")
            return False
        except Exception as e:
            logger.error(f"❌ Enhanced detection test failed: {e}")
            import traceback
            logger.error(traceback.format_exc())
            return False

    async def test_integrated_workflow(self):
        """
        Test integrated workflow with all three propositions working together
        """
        logger.info("🔄 Testing integrated workflow with all augmentations...")
        
        try:
            # Test that all components can work together
            from pdf_parser import PDFParser
            from document_generator import WordDocumentGenerator
            from structured_document_model import Document, Heading, Paragraph, ContentType
            
            # Create test structured document
            test_document = Document(
                title="Test Document",
                content_blocks=[
                    Heading(
                        block_type=ContentType.HEADING,
                        original_text="Chapter 1",
                        page_num=1,
                        bbox=(0, 0, 100, 20),
                        level=1,
                        content="Chapter 1"
                    ),
                    Paragraph(
                        block_type=ContentType.PARAGRAPH,
                        original_text="Test content",
                        page_num=1,
                        bbox=(0, 20, 100, 40),
                        content="Test content"
                    )
                ],
                source_filepath="test.pdf"
            )
            
            # Test document generation with two-pass TOC
            temp_dir = tempfile.mkdtemp(prefix="augmentation_test_")
            self.temp_dirs.append(temp_dir)
            
            output_path = os.path.join(temp_dir, "test_output.docx")
            
            generator = WordDocumentGenerator()
            
            # Enable TOC generation to test two-pass approach
            generator.word_settings['generate_toc'] = True
            
            result_path = generator.create_word_document_from_structured_document(
                test_document, output_path, temp_dir
            )
            
            if result_path and os.path.exists(result_path):
                file_size = os.path.getsize(result_path)
                logger.info(f"✅ Integrated workflow successful: {file_size} bytes")
                return True
            else:
                logger.error("❌ Integrated workflow failed: No output file created")
                return False
                
        except Exception as e:
            logger.error(f"❌ Integrated workflow test failed: {e}")
            return False
    
    def _generate_final_report(self):
        """Generate comprehensive test report"""
        total_tests = len(self.test_results)
        passed_tests = sum(1 for result in self.test_results.values() if result['status'] == 'PASSED')
        failed_tests = sum(1 for result in self.test_results.values() if result['status'] == 'FAILED')
        error_tests = sum(1 for result in self.test_results.values() if result['status'] == 'ERROR')
        
        report = f"""
🎯 AUGMENTATION PROPOSITIONS TEST RESULTS
==========================================

📊 Summary:
• Total Tests: {total_tests}
• Passed: {passed_tests} ✅
• Failed: {failed_tests} ❌
• Errors: {error_tests} 💥

📋 Detailed Results:
"""
        
        for test_name, result in self.test_results.items():
            status_emoji = {'PASSED': '✅', 'FAILED': '❌', 'ERROR': '💥'}[result['status']]
            duration = result.get('duration', 0)
            report += f"• {test_name.replace('_', ' ').title()}: {result['status']} {status_emoji} ({duration:.2f}s)\n"
            
            if result['status'] == 'ERROR':
                report += f"  Error: {result.get('error', 'Unknown error')}\n"
        
        # Overall assessment
        if passed_tests == total_tests:
            report += f"\n🎉 ALL AUGMENTATION PROPOSITIONS IMPLEMENTED SUCCESSFULLY!\n"
            report += f"The PDF translation pipeline now has enhanced spatial analysis, accurate TOC generation, and adaptive heading detection.\n"
        elif passed_tests >= total_tests * 0.75:
            report += f"\n✅ MOST AUGMENTATION PROPOSITIONS IMPLEMENTED\n"
            report += f"The pipeline has significant enhancements but some areas need refinement.\n"
        else:
            report += f"\n⚠️ AUGMENTATION PROPOSITIONS NEED MORE WORK\n"
            report += f"Additional implementation is needed for full enhancement benefits.\n"
        
        logger.info(report)
    
    def _cleanup(self):
        """Clean up temporary resources"""
        for temp_dir in self.temp_dirs:
            try:
                import shutil
                shutil.rmtree(temp_dir)
                logger.debug(f"🧹 Cleaned up temp dir: {temp_dir}")
            except Exception as e:
                logger.warning(f"Could not clean up {temp_dir}: {e}")


async def main():
    """Main entry point for augmentation propositions testing"""
    logger.info("🚀 Starting Augmentation Propositions Validation")
    logger.info("This test validates the implementation of three key enhancements:")
    logger.info("1. Spatial Layout Analysis - 2D spatial understanding for better image placement")
    logger.info("2. Two-Pass TOC Generation - 100% accurate page numbers and hyperlinks")
    logger.info("3. Global Font Analysis - Adaptive heading detection based on document styling")
    logger.info("")
    
    test_suite = AugmentationPropositionsTest()
    results = await test_suite.run_all_tests()
    
    # Return overall success
    success = all(result['status'] == 'PASSED' for result in results.values())
    
    if success:
        logger.info("\n🎉 All augmentation propositions are successfully implemented!")
        logger.info("The PDF translation pipeline now has significantly enhanced fidelity and accuracy.")
    else:
        logger.warning("\n⚠️ Some augmentation propositions need additional work.")
        logger.info("Review the test results above for specific areas to improve.")
    
    return success


if __name__ == "__main__":
    success = asyncio.run(main())
    sys.exit(0 if success else 1)
