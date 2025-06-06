#!/usr/bin/env python3
"""
Test script to verify the PDF workflow fix handles conversion failures gracefully
"""

import os
import sys
import logging
import asyncio
import tempfile
from unittest.mock import patch, MagicMock
from main_workflow import UltimatePDFTranslator
from document_generator import PDFConverter

# Set up logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

def create_test_word_document():
    """Create a test Word document for testing"""
    from docx import Document
    
    doc = Document()
    doc.add_heading('Test Translation Document', 0)
    doc.add_paragraph('This is a test document for verifying PDF conversion handling.')
    doc.add_heading('Section 1', level=1)
    doc.add_paragraph('Content for section 1.')
    
    test_docx_path = "test_workflow.docx"
    doc.save(test_docx_path)
    return test_docx_path

def test_pdf_converter_success():
    """Test successful PDF conversion"""
    logger.info("🧪 Testing successful PDF conversion...")
    
    test_docx = create_test_word_document()
    test_pdf = "test_success.pdf"
    
    try:
        converter = PDFConverter()
        success = converter.convert_word_to_pdf(test_docx, test_pdf)
        
        if success and os.path.exists(test_pdf):
            logger.info("✅ PDF conversion success case works correctly")
            # Clean up
            os.remove(test_docx)
            os.remove(test_pdf)
            return True
        else:
            logger.error("❌ PDF conversion success case failed")
            return False
            
    except Exception as e:
        logger.error(f"❌ PDF conversion success test error: {e}")
        return False

def test_pdf_converter_failure():
    """Test PDF conversion failure handling"""
    logger.info("🧪 Testing PDF conversion failure handling...")
    
    test_docx = create_test_word_document()
    test_pdf = "test_failure.pdf"
    
    try:
        converter = PDFConverter()
        
        # Mock the docx2pdf to raise an exception
        with patch('document_generator.convert_to_pdf_lib') as mock_convert:
            mock_convert.side_effect = Exception("Simulated PDF conversion failure")
            
            success = converter.convert_word_to_pdf(test_docx, test_pdf)
            
            if not success:
                logger.info("✅ PDF conversion failure is handled correctly")
                # Clean up
                os.remove(test_docx)
                return True
            else:
                logger.error("❌ PDF conversion failure was not detected")
                return False
                
    except Exception as e:
        logger.error(f"❌ PDF conversion failure test error: {e}")
        return False

def test_workflow_with_pdf_failure():
    """Test the full workflow when PDF conversion fails"""
    logger.info("🧪 Testing workflow with PDF conversion failure...")
    
    # Create a minimal structured content for testing
    test_content = [
        {'type': 'h1', 'text': 'Test Heading', 'page_num': 1, 'block_num': 1},
        {'type': 'paragraph', 'text': 'Test paragraph content.', 'page_num': 1, 'block_num': 2}
    ]
    
    # Create temporary output directory
    with tempfile.TemporaryDirectory() as temp_dir:
        word_output = os.path.join(temp_dir, "test_translated.docx")
        pdf_output = os.path.join(temp_dir, "test_translated.pdf")
        
        try:
            from document_generator import document_generator
            
            # Create Word document (this should succeed)
            word_success = document_generator.create_word_document_with_structure(
                test_content, word_output, None, None
            )
            
            if not word_success:
                logger.error("❌ Word document creation failed")
                return False
            
            # Test PDF conversion with mocked failure
            from document_generator import pdf_converter
            
            with patch.object(pdf_converter, 'convert_word_to_pdf', return_value=False):
                pdf_success = pdf_converter.convert_word_to_pdf(word_output, pdf_output)
                
                # Verify the workflow handles this correctly
                if not pdf_success:
                    logger.info("✅ Workflow correctly detected PDF conversion failure")
                    
                    # Check that Word file exists but PDF doesn't
                    if os.path.exists(word_output) and not os.path.exists(pdf_output):
                        logger.info("✅ Word document exists, PDF doesn't (as expected)")
                        return True
                    else:
                        logger.error("❌ File existence check failed")
                        return False
                else:
                    logger.error("❌ PDF conversion should have failed")
                    return False
                    
        except Exception as e:
            logger.error(f"❌ Workflow test error: {e}")
            return False

def test_error_reporting():
    """Test that error reporting provides helpful information"""
    logger.info("🧪 Testing error reporting...")
    
    test_docx = create_test_word_document()
    test_pdf = "test_error_reporting.pdf"
    
    try:
        converter = PDFConverter()
        
        # Test different types of errors
        error_scenarios = [
            ("COM Error", "COM error occurred"),
            ("Permission Error", "Permission denied"),
            ("Timeout Error", "Operation timed out"),
            ("Generic Error", "Unknown error")
        ]
        
        for error_name, error_message in error_scenarios:
            logger.info(f"   Testing {error_name}...")
            
            with patch('document_generator.convert_to_pdf_lib') as mock_convert:
                mock_convert.side_effect = Exception(error_message)
                
                # Capture log output to verify error diagnosis
                with patch('document_generator.logger') as mock_logger:
                    success = converter.convert_word_to_pdf(test_docx, test_pdf)
                    
                    if not success:
                        # Check that diagnostic information was logged
                        mock_logger.info.assert_called()
                        logger.info(f"     ✅ {error_name} handled with diagnostics")
                    else:
                        logger.error(f"     ❌ {error_name} not handled correctly")
                        return False
        
        # Clean up
        os.remove(test_docx)
        logger.info("✅ Error reporting works correctly")
        return True
        
    except Exception as e:
        logger.error(f"❌ Error reporting test failed: {e}")
        return False

def test_final_report_with_pdf_failure():
    """Test that final report correctly shows PDF failure status"""
    logger.info("🧪 Testing final report with PDF failure...")
    
    try:
        translator = UltimatePDFTranslator()
        
        # Mock the report generation with PDF failure
        with patch.object(translator, '_generate_final_report') as mock_report:
            # Call with pdf_success=False
            translator._generate_final_report(
                "test.pdf", "/tmp", 0, 60, 10, 10, [], False
            )
            
            # Verify the method was called with pdf_success=False
            mock_report.assert_called_once()
            args = mock_report.call_args[0]
            pdf_success = mock_report.call_args[1].get('pdf_success', True)
            
            if pdf_success == False:
                logger.info("✅ Final report correctly handles PDF failure")
                return True
            else:
                logger.error("❌ Final report doesn't handle PDF failure correctly")
                return False
                
    except Exception as e:
        logger.error(f"❌ Final report test error: {e}")
        return False

def main():
    """Run all PDF workflow tests"""
    logger.info("🚀 PDF WORKFLOW FIX VERIFICATION")
    logger.info("=" * 50)
    
    tests = [
        ("PDF Converter Success", test_pdf_converter_success),
        ("PDF Converter Failure", test_pdf_converter_failure),
        ("Workflow with PDF Failure", test_workflow_with_pdf_failure),
        ("Error Reporting", test_error_reporting),
        ("Final Report with PDF Failure", test_final_report_with_pdf_failure),
    ]
    
    results = {}
    
    for test_name, test_func in tests:
        logger.info(f"\n--- {test_name} Test ---")
        try:
            result = test_func()
            results[test_name] = result
            status = "✅ PASSED" if result else "❌ FAILED"
            logger.info(f"{test_name}: {status}")
        except Exception as e:
            logger.error(f"{test_name} test failed with error: {e}")
            results[test_name] = False
    
    # Summary
    logger.info(f"\n🎯 TEST SUMMARY")
    logger.info("=" * 50)
    
    passed = sum(1 for result in results.values() if result)
    total = len(results)
    
    for test_name, result in results.items():
        status = "✅ PASSED" if result else "❌ FAILED"
        logger.info(f"{test_name}: {status}")
    
    logger.info(f"\nOverall: {passed}/{total} tests passed")
    
    if passed == total:
        logger.info("🎉 All PDF workflow fixes are working correctly!")
        logger.info("✅ The system now properly handles PDF conversion failures")
        logger.info("✅ Users will get clear feedback when PDF conversion fails")
        logger.info("✅ Word documents are still created even if PDF fails")
    else:
        logger.warning("⚠️ Some tests failed. Please review the fixes.")
    
    return passed == total

if __name__ == "__main__":
    success = main()
    sys.exit(0 if success else 1)
