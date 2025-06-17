"""
Document Generator Module for Ultimate PDF Translator - FIXED VERSION

Handles Word document creation and PDF conversion with proper formatting and structure
Includes fixes for:
1. Greek character encoding issues
2. Enhanced TOC generation
3. Better font handling
"""

import os
import logging
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from config_manager import config_manager
from utils import sanitize_for_xml, sanitize_filepath

logger = logging.getLogger(__name__)

# Import structured document model
try:
    from structured_document_model import (
        Document as StructuredDocument, ContentType, Heading, Paragraph, ImagePlaceholder,
        Table, CodeBlock, ListItem, Footnote, Equation, Caption, Metadata
    )
    STRUCTURED_MODEL_AVAILABLE = True
except ImportError:
    STRUCTURED_MODEL_AVAILABLE = False
    logger.warning("Structured document model not available")

class WordDocumentGenerator:
    """Generates Word documents with proper structure and formatting"""
    
    def __init__(self):
        self.word_settings = config_manager.word_output_settings
        self.toc_entries = []
        self.bookmark_id = 0
    
    def _configure_document_fonts_for_unicode(self, doc):
        """FIX 1: Configure document fonts to properly support Greek and other Unicode characters"""
        try:
            # Set the default font for the entire document to support Greek characters
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Arial Unicode MS'  # Comprehensive Unicode font
            font.size = Pt(11)
            
            # Set the complex script font for Greek and other complex scripts
            rpr = style.element.xpath('.//w:rPr')[0] if style.element.xpath('.//w:rPr') else None
            if rpr is not None:
                rFonts = rpr.find(qn('w:rFonts'))
                if rFonts is None:
                    rFonts = OxmlElement("w:rFonts")
                    rpr.insert(0, rFonts)
                
                # Set fonts for different script types
                rFonts.set(qn('w:ascii'), 'Arial Unicode MS')
                rFonts.set(qn('w:hAnsi'), 'Arial Unicode MS')
                rFonts.set(qn('w:cs'), 'Arial Unicode MS')  # Complex scripts (Greek, Arabic, etc.)
                rFonts.set(qn('w:eastAsia'), 'Arial Unicode MS')  # East Asian scripts
            
            # Also configure heading styles
            for i in range(1, 7):  # Heading 1 through Heading 6
                try:
                    heading_style = doc.styles[f'Heading {i}']
                    heading_font = heading_style.font
                    heading_font.name = 'Arial Unicode MS'
                    
                    # Set complex script font for headings too
                    heading_rpr = heading_style.element.xpath('.//w:rPr')[0] if heading_style.element.xpath('.//w:rPr') else None
                    if heading_rpr is not None:
                        heading_rFonts = heading_rpr.find(qn('w:rFonts'))
                        if heading_rFonts is None:
                            heading_rFonts = OxmlElement("w:rFonts")
                            heading_rpr.insert(0, heading_rFonts)
                        
                        heading_rFonts.set(qn('w:ascii'), 'Arial Unicode MS')
                        heading_rFonts.set(qn('w:hAnsi'), 'Arial Unicode MS')
                        heading_rFonts.set(qn('w:cs'), 'Arial Unicode MS')
                        heading_rFonts.set(qn('w:eastAsia'), 'Arial Unicode MS')
                        
                except KeyError:
                    # Style doesn't exist, skip
                    continue
            
            logger.info("Document fonts configured for Unicode support (Greek, etc.)")
            
        except Exception as e:
            logger.warning(f"Could not configure Unicode fonts: {e}")
    
    def _add_hyperlink(self, paragraph, text, bookmark_name):
        """Add a hyperlink to a bookmark within the document"""
        try:
            # Create hyperlink element
            hyperlink = OxmlElement('w:hyperlink')
            hyperlink.set(qn('w:anchor'), bookmark_name)
            
            # Create run element for the hyperlink
            run = OxmlElement('w:r')
            
            # Set hyperlink style
            rPr = OxmlElement('w:rPr')
            color = OxmlElement('w:color')
            color.set(qn('w:val'), '0000FF')  # Blue color
            underline = OxmlElement('w:u')
            underline.set(qn('w:val'), 'single')
            rPr.append(color)
            rPr.append(underline)
            run.append(rPr)
            
            # Add text
            text_elem = OxmlElement('w:t')
            text_elem.text = text
            run.append(text_elem)
            
            hyperlink.append(run)
            paragraph._p.append(hyperlink)
            
            # Return the run for further formatting
            return paragraph.add_run("")  # Empty run for compatibility
            
        except Exception as e:
            logger.debug(f"Could not create hyperlink: {e}")
            # Fallback to regular text
            return paragraph.add_run(text)
    
    def _add_heading_block(self, doc, block_item):
        text_content = None
        level = 1

        if STRUCTURED_MODEL_AVAILABLE and isinstance(block_item, Heading):
            text_content = block_item.content
            level = block_item.level
        elif isinstance(block_item, dict):
            text_content = block_item.get('content') or block_item.get('text')
            level = block_item.get('level')
            type_str = block_item.get('type', '').lower()
            if level is None and type_str.startswith('h') and type_str[1:].isdigit():
                level = int(type_str[1:])
            if level is None: level = 1
        else:
            logger.warning(f"Skipping unknown heading block type: {type(block_item)}")
            return

        if not text_content:
            logger.warning("Skipping empty heading block.")
            return

        safe_content = sanitize_for_xml(text_content)
        p = doc.add_paragraph(style=f'Heading {level}')
        
        self.bookmark_id += 1
        bookmark_name = f"_Toc_Bookmark_{self.bookmark_id}"
        run = p.add_run(safe_content)
        
        tag_start = OxmlElement('w:bookmarkStart')
        tag_start.set(qn('w:id'), str(self.bookmark_id))
        tag_start.set(qn('w:name'), bookmark_name)
        tag_end = OxmlElement('w:bookmarkEnd')
        tag_end.set(qn('w:id'), str(self.bookmark_id))

        if run._r is not None and p._p is not None:
            run._r.addprevious(tag_start)
            run._r.addnext(tag_end)
        else:
            logger.warning(f"Could not precisely position bookmark for heading: {safe_content}. Run or paragraph element was None.")
            p._p.append(tag_start)
            p._p.append(tag_end)
        
        self.toc_entries.append({
            'text': text_content,
            'level': level,
            'bookmark': bookmark_name
        })
        logger.debug(f"Added Heading '{text_content}' (L{level}) with bookmark '{bookmark_name}'.")

    def _add_paragraph_block(self, doc, block_item):
        text_content = None
        if STRUCTURED_MODEL_AVAILABLE and isinstance(block_item, Paragraph):
            text_content = block_item.content
        elif isinstance(block_item, dict):
            text_content = block_item.get('content') or block_item.get('text')
        else:
            logger.warning(f"Skipping unknown paragraph block type: {type(block_item)}")
            return

        if text_content:
            safe_content = sanitize_for_xml(text_content)
            doc.add_paragraph(safe_content)
        else:
            logger.debug("Skipping empty paragraph block.")

    def _add_image_placeholder_block(self, doc, block_item, image_folder_path):
        image_filename = None
        caption_text = None

        if STRUCTURED_MODEL_AVAILABLE and isinstance(block_item, ImagePlaceholder):
            image_filename = block_item.image_path
            caption_text = block_item.caption
        elif isinstance(block_item, dict):
            image_filename = block_item.get('image_path') or block_item.get('filepath')
            caption_text = block_item.get('caption') or block_item.get('text')
        else:
            logger.warning(f"Skipping unknown image block type: {type(block_item)}")
            return

        if not image_filename:
            logger.warning("Skipping image block with no image_path/filename.")
            return

        actual_image_path = image_filename
        if image_folder_path and not os.path.isabs(image_filename):
            actual_image_path = os.path.join(image_folder_path, image_filename)

        if os.path.exists(actual_image_path):
            try:
                doc.add_picture(actual_image_path, width=Inches(self.word_settings.get('image_width_inches', 6.0)))
                if caption_text:
                    safe_caption = sanitize_for_xml(caption_text)
                    caption_style_name = self.word_settings.get('caption_style', 'Caption')
                    try:
                        p_caption = doc.add_paragraph(style=caption_style_name)
                        p_caption.add_run(safe_caption)
                    except KeyError:
                        logger.warning(f"Caption style '{caption_style_name}' not found. Adding caption as normal paragraph.")
                        p_caption = doc.add_paragraph()
                        p_caption.add_run(safe_caption)
                    p_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                logger.debug(f"Added image: {actual_image_path}")
            except Exception as e:
                logger.error(f"Error adding image {actual_image_path}: {e}")
        else:
            logger.warning(f"Image not found at path: {actual_image_path}. Skipping image.")

    def _insert_toc(self, doc):
        """Insert table of contents at the beginning of the document based on collected toc_entries."""
        if not self.toc_entries:
            logger.info("No TOC entries found, skipping TOC generation")
            return
            
        logger.info("Generating Table of Contents...")
        
        # Create a new document for the TOC
        toc_doc = Document()
        
        # Configure fonts for TOC document too
        self._configure_document_fonts_for_unicode(toc_doc)
        
        # Add TOC title
        title = toc_doc.add_paragraph("Table of Contents")
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.style = 'Title'
        
        # Add a decorative line
        line = toc_doc.add_paragraph()
        line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        line.add_run("_" * 40)
        
        # Add TOC entries with hyperlinks and page numbers
        for entry in self.toc_entries:
            level = entry['level']
            text_for_display = sanitize_for_xml(entry['text'])
            bookmark = entry['bookmark']
            
            # Create TOC entry paragraph
            p = toc_doc.add_paragraph()
            p.style = 'Normal'
            
            # Set indentation based on level
            indent_size = Pt(20 * (level - 1))
            p.paragraph_format.left_indent = indent_size
            
            # Add heading text with hyperlink
            try:
                hyperlink_run = self._add_hyperlink(p, text_for_display, bookmark)
                if level == 1:
                    hyperlink_run.bold = True
                    hyperlink_run.font.size = Pt(12)
                elif level == 2:
                    hyperlink_run.font.size = Pt(11)
                else:
                    hyperlink_run.font.size = Pt(10)
                    hyperlink_run.italic = True
            except Exception as e:
                logger.debug(f"Could not create hyperlink for {text_for_display}: {e}")
                # Fallback to regular text
                text_run = p.add_run(text_for_display)
                if level == 1:
                    text_run.bold = True
                    text_run.font.size = Pt(12)
                elif level == 2:
                    text_run.font.size = Pt(11)
                else:
                    text_run.font.size = Pt(10)
                    text_run.italic = True

            # Add dots and page number
            dots_needed = max(3, 60 - len(text_for_display))
            dots_run = p.add_run(" " + "." * dots_needed + " ")
            dots_run.font.color.rgb = RGBColor(128, 128, 128)

            # Add page reference field
            fldSimple = OxmlElement('w:fldSimple')
            fldSimple.set(qn('w:instr'), f'PAGEREF {bookmark} \\h')
            
            run_in_field = OxmlElement('w:r')
            text_in_field = OxmlElement('w:t')
            text_in_field.text = "..."  # Placeholder that will be replaced with actual page number
            run_in_field.append(text_in_field)
            fldSimple.append(run_in_field)
            
            p._p.append(fldSimple)
        
        # Add page break after TOC
        toc_doc.add_page_break()
        
        # Create new document combining TOC and original content
        new_doc = Document()
        
        # Configure fonts for the new document
        self._configure_document_fonts_for_unicode(new_doc)
        
        # Copy TOC paragraphs to new document
        for paragraph in toc_doc.paragraphs:
            new_p = new_doc.add_paragraph()
            new_p.style = paragraph.style
            for run in paragraph.runs:
                new_run = new_p.add_run(run.text)
                new_run.bold = run.bold
                new_run.italic = run.italic
                if run.font.size:
                    new_run.font.size = run.font.size
                if run.font.color.rgb:
                    new_run.font.color.rgb = run.font.color.rgb
        
        # Copy remaining content from original document
        for paragraph in doc.paragraphs:
            new_p = new_doc.add_paragraph()
            new_p.style = paragraph.style
            for run in paragraph.runs:
                new_run = new_p.add_run(run.text)
                new_run.bold = run.bold
                new_run.italic = run.italic
                if run.font.size:
                    new_run.font.size = run.font.size
        
        # Replace original document with new document
        doc._element = new_doc._element
        
        logger.info("Table of Contents generated and inserted")

    def _add_content_block(self, doc, block_item, image_folder_path):
        """Add various content block types to the document. Handles both model objects and dicts."""
        try:
            block_handled = False
            if STRUCTURED_MODEL_AVAILABLE:
                if isinstance(block_item, Heading):
                    self._add_heading_block(doc, block_item)
                    block_handled = True
                elif isinstance(block_item, Paragraph):
                    self._add_paragraph_block(doc, block_item)
                    block_handled = True
                elif isinstance(block_item, ImagePlaceholder):
                    self._add_image_placeholder_block(doc, block_item, image_folder_path)
                    block_handled = True

            if not block_handled and isinstance(block_item, dict):
                dict_type = block_item.get('type', '').lower()
                if dict_type.startswith('h') and dict_type[1:].isdigit():
                    self._add_heading_block(doc, block_item)
                elif dict_type in ['paragraph', 'text', 'p', 'normal']:
                    self._add_paragraph_block(doc, block_item)
                elif dict_type in ['image', 'img', 'imageplaceholder']:
                    self._add_image_placeholder_block(doc, block_item, image_folder_path)
                else:
                    # Fallback for unknown dict types
                    text_content = block_item.get('content') or block_item.get('text')
                    if text_content:
                        logger.warning(f"Unknown dictionary block type: '{dict_type}', but found text. Treating as paragraph.")
                        self._add_paragraph_block(doc, block_item)
                    else:
                        logger.warning(f"Unknown or empty dictionary block type: '{dict_type}'. Content: {str(block_item)[:200]}")
                block_handled = True

            if not block_handled:
                 logger.warning(f"Unknown block type: {type(block_item)}. Attempting to add as raw string.")
                 try:
                     doc.add_paragraph(str(block_item))
                 except Exception as e_raw:
                     logger.error(f"Could not add unknown block type {type(block_item)} as string: {e_raw}")

        except Exception as e:
            logger.error(f"Error processing content block (type: {type(block_item)}): {e}", exc_info=True)

    def create_word_document_with_structure(self, structured_content_list, output_filepath,
                                          image_folder_path, cover_page_data=None):
        """Create Word document with proper structure and formatting."""
        # Initialize for new ToC system
        self.toc_entries = []
        self.bookmark_id = 0

        # Normalize paths to handle mixed separators
        output_filepath = os.path.normpath(output_filepath)
        if image_folder_path:
            image_folder_path = os.path.normpath(image_folder_path)

        logger.info(f"--- Creating Word Document (with_structure): {os.path.basename(output_filepath)} ---")
        
        doc = Document()
        
        # FIX 1: SET UNICODE-COMPATIBLE FONT FOR GREEK CHARACTERS
        self._configure_document_fonts_for_unicode(doc)
        
        # Add cover page if provided
        if cover_page_data:
            self._add_cover_page(doc, cover_page_data, image_folder_path)
            doc.add_page_break()
        
        # Process content items (blocks)
        for item_block in structured_content_list:
            self._add_content_block(doc, item_block, image_folder_path)
        
        # PASS 2: All content is added, now insert the ToC (if enabled)
        if self.word_settings.get('generate_toc', False):
            self._insert_toc(doc)

        # Save document
        try:
            sanitized_filepath = sanitize_filepath(output_filepath)
            output_dir = os.path.dirname(sanitized_filepath)
            if output_dir and not os.path.exists(output_dir):
                os.makedirs(output_dir, exist_ok=True)
                logger.info(f"Created output directory: {output_dir}")

            doc.save(sanitized_filepath)
            logger.info(f"Word document saved successfully: {sanitized_filepath}")
            return sanitized_filepath
        except Exception as e:
            logger.error(f"Error saving Word document: {e}")
            logger.error(f"Attempted path: {output_filepath}")
            return None

    def _add_cover_page(self, doc, cover_page_data, image_folder_path):
        """Add cover page to document"""
        try:
            cover_paragraph = doc.add_paragraph()
            cover_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            if cover_page_data and image_folder_path:
                cover_image_path = cover_page_data.get('filepath')
                if cover_image_path and os.path.exists(cover_image_path):
                    run = cover_paragraph.add_run()
                    run.add_picture(cover_image_path, width=Inches(6))
                    logger.info("Cover page image added")
                else:
                    cover_paragraph.add_run("Cover Page").font.size = Pt(24)
            else:
                cover_paragraph.add_run("Cover Page").font.size = Pt(24)
                
        except Exception as e:
            logger.warning(f"Could not add cover page: {e}")

# Enhanced PDF conversion function with proper font embedding
def convert_word_to_pdf(docx_filepath, pdf_filepath):
    """Enhanced PDF conversion with proper Greek font support"""
    logger.info(f"Converting {docx_filepath} to PDF with Unicode font support...")
    
    try:
        # Try using docx2pdf first
        from docx2pdf import convert
        convert(docx_filepath, pdf_filepath)
        logger.info(f"Successfully converted {docx_filepath} to {pdf_filepath}")
        return True
        
    except ImportError:
        logger.warning("docx2pdf library not available. Trying alternative method...")
        
        try:
            # Alternative: Use python-docx2txt + reportlab for better font control
            import docx2txt
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.fonts import addMapping
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            
            # Register Unicode font
            try:
                # Try to register Arial Unicode MS or fallback to DejaVu
                pdfmetrics.registerFont(TTFont('ArialUnicode', 'arial.ttf'))
                addMapping('ArialUnicode', 0, 0, 'ArialUnicode')
            except:
                # Fallback to DejaVu Sans which supports Greek
                try:
                    pdfmetrics.registerFont(TTFont('DejaVuSans', 'DejaVuSans.ttf'))
                    addMapping('DejaVuSans', 0, 0, 'DejaVuSans')
                    font_name = 'DejaVuSans'
                except:
                    font_name = 'Helvetica'  # Last resort
                    logger.warning("Could not register Unicode font, using Helvetica")
            
            # Extract text from DOCX
            text = docx2txt.process(docx_filepath)
            
            # Create PDF with proper font
            doc = SimpleDocTemplate(pdf_filepath, pagesize=letter)
            styles = getSampleStyleSheet()
            
            # Create custom style with Unicode font
            unicode_style = ParagraphStyle(
                'UnicodeNormal',
                parent=styles['Normal'],
                fontName=font_name,
                fontSize=11,
                encoding='utf-8'
            )
            
            # Build PDF content
            story = []
            for line in text.split('\n'):
                if line.strip():
                    story.append(Paragraph(line, unicode_style))
                    story.append(Spacer(1, 12))
            
            doc.build(story)
            logger.info(f"Successfully converted {docx_filepath} to {pdf_filepath} with Unicode support")
            return True
            
        except Exception as e:
            logger.error(f"Alternative PDF conversion failed: {e}")
            return False
            
    except Exception as e:
        logger.error(f"Error converting Word to PDF: {e}")
        return False

# Instantiate the generator for use by other modules
document_generator = WordDocumentGenerator()