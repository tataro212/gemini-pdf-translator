#!/usr/bin/env python3
"""
Debug script to test file creation and path handling
"""

import os
import logging
from docx import Document

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def test_file_creation():
    """Test file creation in various locations"""
    
    # Test 1: Create a simple Word document in current directory
    print("🔧 Test 1: Creating Word document in current directory...")
    try:
        doc = Document()
        doc.add_paragraph("Test document content")
        test_path = "test_document.docx"
        doc.save(test_path)
        
        if os.path.exists(test_path):
            size = os.path.getsize(test_path)
            print(f"✅ Success: {test_path} created ({size} bytes)")
            os.remove(test_path)  # Clean up
        else:
            print(f"❌ Failed: {test_path} not created")
    except Exception as e:
        print(f"❌ Error: {e}")
    
    # Test 2: Test the exact path from the logs
    print("\n🔧 Test 2: Testing the exact path from logs...")
    test_dir = r"C:\Users\30694\Downloads\words-to-fire-press-betrayal\fel-a-militant-assessment-of-the-experience-of-the-libertarian-student-fed"
    test_file = os.path.join(test_dir, "test_file.docx")
    
    try:
        # Create directory if it doesn't exist
        os.makedirs(test_dir, exist_ok=True)
        print(f"📁 Directory created/exists: {test_dir}")
        
        # Create test file
        doc = Document()
        doc.add_paragraph("Test content for path debugging")
        doc.save(test_file)
        
        if os.path.exists(test_file):
            size = os.path.getsize(test_file)
            print(f"✅ Success: {test_file} created ({size} bytes)")
            
            # List directory contents
            files = os.listdir(test_dir)
            print(f"📂 Directory contents: {files}")
            
            os.remove(test_file)  # Clean up
        else:
            print(f"❌ Failed: {test_file} not created")
            
    except Exception as e:
        print(f"❌ Error: {e}")
    
    # Test 3: Check Downloads folder accessibility
    print("\n🔧 Test 3: Checking Downloads folder...")
    downloads_dir = r"C:\Users\30694\Downloads"
    try:
        if os.path.exists(downloads_dir):
            print(f"✅ Downloads folder exists: {downloads_dir}")
            contents = os.listdir(downloads_dir)
            print(f"📂 Downloads contents (first 10): {contents[:10]}")
        else:
            print(f"❌ Downloads folder not found: {downloads_dir}")
    except Exception as e:
        print(f"❌ Error accessing Downloads: {e}")
    
    # Test 4: Check if the specific output folder exists
    print("\n🔧 Test 4: Checking specific output folder...")
    specific_folder = r"C:\Users\30694\Downloads\words-to-fire-press-betrayal\fel-a-militant-assessment-of-the-experience-of-the-libertarian-student-fed"
    try:
        if os.path.exists(specific_folder):
            print(f"✅ Specific folder exists: {specific_folder}")
            contents = os.listdir(specific_folder)
            print(f"📂 Folder contents: {contents}")
        else:
            print(f"❌ Specific folder not found: {specific_folder}")
            
            # Check parent directories
            parent = os.path.dirname(specific_folder)
            if os.path.exists(parent):
                print(f"📁 Parent exists: {parent}")
                parent_contents = os.listdir(parent)
                print(f"📂 Parent contents: {parent_contents}")
            else:
                print(f"❌ Parent not found: {parent}")
                
    except Exception as e:
        print(f"❌ Error: {e}")

if __name__ == "__main__":
    test_file_creation()
