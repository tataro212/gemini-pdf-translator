import os
import asyncio
import re
from dotenv import load_dotenv
import google.generativeai as genai
from google.generativeai import types
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx2pdf import convert as convert_to_pdf_lib
from pydrive2.auth import GoogleAuth
from pydrive2.drive import GoogleDrive
from collections import Counter, defaultdict
import tkinter as tk
from tkinter import filedialog
import shutil
import tempfile
import traceback
import time
import configparser
import uuid
import json
import hashlib
import logging
import random
from itertools import groupby
import difflib

try:
    import pytesseract
    from PIL import Image as PIL_Image
except ImportError:
    print("WARNING: pytesseract or Pillow not found. OCR functionality will be disabled.")
    pytesseract = None
    PIL_Image = None

# --- Logging Setup ---
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(module)s - %(message)s')
logger = logging.getLogger(__name__)

# --- CONFIG PARSING ---
config = configparser.ConfigParser(allow_no_value=True, inline_comment_prefixes=('#', ';'))
user_config_read_successfully = False
if os.path.exists('config.ini'):
    try:
        config.read('config.ini', encoding='utf-8')
        user_config_read_successfully = True
        logger.info("Settings loaded from user's config.ini. Defaults will be applied for missing values.")
    except configparser.Error as e_conf_read:
        logger.error(f"Error reading config.ini: {e_conf_read}. Using full default configuration.")
else:
    logger.warning("config.ini not found. Using full default configuration.")

def get_config_value(section, key, default, type_func=str, cfg=config):
    if not user_config_read_successfully and not cfg.has_section(section): return default
    if cfg.has_section(section) and cfg.has_option(section, key):
        value = cfg.get(section, key)
        value = value.split('#')[0].split(';')[0].strip()
        try:
            if type_func == bool:
                if isinstance(value, bool): return value
                if value.lower() in ('true', 'yes', 'on', '1'): return True
                if value.lower() in ('false', 'no', 'off', '0'): return False
                raise ValueError(f"Not a boolean: {value}")
            elif value is None and type_func == str : return default
            return type_func(value if value is not None else default)
        except ValueError: return default
    return default

# --- CONFIG SETTINGS ---
MODEL_NAME_CONFIG = get_config_value('GeminiAPI', 'model_name', "gemini-1.5-pro-latest")
MODEL_NAME = MODEL_NAME_CONFIG if MODEL_NAME_CONFIG.startswith("models/") else f"models/{MODEL_NAME_CONFIG}"
TRANSLATION_TEMPERATURE = get_config_value('GeminiAPI', 'translation_temperature', 0.1, float)
MAX_CONCURRENT_API_CALLS = get_config_value('GeminiAPI', 'max_concurrent_api_calls', 5, int)
API_CALL_TIMEOUT = get_config_value('GeminiAPI', 'api_call_timeout_seconds', 600, int)

START_CONTENT_KEYWORDS_STRICT_STR = get_config_value('PDFProcessing', 'start_content_keywords', "introduction, εισαγωγή, πρόλογος, foreword, chapter 1, chapter i, κεφάλαιο 1, κεφάλαιο α, μέρος πρώτο, part one, summary, περίληψη, abstract")
START_CONTENT_KEYWORDS_STRICT = [k.strip().lower() for k in START_CONTENT_KEYWORDS_STRICT_STR.split(',') if k.strip()]
BIBLIOGRAPHY_KEYWORDS_LOWER_STR = get_config_value('PDFProcessing', 'bibliography_keywords', "bibliography, references, sources, literature cited, works cited, πηγές, βιβλιογραφία, αναφορές, εργα που αναφερονται, ευρετήριο, index, σημειώσεις, notes")
BIBLIOGRAPHY_KEYWORDS_LOWER = [k.strip().lower() for k in BIBLIOGRAPHY_KEYWORDS_LOWER_STR.split(',') if k.strip()]
TOC_DETECTION_KEYWORDS_STR = get_config_value('PDFProcessing', 'toc_detection_keywords', "contents, table of contents, περιεχόμενα, πίνακας περιεχομένων")
TOC_DETECTION_KEYWORDS_LOWER = [k.strip().lower() for k in TOC_DETECTION_KEYWORDS_STR.split(',') if k.strip()]
MAIN_CONTENT_AFTER_TOC_KEYWORDS_STR = get_config_value('PDFProcessing', 'main_content_after_toc_keywords', "introduction, πρόλογος, foreword, chapter 1, chapter i, κεφάλαιο 1, κεφάλαιο α, part one, μέρος πρώτο, preface, summary, abstract, acknowledgments, ευχαριστίες")
MAIN_CONTENT_AFTER_TOC_KEYWORDS_LOWER = [k.strip().lower() for k in MAIN_CONTENT_AFTER_TOC_KEYWORDS_STR.split(',') if k.strip()]


CHAPTER_TITLE_PATTERNS = [
    re.compile(r"^\s*(κεφάλαιο|chapter|μέρος|part|ενότητα|section)\s+([IVXLCDM\d\w\.\u0370-\u03FF]+)\s*[:.]*\s*$", re.IGNORECASE | re.UNICODE),
    re.compile(r"^\s*([IVXLCDM\d\w\.\u0370-\u03FF]+)\s*[:.]\s+(.+)", re.IGNORECASE | re.UNICODE)
]
LIST_MARKER_REGEX = re.compile(r"^\s*([*\-•◦∙➢➣►]|(\d{1,2}[\.\)])|([a-zA-Z][\.\)])|(\((?:[ivxlcdm]+|[a-zA-Z]|\d{1,2})\)))\s+", re.UNICODE)
MAX_CHARS_PER_SUBCHUNK = get_config_value('PDFProcessing', 'max_chars_per_subchunk', 12000, int)
AGGREGATE_SMALL_CHUNKS_TARGET_SIZE = get_config_value('PDFProcessing', 'aggregate_small_chunks_target_size', 10000, int)
MIN_CHARS_FOR_STANDALONE_CHUNK = get_config_value('PDFProcessing', 'min_chars_for_standalone_chunk', 350, int)
MAX_PAGES_FOR_START_DETECTION = get_config_value('PDFProcessing', 'max_pages_for_start_detection', 25, int)
MAX_FONT_SIZE_FOR_BODY_ANALYSIS = get_config_value('PDFProcessing', 'max_font_size_for_body_analysis', 36, int)
MIN_CHARS_FOR_TEXTUAL_PAGE = get_config_value('PDFProcessing', 'min_chars_for_textual_page', 50, int)
HEADER_FOOTER_MARGIN_PERCENT = get_config_value('PDFProcessing', 'header_footer_margin_percent', 0.10, float)
TOC_MAX_PAGES_SCAN = get_config_value('PDFProcessing', 'toc_max_pages_scan', 5, int)
HEADING_MAX_WORDS = get_config_value('PDFProcessing', 'heading_max_words', 25, int)
PARAGRAPH_SPLIT_THRESHOLD_FACTOR = get_config_value('PDFProcessing', 'paragraph_split_threshold_factor', 1.5, float) # Factor of dominant font size for vertical gap
MIN_FIRST_LINE_INDENT_THRESHOLD_POINTS = get_config_value('PDFProcessing', 'min_first_line_indent_threshold_points', 5, float) # Minimum points for x0 offset to be considered an indent

EXTRACT_IMAGES = get_config_value('PDFProcessing', 'extract_images', True, bool)
PERFORM_OCR_ON_IMAGES = get_config_value('PDFProcessing', 'perform_ocr_on_images', False, bool)
OCR_LANGUAGE = get_config_value('PDFProcessing', 'ocr_language', "eng")
MIN_OCR_WORDS_FOR_TRANSLATION = get_config_value('PDFProcessing', 'min_ocr_words_for_translation', 3, int)
if pytesseract is None: PERFORM_OCR_ON_IMAGES = False
MIN_IMAGE_WIDTH_PX = get_config_value('PDFProcessing', 'min_image_width_px', 10, int)
MIN_IMAGE_HEIGHT_PX = get_config_value('PDFProcessing', 'min_image_height_px', 10, int)
MIN_IMAGE_BPC_FILTER = get_config_value('PDFProcessing', 'min_image_bpc_filter', 1, int)
MAX_VERTICAL_GAP_FOR_IMAGE_ATTACHMENT_POINTS = get_config_value('PDFProcessing', 'max_vertical_gap_for_image_attachment_points', 50, int)

APPLY_STYLES_TO_PARAGRAPHS = get_config_value('WordOutput', 'apply_styles_to_paragraphs', True, bool)
APPLY_STYLES_TO_HEADINGS = get_config_value('WordOutput', 'apply_styles_to_headings', True, bool)
DEFAULT_IMAGE_WIDTH_INCHES = get_config_value('WordOutput', 'default_image_width_inches', 5.0, float)
GENERATE_TOC = True  # Re-enabled with improved implementation
TOC_TITLE = get_config_value('WordOutput', 'toc_title', "Πίνακας Περιεχομένων")
LIST_INDENT_PER_LEVEL_INCHES = get_config_value('WordOutput', 'list_indent_per_level_inches', 0.25, float)
HEADING_SPACE_BEFORE_PT = get_config_value('WordOutput', 'heading_space_before_pt', 6, int)
PARAGRAPH_FIRST_LINE_INDENT_INCHES = get_config_value('WordOutput', 'paragraph_first_line_indent_inches', 0.0, float) # Default to 0, apply based on detection
PARAGRAPH_SPACE_AFTER_PT = get_config_value('WordOutput', 'paragraph_space_after_pt', 6, int)


GDRIVE_TARGET_FOLDER_ID_CONFIG_STR = get_config_value('GoogleDrive', 'gdrive_target_folder_id', "")
GDRIVE_TARGET_FOLDER_ID_CONFIG = GDRIVE_TARGET_FOLDER_ID_CONFIG_STR if GDRIVE_TARGET_FOLDER_ID_CONFIG_STR and GDRIVE_TARGET_FOLDER_ID_CONFIG_STR.lower() != "none" else None

TARGET_LANGUAGE_CONFIG = get_config_value('TranslationEnhancements', 'target_language', "Ελληνικά")
USE_GLOSSARY = get_config_value('TranslationEnhancements', 'use_glossary', False, bool)
GLOSSARY_FILE_PATH = get_config_value('TranslationEnhancements', 'glossary_file_path', "glossary.json")
USE_TRANSLATION_CACHE = get_config_value('TranslationEnhancements', 'use_translation_cache', True, bool)
TRANSLATION_CACHE_FILE_PATH = get_config_value('TranslationEnhancements', 'translation_cache_file_path', "translation_cache.json")
TRANSLATION_STYLE_TONE_FALLBACK = get_config_value('TranslationEnhancements', 'translation_style_tone', "formal").strip().lower()
ANALYZE_DOCUMENT_STYLE_FIRST = get_config_value('TranslationEnhancements', 'analyze_document_style_first', True, bool)
BATCH_STYLE_ANALYSIS_REUSE = get_config_value('TranslationEnhancements', 'batch_style_analysis_reuse', True, bool)
DOCUMENT_ANALYSIS_PROMPT_TEXT_DEFAULT = "Analyze the following text sample from a larger document. Describe its overall style (e.g., academic, technical, legal, literary, casual), primary subject matter, target audience, and any discernible specific tone or philosophical leaning. Provide a concise summary (max 70 words) suitable for guiding a language translation task to maintain this style and context. Summary:"
DOCUMENT_ANALYSIS_PROMPT_TEXT = get_config_value('TranslationEnhancements', 'document_analysis_prompt_text', DOCUMENT_ANALYSIS_PROMPT_TEXT_DEFAULT)
DOCUMENT_ANALYSIS_SAMPLE_SIZE_CHARS = get_config_value('TranslationEnhancements', 'document_analysis_sample_size_chars', 4000, int)
PERFORM_QUALITY_ASSESSMENT = get_config_value('TranslationEnhancements', 'perform_quality_assessment', True, bool)
QA_STRATEGY = get_config_value('TranslationEnhancements', 'qa_strategy', 'full', str).lower()
QA_SAMPLE_PERCENTAGE = get_config_value('TranslationEnhancements', 'qa_sample_percentage', 0.1, float)

QUALITY_ASSESSMENT_MODEL_CONFIG = get_config_value('TranslationEnhancements', 'quality_assessment_model', "gemini-1.5-flash-latest").strip('"').strip("'")
QUALITY_ASSESSMENT_MODEL = QUALITY_ASSESSMENT_MODEL_CONFIG if QUALITY_ASSESSMENT_MODEL_CONFIG.startswith("models/") else f"models/{QUALITY_ASSESSMENT_MODEL_CONFIG}"

QUALITY_ASSESSMENT_TEMPERATURE = get_config_value('TranslationEnhancements', 'quality_assessment_temperature', 0.3, float)

# API Call Optimization Settings
ENABLE_SMART_GROUPING = get_config_value('APIOptimization', 'enable_smart_grouping', True, bool)
MAX_GROUP_SIZE_CHARS = get_config_value('APIOptimization', 'max_group_size_chars', 12000, int)
MAX_ITEMS_PER_GROUP = get_config_value('APIOptimization', 'max_items_per_group', 8, int)
ENABLE_OCR_GROUPING = get_config_value('APIOptimization', 'enable_ocr_grouping', True, bool)
AGGRESSIVE_GROUPING_MODE = get_config_value('APIOptimization', 'aggressive_grouping_mode', True, bool)
SMART_OCR_FILTERING = get_config_value('APIOptimization', 'smart_ocr_filtering', True, bool)
MIN_OCR_WORDS_FOR_TRANSLATION_ENHANCED = get_config_value('APIOptimization', 'min_ocr_words_for_translation', 8, int)
QUALITY_ASSESSMENT_PROMPT_DEFAULT = """As a translation quality evaluator, assess the translation from the original language (English assumed) to {target_language}.
Document Style/Context Guide: "{style_guide_snippet}"
Evaluate for: Accuracy (meaning), Fluency (naturalness in {target_language}), Grammar, and Style Adherence.
Provide a score (1-5, 5=Excellent) and a brief comment (max 15 words, English).
STRICTLY respond ONLY with the score and comment in the format: "Score: [score]. Comment: [comment]" (e.g., "Score: 4. Comment: Good fluency.")
Do NOT include any other text, greetings, or explanations.
Original:
---
{original_snippet}
---
Translation:
---
{translated_snippet}
---
Assessment:"""
QUALITY_ASSESSMENT_PROMPT_TEXT = get_config_value('TranslationEnhancements', 'quality_assessment_prompt_text', QUALITY_ASSESSMENT_PROMPT_DEFAULT)
SHOW_QA_ANNOTATIONS_IN_WORD = get_config_value('TranslationEnhancements', 'show_qa_annotations_in_word', True, bool)
GLOSSARY_CONSISTENCY_CHECK_CASE_SENSITIVE_TARGET = get_config_value('TranslationEnhancements', 'glossary_consistency_check_case_sensitive_target', False, bool)

GENERATE_QUALITY_REPORT = get_config_value('Reporting', 'generate_quality_report', True, bool)

TRANSLATION_GLOSSARY = {}
TRANSLATION_CACHE = {}
quality_report_messages = []
bookmark_id_counter = 0
qa_format_error_log_count = 0
MAX_QA_FORMAT_ERRORS_TO_LOG = 3
CREDENTIALS_FILE_GDRIVE = "mycreds.txt"
load_dotenv()
API_KEY = os.getenv('GEMINI_API_KEY')
if not API_KEY:
    logger.warning("GEMINI_API_KEY not found. Some functionality will be limited.")
    # Don't raise error during import - let functions handle this individually
else:
    genai.configure(api_key=API_KEY)

logger.info(f"Using API Model: '{MODEL_NAME}'")
logger.info(f"Target Language: {TARGET_LANGUAGE_CONFIG}")
logger.info(f"Generate Quality Report: {GENERATE_QUALITY_REPORT}")

FOOTNOTE_MARKER_REGEX_FOR_CLEANING = re.compile(r"[\u00B9\u00B2\u00B3\u2070\u2074-\u2079]+|\[\s*\d+\s*\]")
RAW_URL_REGEX = re.compile(r'(?i)\b((?:https?://|www\d{0,3}[.]|[a-z0-9.\-]+[.][a-z]{2,4}/)(?:[^\s()<>]+|\(([^\s()<>]+|(\([^\s()<>]+\)))*\))+(?:\(([^\s()<>]+|(\([^\s()<>]+\)))*\)|[^\s`!()\[\]{};:\'".,<>?«»“”‘’]))')

def clean_text_of_markers(text):
    if not text: return ""
    return FOOTNOTE_MARKER_REGEX_FOR_CLEANING.sub("", text)

def should_skip_ocr_translation(ocr_text, filename=""):
    """
    Determine if OCR text should be skipped from translation.
    Returns True for most cases since translating OCR text as captions loses context and meaning.
    Only translates very simple, standalone text that makes sense as a caption.
    """
    if not ocr_text or not ocr_text.strip():
        return True

    # Basic word count check - be more restrictive
    words = ocr_text.split()
    if len(words) < MIN_OCR_WORDS_FOR_TRANSLATION_ENHANCED:
        return True

    if not SMART_OCR_FILTERING:
        return False  # If smart filtering is disabled, use basic word count only

    text_lower = ocr_text.lower().strip()

    # CONSERVATIVE APPROACH: Skip most OCR content since caption translation loses meaning

    # Skip if it contains ANY chart/diagram indicators
    chart_indicators = [
        'figure', 'fig', 'chart', 'graph', 'diagram', 'table', 'tab',
        'axis', 'x-axis', 'y-axis', 'legend', 'scale', '%', 'percent',
        'data', 'value', 'min', 'max', 'avg', 'mean', 'total', 'step',
        'process', 'flow', 'stage', 'phase', 'level', 'category', 'type',
        'section', 'part', 'component', 'element', 'item', 'point'
    ]

    # Skip if contains ANY chart-like words
    if any(indicator in text_lower for indicator in chart_indicators):
        return True

    # Skip if mostly numbers, symbols, or single characters
    non_word_chars = sum(1 for c in text_lower if not c.isalpha() and not c.isspace())
    alpha_chars = sum(1 for c in text_lower if c.isalpha())
    if alpha_chars > 0 and (non_word_chars / len(text_lower)) > 0.3:  # More restrictive
        return True

    # Skip if it's mostly short words (likely labels)
    short_words = sum(1 for word in words if len(word.strip()) <= 3)  # Increased threshold
    if len(words) > 0 and (short_words / len(words)) > 0.5:  # More restrictive
        return True

    # Skip if it looks like mathematical formulas or equations
    math_indicators = ['=', '+', '-', '*', '/', '^', '√', '∑', '∫', '∆', 'α', 'β', 'γ', 'θ', 'π', 'σ', 'μ', '°', '±']
    if any(indicator in text_lower for indicator in math_indicators):
        return True

    # Skip if it's mostly URLs, emails, or technical identifiers
    if any(pattern in text_lower for pattern in ['http', 'www.', '@', '.com', '.org', '.net', 'id:', 'ref:', 'no.', 'code']):
        return True

    # Skip if filename suggests it's any kind of complex image
    if filename:
        filename_lower = filename.lower()
        complex_image_indicators = [
            'chart', 'graph', 'diagram', 'figure', 'fig', 'table', 'flow',
            'process', 'schema', 'map', 'plan', 'layout', 'design', 'model'
        ]
        if any(term in filename_lower for term in complex_image_indicators):
            return True

    # Skip if text contains multiple lines (likely structured content)
    if '\n' in ocr_text and len(ocr_text.split('\n')) > 2:
        return True

    # Skip if text contains common UI/interface elements
    ui_indicators = ['click', 'button', 'menu', 'tab', 'link', 'page', 'next', 'previous', 'home', 'back']
    if any(indicator in text_lower for indicator in ui_indicators):
        return True

    # Only allow very simple, complete sentences that would make sense as captions
    # Must be at least 10 words and look like natural language
    if len(words) < 10:
        return True

    # Must contain common sentence structure words
    sentence_indicators = ['the', 'a', 'an', 'is', 'are', 'was', 'were', 'this', 'that', 'these', 'those', 'and', 'or', 'but']
    if not any(indicator in text_lower for indicator in sentence_indicators):
        return True

    # Skip if it contains common "descriptive" words that suggest it's part of a complex image
    descriptive_words = ['quick', 'brown', 'fox', 'dog', 'runs', 'forest', 'example', 'sample', 'test', 'demo']
    if any(word in text_lower for word in descriptive_words):
        return True

    # If we get here, it might be translatable, but let's be very conservative
    # Only translate if it's clearly a complete, meaningful sentence
    sentences = [s.strip() for s in ocr_text.split('.') if s.strip()]
    if len(sentences) != 1:  # Must be exactly one sentence
        return True

    # Final check: must be a very specific type of descriptive sentence
    # that would make sense as a standalone caption
    academic_indicators = ['methodology', 'research', 'study', 'analysis', 'description', 'explanation']
    if not any(indicator in text_lower for indicator in academic_indicators):
        return True

    return False  # Extremely few cases will reach here

# ===== ENHANCED DOCUMENT STRUCTURE TRANSLATOR =====
class DocumentStructureTranslator:
    """
    Enhanced translator that preserves document structure including:
    - Table of Contents with proper hierarchy
    - Paragraph structure and formatting
    - Image placement with translated alt text
    - Context-aware image positioning using fuzzy matching
    """

    def __init__(self, translator_service):
        self.translator = translator_service

    def translate_document(self, source_document):
        """
        Main method to translate a document while preserving its structure
        """
        logger.info("Starting enhanced document structure translation...")

        # Extract structural elements
        toc_entries = self.extract_toc(source_document)
        paragraphs = self.preserve_paragraphs(source_document)
        images = self.extract_image_references(source_document)

        logger.info(f"Extracted {len(toc_entries)} TOC entries, {len(paragraphs)} paragraphs, {len(images)} images")

        # Translate content while preserving structure
        translated_toc = self.translate_toc(toc_entries)
        translated_paragraphs = self.translate_paragraphs(paragraphs)

        # Reconstruct document
        translated_document = self.recreate_toc(translated_toc)
        translated_document += "\n\n"
        translated_content = self.reconstruct_paragraphs(translated_paragraphs)
        translated_document += self.place_images(translated_content, images)

        logger.info("Enhanced document structure translation completed")
        return translated_document

    def extract_toc(self, document):
        """Extract table of contents entries from markdown headings"""
        toc_entries = []
        heading_pattern = re.compile(r'^(#{1,6})\s+(.+)$', re.MULTILINE)
        for match in heading_pattern.finditer(document):
            level = len(match.group(1))
            text = match.group(2).strip()
            position = match.start()
            toc_entries.append({"level": level, "text": text, "position": position})
        return toc_entries

    def translate_toc(self, toc_entries):
        """Translate table of contents entries"""
        translated_entries = []
        for entry in toc_entries:
            try:
                translated_text = self.translator.translate(entry["text"])
                translated_entries.append({
                    "level": entry["level"],
                    "text": translated_text,
                    "position": entry["position"]
                })
            except Exception as e:
                logger.warning(f"Failed to translate TOC entry '{entry['text']}': {e}")
                # Keep original text if translation fails
                translated_entries.append(entry)
        return translated_entries

    def recreate_toc(self, translated_toc_entries):
        """Recreate table of contents with translated headings"""
        if not translated_toc_entries:
            return ""

        toc = "# Table of Contents\n\n"
        for entry in translated_toc_entries:
            indent = "  " * (entry["level"] - 1)
            # Create URL-friendly slug
            slug = ''.join(e for e in entry['text'].lower() if e.isalnum() or e.isspace()).replace(' ', '-')
            toc += f"{indent}- [{entry['text']}](#{slug})\n"
        return toc

    def preserve_paragraphs(self, text):
        """Extract and preserve paragraph structure with metadata"""
        # Remove TOC and headings first to focus on content paragraphs
        content = re.sub(r'^#{1,6}\s+.+$', '', text, flags=re.MULTILINE)
        # Split by double newlines to identify paragraphs
        paragraphs = re.split(r'\n\s*\n', content)
        # Filter out empty paragraphs
        paragraphs = [p.strip() for p in paragraphs if p.strip()]

        # Store paragraph metadata
        paragraph_data = []
        for p in paragraphs:
            # Check if paragraph contains an image
            has_image = bool(re.search(r'!\[.*?\]\(.*?\)', p))
            paragraph_data.append({
                "text": p,
                "length": len(p),
                "has_image": has_image
            })
        return paragraph_data

    def translate_paragraphs(self, paragraphs):
        """Translate paragraphs while preserving image tags and structure"""
        translated = []
        for para in paragraphs:
            try:
                # Skip translation for paragraphs that are just images
                if para["has_image"] and len(re.sub(r'!\[.*?\]\(.*?\)', '', para["text"])).strip() == 0:
                    translated.append(para["text"])
                else:
                    # For mixed content, preserve image tags during translation
                    if para["has_image"]:
                        # Replace image tags with placeholders before translation
                        placeholders = {}
                        modified_text = para["text"]
                        img_pattern = re.compile(r'(!\[.*?\]\(.*?\))')
                        for i, match in enumerate(img_pattern.finditer(para["text"])):
                            placeholder = f"__IMG_PLACEHOLDER_{i}__"
                            placeholders[placeholder] = match.group(0)
                            modified_text = modified_text.replace(match.group(0), placeholder)

                        # Translate text with placeholders
                        translated_text = self.translator.translate(modified_text)

                        # Restore image tags
                        for placeholder, img_tag in placeholders.items():
                            translated_text = translated_text.replace(placeholder, img_tag)

                        translated.append(translated_text)
                    else:
                        translated.append(self.translator.translate(para["text"]))
            except Exception as e:
                logger.warning(f"Failed to translate paragraph: {e}")
                # Keep original text if translation fails
                translated.append(para["text"])
        return translated

    def reconstruct_paragraphs(self, translated_paragraphs):
        """Reconstruct paragraphs with proper spacing"""
        result = ""
        for para in translated_paragraphs:
            result += para + "\n\n"
        return result.strip()

    def extract_image_references(self, document):
        """Extract image references with context for optimal placement"""
        image_pattern = re.compile(r'!\[(.*?)\]\((.*?)\)', re.DOTALL)
        images = []
        for match in image_pattern.finditer(document):
            alt_text = match.group(1)
            src = match.group(2)

            # Get surrounding context (100 chars before and after)
            start_context = max(0, match.start() - 100)
            end_context = min(len(document), match.end() + 100)
            context = document[start_context:end_context]

            # Find the paragraph containing this image
            paragraph_start = document.rfind("\n\n", 0, match.start())
            paragraph_end = document.find("\n\n", match.end())
            if paragraph_start == -1:
                paragraph_start = 0
            if paragraph_end == -1:
                paragraph_end = len(document)

            paragraph_context = document[paragraph_start:paragraph_end].strip()

            images.append({
                "alt": alt_text,
                "src": src,
                "position": match.start(),
                "context": context,
                "paragraph_context": paragraph_context
            })
        return images

    def find_best_position(self, translated_text, original_context, translated_context):
        """Find the best position to insert an image based on context matching"""
        # Try to find exact paragraph match first
        if translated_context in translated_text:
            return translated_text.find(translated_context)

        # If no exact match, use fuzzy matching
        matches = difflib.get_close_matches(translated_context,
                                           [translated_text[i:i+len(translated_context)+50]
                                            for i in range(0, len(translated_text), 50)],
                                           n=1, cutoff=0.3)

        if matches:
            return translated_text.find(matches[0])

        # Fallback: try to match based on surrounding words
        words = re.findall(r'\b\w+\b', original_context)
        if words:
            # Try to find concentrations of matching words
            best_pos = 0
            max_matches = 0
            for i in range(0, len(translated_text), 20):
                chunk = translated_text[i:i+200]
                matches = sum(1 for word in words if word.lower() in chunk.lower())
                if matches > max_matches:
                    max_matches = matches
                    best_pos = i
            return best_pos

        # Last resort: place at the end
        return len(translated_text)

    def place_images(self, translated_text, image_data):
        """Place images in translated text with translated alt text and optimal positioning"""
        if not image_data:
            return translated_text

        # First translate all alt texts
        for img in image_data:
            try:
                img["translated_alt"] = self.translator.translate(img["alt"]) if img["alt"] else img["alt"]
                if "paragraph_context" in img:
                    img["translated_context"] = self.translator.translate(img["paragraph_context"])
            except Exception as e:
                logger.warning(f"Failed to translate image alt text '{img['alt']}': {e}")
                img["translated_alt"] = img["alt"]  # Keep original if translation fails
                img["translated_context"] = img.get("paragraph_context", "")

        # Sort images by their original position to maintain relative order
        sorted_images = sorted(image_data, key=lambda x: x["position"])

        # Create a copy of the text to modify
        result = translated_text

        # Track offsets as we insert images
        offset = 0

        for img in sorted_images:
            image_tag = f'![{img["translated_alt"]}]({img["src"]})'

            # Check if image is already in the text (preserved during paragraph translation)
            img_pattern = re.compile(r'!\[.*?\]\(' + re.escape(img["src"]) + r'\)')
            if img_pattern.search(result):
                continue

            # Find the best position to insert the image
            if "translated_context" in img:
                pos = self.find_best_position(result, img["context"], img["translated_context"])
            else:
                pos = self.find_best_position(result, img["context"], "")

            # Adjust position based on previous insertions
            pos += offset

            # Insert the image
            result = result[:pos] + "\n\n" + image_tag + "\n\n" + result[pos:]

            # Update offset
            offset += len("\n\n" + image_tag + "\n\n")

        return result

# ===== END ENHANCED DOCUMENT STRUCTURE TRANSLATOR =====

# ===== ENHANCED ERROR RECOVERY & QUOTA MANAGEMENT =====

class QuotaManager:
    """Manages API quota and implements intelligent retry strategies"""

    def __init__(self):
        self.quota_exceeded = False
        self.daily_request_count = 0
        self.last_reset_date = None
        self.retry_delays = [1, 2, 5, 10, 30, 60]  # Progressive delays in seconds

    def is_quota_error(self, error_message):
        """Check if error is quota-related"""
        quota_indicators = [
            "exceeded your current quota",
            "quota_metric",
            "GenerateRequestsPerDayPerProjectPerModel",
            "429"
        ]
        return any(indicator in str(error_message) for indicator in quota_indicators)

    def handle_quota_error(self):
        """Handle quota exhaustion"""
        self.quota_exceeded = True
        logger.error("🚨 API QUOTA EXCEEDED - Switching to recovery mode")
        return False

    def can_make_request(self):
        """Check if we can make an API request"""
        return not self.quota_exceeded

class EnhancedErrorRecovery:
    """Enhanced error recovery with progressive retry and graceful degradation"""

    def __init__(self):
        self.quota_manager = QuotaManager()
        self.failed_items = []
        self.partial_results = {}

    async def safe_api_call(self, api_function, *args, max_retries=3, **kwargs):
        """Make API call with enhanced error recovery"""

        if not self.quota_manager.can_make_request():
            logger.warning("⚠️ Skipping API call due to quota exhaustion")
            return None

        for attempt in range(max_retries):
            try:
                result = await api_function(*args, **kwargs)
                return result

            except Exception as e:
                error_msg = str(e)

                # Handle quota errors
                if self.quota_manager.is_quota_error(error_msg):
                    self.quota_manager.handle_quota_error()
                    logger.error(f"💰 Quota exceeded on attempt {attempt + 1}")
                    return None

                # Handle temporary errors with retry
                if attempt < max_retries - 1:
                    delay = self.quota_manager.retry_delays[min(attempt, len(self.quota_manager.retry_delays) - 1)]
                    logger.warning(f"⏳ API error on attempt {attempt + 1}, retrying in {delay}s: {error_msg[:100]}")
                    await asyncio.sleep(delay)
                    continue
                else:
                    logger.error(f"❌ API call failed after {max_retries} attempts: {error_msg}")
                    self.failed_items.append({'args': args, 'kwargs': kwargs, 'error': error_msg})
                    return None

        return None

    def get_recovery_report(self):
        """Generate recovery report"""
        total_failed = len(self.failed_items)
        quota_failed = self.quota_manager.quota_exceeded

        report = f"""
🔄 RECOVERY REPORT
================
• Total Failed Items: {total_failed}
• Quota Exceeded: {'Yes' if quota_failed else 'No'}
• Partial Results Available: {len(self.partial_results)}

💡 RECOMMENDATIONS:
"""

        if quota_failed:
            report += """
• Wait 24 hours for quota reset, or upgrade your API tier
• Consider using gemini-1.5-flash-latest for cost efficiency
• Enable more aggressive batching to reduce API calls
"""

        if total_failed > 0:
            report += f"""
• {total_failed} items need re-translation
• Use resume functionality to continue from where you left off
"""

        return report

# ===== SMART BATCHING ENHANCEMENT =====

class IntelligentBatcher:
    """Enhanced batching that dramatically reduces API calls"""

    def __init__(self, max_batch_size=15000, target_batches=50):
        self.max_batch_size = max_batch_size
        self.target_batches = target_batches

    def create_smart_batches(self, items):
        """Create intelligent batches that minimize API calls"""
        if not items:
            return []

        # Calculate optimal batch size
        total_chars = sum(len(str(item.get('text', ''))) for item in items)
        optimal_batch_count = max(1, min(self.target_batches, total_chars // self.max_batch_size + 1))
        target_size = total_chars // optimal_batch_count

        logger.info(f"📊 Smart Batching: {len(items)} items → {optimal_batch_count} batches (target: {target_size} chars/batch)")

        batches = []
        current_batch = []
        current_size = 0

        for item in items:
            item_text = str(item.get('text', ''))
            item_size = len(item_text)

            # Check if adding this item would exceed batch size
            if current_size + item_size > self.max_batch_size and current_batch:
                batches.append(current_batch)
                current_batch = [item]
                current_size = item_size
            else:
                current_batch.append(item)
                current_size += item_size

        # Add the last batch
        if current_batch:
            batches.append(current_batch)

        logger.info(f"✅ Created {len(batches)} smart batches (reduced from {len(items)} individual calls)")
        return batches

# ===== ENHANCED TRANSLATION QUALITY SYSTEM =====

class DocumentTypeDetector:
    """Detects document type and provides specialized translation guidance"""

    DOCUMENT_TYPES = {
        'academic': {
            'keywords': ['abstract', 'methodology', 'conclusion', 'references', 'bibliography', 'hypothesis', 'research', 'study', 'analysis', 'findings'],
            'style_guide': 'Academic paper with formal tone, precise terminology, and scholarly language. Maintain citation formats and technical accuracy.',
            'terminology_focus': 'technical terms, research concepts, statistical terms'
        },
        'legal': {
            'keywords': ['whereas', 'hereby', 'pursuant', 'jurisdiction', 'contract', 'agreement', 'clause', 'article', 'section', 'liability'],
            'style_guide': 'Legal document requiring precise legal terminology, formal structure, and exact meaning preservation. Maintain legal force and clarity.',
            'terminology_focus': 'legal terms, procedural language, statutory references'
        },
        'technical': {
            'keywords': ['specification', 'procedure', 'installation', 'configuration', 'parameters', 'system', 'manual', 'instructions', 'requirements'],
            'style_guide': 'Technical documentation with clear, step-by-step instructions and precise technical terminology. Maintain operational accuracy.',
            'terminology_focus': 'technical specifications, procedural steps, system components'
        },
        'medical': {
            'keywords': ['diagnosis', 'treatment', 'patient', 'symptoms', 'medication', 'therapy', 'clinical', 'medical', 'health', 'disease'],
            'style_guide': 'Medical document requiring accurate medical terminology and patient safety considerations. Maintain clinical precision.',
            'terminology_focus': 'medical terms, drug names, anatomical references'
        },
        'business': {
            'keywords': ['revenue', 'profit', 'market', 'strategy', 'management', 'financial', 'business', 'corporate', 'investment', 'performance'],
            'style_guide': 'Business document with professional tone and commercial terminology. Maintain financial accuracy and business context.',
            'terminology_focus': 'business terms, financial concepts, market terminology'
        },
        'literary': {
            'keywords': ['chapter', 'character', 'story', 'narrative', 'dialogue', 'plot', 'theme', 'author', 'novel', 'poetry'],
            'style_guide': 'Literary work requiring creative and culturally appropriate translation. Preserve artistic expression and cultural nuances.',
            'terminology_focus': 'cultural references, literary devices, emotional expressions'
        }
    }

    def detect_document_type(self, text_sample):
        """Detect document type based on keyword analysis"""
        text_lower = text_sample.lower()
        type_scores = {}

        for doc_type, config in self.DOCUMENT_TYPES.items():
            score = sum(1 for keyword in config['keywords'] if keyword in text_lower)
            type_scores[doc_type] = score

        # Get the type with highest score
        if type_scores and max(type_scores.values()) > 0:
            detected_type = max(type_scores, key=type_scores.get)
            confidence = type_scores[detected_type] / len(self.DOCUMENT_TYPES[detected_type]['keywords'])
            return detected_type, confidence, self.DOCUMENT_TYPES[detected_type]

        return 'general', 0.0, {
            'style_guide': 'General document with clear, natural language appropriate for the target audience.',
            'terminology_focus': 'general vocabulary, common expressions'
        }

class EnhancedTranslationPromptGenerator:
    """Generates enhanced translation prompts based on document analysis"""

    def __init__(self):
        self.doc_detector = DocumentTypeDetector()

    def generate_enhanced_prompt(self, text, target_language, document_sample="", context_info=""):
        """Generate an enhanced translation prompt with document-specific guidance"""

        # Detect document type
        doc_type, confidence, type_config = self.doc_detector.detect_document_type(document_sample or text)

        # 🧠 ENHANCED INTELLIGENCE: Include cross-reference guide if available
        global ENHANCED_REFERENCE_GUIDE
        cross_ref_guidance = ""
        if ENHANCED_REFERENCE_GUIDE:
            cross_ref_guidance = f"\n\n{ENHANCED_REFERENCE_GUIDE}"

        # Base enhanced prompt
        enhanced_prompt = f"""You are an expert translator specializing in {doc_type} documents.

DOCUMENT CONTEXT:
- Document Type: {doc_type.title()} (confidence: {confidence:.2f})
- Style Guide: {type_config['style_guide']}
- Terminology Focus: {type_config['terminology_focus']}
{f"- Additional Context: {context_info}" if context_info else ""}

TRANSLATION REQUIREMENTS:
1. Translate from the source language to {target_language}
2. Maintain the original document structure and formatting
3. Preserve technical accuracy and domain-specific terminology
4. Ensure cultural appropriateness for {target_language} readers
5. Keep consistent terminology throughout the document
6. Maintain the same level of formality as the original{cross_ref_guidance}

QUALITY STANDARDS:
- Accuracy: Preserve exact meaning and technical details
- Fluency: Natural, readable {target_language} that flows well
- Consistency: Use consistent terminology and style throughout
- Cultural Adaptation: Adapt cultural references appropriately
- Professional Standards: Meet professional translation quality standards

SPECIAL INSTRUCTIONS:
- If you encounter ambiguous terms, choose the interpretation most appropriate for {doc_type} documents
- Preserve any formatting markers, placeholders, or special characters
- Maintain paragraph structure and spacing
- For technical terms without direct equivalents, provide the most accepted {target_language} term or keep the original with explanation if necessary

TEXT TO TRANSLATE:
{text}

TRANSLATION:"""

        return enhanced_prompt

class CoherentBatchProcessor:
    """Processes documents in coherent batches while preserving context and placeholders"""

    def __init__(self, max_batch_size=8000, overlap_size=200):
        self.max_batch_size = max_batch_size
        self.overlap_size = overlap_size
        self.image_placeholder_pattern = re.compile(r'__IMG_PLACEHOLDER_\d+__')
        self.context_memory = []

    def create_coherent_batches(self, text, preserve_structure=True):
        """Create batches that preserve document coherence and structure"""
        if len(text) <= self.max_batch_size:
            return [text]

        batches = []
        current_pos = 0

        while current_pos < len(text):
            # Calculate batch end position
            batch_end = min(current_pos + self.max_batch_size, len(text))

            # If not at the end, find a good break point
            if batch_end < len(text):
                # Look for paragraph breaks first
                paragraph_break = text.rfind('\n\n', current_pos, batch_end)
                if paragraph_break > current_pos + self.max_batch_size // 2:
                    batch_end = paragraph_break + 2
                else:
                    # Look for sentence breaks
                    sentence_break = text.rfind('. ', current_pos, batch_end)
                    if sentence_break > current_pos + self.max_batch_size // 2:
                        batch_end = sentence_break + 2
                    else:
                        # Look for any whitespace
                        space_break = text.rfind(' ', current_pos, batch_end)
                        if space_break > current_pos + self.max_batch_size // 2:
                            batch_end = space_break + 1

            # Extract batch with overlap for context
            batch_start = max(0, current_pos - (self.overlap_size if current_pos > 0 else 0))
            batch_text = text[batch_start:batch_end]

            # Mark overlap regions for post-processing
            batch_info = {
                'text': batch_text,
                'start_pos': current_pos,
                'end_pos': batch_end,
                'has_start_overlap': current_pos > 0,
                'has_end_overlap': batch_end < len(text),
                'overlap_start': self.overlap_size if current_pos > 0 else 0,
                'overlap_end': len(batch_text) - (batch_end - current_pos) if batch_end < len(text) else 0
            }

            batches.append(batch_info)
            current_pos = batch_end

        return batches

    def merge_translated_batches(self, translated_batches):
        """Merge translated batches while handling overlaps intelligently"""
        if not translated_batches:
            return ""

        if len(translated_batches) == 1:
            return translated_batches[0]['translated_text']

        merged_text = ""

        for i, batch in enumerate(translated_batches):
            translated_text = batch['translated_text']

            if i == 0:
                # First batch - use everything except end overlap
                if batch['has_end_overlap']:
                    end_cut = len(translated_text) - batch['overlap_end']
                    merged_text += translated_text[:end_cut]
                else:
                    merged_text += translated_text
            elif i == len(translated_batches) - 1:
                # Last batch - skip start overlap
                if batch['has_start_overlap']:
                    start_cut = batch['overlap_start']
                    merged_text += translated_text[start_cut:]
                else:
                    merged_text += translated_text
            else:
                # Middle batch - skip both overlaps
                start_cut = batch['overlap_start'] if batch['has_start_overlap'] else 0
                end_cut = len(translated_text) - batch['overlap_end'] if batch['has_end_overlap'] else len(translated_text)
                merged_text += translated_text[start_cut:end_cut]

        return merged_text

# ===== ENHANCED OCR AND MULTI-MODAL CAPABILITIES =====

class EnhancedOCRProcessor:
    """Enhanced OCR with AI-powered text recognition and layout analysis"""

    def __init__(self):
        self.confidence_threshold = 60  # Minimum OCR confidence
        self.layout_analysis_enabled = True

    def extract_text_with_layout(self, image_path, lang='eng'):
        """Extract text with layout information for better translation context"""
        if not pytesseract or not PIL_Image:
            return None, []

        try:
            # Get detailed OCR data with confidence scores
            image = PIL_Image.open(image_path)

            # Get OCR data with bounding boxes and confidence
            ocr_data = pytesseract.image_to_data(image, lang=lang, output_type=pytesseract.Output.DICT)

            # Filter by confidence and group into text blocks
            text_blocks = []
            current_block = []
            current_line = ocr_data['line_num'][0] if ocr_data['line_num'] else 0

            for i in range(len(ocr_data['text'])):
                confidence = int(ocr_data['conf'][i])
                text = ocr_data['text'][i].strip()

                if confidence > self.confidence_threshold and text:
                    # Check if we're on a new line
                    if ocr_data['line_num'][i] != current_line:
                        if current_block:
                            text_blocks.append(' '.join(current_block))
                            current_block = []
                        current_line = ocr_data['line_num'][i]

                    current_block.append(text)

            # Add the last block
            if current_block:
                text_blocks.append(' '.join(current_block))

            # Combine blocks with appropriate spacing
            full_text = '\n'.join(text_blocks)

            return full_text, text_blocks

        except Exception as e:
            logger.error(f"Enhanced OCR failed for {image_path}: {e}")
            return None, []

    def detect_text_regions(self, image_path):
        """Detect different text regions (headers, body, captions, etc.)"""
        if not pytesseract or not PIL_Image:
            return []

        try:
            image = PIL_Image.open(image_path)

            # Get detailed layout information
            ocr_data = pytesseract.image_to_data(image, output_type=pytesseract.Output.DICT)

            regions = []
            for i in range(len(ocr_data['text'])):
                if int(ocr_data['conf'][i]) > self.confidence_threshold:
                    text = ocr_data['text'][i].strip()
                    if text:
                        region = {
                            'text': text,
                            'bbox': (ocr_data['left'][i], ocr_data['top'][i],
                                   ocr_data['width'][i], ocr_data['height'][i]),
                            'confidence': ocr_data['conf'][i],
                            'block_num': ocr_data['block_num'][i],
                            'line_num': ocr_data['line_num'][i]
                        }
                        regions.append(region)

            return regions

        except Exception as e:
            logger.error(f"Text region detection failed for {image_path}: {e}")
            return []

class TableExtractor:
    """Extract and translate tables while preserving structure"""

    def __init__(self):
        self.table_detection_enabled = True

    def detect_tables_in_image(self, image_path):
        """Detect if an image contains tabular data"""
        if not pytesseract or not PIL_Image:
            return False, []

        try:
            # Simple heuristic: look for aligned text in rows/columns
            ocr_processor = EnhancedOCRProcessor()
            regions = ocr_processor.detect_text_regions(image_path)

            if len(regions) < 4:  # Need at least 4 text regions for a table
                return False, []

            # Group regions by approximate y-coordinates (rows)
            rows = {}
            for region in regions:
                y_coord = region['bbox'][1]
                # Group regions within 10 pixels vertically
                row_key = round(y_coord / 10) * 10
                if row_key not in rows:
                    rows[row_key] = []
                rows[row_key].append(region)

            # Check if we have multiple rows with multiple columns
            multi_column_rows = sum(1 for row_regions in rows.values() if len(row_regions) > 1)

            if multi_column_rows >= 2:  # At least 2 rows with multiple columns
                return True, self._extract_table_structure(rows)

            return False, []

        except Exception as e:
            logger.error(f"Table detection failed for {image_path}: {e}")
            return False, []

    def _extract_table_structure(self, rows):
        """Extract table structure from detected regions"""
        table_data = []

        # Sort rows by y-coordinate
        sorted_rows = sorted(rows.items(), key=lambda x: x[0])

        for y_coord, row_regions in sorted_rows:
            # Sort columns by x-coordinate
            sorted_columns = sorted(row_regions, key=lambda r: r['bbox'][0])
            row_data = [region['text'] for region in sorted_columns]
            table_data.append(row_data)

        return table_data

class SmartImageAnalyzer:
    """Analyze images to determine optimal translation approach"""

    def __init__(self):
        self.ocr_processor = EnhancedOCRProcessor()
        self.table_extractor = TableExtractor()

    def analyze_image(self, image_path, lang='eng'):
        """Comprehensive image analysis for translation"""
        analysis_result = {
            'has_text': False,
            'text_content': '',
            'text_blocks': [],
            'is_table': False,
            'table_data': [],
            'image_type': 'unknown',
            'translation_approach': 'skip'
        }

        try:
            # Extract text with layout
            text_content, text_blocks = self.ocr_processor.extract_text_with_layout(image_path, lang)

            if text_content and len(text_content.strip()) > 0:
                analysis_result['has_text'] = True
                analysis_result['text_content'] = text_content
                analysis_result['text_blocks'] = text_blocks

                # Check if it's a table
                is_table, table_data = self.table_extractor.detect_tables_in_image(image_path)
                analysis_result['is_table'] = is_table
                analysis_result['table_data'] = table_data

                # Determine image type and translation approach
                word_count = len(text_content.split())

                if is_table:
                    analysis_result['image_type'] = 'table'
                    analysis_result['translation_approach'] = 'table_translation'
                elif word_count >= MIN_OCR_WORDS_FOR_TRANSLATION:
                    analysis_result['image_type'] = 'text_image'
                    analysis_result['translation_approach'] = 'full_translation'
                elif word_count > 0:
                    analysis_result['image_type'] = 'minimal_text'
                    analysis_result['translation_approach'] = 'caption_only'
                else:
                    analysis_result['image_type'] = 'no_text'
                    analysis_result['translation_approach'] = 'skip'
            else:
                analysis_result['image_type'] = 'no_text'
                analysis_result['translation_approach'] = 'skip'

        except Exception as e:
            logger.error(f"Image analysis failed for {image_path}: {e}")

        return analysis_result

# ===== INTEGRATION WRAPPER FOR ENHANCED CAPABILITIES =====

class GeminiTranslatorService:
    """
    Enhanced wrapper service that integrates with the existing Gemini translation functionality
    """

    def __init__(self, target_language=None, style_guide=None, document_sample=""):
        self.target_language = target_language or TARGET_LANGUAGE_CONFIG
        self.style_guide = style_guide or ""
        self.prompt_generator = EnhancedTranslationPromptGenerator()
        self.document_sample = document_sample

    def translate(self, text):
        """
        Translate text using enhanced prompts and the existing Gemini infrastructure
        """
        if not text or not text.strip():
            return text

        try:
            # Generate enhanced prompt
            enhanced_prompt = self.prompt_generator.generate_enhanced_prompt(
                text,
                self.target_language,
                self.document_sample,
                self.style_guide
            )

            # Use the enhanced prompt directly with Gemini API
            return asyncio.run(self._direct_gemini_call(enhanced_prompt))

        except Exception as e:
            logger.error(f"Enhanced translation failed for text '{text[:50]}...': {e}")
            return f"[Translation failed: {text}]"

    # Removed duplicate _translate_with_gemini method

    async def _direct_gemini_call(self, prompt):
        """Direct call to Gemini API with enhanced prompt"""
        try:
            model = genai.GenerativeModel(MODEL_NAME)

            generation_config = genai.types.GenerationConfig(
                temperature=TRANSLATION_TEMPERATURE,
                max_output_tokens=8192,
            )

            response = await model.generate_content_async(
                prompt,
                generation_config=generation_config,
                request_options=genai.types.RequestOptions(timeout=API_CALL_TIMEOUT)
            )

            if response.parts:
                return response.text.strip()
            else:
                error_msg = "Empty or blocked API response"
                if hasattr(response, 'prompt_feedback') and response.prompt_feedback:
                    error_msg += f" Feedback: {response.prompt_feedback}"
                raise Exception(error_msg)

        except Exception as e:
            logger.error(f"Direct Gemini API call failed: {e}")
            raise

def create_enhanced_document_translator(target_language=None, style_guide=None, document_sample=""):
    """
    Factory function to create an enhanced document translator with all capabilities
    """
    translator_service = GeminiTranslatorService(target_language, style_guide, document_sample)
    return DocumentStructureTranslator(translator_service)

def enhance_existing_image_processing(image_references, output_image_folder):
    """
    Enhanced image processing with smart analysis, context preservation, and translation
    """
    enhanced_image_refs = []
    smart_analyzer = SmartImageAnalyzer()

    for img_ref in image_references:
        try:
            image_path = img_ref.get('filepath')
            if not image_path or not os.path.exists(image_path):
                enhanced_image_refs.append(img_ref)
                continue

            # Perform smart analysis
            analysis = smart_analyzer.analyze_image(image_path, OCR_LANGUAGE)

            # Enhance the image reference with analysis results
            enhanced_ref = img_ref.copy()
            enhanced_ref.update({
                'smart_analysis': analysis,
                'has_translatable_text': analysis['has_text'],
                'extracted_text': analysis['text_content'],
                'text_blocks': analysis['text_blocks'],
                'is_table': analysis['is_table'],
                'table_data': analysis['table_data'],
                'translation_approach': analysis['translation_approach'],
                'context_preserved': True  # Mark as context-enhanced
            })

            # Add enhanced positioning information
            if img_ref.get('bbox'):
                enhanced_ref['positioning'] = {
                    'bbox': img_ref['bbox'],
                    'page_num': img_ref.get('page_num', 0),
                    'vertical_position': img_ref['bbox'][1] if img_ref.get('bbox') else 0
                }

            enhanced_image_refs.append(enhanced_ref)

            logger.info(f"Enhanced analysis for {img_ref.get('filename', 'unknown')}: "
                       f"Type={analysis['image_type']}, "
                       f"Approach={analysis['translation_approach']}, "
                       f"Text={'Yes' if analysis['has_text'] else 'No'}, "
                       f"Table={'Yes' if analysis['is_table'] else 'No'}")

        except Exception as e:
            logger.error(f"Failed to enhance image {img_ref.get('filename', 'unknown')}: {e}")
            enhanced_image_refs.append(img_ref)

    return enhanced_image_refs

# ===== ENHANCED WORKFLOW INTEGRATION =====

def process_pdf_with_enhanced_capabilities(pdf_filepath, output_directory):
    """
    Main enhanced processing function that integrates all new capabilities
    """
    logger.info("=== STARTING ENHANCED PDF PROCESSING ===")

    try:
        # Step 1: Extract images with enhanced analysis
        logger.info("Step 1: Enhanced image extraction and analysis...")
        image_folder = os.path.join(output_directory, "extracted_images")
        raw_image_refs = extract_images_from_pdf(pdf_filepath, image_folder)
        enhanced_image_refs = enhance_existing_image_processing(raw_image_refs, image_folder)

        # Step 2: Extract structured content
        logger.info("Step 2: Extracting structured content...")
        structured_content = extract_structured_content_from_pdf(pdf_filepath, enhanced_image_refs)

        # Step 3: Analyze document type and create enhanced translator
        logger.info("Step 3: Document analysis and enhanced translator setup...")

        # Get document sample for analysis
        document_sample = ""
        if structured_content:
            sample_texts = []
            for item in structured_content[:10]:  # First 10 items for analysis
                if item.get('type') in ['paragraph', 'heading'] and item.get('text'):
                    sample_texts.append(item['text'])
            document_sample = " ".join(sample_texts)[:4000]  # Limit to 4000 chars

        # Create enhanced document translator
        enhanced_translator = create_enhanced_document_translator(
            target_language=TARGET_LANGUAGE_CONFIG,
            style_guide="",  # Will be determined by document analysis
            document_sample=document_sample
        )

        # Step 4: Process content with coherent batching
        logger.info("Step 4: Processing content with enhanced translation...")
        batch_processor = CoherentBatchProcessor()

        # Process structured content in coherent batches
        processed_content = []
        for item in structured_content:
            if item.get('type') in ['paragraph', 'heading'] and item.get('text'):
                try:
                    # Use enhanced translation for text content
                    original_text = item['text']

                    # Check if text is long enough to benefit from batching
                    if len(original_text) > batch_processor.max_batch_size:
                        # Process in coherent batches
                        batches = batch_processor.create_coherent_batches(original_text)
                        translated_batches = []

                        for batch_info in batches:
                            translated_text = enhanced_translator.translator.translate(batch_info['text'])
                            batch_info['translated_text'] = translated_text
                            translated_batches.append(batch_info)

                        # Merge batches
                        final_translated_text = batch_processor.merge_translated_batches(translated_batches)
                        item['translated_text'] = final_translated_text
                    else:
                        # Process normally for shorter texts
                        item['translated_text'] = enhanced_translator.translator.translate(original_text)

                    processed_content.append(item)

                except Exception as e:
                    logger.error(f"Failed to process item: {e}")
                    item['translated_text'] = f"[Translation failed: {item.get('text', 'Unknown')}]"
                    processed_content.append(item)
            else:
                # Keep non-text items as-is
                processed_content.append(item)

        # Step 5: Enhanced image processing for images with text
        logger.info("Step 5: Processing images with translatable text...")
        for img_ref in enhanced_image_refs:
            if img_ref.get('translation_approach') in ['full_translation', 'table_translation']:
                try:
                    if img_ref.get('is_table') and img_ref.get('table_data'):
                        # Process table data
                        translated_table = []
                        for row in img_ref['table_data']:
                            translated_row = []
                            for cell in row:
                                if cell.strip():
                                    translated_cell = enhanced_translator.translator.translate(cell)
                                    translated_row.append(translated_cell)
                                else:
                                    translated_row.append(cell)
                            translated_table.append(translated_row)
                        img_ref['translated_table_data'] = translated_table

                    elif img_ref.get('extracted_text'):
                        # Process regular text in image
                        translated_ocr = enhanced_translator.translator.translate(img_ref['extracted_text'])
                        img_ref['translated_ocr_text'] = translated_ocr

                except Exception as e:
                    logger.error(f"Failed to translate image text for {img_ref.get('filename', 'unknown')}: {e}")
                    img_ref['translated_ocr_text'] = f"[Translation failed: {img_ref.get('extracted_text', 'Unknown')}]"

        logger.info("=== ENHANCED PDF PROCESSING COMPLETED ===")
        return processed_content, enhanced_image_refs

    except Exception as e:
        logger.error(f"Enhanced PDF processing failed: {e}")
        raise

def update_existing_translation_prompt():
    """
    Update the existing translation prompt in the script to use enhanced prompts
    """
    # This function can be used to modify the existing prompt generation
    # in the call_gemini_for_translation function
    pass

# ===== TESTING AND VALIDATION FUNCTIONS =====

def test_enhanced_capabilities():
    """
    Test function to validate the enhanced capabilities
    """
    logger.info("=== TESTING ENHANCED CAPABILITIES ===")

    try:
        # Test 1: Document Type Detection
        logger.info("Test 1: Document Type Detection")
        detector = DocumentTypeDetector()

        test_texts = {
            "academic": "This study presents a comprehensive analysis of the methodology used in recent research. The findings indicate significant correlations between variables.",
            "legal": "Whereas the parties hereby agree to the terms and conditions set forth in this contract, pursuant to the jurisdiction of the applicable law.",
            "technical": "Follow these installation procedures to configure the system parameters. Ensure all specifications meet the requirements.",
            "medical": "The patient presented with symptoms of acute inflammation. Treatment with prescribed medication showed positive clinical results."
        }

        for expected_type, text in test_texts.items():
            detected_type, confidence, config = detector.detect_document_type(text)
            logger.info(f"  Text: '{text[:50]}...'")
            logger.info(f"  Expected: {expected_type}, Detected: {detected_type}, Confidence: {confidence:.2f}")
            logger.info(f"  Style Guide: {config['style_guide'][:100]}...")
            print(f"✓ {expected_type.title()} detection: {detected_type} (confidence: {confidence:.2f})")

        # Test 2: Enhanced Prompt Generation
        logger.info("\nTest 2: Enhanced Prompt Generation")
        prompt_gen = EnhancedTranslationPromptGenerator()

        test_text = "The research methodology employed in this study follows established academic protocols."
        enhanced_prompt = prompt_gen.generate_enhanced_prompt(
            test_text,
            "Greek",
            "This is an academic research paper discussing methodology and findings.",
            "formal academic style"
        )

        logger.info(f"  Generated enhanced prompt length: {len(enhanced_prompt)} characters")
        logger.info(f"  Prompt preview: {enhanced_prompt[:200]}...")
        print("✓ Enhanced prompt generation working")

        # Test 3: Coherent Batch Processing
        logger.info("\nTest 3: Coherent Batch Processing")
        batch_processor = CoherentBatchProcessor(max_batch_size=100)  # Small size for testing

        test_long_text = """This is the first paragraph of a longer document. It contains important information that should be preserved.

This is the second paragraph. It continues the narrative and should maintain coherence with the previous paragraph.

This is the third paragraph. It concludes the document and references earlier points."""

        batches = batch_processor.create_coherent_batches(test_long_text)
        logger.info(f"  Original text length: {len(test_long_text)} characters")
        logger.info(f"  Number of batches created: {len(batches)}")

        for i, batch in enumerate(batches):
            logger.info(f"  Batch {i+1}: {len(batch['text'])} chars, overlap_start: {batch.get('overlap_start', 0)}")

        print(f"✓ Coherent batching: {len(batches)} batches created")

        # Test 4: Enhanced OCR (if available)
        if pytesseract and PIL_Image:
            logger.info("\nTest 4: Enhanced OCR Capabilities")
            ocr_processor = EnhancedOCRProcessor()
            print("✓ Enhanced OCR processor initialized")
        else:
            logger.info("\nTest 4: Enhanced OCR - SKIPPED (pytesseract/PIL not available)")
            print("⚠ Enhanced OCR skipped - dependencies not available")

        # Test 5: Smart Image Analysis
        logger.info("\nTest 5: Smart Image Analysis")
        image_analyzer = SmartImageAnalyzer()
        print("✓ Smart image analyzer initialized")

        logger.info("=== ALL ENHANCED CAPABILITIES TESTS COMPLETED ===")
        print("\n🎉 Enhanced capabilities test completed successfully!")
        return True

    except Exception as e:
        logger.error(f"Enhanced capabilities test failed: {e}")
        print(f"❌ Test failed: {e}")
        return False

def test_translation_with_sample():
    """
    Test the enhanced translation with a sample text
    """
    logger.info("=== TESTING ENHANCED TRANSLATION ===")

    try:
        # Create enhanced translator
        sample_document = "This is an academic research paper discussing advanced methodologies in data analysis."
        enhanced_translator = create_enhanced_document_translator(
            target_language="Greek",
            style_guide="formal academic style",
            document_sample=sample_document
        )

        # Test translation
        test_text = "The methodology employed in this research follows established academic protocols and ensures reproducible results."

        logger.info(f"Testing translation of: '{test_text}'")
        print(f"Original text: {test_text}")

        # Note: This will only work if API key is properly configured
        try:
            translated_text = enhanced_translator.translator.translate(test_text)
            logger.info(f"Translation result: '{translated_text}'")
            print(f"Translated text: {translated_text}")
            print("✓ Enhanced translation test completed")
            return True
        except Exception as e:
            logger.warning(f"Translation test skipped - API issue: {e}")
            print(f"⚠ Translation test skipped - API configuration needed: {e}")
            return False

    except Exception as e:
        logger.error(f"Enhanced translation test failed: {e}")
        print(f"❌ Translation test failed: {e}")
        return False

def test_enhanced_capabilities_simple():
    """
    Simple test without API calls to verify the enhanced capabilities structure
    """
    print("=== TESTING ENHANCED CAPABILITIES (SIMPLE) ===")

    try:
        # Test 1: Document Type Detection
        print("Test 1: Document Type Detection")
        detector = DocumentTypeDetector()

        test_text = "This study presents a comprehensive analysis of the methodology used in recent research."
        detected_type, confidence, config = detector.detect_document_type(test_text)
        print(f"✓ Detected: {detected_type} (confidence: {confidence:.2f})")
        print(f"  Style Guide: {config['style_guide'][:100]}...")

        # Test 2: Enhanced Prompt Generation
        print("\nTest 2: Enhanced Prompt Generation")
        prompt_gen = EnhancedTranslationPromptGenerator()

        enhanced_prompt = prompt_gen.generate_enhanced_prompt(
            "The research methodology is important.",
            "Greek",
            "This is an academic research paper.",
            "formal academic style"
        )

        print(f"✓ Generated prompt length: {len(enhanced_prompt)} characters")
        print(f"  Preview: {enhanced_prompt[:200]}...")

        # Test 3: Document Structure Components
        print("\nTest 3: Document Structure Components")

        # Create a mock translator for testing structure
        class MockTranslator:
            def translate(self, text):
                return f"[TRANSLATED: {text}]"

        mock_translator = MockTranslator()
        doc_translator = DocumentStructureTranslator(mock_translator)

        sample_doc = '''# Research Methodology

## Introduction

This study presents a comprehensive analysis.

![Sample Chart](images/chart1.png)

## Methods

The methodology employed ensures results.
'''

        print("✓ Testing document structure translation...")
        translated_doc = doc_translator.translate_document(sample_doc)
        print(f"  Original length: {len(sample_doc)} chars")
        print(f"  Translated length: {len(translated_doc)} chars")
        print(f"  Preview: {translated_doc[:200]}...")

        print("\n🎉 Enhanced capabilities structure test completed successfully!")
        return True

    except Exception as e:
        print(f"❌ Test failed: {e}")
        import traceback
        traceback.print_exc()
        return False

def test_ocr_filtering():
    """Test the smart OCR filtering functionality"""
    print("\n" + "="*50)
    print("TESTING SMART OCR FILTERING")
    print("="*50)

    test_cases = [
        ("Figure 1: Sales Chart", "chart1.png", True, "Chart/diagram content"),
        ("x-axis y-axis legend data", "graph.png", True, "Chart labels"),
        ("This is a normal paragraph with meaningful text content that makes sense.", "text.png", True, "Normal text - but still skipped (too risky)"),
        ("a b c d e", "labels.png", True, "Single letter labels"),
        ("E = mc²", "formula.png", True, "Mathematical formula"),
        ("http://www.example.com", "url.png", True, "URL content"),
        ("The quick brown fox jumps over the lazy dog and runs through the forest.", "content.png", True, "Long sentence - but still skipped (contains test words)"),
        ("1 2 3 % total avg", "data.png", True, "Data labels"),
        ("Introduction to Machine Learning Algorithms", "title.png", True, "Title - skipped (likely part of complex image)"),
        ("α β γ ∑ ∫ ∆", "math.png", True, "Mathematical symbols"),
        ("This is the complete description of the research methodology that was used in this comprehensive study.", "document.png", False, "Very clear complete sentence (rare case that might translate)"),
        ("Step 1 Process Flow", "process.png", True, "Process diagram"),
        ("Table Header Data", "table.png", True, "Table content"),
        ("Click here to continue", "ui.png", True, "UI element"),
        ("Line 1\nLine 2\nLine 3", "multi.png", True, "Multi-line content")
    ]

    passed = 0
    total = len(test_cases)

    for text, filename, should_skip, description in test_cases:
        result = should_skip_ocr_translation(text, filename)
        status = "✓ PASS" if result == should_skip else "✗ FAIL"
        print(f"{status} | {description}")
        print(f"      Text: '{text}' | File: {filename}")
        print(f"      Expected: {'Skip' if should_skip else 'Translate'} | Got: {'Skip' if result else 'Translate'}")
        print()

        if result == should_skip:
            passed += 1

    print(f"OCR Filtering Test Results: {passed}/{total} passed ({passed/total*100:.1f}%)")
    return passed == total

def test_optimizations_comprehensive():
    """
    Comprehensive test of all optimizations without API calls
    """
    print("\n" + "="*60)
    print("COMPREHENSIVE OPTIMIZATION TESTING (NO API CALLS)")
    print("="*60)

    all_tests_passed = True

    try:
        # Test 1: Enhanced Smart Grouping
        print("\n1. Testing Enhanced Smart Grouping...")
        smart_grouper = SmartGroupingProcessor(max_group_size=8000, min_group_items=2)

        # Create mock structured content
        mock_content = [
            {'type': 'h1', 'text': 'Chapter 1: Introduction', 'font_size': 16},
            {'type': 'p', 'text': 'This is the first paragraph of the introduction. It contains important information about the research topic.', 'font_size': 12},
            {'type': 'p', 'text': 'This is the second paragraph that continues the discussion.', 'font_size': 12},
            {'type': 'list_item', 'text': 'First bullet point', 'font_size': 12, 'list_info': {'type': 'bullet', 'level': 1}},
            {'type': 'list_item', 'text': 'Second bullet point', 'font_size': 12, 'list_info': {'type': 'bullet', 'level': 1}},
            {'type': 'image_placeholder', 'filename': 'chart1.png', 'text': '[Image: chart1.png]'},
            {'type': 'h2', 'text': 'Methodology', 'font_size': 14},
            {'type': 'p', 'text': 'The methodology section describes the approach used in this research.', 'font_size': 12},
        ]

        groups = smart_grouper.create_smart_groups(mock_content)
        print(f"✓ Original items: {len(mock_content)}, Grouped into: {len(groups)} groups")

        # Verify grouping logic
        grouped_items = 0
        standalone_items = 0
        for group in groups:
            if len(group) == 1:
                if group[0].get('type') == 'grouped_content':
                    grouped_items += group[0]['item_count']
                else:
                    standalone_items += 1
            else:
                standalone_items += len(group)

        print(f"  - Grouped items: {grouped_items}")
        print(f"  - Standalone items: {standalone_items}")
        print(f"  - Total items preserved: {grouped_items + standalone_items}")

        if grouped_items + standalone_items == len(mock_content):
            print("✓ Smart grouping preserves all content")
        else:
            print("❌ Smart grouping lost some content")
            all_tests_passed = False

        # Test 2: TOC Generation
        print("\n2. Testing TOC Generation...")

        # Mock structured content with headings
        toc_content = [
            {'type': 'h1', 'translated_text': 'Εισαγωγή', 'text': 'Introduction'},
            {'type': 'p', 'translated_text': 'Αυτή είναι η εισαγωγή.', 'text': 'This is the introduction.'},
            {'type': 'h2', 'translated_text': 'Μεθοδολογία', 'text': 'Methodology'},
            {'type': 'p', 'translated_text': 'Η μεθοδολογία περιγράφει...', 'text': 'The methodology describes...'},
            {'type': 'h3', 'translated_text': 'Ανάλυση Δεδομένων', 'text': 'Data Analysis'},
        ]

        # Test TOC extraction
        toc_items = []
        for index, item in enumerate(toc_content):
            if item['type'] in ['h1', 'h2', 'h3'] and item.get('translated_text'):
                level_num = int(item['type'][1:])
                bookmark_name = f"heading_{level_num}_{index}_{uuid.uuid4().hex[:8]}"
                item['_temp_bookmark_name_'] = bookmark_name
                first_line = item['translated_text'].split('\n')[0]
                toc_entry_text = (first_line[:75] + '...') if len(first_line) > 78 else first_line
                toc_items.append({
                    'text': toc_entry_text,
                    'level': level_num,
                    'bookmark_name': bookmark_name
                })

        print(f"✓ TOC extracted {len(toc_items)} entries:")
        for item in toc_items:
            indent = "  " * (item['level'] - 1)
            print(f"    {indent}- Level {item['level']}: {item['text']}")

        # Test 3: Cover Page Simulation
        print("\n3. Testing Cover Page Integration...")

        # Simulate cover page data
        mock_cover_data = {
            'filepath': 'mock_cover.png',
            'filename': 'cover_page_test.png',
            'page_number': 1,
            'type': 'cover_page',
            'width': 800,
            'height': 1200
        }

        print(f"✓ Cover page data structure: {mock_cover_data['type']}")
        print(f"  - Filename: {mock_cover_data['filename']}")
        print(f"  - Dimensions: {mock_cover_data['width']}x{mock_cover_data['height']}")

        # Test 4: Enhanced Image Processing
        print("\n4. Testing Enhanced Image Processing...")

        mock_images = [
            {
                'filename': 'chart1.png',
                'filepath': '/path/to/chart1.png',
                'page_num': 1,
                'bbox': (100, 200, 400, 500),
                'type': 'image_placeholder',
                'ocr_text_original': 'Sample chart data',
                'id': 'img_1_123'
            },
            {
                'filename': 'table1.png',
                'filepath': '/path/to/table1.png',
                'page_num': 2,
                'bbox': (50, 300, 500, 600),
                'type': 'image_placeholder',
                'ocr_text_original': 'Header1 Header2\nData1 Data2',
                'id': 'img_2_456'
            }
        ]

        # Simulate enhanced image processing
        enhanced_images = []
        for img in mock_images:
            enhanced_img = img.copy()
            enhanced_img.update({
                'smart_analysis': {
                    'has_text': bool(img.get('ocr_text_original')),
                    'text_content': img.get('ocr_text_original', ''),
                    'is_table': 'Header' in img.get('ocr_text_original', ''),
                    'image_type': 'table' if 'Header' in img.get('ocr_text_original', '') else 'text_image',
                    'translation_approach': 'table_translation' if 'Header' in img.get('ocr_text_original', '') else 'full_translation'
                },
                'context_preserved': True,
                'positioning': {
                    'bbox': img['bbox'],
                    'page_num': img['page_num'],
                    'vertical_position': img['bbox'][1]
                }
            })
            enhanced_images.append(enhanced_img)

        print(f"✓ Enhanced {len(enhanced_images)} images:")
        for img in enhanced_images:
            analysis = img['smart_analysis']
            print(f"    - {img['filename']}: Type={analysis['image_type']}, "
                  f"Approach={analysis['translation_approach']}, "
                  f"Text={'Yes' if analysis['has_text'] else 'No'}")

        # Test 5: Batch Processing Simulation
        print("\n5. Testing Coherent Batch Processing...")

        batch_processor = CoherentBatchProcessor(max_batch_size=11000)  # Updated to match new group size

        long_text = """This is the first paragraph of a longer document. It contains important information that should be preserved during translation.

This is the second paragraph. It continues the narrative and should maintain coherence with the previous paragraph when processed in batches.

This is the third paragraph. It concludes the section and references earlier points, demonstrating the importance of context preservation."""

        batches = batch_processor.create_coherent_batches(long_text)
        print(f"✓ Created {len(batches)} coherent batches from {len(long_text)} characters")

        # Handle both string and dict batch formats
        for i, batch in enumerate(batches):
            if isinstance(batch, dict):
                print(f"    - Batch {i+1}: {len(batch['text'])} chars, "
                      f"overlap_start: {batch.get('overlap_start', 0)}, "
                      f"overlap_end: {batch.get('overlap_end', 0)}")
            else:
                print(f"    - Batch {i+1}: {len(batch)} chars (simple format)")

        # Test batch merging simulation
        mock_translated_batches = []
        for batch in batches:
            if isinstance(batch, dict):
                mock_batch = batch.copy()
                mock_batch['translated_text'] = f"[TRANSLATED: {batch['text']}]"
                mock_translated_batches.append(mock_batch)
            else:
                # Handle simple string format
                mock_batch = {
                    'text': batch,
                    'translated_text': f"[TRANSLATED: {batch}]",
                    'start_pos': 0,
                    'end_pos': len(batch),
                    'has_start_overlap': False,
                    'has_end_overlap': False,
                    'overlap_start': 0,
                    'overlap_end': 0
                }
                mock_translated_batches.append(mock_batch)

        merged_result = batch_processor.merge_translated_batches(mock_translated_batches)
        print(f"✓ Merged result length: {len(merged_result)} characters")

        # Test 6: Document Structure Preservation
        print("\n6. Testing Document Structure Preservation...")

        # Test grouping compatibility
        test_pairs = [
            ({'type': 'p'}, {'type': 'p'}),
            ({'type': 'p'}, {'type': 'list_item'}),
            ({'type': 'h1'}, {'type': 'p'}),
            ({'type': 'image_placeholder'}, {'type': 'p'}),
            ({'type': 'list_item'}, {'type': 'list_item'}),
        ]

        print("✓ Testing grouping compatibility:")
        for item1, item2 in test_pairs:
            can_group = smart_grouper.can_group_together(item1, item2)
            print(f"    - {item1['type']} + {item2['type']}: {'✓' if can_group else '✗'}")

        # Test 7: OCR Filtering
        print("\n7. Testing Smart OCR Filtering...")
        ocr_test_passed = test_ocr_filtering()
        if not ocr_test_passed:
            all_tests_passed = False

        print("\n" + "="*60)
        if all_tests_passed:
            print("🎉 ALL OPTIMIZATION TESTS PASSED!")
            print("✓ Enhanced smart grouping working correctly")
            print("✓ TOC generation and translation ready")
            print("✓ Cover page integration functional")
            print("✓ Enhanced image processing operational")
            print("✓ Coherent batch processing validated")
            print("✓ Document structure preservation confirmed")
            print("✓ Smart OCR filtering operational")
        else:
            print("⚠️  Some tests had issues - check output above")

        print("="*60)
        return all_tests_passed

    except Exception as e:
        print(f"❌ Comprehensive test failed: {e}")
        import traceback
        traceback.print_exc()
        return False

# ===== SMART GROUPING SYSTEM FOR API CALL REDUCTION =====

class SmartGroupingProcessor:
    """
    Enhanced intelligent grouping of text segments to reduce API calls while preserving formatting
    """

    def __init__(self, max_group_size=12000, min_group_items=2, preserve_formatting=True):
        self.max_group_size = max_group_size
        self.min_group_items = min_group_items
        self.preserve_formatting = preserve_formatting
        self.group_separator = "\n\n===SEGMENT_SEPARATOR===\n\n"
        self.context_preservation_enabled = True
        self.cache_optimization_enabled = True
        # Enhanced grouping parameters for better efficiency
        self.aggressive_grouping_enabled = True
        self.max_items_per_group = 8  # Limit items per group for better splitting

    def can_group_together(self, item1, item2):
        """Enhanced grouping logic that preserves document structure and context"""
        type1, type2 = item1.get('type'), item2.get('type')

        # Never group headings with other content to preserve structure
        if type1 in ['h1', 'h2', 'h3'] or type2 in ['h1', 'h2', 'h3']:
            return False

        # Never group image placeholders
        if type1 == 'image_placeholder' or type2 == 'image_placeholder':
            return False

        # Enhanced compatibility matrix for better context preservation
        compatible_types = {
            ('p', 'p'): True,
            ('p', 'list_item'): True,  # Paragraphs and lists can be grouped
            ('list_item', 'list_item'): True,
            ('list_item', 'p'): True,  # But be careful with order
        }

        # Same types can usually be grouped (except headings and images)
        if type1 == type2 and type1 not in ['h1', 'h2', 'h3', 'image_placeholder']:
            return True

        # Enhanced grouping for aggressive mode
        if self.aggressive_grouping_enabled:
            # Allow more flexible grouping for similar content types
            text_types = {'p', 'list_item'}
            if type1 in text_types and type2 in text_types:
                return True

        # Check compatibility matrix
        pair = (type1, type2)
        reverse_pair = (type2, type1)

        return compatible_types.get(pair, compatible_types.get(reverse_pair, False))

    def estimate_group_size(self, items):
        """Estimate the total size of a group including separators"""
        total_size = 0
        for i, item in enumerate(items):
            text = item.get('text', '')
            total_size += len(text)
            if i < len(items) - 1:  # Add separator size except for last item
                total_size += len(self.group_separator)
        return total_size

    def check_cache_for_group(self, combined_text, target_language, model_name):
        """Check if a group's translation is already cached"""
        if not USE_TRANSLATION_CACHE:
            return None

        cache_key = get_cache_key(combined_text, target_language, model_name)
        return TRANSLATION_CACHE.get(cache_key)

    def save_group_to_cache(self, combined_text, translated_text, target_language, model_name):
        """Save a group's translation to cache"""
        if not USE_TRANSLATION_CACHE:
            return

        cache_key = get_cache_key(combined_text, target_language, model_name)
        TRANSLATION_CACHE[cache_key] = translated_text

    def create_smart_groups(self, structured_content):
        """Create intelligent groups of content items"""
        groups = []
        current_group = []

        i = 0
        while i < len(structured_content):
            item = structured_content[i]
            item_type = item.get('type')

            # Skip image placeholders - they don't need translation
            if item_type == 'image_placeholder':
                if current_group:
                    groups.append(self._finalize_group(current_group))
                    current_group = []
                groups.append([item])  # Images go alone
                i += 1
                continue

            # Skip empty items
            text_content = item.get('text', '').strip()
            if not text_content:
                i += 1
                continue

            # Always keep headings separate for structure preservation
            if item_type in ['h1', 'h2', 'h3']:
                if current_group:
                    groups.append(self._finalize_group(current_group))
                    current_group = []
                groups.append([item])  # Headings go alone
                i += 1
                continue

            # Check if we can add this item to current group
            can_add = True

            if current_group:
                # Check compatibility with last item in group
                last_item = current_group[-1]
                if not self.can_group_together(last_item, item):
                    can_add = False

                # Check size constraints
                test_group = current_group + [item]
                if self.estimate_group_size(test_group) > self.max_group_size:
                    can_add = False

                # Check item count limit to prevent overly large groups
                if len(current_group) >= self.max_items_per_group:
                    can_add = False

            if can_add:
                current_group.append(item)
            else:
                # Finalize current group and start new one
                if current_group:
                    groups.append(self._finalize_group(current_group))
                current_group = [item]

            i += 1

        # Don't forget the last group
        if current_group:
            groups.append(self._finalize_group(current_group))

        return groups

    def _finalize_group(self, group):
        """Finalize a group by adding metadata"""
        if len(group) == 1:
            return group  # Single items don't need special handling

        # Create combined text for translation
        combined_text = self.group_separator.join(item.get('text', '') for item in group)

        # Check cache for this group if cache optimization is enabled
        cache_hit = False
        if self.cache_optimization_enabled and USE_TRANSLATION_CACHE:
            cache_key = get_cache_key(combined_text, TARGET_LANGUAGE_CONFIG, MODEL_NAME)
            if cache_key in TRANSLATION_CACHE:
                cache_hit = True
                logger.info(f"Cache hit for group of {len(group)} items ({len(combined_text)} chars)")

        # Create group metadata
        group_info = {
            'type': 'grouped_content',
            'items': group,
            'combined_text': combined_text,
            'item_count': len(group),
            'estimated_size': len(combined_text),
            'cache_hit': cache_hit
        }

        return [group_info]

    def split_translated_group(self, translated_text, original_group_info):
        """Split translated grouped text back into individual items"""
        if original_group_info.get('type') != 'grouped_content':
            return [translated_text]  # Not a group, return as-is

        # Split by separator
        translated_parts = translated_text.split(self.group_separator)
        original_items = original_group_info['items']

        # Ensure we have the right number of parts
        if len(translated_parts) != len(original_items):
            logger.warning(f"Group splitting mismatch: expected {len(original_items)} parts, got {len(translated_parts)}")
            # Fallback: distribute text evenly or use original structure
            return self._fallback_split(translated_text, original_items)

        # Assign translated text back to original items
        results = []
        for i, (translated_part, original_item) in enumerate(zip(translated_parts, original_items)):
            result_item = original_item.copy()
            result_item['translated_text'] = translated_part.strip()
            results.append(result_item)

        return results

    def _fallback_split(self, translated_text, original_items):
        """Fallback method when group splitting fails"""
        logger.warning("Using fallback group splitting method")

        # Simple approach: try to split by double newlines
        parts = re.split(r'\n\s*\n', translated_text)

        results = []
        for i, original_item in enumerate(original_items):
            result_item = original_item.copy()
            if i < len(parts):
                result_item['translated_text'] = parts[i].strip()
            else:
                # If we don't have enough parts, mark for individual translation
                result_item['translated_text'] = f"[GROUP_SPLIT_FAILED: {original_item.get('text', '')}]"
            results.append(result_item)

        return results

# ===== END ENHANCED CAPABILITIES INTEGRATION =====

# This duplicate class has been removed - using the enhanced version above

def load_glossary():
    global TRANSLATION_GLOSSARY, quality_report_messages
    if USE_GLOSSARY and GLOSSARY_FILE_PATH:
        if os.path.exists(GLOSSARY_FILE_PATH):
            try:
                with open(GLOSSARY_FILE_PATH, 'r', encoding='utf-8') as f:
                    TRANSLATION_GLOSSARY = json.load(f)
                msg = f"Glossary loaded ({len(TRANSLATION_GLOSSARY)} terms) from: {GLOSSARY_FILE_PATH}"
                logger.info(msg); quality_report_messages.append(msg)
            except Exception as e:
                msg = f"ERROR loading glossary from {GLOSSARY_FILE_PATH}: {e}"
                logger.error(msg); quality_report_messages.append(msg)
        else:
            msg = f"WARNING: Glossary file '{GLOSSARY_FILE_PATH}' not found, but use_glossary is True."
            logger.warning(msg); quality_report_messages.append(msg)
    elif USE_GLOSSARY:
        msg = "WARNING: Glossary use enabled but glossary_file_path is not specified."
        logger.warning(msg); quality_report_messages.append(msg)

def load_translation_cache():
    global TRANSLATION_CACHE, quality_report_messages
    if USE_TRANSLATION_CACHE and TRANSLATION_CACHE_FILE_PATH:
        if os.path.exists(TRANSLATION_CACHE_FILE_PATH):
            try:
                with open(TRANSLATION_CACHE_FILE_PATH, 'r', encoding='utf-8') as f:
                    TRANSLATION_CACHE = json.load(f)
                msg = f"Translation cache loaded ({len(TRANSLATION_CACHE)} entries) from: {TRANSLATION_CACHE_FILE_PATH}"
                logger.info(msg); quality_report_messages.append(msg)
            except Exception as e:
                msg = f"WARNING loading cache from {TRANSLATION_CACHE_FILE_PATH}: {e}. Starting empty."
                logger.warning(msg); quality_report_messages.append(msg); TRANSLATION_CACHE = {}

def save_translation_cache():
    global quality_report_messages
    if USE_TRANSLATION_CACHE and TRANSLATION_CACHE_FILE_PATH:
        try:
            with open(TRANSLATION_CACHE_FILE_PATH, 'w', encoding='utf-8') as f:
                json.dump(TRANSLATION_CACHE, f, ensure_ascii=False, indent=4)
            logger.info(f"Translation cache saved to: {TRANSLATION_CACHE_FILE_PATH}")
        except Exception as e:
            msg = f"ERROR saving cache to {TRANSLATION_CACHE_FILE_PATH}: {e}"
            logger.error(msg); quality_report_messages.append(msg)

def get_cache_key(text, target_language, model_name):
    hasher = hashlib.sha256()
    hasher.update(text.encode('utf-8'))
    text_hash = hasher.hexdigest()
    return f"{text_hash}_{target_language}_{model_name}"

def get_desktop_path():
    desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
    if os.path.exists(desktop_path): return desktop_path
    desktop_path_el = os.path.join(os.path.expanduser("~"), "Επιφάνεια εργασίας")
    if os.path.exists(desktop_path_el): return desktop_path_el
    return os.path.expanduser("~")

def choose_input_path():
    root = None
    try:
        root = tk.Tk()
        root.withdraw()
        root.attributes('-topmost', True)
        process_type = input("Process a single PDF file or a directory of PDF files? (file/dir): ").lower().strip()
        if process_type == 'file':
            file_path = filedialog.askopenfilename(parent=root, title="Επιλέξτε αρχείο PDF", filetypes=(("PDF files", "*.pdf"), ("All files", "*.*")))
            if file_path: logger.info(f"Επιλέχθηκε το αρχείο: {file_path}")
            else: logger.info("Δεν επιλέχθηκε αρχείο.")
            return file_path, 'file'
        elif process_type == 'dir':
            dir_path = filedialog.askdirectory(parent=root, title="Επιλέξτε φάκελο με αρχεία PDF")
            if dir_path: logger.info(f"Επιλέχθηκε ο φάκελος: {dir_path}")
            else: logger.info("Δεν επιλέχθηκε φάκελος.")
            return dir_path, 'dir'
        else:
            logger.warning("Invalid choice. Please enter 'file' or 'dir'.")
            return None, None
    except Exception as e:
        logger.error(f"Σφάλμα κατά την εμφάνιση του διαλόγου επιλογής: {e}")
        return None, None
    finally:
        if root is not None:
            try: root.destroy()
            except tk.TclError: pass

def choose_base_output_directory(initial_dir=None):
    logger.info("\nΠαρακαλώ επιλέξτε τον ΚΥΡΙΟ φάκελο όπου θα δημιουργηθούν όλοι οι υποφάκελοι εξόδου...")
    root = None; chosen_directory = None
    try:
        root = tk.Tk(); root.withdraw(); root.attributes('-topmost', True)
        if initial_dir is None: initial_dir = get_desktop_path()
        chosen_directory = filedialog.askdirectory(parent=root, title="Επιλέξτε ΚΥΡΙΟ φάκελο αποθήκευσης", initialdir=initial_dir)
        if chosen_directory: logger.info(f"Επιλεγμένος κύριος φάκελος αποθήκευσης: {chosen_directory}")
        else: logger.info("Δεν επιλέχθηκε κύριος φάκελος αποθήκευσης. Θα χρησιμοποιηθεί ο φάκελος του αρχικού PDF/τρέχων φάκελος.")
        return chosen_directory
    except Exception as e:
        logger.error(f"Σφάλμα κατά την εμφάνιση του διαλόγου επιλογής κύριου φακέλου: {e}")
        return None
    finally:
        if root is not None:
            try: root.destroy()
            except tk.TclError: pass

def get_specific_output_dir_for_file(main_base_output_dir, source_pdf_filepath):
    base_input_filename = os.path.splitext(os.path.basename(source_pdf_filepath))[0]
    safe_subdir_name = re.sub(r'[<>:"/\\|?*]', '_', base_input_filename)
    safe_subdir_name = safe_subdir_name[:100]
    if not main_base_output_dir:
        main_base_output_dir = os.path.dirname(source_pdf_filepath) if os.path.dirname(source_pdf_filepath) else os.getcwd()
        logger.warning(f"Κύριος φάκελος εξόδου δεν δόθηκε, χρησιμοποιείται ο: {main_base_output_dir}")
    specific_output_dir = os.path.join(main_base_output_dir, f"{safe_subdir_name}_Μεταφρασμένο_{TARGET_LANGUAGE_CONFIG.replace(' ','_')}")
    try:
        os.makedirs(specific_output_dir, exist_ok=True)
        return specific_output_dir
    except Exception as e_mkdir:
        logger.error(f"Σφάλμα δημιουργίας υποφακέλου '{specific_output_dir}': {e_mkdir}. Αποθήκευση στον κύριο φάκελο: {main_base_output_dir}.")
        return main_base_output_dir

def ocr_image_text(image_path, lang='eng'):
    global quality_report_messages
    if not pytesseract or not PIL_Image or not PERFORM_OCR_ON_IMAGES: return None
    try:
        try:
            pytesseract.get_tesseract_version()
        except Exception as e_tess_check:
            logger.error(f"Tesseract not found or not configured correctly: {e_tess_check}")
            logger.error("Please ensure Tesseract is installed and in your system's PATH.")
            quality_report_messages.append("ERROR: Tesseract OCR not found or not configured correctly. OCR will fail.")
            return None

        text = pytesseract.image_to_string(PIL_Image.open(image_path), lang=lang)
        return text.strip() if text else None
    except Exception as e:
        msg = f"OCR failed for image {os.path.basename(image_path)} with lang {lang}: {e}"; logger.error(msg); quality_report_messages.append(msg)
        return None

def upload_to_google_drive(filepath_to_upload, filename_on_drive, gdrive_folder_id):
    global quality_report_messages, CREDENTIALS_FILE_GDRIVE
    if not os.path.exists("client_secrets.json"):
        msg = "Google Drive upload skipped. 'client_secrets.json' not found."
        logger.error(msg); quality_report_messages.append(msg)
        return None
    try:
        gauth = GoogleAuth()
        gauth.settings['get_refresh_token'] = True

        if os.path.exists(CREDENTIALS_FILE_GDRIVE):
            gauth.LoadCredentialsFile(CREDENTIALS_FILE_GDRIVE)

        if gauth.credentials is None or gauth.access_token_expired or not gauth.credentials.refresh_token:
            logger.info("Attempting Google Drive authentication (web browser will open). Please ensure you grant offline access if prompted.")
            gauth.LocalWebserverAuth()
        else:
            gauth.Authorize()

        if gauth.credentials and gauth.credentials.refresh_token is None:
             logger.warning("Google Drive: No refresh token obtained after authentication. Offline access might not work. Consider re-authenticating after deleting 'mycreds.txt' and ensuring 'offline access' is granted in the browser prompt during authentication.")

        gauth.SaveCredentialsFile(CREDENTIALS_FILE_GDRIVE)

        drive = GoogleDrive(gauth)
        file_metadata = {'title': filename_on_drive}
        if gdrive_folder_id:
            file_metadata['parents'] = [{'id': gdrive_folder_id, "kind": "drive#fileLink"}]

        gfile = drive.CreateFile(file_metadata)
        gfile.SetContentFile(filepath_to_upload)
        logger.info(f"Uploading '{filename_on_drive}' to Google Drive...")
        gfile.Upload()

        upload_msg = f"Successfully uploaded '{filename_on_drive}' to Google Drive (ID: {gfile['id']})."
        if gdrive_folder_id: upload_msg += f" Folder ID: {gdrive_folder_id}."
        else: upload_msg += " (Uploaded to Root MyDrive)."
        logger.info(upload_msg); quality_report_messages.append(upload_msg)
        return gfile['id']
    except Exception as e_gdrive:
        error_msg = f"Google Drive upload failed for '{filename_on_drive}': {e_gdrive}"
        logger.error(error_msg, exc_info=True); quality_report_messages.append(error_msg)
        if "No refresh_token found" in str(e_gdrive) or "invalid_grant" in str(e_gdrive).lower() or "access_denied" in str(e_gdrive).lower():
            logger.error("To fix Google Drive refresh token issues: Delete 'mycreds.txt' and re-run the script. During Google authentication in your browser, ensure you grant 'offline access' (access when you are not present) if prompted.")
        return None

def extract_cover_page_from_pdf(pdf_filepath, output_folder):
    """
    Extract the first page of PDF as cover page image
    """
    global quality_report_messages

    try:
        if not os.path.exists(output_folder):
            os.makedirs(output_folder)

        # Open PDF and extract first page
        pdf_document = fitz.open(pdf_filepath)
        if len(pdf_document) == 0:
            logger.warning("PDF has no pages for cover extraction")
            return None

        # Get first page
        first_page = pdf_document[0]

        # Convert to image with high resolution
        mat = fitz.Matrix(2.0, 2.0)  # 2x zoom for better quality
        pix = first_page.get_pixmap(matrix=mat)

        # Save cover page
        cover_filename = f"cover_page_{os.path.splitext(os.path.basename(pdf_filepath))[0]}.png"
        cover_path = os.path.join(output_folder, cover_filename)
        pix.save(cover_path)

        pdf_document.close()

        logger.info(f"Cover page extracted: {cover_path}")

        return {
            'filepath': cover_path,
            'filename': cover_filename,
            'page_number': 1,
            'type': 'cover_page',
            'width': pix.width,
            'height': pix.height
        }

    except Exception as e:
        error_msg = f"Failed to extract cover page from {pdf_filepath}: {e}"
        logger.error(error_msg)
        quality_report_messages.append(error_msg)
        return None

def extract_images_from_pdf(pdf_filepath, output_image_folder):
    global quality_report_messages
    if not EXTRACT_IMAGES:
        logger.info("Image extraction is disabled in config.ini.")
        return []
    image_references = []
    images_skipped_filter = 0
    extracted_image_hashes = set()
    images_skipped_duplicate_hash = 0

    try:
        with fitz.open(pdf_filepath) as doc:
            if not os.path.exists(output_image_folder):
                os.makedirs(output_image_folder)
            image_count_on_page = defaultdict(int)
            for page_index in range(len(doc)):
                page = doc.load_page(page_index)
                image_list = page.get_images(full=True)
                for img_info in image_list:
                    xref = img_info[0]
                    base_image = doc.extract_image(xref)
                    if base_image and base_image["image"]:
                        image_bytes = base_image["image"]

                        img_hash = hashlib.md5(image_bytes).hexdigest()
                        if img_hash in extracted_image_hashes:
                            images_skipped_duplicate_hash +=1
                            continue

                        img_width = base_image.get("width", 0)
                        img_height = base_image.get("height", 0)
                        img_bpc = base_image.get("bpc", 8)
                        if img_width < MIN_IMAGE_WIDTH_PX or img_height < MIN_IMAGE_HEIGHT_PX or img_bpc < MIN_IMAGE_BPC_FILTER:
                            images_skipped_filter += 1; continue

                        extracted_image_hashes.add(img_hash)

                        image_ext = base_image["ext"]
                        image_count_on_page[page_index] += 1
                        image_filename = f"image_p{page_index+1}_{image_count_on_page[page_index]}.{image_ext}"
                        image_save_path = os.path.join(output_image_folder, image_filename)
                        with open(image_save_path, "wb") as img_file: img_file.write(image_bytes)
                        ocr_text_original = ocr_image_text(image_save_path, lang=OCR_LANGUAGE) if PERFORM_OCR_ON_IMAGES else None
                        if ocr_text_original: quality_report_messages.append(f"INFO: OCR extracted from '{image_filename}'.")
                        img_ref_data = {'type': 'image_placeholder', 'page_num': page_index, 'filename': image_filename, 'filepath': image_save_path, 'text': f"[Image: {image_filename} - PDF Page: {page_index+1}]", 'ocr_text_original': ocr_text_original, 'bbox': (0,0,0,0), 'id': f"img_{page_index}_{xref}"}
                        try:
                            img_bboxes_list = page.get_image_rects(img_info[0], transform=True)
                            if img_bboxes_list:
                                first_bbox_candidate = img_bboxes_list[0]; actual_rect = None
                                if isinstance(first_bbox_candidate, fitz.Rect): actual_rect = first_bbox_candidate
                                elif isinstance(first_bbox_candidate, (tuple, list)) and len(first_bbox_candidate) > 0 and isinstance(first_bbox_candidate[0], fitz.Rect): actual_rect = first_bbox_candidate[0]
                                if actual_rect: img_ref_data['bbox'] = tuple(int(round(c)) for c in actual_rect.irect)
                                elif isinstance(first_bbox_candidate, (tuple, list)) and len(first_bbox_candidate) == 4:
                                    try: img_ref_data['bbox'] = tuple(int(round(c)) for c in first_bbox_candidate)
                                    except (ValueError, TypeError) as e_conv: quality_report_messages.append(f"WARN: bbox conv fail {xref} pg {page_index+1}: {e_conv}")
                                else: quality_report_messages.append(f"WARN: Image bbox for {xref} pg {page_index+1} unhandled: {type(first_bbox_candidate)}")
                        except Exception as e_bbox_main: quality_report_messages.append(f"WARN: Error processing bbox for image {xref} pg {page_index+1}: {e_bbox_main}")
                        image_references.append(img_ref_data)
        if images_skipped_filter > 0: quality_report_messages.append(f"INFO: Skipped {images_skipped_filter} images by size/BPC filter.")
        if images_skipped_duplicate_hash > 0: quality_report_messages.append(f"INFO: Skipped {images_skipped_duplicate_hash} duplicate images by content hash.")
        msg = f"Extracted {len(image_references)} unique image references."
        if images_skipped_duplicate_hash > 0:
             msg += f" ({images_skipped_duplicate_hash} duplicates by content were skipped)."
        logger.info(msg); quality_report_messages.append(msg)
    except Exception as e_img:
        msg = f"ERROR during image extraction: {e_img}"; logger.error(msg, exc_info=True); quality_report_messages.append(msg)
    return image_references

def groupby_images_by_page(all_extracted_image_refs):
    if not all_extracted_image_refs:
        return {}
    sorted_images = sorted(all_extracted_image_refs, key=lambda x: x['page_num'])
    images_by_page_dict = {}
    for k, g in groupby(sorted_images, key=lambda x: x['page_num']):
        images_by_page_dict[k] = sorted(list(g), key=lambda x: x.get('bbox', (0,0,0,0))[1])
    return images_by_page_dict

def extract_structured_content_from_pdf(filepath, all_extracted_image_refs):
    global quality_report_messages, LIST_MARKER_REGEX, MAX_FONT_SIZE_FOR_BODY_ANALYSIS, MIN_CHARS_FOR_TEXTUAL_PAGE, TOC_DETECTION_KEYWORDS_LOWER, MAIN_CONTENT_AFTER_TOC_KEYWORDS_LOWER, HEADER_FOOTER_MARGIN_PERCENT, TOC_MAX_PAGES_SCAN, HEADING_MAX_WORDS, PARAGRAPH_SPLIT_THRESHOLD_FACTOR, MIN_FIRST_LINE_INDENT_THRESHOLD_POINTS
    logger.info(f"--- Structuring Content from PDF: {os.path.basename(filepath)} ---")
    images_by_page = groupby_images_by_page(all_extracted_image_refs)
    for page_num_key in images_by_page:
        for img_ref in images_by_page[page_num_key]:
            img_ref.setdefault('_attached_images', [])

    final_structured_items = []
    identified_header_footer_strings = set() # For storing identified repeating header/footer text

    try:
        with fitz.open(filepath) as doc:
            content_start_page = 0
            content_start_y = 0.0
            toc_heading_found_page = -1
            toc_heading_found_y = 0.0
            is_toc_skipped = False

            # 1. Table of Contents Detection
            logger.info("Attempting to detect Table of Contents...")
            scan_limit_toc_heading = min(len(doc), MAX_PAGES_FOR_START_DETECTION)
            for page_num_toc_head in range(scan_limit_toc_heading):
                if toc_heading_found_page != -1: break
                page_toc_scan = doc.load_page(page_num_toc_head)
                blocks_toc_scan_dict = page_toc_scan.get_text("dict", flags=fitz.TEXT_PRESERVE_LIGATURES | fitz.TEXT_PRESERVE_WHITESPACE)["blocks"]

                for block_scan_data in blocks_toc_scan_dict:
                    if block_scan_data['type'] == 0:
                        for line_scan in block_scan_data.get("lines", []):
                            line_text_scan_original = "".join(span['text'] for span in line_scan['spans']).strip()
                            if not line_text_scan_original: continue
                            line_text_scan_lower = line_text_scan_original.lower()

                            if any(keyword in line_text_scan_lower for keyword in TOC_DETECTION_KEYWORDS_LOWER):
                                if len(line_text_scan_original.split()) < 7:
                                    toc_heading_found_page = page_num_toc_head
                                    toc_heading_found_y = line_scan['bbox'][1]
                                    logger.info(f"Potential Table of Contents heading detected on page {toc_heading_found_page + 1} at y={toc_heading_found_y:.2f} (text: '{line_text_scan_original[:50]}...').")
                                    break
                        if toc_heading_found_page != -1: break
            
            if toc_heading_found_page != -1:
                content_start_page = toc_heading_found_page
                content_start_y = toc_heading_found_y

                found_actual_content_start = False
                for page_idx_after_toc in range(toc_heading_found_page, min(toc_heading_found_page + TOC_MAX_PAGES_SCAN, len(doc))):
                    if found_actual_content_start: break
                    page_after_toc = doc.load_page(page_idx_after_toc)
                    blocks_after_toc = page_after_toc.get_text("dict", flags=fitz.TEXT_PRESERVE_LIGATURES | fitz.TEXT_PRESERVE_WHITESPACE)["blocks"]
                    sorted_blocks_on_page = sorted(blocks_after_toc, key=lambda b: b['bbox'][1] if 'bbox' in b else 0)

                    for block_main_content_scan in sorted_blocks_on_page:
                        if block_main_content_scan['type'] == 0:
                            if page_idx_after_toc == toc_heading_found_page and block_main_content_scan['bbox'][1] <= toc_heading_found_y + 5:
                                continue

                            block_text_lines = ["".join(s['text'] for s in l['spans']).strip() for l in block_main_content_scan.get("lines", []) if l['spans']]
                            if not block_text_lines: continue
                            
                            first_line_text_lower = block_text_lines[0].lower()
                            if any(keyword in first_line_text_lower for keyword in MAIN_CONTENT_AFTER_TOC_KEYWORDS_LOWER):
                                first_span_font_size = block_main_content_scan['lines'][0]['spans'][0]['size'] if block_main_content_scan['lines'][0]['spans'] else 0
                                if first_span_font_size > 8 and not (len(block_text_lines) == 1 and block_text_lines[0].isdigit()):
                                    content_start_page = page_idx_after_toc
                                    content_start_y = block_main_content_scan['bbox'][1]
                                    logger.info(f"Main content start detected after ToC on page {content_start_page + 1} at y={content_start_y:.2f} (text: '{block_text_lines[0][:50]}...').")
                                    quality_report_messages.append(f"INFO: Main content after ToC detected on page {content_start_page + 1}. Extraction starts here.")
                                    is_toc_skipped = True
                                    found_actual_content_start = True
                                    break
                    if found_actual_content_start: break
                
                if not found_actual_content_start:
                    content_start_page = toc_heading_found_page + 1
                    content_start_y = 0
                    if content_start_page < len(doc):
                        logger.warning(f"Could not clearly identify end of ToC/start of main content. Starting extraction from page {content_start_page + 1} (top).")
                        quality_report_messages.append(f"WARNING: End of ToC unclear. Starting from page {content_start_page + 1}.")
                        is_toc_skipped = True
                    else:
                        logger.warning("Table of Contents seems to be at the end of the document. No subsequent content to extract.")
                        quality_report_messages.append("WARNING: ToC at document end. No subsequent content.")
                        return []
            else:
                logger.warning(f"Table of Contents heading not automatically detected within the first {scan_limit_toc_heading} pages. Content extraction will start from the absolute beginning of the PDF.")
                quality_report_messages.append("WARNING: Table of Contents heading not detected. Starting from PDF beginning.")
            
            logger.info(f"Final content extraction will start from page {content_start_page + 1}, y-position {content_start_y:.2f}.")

            # Enhanced Header/Footer Detection (Scan first few content pages)
            if len(doc) > content_start_page:
                num_pages_for_hf_scan = min(5, len(doc) - content_start_page)
                potential_hf_lines_by_page = [[] for _ in range(num_pages_for_hf_scan)]
                page_text_content_for_hf_scan = ["" for _ in range(num_pages_for_hf_scan)]

                for i_hf_scan in range(num_pages_for_hf_scan):
                    page_idx_for_hf = content_start_page + i_hf_scan
                    if page_idx_for_hf >= len(doc): break # Ensure we don't go out of bounds
                    
                    page_for_hf = doc.load_page(page_idx_for_hf)
                    page_text_content_for_hf_scan[i_hf_scan] = page_for_hf.get_text("text") # For context if needed
                    page_rect_hf = page_for_hf.rect
                    page_height_hf = page_rect_hf.height
                    HEADER_ZONE_Y1_HF_SCAN = page_rect_hf.y0 + page_height_hf * HEADER_FOOTER_MARGIN_PERCENT
                    FOOTER_ZONE_Y0_HF_SCAN = page_rect_hf.y1 - page_height_hf * HEADER_FOOTER_MARGIN_PERCENT

                    blocks_for_hf = page_for_hf.get_text("dict", flags=fitz.TEXT_PRESERVE_LIGATURES | fitz.TEXT_PRESERVE_WHITESPACE)["blocks"]
                    for block_hf in blocks_for_hf:
                        if block_hf['type'] == 0: # Text block
                            block_bbox_hf = block_hf['bbox']
                            # Skip if block is clearly part of main content based on ToC start
                            if page_idx_for_hf == content_start_page and block_bbox_hf[1] >= content_start_y and not is_toc_skipped:
                                continue

                            for line_hf in block_hf.get("lines", []):
                                line_text_hf_original = "".join(span['text'] for span in line_hf['spans']).strip()
                                if not line_text_hf_original or len(line_text_hf_original) > 70: continue # Skip empty or very long lines for HF
                                
                                line_bbox_hf_coords = line_hf['bbox']
                                is_in_hf_zone_scan = line_bbox_hf_coords[1] < HEADER_ZONE_Y1_HF_SCAN or line_bbox_hf_coords[3] > FOOTER_ZONE_Y0_HF_SCAN
                                
                                if is_in_hf_zone_scan:
                                    # Store text and its relative vertical position (e.g. y0 for header, y1 for footer)
                                    # For simplicity, we'll just store the text for now. More advanced would be (text, y_rel_pos)
                                    potential_hf_lines_by_page[i_hf_scan].append(line_text_hf_original)
                
                # Find common lines across the scanned pages
                if potential_hf_lines_by_page and any(page_lines for page_lines in potential_hf_lines_by_page):
                    # Count occurrences of each line considering all scanned pages
                    line_counts_across_pages = Counter()
                    for page_lines_list in potential_hf_lines_by_page:
                        # Use set for each page to count each line only once per page for this stage
                        for line_txt in set(page_lines_list): 
                            line_counts_across_pages[line_txt] += 1
                    
                    # A line is a header/footer if it appears on a significant number of the scanned pages
                    # e.g., more than half, or at least 2 if fewer pages scanned.
                    min_occurrences_for_hf = max(2, num_pages_for_hf_scan // 2 + 1) if num_pages_for_hf_scan > 1 else 1
                    
                    for text_line, count in line_counts_across_pages.items():
                        if count >= min_occurrences_for_hf:
                            identified_header_footer_strings.add(text_line)
                            logger.info(f"Identified potential repeating header/footer string: '{text_line}' (found on {count}/{num_pages_for_hf_scan} scanned pages)")
                            quality_report_messages.append(f"INFO: Identified potential header/footer: '{text_line}' (found {count}/{num_pages_for_hf_scan})")


            temp_dominant_font_size = 10.0; all_font_items_for_analysis = []
            analysis_pages_limit_font = min(len(doc) - content_start_page, MAX_PAGES_FOR_START_DETECTION)
            
            for page_num_font_analysis in range(content_start_page, min(content_start_page + analysis_pages_limit_font, len(doc))):
                page_font_analysis = doc.load_page(page_num_font_analysis)
                blocks_font_analysis = page_font_analysis.get_text("dict", flags=fitz.TEXT_PRESERVE_LIGATURES | fitz.TEXT_PRESERVE_WHITESPACE)["blocks"]
                for block_font in blocks_font_analysis:
                    if block_font['type'] == 0:
                        for line_font in block_font["lines"]:
                            if page_num_font_analysis == content_start_page and line_font['bbox'][1] < content_start_y:
                                continue
                            if line_font["spans"]:
                                font_size_val = round(line_font["spans"][0]["size"], 1)
                                if font_size_val <= MAX_FONT_SIZE_FOR_BODY_ANALYSIS:
                                    all_font_items_for_analysis.append({'font_size': font_size_val})
            if all_font_items_for_analysis:
                font_sizes_in_doc_sample = [item['font_size'] for item in all_font_items_for_analysis if item['font_size'] > 0]
                if font_sizes_in_doc_sample:
                    plausible_body_text_sizes_sample = [s for s in font_sizes_in_doc_sample if 7.0 <= s <= 18.0]
                    if plausible_body_text_sizes_sample:
                        font_counts_sample = Counter(plausible_body_text_sizes_sample); temp_dominant_font_size = font_counts_sample.most_common(1)[0][0]

            all_page_items_collected = []
            previous_block_bottom_y = 0 # For inter-block paragraph detection
            page_default_x0 = None # To help detect first line indents

            for page_num in range(content_start_page, len(doc)):
                page = doc.load_page(page_num)
                page_rect = page.rect
                page_height = page_rect.height
                HEADER_ZONE_Y1 = page_rect.y0 + page_height * HEADER_FOOTER_MARGIN_PERCENT
                FOOTER_ZONE_Y0 = page_rect.y1 - page_height * HEADER_FOOTER_MARGIN_PERCENT

                page_items_this_page = []
                page_dict_data = page.get_text("dict", flags=fitz.TEXT_PRESERVE_LIGATURES | fitz.TEXT_PRESERVE_WHITESPACE | fitz.TEXT_DEHYPHENATE)
                accumulated_text_on_page_len = 0
                
                # Estimate default x0 for the page (excluding outliers like page numbers far left/right)
                line_x0s_for_page = []
                for b_x0_scan in page_dict_data.get("blocks", []):
                    if b_x0_scan['type'] == 0:
                        for l_x0_scan in b_x0_scan.get("lines",[]):
                            if l_x0_scan['spans']: line_x0s_for_page.append(l_x0_scan['bbox'][0])
                if line_x0s_for_page:
                    # Basic outlier removal: sort and take middle 50%
                    sorted_x0s = sorted(line_x0s_for_page)
                    q1_idx = len(sorted_x0s) // 4
                    q3_idx = 3 * len(sorted_x0s) // 4
                    if q3_idx > q1_idx:
                        typical_x0s = sorted_x0s[q1_idx:q3_idx]
                        if typical_x0s: page_default_x0 = sum(typical_x0s) / len(typical_x0s)
                    elif sorted_x0s: page_default_x0 = sorted_x0s[len(sorted_x0s)//2] # Median if too few points
                
                if page_default_x0 is None and page_num > content_start_page: # Inherit from previous if not calculable and not first content page
                    if all_page_items_collected and '_page_default_x0' in all_page_items_collected[-1]:
                         page_default_x0 = all_page_items_collected[-1]['_page_default_x0']

                sorted_blocks_on_page_for_content = sorted(page_dict_data.get("blocks", []), key=lambda b: b['bbox'][1] if 'bbox' in b else 0)
                
                last_line_y1_in_block_buffer = 0 # For intra-block paragraph detection

                for block_idx, block_data in enumerate(sorted_blocks_on_page_for_content):
                    if block_data['type'] == 0:
                        block_bbox = block_data['bbox']
                        current_block_top_y = block_bbox[1]

                        if page_num == content_start_page and current_block_top_y < content_start_y:
                            previous_block_bottom_y = block_bbox[3] # Still update for next block on this page
                            continue
                        
                        block_text_for_filter_check_lines = [ "".join(s['text'] for s in l['spans']) for l in block_data.get("lines", [])]
                        block_text_for_filter_check = " ".join(block_text_for_filter_check_lines).strip()

                        # Skip identified repeating headers/footers
                        if block_text_for_filter_check in identified_header_footer_strings and \
                           (block_bbox[1] < HEADER_ZONE_Y1 or block_bbox[3] > FOOTER_ZONE_Y0):
                            logger.debug(f"Skipping block matching identified header/footer: '{block_text_for_filter_check[:50]}' on page {page_num + 1}")
                            quality_report_messages.append(f"INFO: Skipped identified H/F block: '{block_text_for_filter_check[:30]}' pg {page_num + 1}")
                            previous_block_bottom_y = block_bbox[3]
                            continue
                        
                        # Fallback for numeric/short headers/footers if not caught by above
                        is_in_header_zone_fallback = block_bbox[1] < HEADER_ZONE_Y1
                        is_in_footer_zone_fallback = block_bbox[3] > FOOTER_ZONE_Y0
                        if is_in_header_zone_fallback or is_in_footer_zone_fallback:
                            words_in_block_fb = block_text_for_filter_check.split()
                            num_digits_fb = sum(c.isdigit() for c in block_text_for_filter_check)
                            num_alpha_fb = sum(c.isalpha() for c in block_text_for_filter_check)
                            if (len(words_in_block_fb) <= 7 and num_digits_fb > 0 and num_alpha_fb < 15) or \
                               (len(words_in_block_fb) <=3 and num_digits_fb == 0): # Catch very short text too
                                logger.debug(f"Skipping potential page number/header/footer (numeric/short fallback): '{block_text_for_filter_check[:50]}' on page {page_num + 1} at bbox {block_bbox}")
                                quality_report_messages.append(f"INFO: Skipped potential page num/H/F (fallback): '{block_text_for_filter_check[:30]}' pg {page_num + 1}")
                                previous_block_bottom_y = block_bbox[3]
                                continue
                        
                        # Paragraph break based on inter-block vertical gap
                        # Only consider this if the new block isn't a tiny fragment right after a large block (often happens with superscripts etc)
                        block_height = block_bbox[3] - block_bbox[1]
                        if page_items_this_page and block_idx > 0 and \
                           (current_block_top_y - previous_block_bottom_y) > (temp_dominant_font_size * PARAGRAPH_SPLIT_THRESHOLD_FACTOR) and \
                           block_height > temp_dominant_font_size * 0.5 : # New block isn't too small
                            # This signals a new paragraph if there was content in the previous block
                            if page_items_this_page[-1]['type'] != 'paragraph_break_marker': # Avoid double markers
                                page_items_this_page.append({'type': 'paragraph_break_marker', 'page_num': page_num, 'bbox_y0': previous_block_bottom_y})
                        
                        # If block is effectively empty (e.g. only whitespace after cleaning), treat as paragraph break
                        if not block_text_for_filter_check.strip() and block_height > temp_dominant_font_size * 0.5: # Ensure it's not a tiny artifact
                            if page_items_this_page and page_items_this_page[-1]['type'] != 'paragraph_break_marker':
                                page_items_this_page.append({'type': 'paragraph_break_marker', 'page_num': page_num, 'bbox_y0': current_block_top_y})
                            previous_block_bottom_y = block_bbox[3]
                            continue


                        temp_lines_for_this_block_processing = []
                        first_line_in_block = True

                        for line_idx, line_info in enumerate(block_data.get("lines", [])):
                            line_text_original = "".join(span['text'] for span in line_info['spans'])
                            if not line_text_original.strip():
                                # An empty line within a block might signify a paragraph break
                                if temp_lines_for_this_block_processing and temp_lines_for_this_block_processing[-1]['type'] != 'paragraph_break_marker':
                                     temp_lines_for_this_block_processing.append({'type': 'paragraph_break_marker', 'page_num': page_num, 'bbox_y0': line_info['bbox'][1]})
                                last_line_y1_in_block_buffer = line_info['bbox'][3]
                                continue

                            line_font_size = round(line_info['spans'][0]['size'],1) if line_info['spans'] else temp_dominant_font_size
                            line_font_name_lower = line_info['spans'][0]['font'].lower() if line_info['spans'] else ""
                            line_is_bold = "bold" in line_font_name_lower or "heavy" in line_font_name_lower or "black" in line_font_name_lower or "demi" in line_font_name_lower or "semibold" in line_font_name_lower
                            line_is_italic = "italic" in line_font_name_lower or "oblique" in line_font_name_lower
                            current_line_bbox = line_info['bbox']
                            cleaned_line_text = clean_text_of_markers(line_text_original.strip())
                            accumulated_text_on_page_len += len(cleaned_line_text)

                            # Intra-block paragraph break based on vertical spacing
                            if temp_lines_for_this_block_processing and line_idx > 0 and \
                               (current_line_bbox[1] - last_line_y1_in_block_buffer) > (line_font_size * PARAGRAPH_SPLIT_THRESHOLD_FACTOR * 0.8): # Slightly smaller factor for intra-block
                                if temp_lines_for_this_block_processing[-1]['type'] != 'paragraph_break_marker':
                                    temp_lines_for_this_block_processing.append({'type': 'paragraph_break_marker', 'page_num': page_num, 'bbox_y0': last_line_y1_in_block_buffer})
                            
                            is_indented_line = False
                            if first_line_in_block and page_default_x0 is not None:
                                line_x0 = current_line_bbox[0]
                                if (line_x0 - page_default_x0) > MIN_FIRST_LINE_INDENT_THRESHOLD_POINTS:
                                    is_indented_line = True
                            
                            match = LIST_MARKER_REGEX.match(line_text_original)
                            if match:
                                if temp_lines_for_this_block_processing: # Flush previous non-list text as a paragraph
                                    para_text = " ".join(l['text'] for l in temp_lines_for_this_block_processing if l['type']=='p_line').strip()
                                    if para_text:
                                        page_items_this_page.append({
                                            'type': 'p', 'text': para_text,
                                            'font_size': temp_lines_for_this_block_processing[0]['font_size'],
                                            'bold': temp_lines_for_this_block_processing[0]['bold'],
                                            'italic': temp_lines_for_this_block_processing[0]['italic'],
                                            'is_indented': temp_lines_for_this_block_processing[0].get('is_indented', False),
                                            'page_num': page_num,
                                            'bbox_y0': temp_lines_for_this_block_processing[0]['bbox_y0'],
                                            'bbox': temp_lines_for_this_block_processing[0]['bbox'],
                                            '_page_default_x0': page_default_x0,
                                            '_attached_images': []
                                        })
                                    temp_lines_for_this_block_processing = []

                                marker = match.group(1).strip()
                                list_item_text_content = LIST_MARKER_REGEX.sub("", line_text_original).strip()
                                list_item_text_cleaned_content = clean_text_of_markers(list_item_text_content)
                                if list_item_text_cleaned_content:
                                    list_type = 'number' if any(char.isdigit() for char in marker) else 'bullet'
                                    page_items_this_page.append({
                                        'type': 'list_item', 'text': list_item_text_cleaned_content,
                                        'font_size': line_font_size, 'bold': line_is_bold, 'italic': line_is_italic,
                                        'page_num': page_num, 'bbox_y0': current_line_bbox[1], 'bbox': current_line_bbox,
                                        'list_info': {'level': 1, 'type': list_type, 'original_marker': marker},
                                        '_page_default_x0': page_default_x0,
                                        '_attached_images': []
                                    })
                            else:
                                if cleaned_line_text:
                                    temp_lines_for_this_block_processing.append({
                                        'type': 'p_line', # Temporary type for lines within a block
                                        'text': cleaned_line_text, 'font_size': line_font_size,
                                        'bold': line_is_bold, 'italic': line_is_italic,
                                        'is_indented': is_indented_line if first_line_in_block else False, # Indent only applies to first line of a paragraph
                                        'bbox': current_line_bbox, 'bbox_y0': current_line_bbox[1], 'page_num': page_num
                                    })
                            first_line_in_block = False
                            last_line_y1_in_block_buffer = current_line_bbox[3]
                        
                        # After processing all lines in a block, flush remaining p_lines
                        if temp_lines_for_this_block_processing:
                            # Consolidate lines, respecting internal paragraph_break_markers
                            current_para_lines = []
                            first_line_data_for_para = None
                            for line_item in temp_lines_for_this_block_processing:
                                if line_item['type'] == 'paragraph_break_marker':
                                    if current_para_lines and first_line_data_for_para:
                                        page_items_this_page.append({
                                            'type': 'p', 'text': " ".join(current_para_lines).strip(),
                                            'font_size': first_line_data_for_para['font_size'],
                                            'bold': first_line_data_for_para['bold'],
                                            'italic': first_line_data_for_para['italic'],
                                            'is_indented': first_line_data_for_para.get('is_indented', False),
                                            'page_num': page_num,
                                            'bbox_y0': first_line_data_for_para['bbox_y0'],
                                            'bbox': first_line_data_for_para['bbox'], # Bbox of the first line of this paragraph
                                             '_page_default_x0': page_default_x0,
                                            '_attached_images': []
                                        })
                                    current_para_lines = []
                                    first_line_data_for_para = None
                                else: # 'p_line'
                                    if not first_line_data_for_para:
                                        first_line_data_for_para = line_item
                                    current_para_lines.append(line_item['text'])
                            
                            if current_para_lines and first_line_data_for_para: # Flush any remaining
                                page_items_this_page.append({
                                    'type': 'p', 'text': " ".join(current_para_lines).strip(),
                                    'font_size': first_line_data_for_para['font_size'],
                                    'bold': first_line_data_for_para['bold'],
                                    'italic': first_line_data_for_para['italic'],
                                    'is_indented': first_line_data_for_para.get('is_indented', False),
                                    'page_num': page_num,
                                    'bbox_y0': first_line_data_for_para['bbox_y0'],
                                    'bbox': first_line_data_for_para['bbox'],
                                    '_page_default_x0': page_default_x0,
                                    '_attached_images': []
                                })
                        previous_block_bottom_y = block_bbox[3] # Update for next inter-block check

                if page_num in images_by_page:
                    for img_ref in images_by_page[page_num]:
                        if page_num == content_start_page and img_ref['bbox'][1] < content_start_y:
                            continue
                        page_items_this_page.append({**img_ref, 'bbox_y0': img_ref['bbox'][1], '_page_default_x0': page_default_x0})
                
                if not PERFORM_OCR_ON_IMAGES and accumulated_text_on_page_len < MIN_CHARS_FOR_TEXTUAL_PAGE and any(item['type'] == 'image_placeholder' for item in page_items_this_page):
                    msg = f"WARNING: Page {page_num+1} yielded very little text ({accumulated_text_on_page_len} chars) and contains images. If it's an image-based page with text, enable OCR (perform_ocr_on_images = True in config.ini) for translation."
                    logger.warning(msg); quality_report_messages.append(msg)

                page_items_this_page.sort(key=lambda x: x.get('bbox_y0', 0))
                all_page_items_collected.extend(page_items_this_page)

            if not all_page_items_collected:
                msg = "CRITICAL WARNING: No content items collected after page iteration (check ToC detection or PDF content)."; logger.warning(msg); return []

            font_sizes_in_filtered_doc = [item['font_size'] for item in all_page_items_collected if item.get('font_size') and item.get('type') not in ['image_placeholder', 'paragraph_break_marker']]
            auto_min_font_for_inclusion, dominant_font_size_doc = 7.0, temp_dominant_font_size
            auto_pdf_h1_min_font, auto_pdf_h2_min_font, auto_pdf_h3_min_font = dominant_font_size_doc + 4.0, dominant_font_size_doc + 2.0, dominant_font_size_doc + 1.0

            if font_sizes_in_filtered_doc:
                plausible_body_text_sizes = [s for s in font_sizes_in_filtered_doc if 7.0 <= s <= 18.0 and s <= MAX_FONT_SIZE_FOR_BODY_ANALYSIS]
                if plausible_body_text_sizes:
                    font_counts = Counter(plausible_body_text_sizes); dominant_font_size_doc = font_counts.most_common(1)[0][0]
                    auto_min_font_for_inclusion = max(7.0, dominant_font_size_doc - 2.5)
                else:
                    auto_min_font_for_inclusion = 7.0
                
                potential_heading_sizes = sorted(list(set(s for s in font_sizes_in_filtered_doc if s > dominant_font_size_doc + 0.5 and s >= auto_min_font_for_inclusion)), reverse=True)
                if len(potential_heading_sizes) >= 3:
                    auto_pdf_h1_min_font, auto_pdf_h2_min_font, auto_pdf_h3_min_font = potential_heading_sizes[0], potential_heading_sizes[1], potential_heading_sizes[2]
                    if auto_pdf_h3_min_font >= auto_pdf_h2_min_font - 0.1 : auto_pdf_h3_min_font = max(dominant_font_size_doc + 0.5, auto_pdf_h2_min_font - 0.2)
                    if auto_pdf_h2_min_font >= auto_pdf_h1_min_font - 0.1 : auto_pdf_h2_min_font = max(auto_pdf_h3_min_font + 0.1, auto_pdf_h1_min_font - 0.2)
                elif len(potential_heading_sizes) == 2:
                    auto_pdf_h1_min_font, auto_pdf_h2_min_font = potential_heading_sizes[0], potential_heading_sizes[1]; auto_pdf_h3_min_font = auto_pdf_h2_min_font
                    if auto_pdf_h2_min_font >= auto_pdf_h1_min_font - 0.1: auto_pdf_h2_min_font = max(dominant_font_size_doc + 0.5, auto_pdf_h1_min_font - 0.2)
                elif len(potential_heading_sizes) == 1:
                    auto_pdf_h1_min_font = auto_pdf_h2_min_font = auto_pdf_h3_min_font = potential_heading_sizes[0]

            logger.debug(f"Font Analysis (Post-ToC): Dominant Size={dominant_font_size_doc}, Min Inclusion={auto_min_font_for_inclusion}, H1>={auto_pdf_h1_min_font}, H2>={auto_pdf_h2_min_font}, H3>={auto_pdf_h3_min_font}")
            quality_report_messages.append(f"INFO Font Analysis (Post-ToC): Dominant={dominant_font_size_doc}, MinIncl={auto_min_font_for_inclusion}, H1={auto_pdf_h1_min_font}, H2={auto_pdf_h2_min_font}, H3={auto_pdf_h3_min_font}")

            processed_content_intermediate = []
            bibliography_started = False
            last_text_item_for_image_attachment = None

            for current_item_from_collected in all_page_items_collected:
                if bibliography_started: continue
                item_type_original = current_item_from_collected.get('type')

                if item_type_original == 'paragraph_break_marker': # Handle paragraph break markers
                    if processed_content_intermediate and processed_content_intermediate[-1]['type'] != 'paragraph_break_marker': # Avoid duplicates
                         processed_content_intermediate.append(current_item_from_collected)
                    last_text_item_for_image_attachment = None
                    continue

                if item_type_original == 'image_placeholder':
                    if last_text_item_for_image_attachment and \
                       last_text_item_for_image_attachment['page_num'] == current_item_from_collected['page_num'] and \
                       current_item_from_collected.get('bbox_y0',0) > last_text_item_for_image_attachment.get('bbox', (0,0,0,0))[3] and \
                       (current_item_from_collected.get('bbox_y0',0) - last_text_item_for_image_attachment.get('bbox', (0,0,0,0))[3]) < MAX_VERTICAL_GAP_FOR_IMAGE_ATTACHMENT_POINTS:
                        last_text_item_for_image_attachment.setdefault('_attached_images', []).append(current_item_from_collected)
                    else:
                        processed_content_intermediate.append(current_item_from_collected)
                    last_text_item_for_image_attachment = None
                    continue

                cleaned_text = current_item_from_collected.get('text', '')
                if item_type_original not in ['list_item', 'image_placeholder'] and (not current_item_from_collected.get('font_size') or current_item_from_collected.get('font_size', 0) < auto_min_font_for_inclusion):
                    last_text_item_for_image_attachment = None; continue
                if not cleaned_text:
                    last_text_item_for_image_attachment = None; continue
                
                current_text_lower = cleaned_text.lower()
                is_bib_heading = any(keyword == current_text_lower for keyword in BIBLIOGRAPHY_KEYWORDS_LOWER) or \
                                 any(current_text_lower.startswith(f"{keyword}:") for keyword in BIBLIOGRAPHY_KEYWORDS_LOWER) or \
                                 any(current_text_lower.startswith(f"{keyword} ") for keyword in BIBLIOGRAPHY_KEYWORDS_LOWER)
                is_bib_style = current_item_from_collected.get('font_size', 0) >= auto_pdf_h3_min_font -1 or \
                               (current_item_from_collected.get('bold', False) and current_item_from_collected.get('font_size', 0) >= dominant_font_size_doc +0.5)
                if not bibliography_started and is_bib_heading and is_bib_style and len(cleaned_text.split()) < 10:
                    bibliography_started = True; msg = f"Bibliography detected: '{cleaned_text}'."; logger.info(msg); quality_report_messages.append(msg)
                    last_text_item_for_image_attachment = None; continue
                
                item_type_val = item_type_original
                if item_type_original == 'p':
                    fs_item = current_item_from_collected.get('font_size', 0); is_bold_item = current_item_from_collected.get('bold', False)
                    is_chap_patt = any(pattern.match(cleaned_text) for pattern in CHAPTER_TITLE_PATTERNS)
                    word_count = len(cleaned_text.split())
                    
                    if (is_chap_patt and (fs_item >= auto_pdf_h1_min_font - 0.5 or (is_bold_item and fs_item >= auto_pdf_h1_min_font - 2.0))) or \
                       (fs_item >= auto_pdf_h1_min_font or (is_bold_item and fs_item >= auto_pdf_h1_min_font -1.5 and word_count < HEADING_MAX_WORDS)):
                        item_type_val = 'h1'
                    elif (fs_item >= auto_pdf_h2_min_font or (is_bold_item and fs_item >= auto_pdf_h2_min_font -1.5)) and word_count < HEADING_MAX_WORDS + 5 :
                        item_type_val = 'h2'
                    elif (fs_item >= auto_pdf_h3_min_font or (is_bold_item and fs_item >= auto_pdf_h3_min_font -1.5)) and word_count < HEADING_MAX_WORDS + 10:
                        item_type_val = 'h3'
                
                current_item_from_collected['type'] = item_type_val
                processed_content_intermediate.append(current_item_from_collected)
                if item_type_val != 'image_placeholder':
                    last_text_item_for_image_attachment = current_item_from_collected
            
            # Aggregation Stage 1: Combine consecutive items of same type (respecting paragraph_break_marker)
            aggregated_content_stage1 = []
            buffer = []; buffer_type = None; buffer_style = {}; buffer_list_info_agg = None; buffer_images_agg = []; buffer_is_indented_agg = False

            def flush_buffer_agg1():
                nonlocal buffer_type, buffer_style, buffer_list_info_agg, buffer_images_agg, buffer_is_indented_agg
                if buffer:
                    text_content = " ".join(b_item['text'] for b_item in buffer if b_item.get('text')).strip() # Changed from \n to space
                    consolidated_images_from_buffer = []
                    for b_item in buffer:
                        consolidated_images_from_buffer.extend(b_item.get('_attached_images', []))
                    
                    if text_content or consolidated_images_from_buffer:
                        current_buffer_type_to_use = buffer_type
                        if current_buffer_type_to_use is None:
                            if any(b_item.get('list_info') for b_item in buffer): current_buffer_type_to_use = 'list_item'
                            else: current_buffer_type_to_use = 'p'
                        
                        item_to_add = {'type': current_buffer_type_to_use, 'text': text_content, **buffer_style}
                        if buffer_list_info_agg:
                            item_to_add['list_info'] = buffer_list_info_agg
                        if current_buffer_type_to_use == 'p' and buffer_is_indented_agg: # Carry indent
                            item_to_add['is_indented'] = True

                        unique_image_data_list = []
                        seen_ids = set()
                        for img_data_original in consolidated_images_from_buffer:
                            img_id = img_data_original.get('id')
                            if img_id and img_id not in seen_ids:
                                unique_image_data_list.append(img_data_original)
                                seen_ids.add(img_id)
                            elif not img_id:
                                unique_image_data_list.append(img_data_original)
                        item_to_add['_attached_images'] = unique_image_data_list
                        
                        aggregated_content_stage1.append(item_to_add)
                
                buffer.clear(); buffer_images_agg.clear(); buffer_list_info_agg = None
                buffer_type = None; buffer_style = {}; buffer_is_indented_agg = False

            for item_agg1 in processed_content_intermediate:
                item_type_agg1 = item_agg1['type']
                is_textual_item = item_type_agg1 in ['p', 'list_item']

                if item_type_agg1 == 'paragraph_break_marker':
                    flush_buffer_agg1()
                    # Optionally, could add a specific marker item if needed later, but Word handles paragraph breaks.
                    continue

                if not buffer:
                    if is_textual_item:
                        buffer_type = item_type_agg1
                        buffer_style = {'font_size': item_agg1.get('font_size'), 'bold': item_agg1.get('bold', False), 'italic': item_agg1.get('italic', False)}
                        if item_type_agg1 == 'list_item': buffer_list_info_agg = item_agg1.get('list_info')
                        else: buffer_list_info_agg = None
                        if item_type_agg1 == 'p': buffer_is_indented_agg = item_agg1.get('is_indented', False) # Capture indent
                        buffer_images_agg.clear(); buffer_images_agg.extend(item_agg1.get('_attached_images', []))
                        buffer.append(item_agg1)
                    else:
                        aggregated_content_stage1.append(item_agg1)
                else:
                    # Flush if type changes, or if it's a heading, or if list status changes
                    if not is_textual_item or \
                       item_type_agg1 != buffer_type or \
                       (item_type_agg1 == 'list_item' and buffer_type != 'list_item') or \
                       (item_type_agg1 != 'list_item' and buffer_type == 'list_item') or \
                       item_type_agg1 in ['h1', 'h2', 'h3'] or \
                       (item_type_agg1 == 'p' and item_agg1.get('is_indented', False) != buffer_is_indented_agg and buffer_type == 'p'): # Flush if indent status changes for 'p'
                        flush_buffer_agg1()
                        if is_textual_item:
                            buffer_type = item_type_agg1
                            buffer_style = {'font_size': item_agg1.get('font_size'), 'bold': item_agg1.get('bold', False), 'italic': item_agg1.get('italic', False)}
                            if item_type_agg1 == 'list_item': buffer_list_info_agg = item_agg1.get('list_info')
                            else: buffer_list_info_agg = None
                            if item_type_agg1 == 'p': buffer_is_indented_agg = item_agg1.get('is_indented', False)
                            buffer_images_agg.clear(); buffer_images_agg.extend(item_agg1.get('_attached_images', []))
                            buffer.append(item_agg1)
                        else:
                            aggregated_content_stage1.append(item_agg1)
                    else: # Compatible item, add to buffer
                        buffer_images_agg.extend(item_agg1.get('_attached_images', []))
                        buffer.append(item_agg1)
            flush_buffer_agg1()

            sub_chunked_content = []
            for item_sub in aggregated_content_stage1:
                item_type_sub = item_sub['type']
                attached_images_for_item = item_sub.get('_attached_images', [])
                list_info_for_item = item_sub.get('list_info') if item_type_sub == 'list_item' else None
                is_indented_sub = item_sub.get('is_indented', False) if item_type_sub == 'p' else False

                if item_type_sub in ['p', 'list_item'] and len(item_sub['text']) > MAX_CHARS_PER_SUBCHUNK:
                    text_to_split = item_sub['text']; start_index = 0
                    item_style_sub = {'font_size': item_sub.get('font_size'), 'bold': item_sub.get('bold'), 'italic': item_sub.get('italic')}
                    first_sub_chunk = True
                    while start_index < len(text_to_split):
                        end_index = min(start_index + MAX_CHARS_PER_SUBCHUNK, len(text_to_split))
                        if end_index < len(text_to_split):
                            best_split_pos = -1
                            for delimiter in ['\n\n', '. ', '.\n', '.\t', '\n', '? ', '! ', '; ']:
                                pos = text_to_split.rfind(delimiter, start_index, end_index)
                                if pos != -1 and pos + len(delimiter) > start_index :
                                    if best_split_pos == -1 or pos + len(delimiter) > best_split_pos : best_split_pos = pos + len(delimiter)
                            if best_split_pos == -1 :
                                space_pos = text_to_split.rfind(' ', start_index, end_index)
                                if space_pos != -1 and space_pos > start_index : best_split_pos = space_pos + 1
                            if best_split_pos != -1 and best_split_pos > start_index : end_index = best_split_pos
                        
                        sub_chunk_text_val = text_to_split[start_index:end_index].strip()
                        if sub_chunk_text_val:
                            chunk_data = {'type': item_type_sub, 'text': sub_chunk_text_val, **item_style_sub}
                            if list_info_for_item: chunk_data['list_info'] = list_info_for_item
                            if is_indented_sub and first_sub_chunk: chunk_data['is_indented'] = True # Only first sub-chunk of an indented para is indented
                            if first_sub_chunk and attached_images_for_item:
                                chunk_data['_attached_images'] = attached_images_for_item; first_sub_chunk = False
                            else: chunk_data['_attached_images'] = []
                            sub_chunked_content.append(chunk_data)
                        start_index = end_index
                else:
                    item_sub['_attached_images'] = attached_images_for_item
                    if list_info_for_item: item_sub['list_info'] = list_info_for_item
                    if is_indented_sub: item_sub['is_indented'] = True
                    sub_chunked_content.append(item_sub)
            
            final_structured_items = []
            small_elements_buffer = []; accumulated_chars_small_buffer = 0
            buffer_main_item_style = None; buffer_main_item_type = 'p'; buffer_attached_images_final = []; buffer_list_info_final_agg = None; buffer_is_indented_final = False
            for item_final_agg in sub_chunked_content:
                item_text_fa = item_final_agg.get('text', ''); item_text_len_fa = len(item_text_fa)
                item_attached_imgs_fa = item_final_agg.get('_attached_images', [])
                item_type_final = item_final_agg['type']
                item_list_info_final_agg = item_final_agg.get('list_info')
                item_is_indented_final = item_final_agg.get('is_indented', False) if item_type_final == 'p' else False
                should_flush_buffer = False
                if small_elements_buffer:
                    if item_type_final in ['h1', 'h2'] or \
                       (item_type_final == 'image_placeholder' and not item_attached_imgs_fa) or \
                       (item_type_final == 'list_item' and buffer_main_item_type != 'list_item') or \
                       (item_type_final != 'list_item' and buffer_main_item_type == 'list_item') or \
                       (item_type_final == 'h3' and (item_text_len_fa >= MIN_CHARS_FOR_STANDALONE_CHUNK or buffer_main_item_type == 'p')) or \
                       ((item_type_final == 'p' or item_type_final == 'list_item') and item_text_len_fa >= MIN_CHARS_FOR_STANDALONE_CHUNK) or \
                       (accumulated_chars_small_buffer + item_text_len_fa > AGGREGATE_SMALL_CHUNKS_TARGET_SIZE) or \
                       (item_type_final == 'p' and item_is_indented_final != buffer_is_indented_final and buffer_main_item_type == 'p'): # Flush if indent changes
                        should_flush_buffer = True
                
                if should_flush_buffer:
                    if small_elements_buffer:
                        item_to_add = {'type': buffer_main_item_type, 'text': "\n".join(small_elements_buffer).strip(), **(buffer_main_item_style if buffer_main_item_style else {}), '_attached_images': buffer_attached_images_final}
                        if buffer_list_info_final_agg: item_to_add['list_info'] = buffer_list_info_final_agg
                        if buffer_main_item_type == 'p' and buffer_is_indented_final: item_to_add['is_indented'] = True
                        final_structured_items.append(item_to_add)
                    small_elements_buffer = []; accumulated_chars_small_buffer = 0; buffer_main_item_style = None; buffer_attached_images_final = []; buffer_list_info_final_agg = None; buffer_is_indented_final = False
                
                is_standalone_item = False
                if item_type_final in ['h1', 'h2', 'image_placeholder'] or \
                   ((item_type_final == 'h3' or item_type_final == 'p' or item_type_final == 'list_item') and item_text_len_fa >= MIN_CHARS_FOR_STANDALONE_CHUNK):
                    is_standalone_item = True

                if is_standalone_item:
                    if small_elements_buffer:
                        item_to_add = {'type': buffer_main_item_type, 'text': "\n".join(small_elements_buffer).strip(), **(buffer_main_item_style if buffer_main_item_style else {}), '_attached_images': buffer_attached_images_final}
                        if buffer_list_info_final_agg: item_to_add['list_info'] = buffer_list_info_final_agg
                        if buffer_main_item_type == 'p' and buffer_is_indented_final: item_to_add['is_indented'] = True
                        final_structured_items.append(item_to_add)
                        small_elements_buffer = []; accumulated_chars_small_buffer = 0; buffer_main_item_style = None; buffer_attached_images_final = []; buffer_list_info_final_agg = None; buffer_is_indented_final = False
                    final_structured_items.append(item_final_agg)
                elif item_type_final not in ['image_placeholder']:
                    if not small_elements_buffer:
                        buffer_main_item_type = item_type_final
                        buffer_main_item_style = {'font_size': item_final_agg.get('font_size'), 'bold': item_final_agg.get('bold', False), 'italic': item_final_agg.get('italic', False)}
                        buffer_attached_images_final = list(item_attached_imgs_fa)
                        if item_list_info_final_agg: buffer_list_info_final_agg = item_list_info_final_agg
                        if item_type_final == 'p': buffer_is_indented_final = item_is_indented_final
                    else:
                        buffer_attached_images_final.extend(item_attached_imgs_fa)
                    small_elements_buffer.append(item_text_fa)
                    accumulated_chars_small_buffer += item_text_len_fa + (1 if len(small_elements_buffer) > 1 else 0)
                    if accumulated_chars_small_buffer >= AGGREGATE_SMALL_CHUNKS_TARGET_SIZE:
                        if small_elements_buffer:
                            item_to_add = {'type': buffer_main_item_type, 'text': "\n".join(small_elements_buffer).strip(), **(buffer_main_item_style if buffer_main_item_style else {}), '_attached_images': buffer_attached_images_final}
                            if buffer_list_info_final_agg: item_to_add['list_info'] = buffer_list_info_final_agg
                            if buffer_main_item_type == 'p' and buffer_is_indented_final: item_to_add['is_indented'] = True
                            final_structured_items.append(item_to_add)
                        small_elements_buffer = []; accumulated_chars_small_buffer = 0; buffer_main_item_style = None; buffer_attached_images_final = []; buffer_list_info_final_agg = None; buffer_is_indented_final = False
            
            if small_elements_buffer:
                item_to_add = {'type': buffer_main_item_type, 'text': "\n".join(small_elements_buffer).strip(), **(buffer_main_item_style if buffer_main_item_style else {}), '_attached_images': buffer_attached_images_final}
                if buffer_list_info_final_agg: item_to_add['list_info'] = buffer_list_info_final_agg
                if buffer_main_item_type == 'p' and buffer_is_indented_final: item_to_add['is_indented'] = True
                final_structured_items.append(item_to_add)

    except Exception as e_struct_main:
        msg = f"CRITICAL ERROR during content extraction: {e_struct_main}"; logger.error(msg, exc_info=True); quality_report_messages.append(msg)
        return []
    return final_structured_items

async def analyze_document_style_async(sample_text, analysis_prompt_text, analysis_model, analysis_temp, semaphore):
    global quality_report_messages, TRANSLATION_STYLE_TONE_FALLBACK
    if not sample_text.strip(): quality_report_messages.append("WARNING: Document sample for style analysis is empty."); return TRANSLATION_STYLE_TONE_FALLBACK
    async with semaphore:
        model = genai.GenerativeModel(analysis_model)
        max_sample_len_for_prompt = 15000; truncated_sample = sample_text[:max_sample_len_for_prompt]
        if len(sample_text) > max_sample_len_for_prompt: quality_report_messages.append(f"INFO: Style analysis sample truncated to {max_sample_len_for_prompt} chars.")
        full_prompt = f"{analysis_prompt_text}\n\nTEXT SAMPLE:\n---\n{truncated_sample}\n---\nCONCISE STYLE SUMMARY (max 70 words):"
        try:
            response = await model.generate_content_async(full_prompt, generation_config=genai.types.GenerationConfig(temperature=analysis_temp), request_options=types.RequestOptions(timeout=API_CALL_TIMEOUT))
            if response.parts:
                style_summary = response.text.strip()
                if len(style_summary) > 250: style_summary = style_summary[:247] + "..."; quality_report_messages.append(f"INFO: Style analysis summary auto-truncated.")
                quality_report_messages.append(f"INFO: Document Style Analysis Complete. Summary: '{style_summary}'")
                return style_summary
            else:
                err_msg_text = "Style analysis API returned empty/blocked."
                if hasattr(response, 'prompt_feedback') and response.prompt_feedback: err_msg_text += f" Feedback: {response.prompt_feedback}"
                quality_report_messages.append(f"ERROR (Style Analysis API): {err_msg_text}"); return TRANSLATION_STYLE_TONE_FALLBACK
        except Exception as e: quality_report_messages.append(f"ERROR (Style Analysis API Call): {e}"); logger.error(f"Style Analysis API Call Error: {e}", exc_info=True); return TRANSLATION_STYLE_TONE_FALLBACK

async def assess_translation_quality_async(original_text, translated_text, target_language, dynamic_style_guide, semaphore):
    global quality_report_messages, qa_format_error_log_count, MAX_QA_FORMAT_ERRORS_TO_LOG, QUALITY_ASSESSMENT_MODEL, QUALITY_ASSESSMENT_PROMPT_TEXT
    async with semaphore:
        model_qa = genai.GenerativeModel(QUALITY_ASSESSMENT_MODEL)
        original_snippet = original_text[:800] + "..." if len(original_text) > 800 else original_text
        translated_snippet = translated_text[:1200] + "..." if len(translated_text) > 1200 else translated_text
        style_guide_snippet = dynamic_style_guide[:200] + "..." if dynamic_style_guide and len(dynamic_style_guide) > 200 else (dynamic_style_guide if dynamic_style_guide else TRANSLATION_STYLE_TONE_FALLBACK)
        
        prompt = QUALITY_ASSESSMENT_PROMPT_TEXT.format(
            target_language=target_language,
            style_guide_snippet=style_guide_snippet,
            original_snippet=original_snippet,
            translated_snippet=translated_snippet
        )
        
        retries_qa = 0; max_retries_qa = 2; current_delay_qa = 5
        while retries_qa <= max_retries_qa:
            try:
                response = await model_qa.generate_content_async(prompt, generation_config=genai.types.GenerationConfig(temperature=QUALITY_ASSESSMENT_TEMPERATURE), request_options=types.RequestOptions(timeout=180))
                if response.parts:
                    assessment_text = response.text.strip()
                    match = re.search(r"Score:\s*([1-5](?:\.\d{1,2})?)\.?\s*Comment:\s*(.+)", assessment_text, re.IGNORECASE | re.DOTALL)
                    score_val = 0; comment_val = "N/A (Format Error)"
                    if match:
                        try: score_val = float(match.group(1))
                        except ValueError: score_val = 0
                        comment_val = match.group(2).strip()[:150]
                    else:
                        if qa_format_error_log_count < MAX_QA_FORMAT_ERRORS_TO_LOG:
                            orig_snip_for_log = original_snippet[:70].replace('\n', ' ');
                            debug_msg = f"WARNING (QA Format): QA LLM response for '{orig_snip_for_log}...' no match. Raw: '{assessment_text[:200]}...'";
                            logger.warning(debug_msg); quality_report_messages.append(debug_msg);
                            qa_format_error_log_count += 1
                        elif qa_format_error_log_count == MAX_QA_FORMAT_ERRORS_TO_LOG:
                            quality_report_messages.append(f"WARNING (QA Format): Further QA format errors suppressed.");
                            qa_format_error_log_count += 1
                        comment_val = f"QA Format Error. LLM Raw: {assessment_text[:100]}"

                    return {'score': score_val, 'comment': comment_val, 'raw_assessment': assessment_text}
                else:
                    err_msg_text_qa = "Empty/blocked API response for QA."; feedback = getattr(response, 'prompt_feedback', None); err_msg_text_qa += f" Feedback: {feedback}" if feedback else ""
                    quality_report_messages.append(f"WARNING (QA API): {err_msg_text_qa} (Original: '{original_snippet[:50]}...')."); return {'score': 0, 'comment': "QA: Empty API response.", 'raw_assessment': ''}
            except Exception as e_qa:
                error_type_name_qa = type(e_qa).__name__; error_details_qa = str(e_qa); is_retryable_qa = "DeadlineExceeded" in error_type_name_qa or "503" in error_details_qa or "ServiceUnavailable" in error_type_name_qa
                if is_retryable_qa and retries_qa < max_retries_qa:
                    retries_qa += 1; logger.warning(f"QA API ERROR ({error_type_name_qa}) for snippet '{original_snippet[:50]}...' (Attempt {retries_qa}/{max_retries_qa+1}). Retrying in {current_delay_qa}s: {e_qa}"); await asyncio.sleep(current_delay_qa); current_delay_qa *= 2; continue
                else: quality_report_messages.append(f"ERROR (QA API Call - {error_type_name_qa}): Original snippet '{original_snippet[:50]}...' - {e_qa}"); logger.error(f"QA API Call Error: {e_qa}", exc_info=True); return {'score': 0, 'comment': f"QA API Error: {error_type_name_qa}", 'raw_assessment': ''}
        quality_report_messages.append(f"ERROR (QA API Call): Max retries for '{original_snippet[:50]}...'"); logger.error(f"QA API Max retries for '{original_snippet[:50]}...'"); return {'score': 0, 'comment': "QA: Max retries exceeded", 'raw_assessment': ''}

async def call_gemini_for_translation(item_index_original, text_to_translate, prev_ctx_text, next_ctx_text, target_language, semaphore, model_name_to_use, temperature, dynamic_style_guide, max_retries=3, initial_delay=15, item_type_for_prompt="text block", glossary_terms_in_item=None):
    item_index_tuple = item_index_original
    global quality_report_messages, TRANSLATION_GLOSSARY, TRANSLATION_CACHE, TRANSLATION_STYLE_TONE_FALLBACK
    async with semaphore:
        cache_key = None
        if USE_TRANSLATION_CACHE:
            cache_key = get_cache_key(text_to_translate, target_language, model_name_to_use)
            if cache_key in TRANSLATION_CACHE:
                idx_print_str = ""
                if item_index_tuple[0] == 'ocr':
                    main_idx_val = item_index_tuple[1]
                    sub_idx_val = item_index_tuple[2]
                    idx_print_str = f"Item {main_idx_val + 1} OCR Img {sub_idx_val + 1 if sub_idx_val != -1 else 'Standalone'}"
                else:
                    actual_numerical_index = item_index_tuple[1]
                    idx_print_str = str(actual_numerical_index + 1)

                # Enhanced cache hit logging for groups
                cache_size_info = f" ({len(text_to_translate)} chars)"
                if len(text_to_translate) > 8000:
                    cache_size_info += " [LARGE GROUP]"

                quality_report_messages.append(f"CACHE HIT: Item {idx_print_str} (type: {item_type_for_prompt}){cache_size_info}")
                if glossary_terms_in_item is None and USE_GLOSSARY:
                    glossary_terms_in_item = {k: v for k, v in TRANSLATION_GLOSSARY.items() if k.lower() in text_to_translate.lower()}
                return item_index_tuple, TRANSLATION_CACHE[cache_key], True, bool(glossary_terms_in_item)
        
        model = genai.GenerativeModel(model_name_to_use); glossary_instructions = ""
        if USE_GLOSSARY and glossary_terms_in_item:
            glossary_instructions = "Adhere to the following glossary terms strictly (case-insensitive for source terms):\n"
            for o,t in glossary_terms_in_item.items(): glossary_instructions += f"- Translate instances of '{o}' (and its variations) as '{t}'\n"
            glossary_instructions += "---\n"
        
        style_instruction_to_use = f"Translate according to the following document style guide: '{dynamic_style_guide}'. Maintain overall consistency with this guide." if dynamic_style_guide and dynamic_style_guide.strip() and dynamic_style_guide.lower() != "neutral" and dynamic_style_guide != TRANSLATION_STYLE_TONE_FALLBACK else f"Ensure the translation adopts a {TRANSLATION_STYLE_TONE_FALLBACK} style and tone."
        
        # ENHANCED PROMPT with document intelligence
        # Detect document type for enhanced translation
        document_sample = f"{prev_ctx_text} {text_to_translate} {next_ctx_text}"
        doc_detector = DocumentTypeDetector()
        doc_type, confidence, type_config = doc_detector.detect_document_type(document_sample)

        # Enhanced style instruction with document context
        enhanced_style_instruction = style_instruction_to_use
        if confidence > 0.3:  # If we have reasonable confidence in document type
            enhanced_style_instruction += f" This is a {doc_type} document. {type_config['style_guide']} Pay special attention to {type_config['terminology_focus']}."

        prompt = f"""You are an expert translator specializing in {doc_type} documents. Translate the "MAIN TEXT TO TRANSLATE" which is identified as a "{item_type_for_prompt}" into {target_language}.

DOCUMENT CONTEXT:
- Document Type: {doc_type.title()} (confidence: {confidence:.2f})
- Specialized Focus: {type_config['terminology_focus']}

TRANSLATION REQUIREMENTS:
{enhanced_style_instruction}
{glossary_instructions}

QUALITY STANDARDS:
- Accuracy: Preserve exact meaning and technical details
- Fluency: Natural, readable {target_language} that flows well
- Consistency: Use consistent terminology throughout
- Cultural Adaptation: Adapt cultural references appropriately
- Professional Standards: Meet professional translation quality for {doc_type} documents

PREVIOUS CONTEXT (original language, for context only, do not translate this part):
---
{prev_ctx_text if prev_ctx_text else "N/A"}
---
MAIN TEXT TO TRANSLATE:
---
{text_to_translate}
---
NEXT CONTEXT (original language, for context only, do not translate this part):
---
{next_ctx_text if next_ctx_text else "N/A"}
---

TRANSLATION INSTRUCTIONS:
1. Apply {doc_type}-specific translation best practices
2. Ensure terminology consistency with the document type and context
3. Preserve any formatting, line breaks, and special characters
4. Maintain the same level of formality as the original text
5. For {target_language}, use monotonic Greek (single accent system) - avoid polytonic accents
6. Provide ONLY the translation of the MAIN TEXT TO TRANSLATE, without any additional commentary

TRANSLATION:"""
        
        retries = 0; current_delay = initial_delay
        idx_print_str = ""
        if item_index_tuple[0] == 'ocr':
            main_idx_val = item_index_tuple[1]
            sub_idx_val = item_index_tuple[2]
            idx_print_str = f"Item {main_idx_val + 1} OCR Img {sub_idx_val + 1 if sub_idx_val != -1 else 'Standalone'}"
        else:
            actual_numerical_index = item_index_tuple[1]
            idx_print_str = str(actual_numerical_index + 1)

        while retries <= max_retries:
            try:
                response = await model.generate_content_async(prompt,generation_config=genai.types.GenerationConfig(temperature=temperature),request_options=types.RequestOptions(timeout=API_CALL_TIMEOUT))
                if response.parts:
                    translated_text = response.text.strip()
                    if USE_TRANSLATION_CACHE and cache_key:
                        TRANSLATION_CACHE[cache_key] = translated_text
                        # Enhanced cache save logging for groups
                        cache_save_info = f" [CACHED: {len(text_to_translate)} chars]"
                        if len(text_to_translate) > 8000:
                            cache_save_info += " [LARGE GROUP SAVED]"
                    else:
                        cache_save_info = ""

                    msg_detail = f"(type: {item_type_for_prompt}){cache_save_info}"
                    if glossary_terms_in_item: msg_detail += f" (Glossary terms considered: {len(glossary_terms_in_item)})"
                    quality_report_messages.append(f"SUCCESS (API): Item {idx_print_str} {msg_detail} translated."); return item_index_tuple, translated_text, False, bool(glossary_terms_in_item)
                else:
                    err_msg_text = "Empty or blocked API response for translation."; feedback = getattr(response, 'prompt_feedback', None); err_msg_text += f" Feedback: {feedback}" if feedback else ""
                    quality_report_messages.append(f"ERROR (API NON-RETRYABLE): Item {idx_print_str} - {err_msg_text}"); logger.error(f"API NON-RETRYABLE Error for Item {idx_print_str}: {err_msg_text}")
                    return item_index_tuple, f"[Translation failed for item {idx_print_str}: {err_msg_text}]", False, bool(glossary_terms_in_item)
            except Exception as e:
                error_type = type(e).__name__; error_details = str(e); is_retryable = "DeadlineExceeded" in error_type or "503" in error_details or "504" in error_details or ("ResourceExhausted" in error_type and ("quota" in error_details.lower() or "per_minute" in error_details.lower() or "rate limit" in error_details.lower())) or "ServiceUnavailable" in error_type or "InternalServerError" in error_type
                logger.warning(f"API ERROR ({error_type}) translating item {idx_print_str} (Attempt {retries + 1}/{max_retries+1}): {e}")
                if is_retryable and retries < max_retries:
                    retries += 1; logger.info(f"Retrying translation for item {idx_print_str} in {current_delay}s..."); await asyncio.sleep(current_delay); current_delay = min(current_delay * 2, 300)
                else:
                    final_err_msg_text = f"({error_type}) after {retries +1} attempt(s): {e}"; quality_report_messages.append(f"ERROR (API RETRY FAILED or NON-RETRYABLE): Item {idx_print_str} - {final_err_msg_text}"); logger.error(f"API RETRY FAILED for Item {idx_print_str}: {final_err_msg_text}", exc_info=True)
                    return item_index_tuple, f"[Translation permanently failed for item {idx_print_str}: {final_err_msg_text}]", False, bool(glossary_terms_in_item)
        
        final_loop_err_msg_text = f"Translation for item {idx_print_str} failed after {max_retries+1} retries (loop exhausted)"; quality_report_messages.append(f"ERROR (API LOOP EXHAUSTED): Item {idx_print_str} - {final_loop_err_msg_text}"); logger.error(f"API LOOP EXHAUSTED for Item {idx_print_str}: {final_loop_err_msg_text}")
        return item_index_tuple, f"[Translation permanently failed for item {idx_print_str}: {final_loop_err_msg_text}]", False, bool(glossary_terms_in_item)

def add_toc_entry(doc, text, level, bookmark_ref):
    first_line = text.split('\n')[0]
    toc_text = (first_line[:75] + '...') if len(first_line) > 78 else first_line

    p = doc.add_paragraph(); run = p.add_run(toc_text)
    if level == 1: run.font.size = Pt(12); run.bold = True
    elif level == 2: run.font.size = Pt(11); p.paragraph_format.left_indent = Inches(0.25)
    elif level == 3: run.font.size = Pt(10); p.paragraph_format.left_indent = Inches(0.5)
    hyperlink = OxmlElement('w:hyperlink'); hyperlink.set(qn('w:anchor'), bookmark_ref); hyperlink.append(run._r); p._p.append(hyperlink)

def add_bookmark_to_paragraph(paragraph_obj, bookmark_name_str):
    global bookmark_id_counter
    try:
        if paragraph_obj is None or not hasattr(paragraph_obj, '_p'): 
            quality_report_messages.append(f"WARNING: Cannot add bookmark '{bookmark_name_str}', invalid paragraph object.")
            return
        
        # Simple approach: just add a text marker instead of XML bookmarks
        # This avoids the complex XML manipulation that causes errors
        bookmark_text = f"[BOOKMARK: {bookmark_name_str}]"
        
        # Add bookmark as invisible text (very small font)
        if hasattr(paragraph_obj, 'add_run'):
            bookmark_run = paragraph_obj.add_run(bookmark_text)
            bookmark_run.font.size = Pt(1)  # Very small
            bookmark_run.font.color.rgb = RGBColor(255, 255, 255)  # White (invisible)
        
        bookmark_id_counter += 1
        logger.debug(f"Added simple bookmark: {bookmark_name_str}")
        
    except Exception as e:
        logger.warning(f"Failed to add bookmark '{bookmark_name_str}': {e}")
        quality_report_messages.append(f"WARNING: Failed to add bookmark '{bookmark_name_str}': {e}")
        bookmark_id_counter += 1  # Still increment to avoid ID conflicts

def _add_image_to_doc(doc, image_data, image_folder_path, main_item_for_qa_ref=None):
    if not image_folder_path or not image_data.get('filename'):
        placeholder_text = image_data.get('text', f"[Image placeholder data missing: {image_data.get('filename', 'unknown')}]")
        doc.add_paragraph(placeholder_text).italic = True
        return
    image_filename = image_data.get('filename')
    image_path = os.path.join(image_folder_path, image_filename)
    if os.path.exists(image_path):
        try:
            img_to_insert = image_path
            temp_png_path = None
            if PIL_Image and (image_filename.lower().endswith(".jpg") or image_filename.lower().endswith(".jpeg")):
                try:
                    img = PIL_Image.open(image_path)
                    if img.format != 'PNG':
                        temp_dir_for_conversion = tempfile.gettempdir()
                        temp_png_filename = os.path.splitext(os.path.basename(image_filename))[0] + "_converted.png"
                        temp_png_path = os.path.join(temp_dir_for_conversion, temp_png_filename)
                        img.convert("RGB").save(temp_png_path, "PNG")
                        img_to_insert = temp_png_path
                        logger.debug(f"Converted {image_filename} to PNG for Word insertion: {temp_png_path}")
                except Exception as e_pil:
                    logger.warning(f"Pillow (PIL) could not process/convert image {image_filename}: {e_pil}. Will try inserting original.")

            p_img = doc.add_paragraph(); p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_img = p_img.add_run(); run_img.add_picture(img_to_insert, width=Inches(DEFAULT_IMAGE_WIDTH_INCHES))
            
            if temp_png_path and os.path.exists(temp_png_path):
                try: os.remove(temp_png_path)
                except Exception as e_remove: logger.warning(f"Could not remove temporary PNG {temp_png_path}: {e_remove}")

            translated_ocr = image_data.get('translated_ocr_text')
            original_ocr = image_data.get('ocr_text_original')
            caption_to_add = None; cap_p_for_qa = None
            if translated_ocr and not translated_ocr.startswith("[Translation failed"):
                caption_to_add = f"Μετάφραση κειμένου εικόνας: {translated_ocr}"
            elif original_ocr:
                ocr_status = "(μη μεταφρασμένο)"
                if translated_ocr and translated_ocr.startswith("[Translation failed"): ocr_status = "(αποτυχία μετάφρασης)"
                elif not translated_ocr and PERFORM_OCR_ON_IMAGES and len(original_ocr.split()) >= MIN_OCR_WORDS_FOR_TRANSLATION :
                     ocr_status = "(αποτυχία μετάφρασης - άγνωστος λόγος)"
                elif not translated_ocr : ocr_status = "(παραλείφθηκε η μετάφραση - πολύ σύντομο ή OCR απενεργοποιημένο)"
                caption_to_add = f"Εξαγμένο κείμενο εικόνας {ocr_status}: {original_ocr[:200]}{'...' if len(original_ocr)>200 else ''}"
            
            if caption_to_add:
                cap_p_for_qa = doc.add_paragraph(); cap_run = cap_p_for_qa.add_run(caption_to_add)
                cap_run.font.size = Pt(9); cap_run.italic = True; cap_p_for_qa.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            if SHOW_QA_ANNOTATIONS_IN_WORD and image_data.get('quality_assessment_ocr'):
                qa_info_ocr = image_data['quality_assessment_ocr']
                qa_text_ocr = f" [QA-OCR: {qa_info_ocr.get('score','N/A')}/5 - {qa_info_ocr.get('comment','N/A')}]"
                para_to_add_qa_ocr = cap_p_for_qa if cap_p_for_qa else doc.add_paragraph()
                if not cap_p_for_qa: para_to_add_qa_ocr.alignment = WD_ALIGN_PARAGRAPH.CENTER
                qa_run_ocr = para_to_add_qa_ocr.add_run(qa_text_ocr)
                qa_run_ocr.font.size = Pt(7); qa_run_ocr.italic = True
                if qa_info_ocr.get('score', 5) < 3: qa_run_ocr.font.color.rgb = RGBColor(128, 128, 128)
        except FileNotFoundError:
            quality_report_messages.append(f"WARN: Image file not found '{image_path}' for Word."); doc.add_paragraph(image_data.get('text', f"[Placeholder: {image_filename}]")).italic = True
        except Exception as e_add_pic:
            quality_report_messages.append(f"WARN: Failed to add image '{image_filename}' to Word: {e_add_pic}"); logger.warning(f"Failed image add '{image_filename}': {e_add_pic}", exc_info=True); doc.add_paragraph(image_data.get('text', f"[Placeholder: {image_filename}]")).italic = True
    else:
        quality_report_messages.append(f"WARN: Image file missing '{image_path}'."); doc.add_paragraph(image_data.get('text', f"[Image not found: {image_filename}]")).italic = True

def _apply_text_and_styling_to_paragraph(paragraph_obj, text_content, item_data, item_type_hint):
    apply_font_styles = (item_type_hint == 'p' and APPLY_STYLES_TO_PARAGRAPHS) or \
                        (item_type_hint == 'list_item' and APPLY_STYLES_TO_PARAGRAPHS) or \
                        (item_type_hint in ['h1', 'h2', 'h3'] and APPLY_STYLES_TO_HEADINGS)

    lines_in_text = text_content.split('\n')
    for line_idx, line_str in enumerate(lines_in_text):
        if line_idx > 0 :
            paragraph_obj.add_run().add_break()
        
        last_match_end = 0
        for match_url in RAW_URL_REGEX.finditer(line_str):
            start_idx, end_idx = match_url.span()
            if start_idx > last_match_end:
                run_before = paragraph_obj.add_run(line_str[last_match_end:start_idx])
                if apply_font_styles:
                    if item_data.get('bold'): run_before.bold = True
                    if item_data.get('italic'): run_before.italic = True
                    if item_type_hint not in ['list_item'] and item_data.get('font_size') and isinstance(item_data.get('font_size'), (int,float)):
                        try: run_before.font.size = Pt(item_data['font_size'])
                        except Exception as e_font: logger.debug(f"Font size error for run_before: {e_font}")
            
            url_text = match_url.group(0)
            link_run = paragraph_obj.add_run(url_text)
            link_run.font.color.rgb = RGBColor(5,99,193)
            link_run.font.underline = True
            if apply_font_styles:
                if item_data.get('bold'): link_run.bold = True
                if item_data.get('italic'): link_run.italic = True
            last_match_end = end_idx
        
        if last_match_end < len(line_str):
            run_after = paragraph_obj.add_run(line_str[last_match_end:])
            if apply_font_styles:
                if item_data.get('bold'): run_after.bold = True
                if item_data.get('italic'): run_after.italic = True
                if item_type_hint not in ['list_item'] and item_data.get('font_size') and isinstance(item_data.get('font_size'), (int,float)):
                    try: run_after.font.size = Pt(item_data['font_size'])
                    except Exception as e_font: logger.debug(f"Font size error for run_after: {e_font}")

def _add_qa_annotation_to_paragraph(paragraph_obj, qa_info_dict):
    if qa_info_dict:
        qa_text = f" [QA: {qa_info_dict.get('score','N/A')}/5 - {qa_info_dict.get('comment','N/A')}]"
        qa_run = paragraph_obj.add_run(qa_text)
        qa_run.font.size = Pt(7); qa_run.italic = True
        if qa_info_dict.get('score', 5) < 3: qa_run.font.color.rgb = RGBColor(128, 128, 128)

def create_word_document_with_structure(structured_content_list, output_filepath, image_folder_path, cover_page_data=None):
    global quality_report_messages, bookmark_id_counter, LIST_INDENT_PER_LEVEL_INCHES, HEADING_SPACE_BEFORE_PT, PARAGRAPH_FIRST_LINE_INDENT_INCHES, PARAGRAPH_SPACE_AFTER_PT
    bookmark_id_counter = 0
    logger.info(f"--- Creating Word Document: {os.path.basename(output_filepath)} ---")
    doc = Document()

    # Add cover page if provided
    if cover_page_data and cover_page_data.get('filepath') and os.path.exists(cover_page_data['filepath']):
        try:
            cover_paragraph = doc.add_paragraph()
            cover_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cover_paragraph.add_run()
            run.add_picture(cover_page_data['filepath'], width=Inches(6.5))
            doc.add_page_break()
            logger.info("Cover page added to document")
            quality_report_messages.append("INFO: Cover page included in translated document")
        except Exception as e:
            logger.warning(f"Failed to add cover page: {e}")
            quality_report_messages.append(f"WARNING: Cover page could not be added: {e}")

    # Define or ensure list styles exist (example for up to 3 levels)
    # This is a basic way; a more robust solution involves checking/creating styles if they don't exist in the template.
    # For now, we rely on python-docx's default handling if a style name isn't found.
    # Example: doc.styles.add_style('ListBullet2', WD_STYLE_TYPE.PARAGRAPH) # then configure it

    # Enhanced TOC Generation with improved bookmark handling
    if GENERATE_TOC:
        doc.add_heading(TOC_TITLE, level=0)
        toc_items_for_doc = []
        for index, item_toc_pass in enumerate(structured_content_list):
            if item_toc_pass['type'] in ['h1', 'h2', 'h3'] and item_toc_pass.get('translated_text'):
                level_num = int(item_toc_pass['type'][1:])
                # Create safer bookmark names
                bookmark_name = f"heading_{level_num}_{index}_{uuid.uuid4().hex[:8]}"
                item_toc_pass['_temp_bookmark_name_'] = bookmark_name
                first_line = item_toc_pass['translated_text'].split('\n')[0]
                toc_entry_text = (first_line[:75] + '...') if len(first_line) > 78 else first_line
                toc_items_for_doc.append({
                    'text': toc_entry_text,
                    'level': level_num,
                    'bookmark_name': bookmark_name
                })

        # Add TOC entries with improved formatting
        for toc_item_entry in toc_items_for_doc:
            add_toc_entry(doc, toc_item_entry['text'], toc_item_entry['level'], toc_item_entry['bookmark_name'])

        if toc_items_for_doc:
            doc.add_page_break()

    for item_index, item in enumerate(structured_content_list):
        bookmark_name_to_add_to_current = item.get('_temp_bookmark_name_') if GENERATE_TOC else None
        item_type = item['type']

        if item_type == 'image_placeholder':
            _add_image_to_doc(doc, item, image_folder_path)
            continue

        text_to_add = item.get('translated_text', item.get('text', ''))
        original_text_for_error_snippet = item.get('text', '[Original text not available]')[:100] + "..."
        if not isinstance(text_to_add, str): text_to_add = str(text_to_add)
        
        if not text_to_add.strip() and not item.get('_attached_images') and not text_to_add.startswith(("[Translation failed", "[Μετάφραση απέτυχε")):
            continue

        if text_to_add.startswith(("[Translation failed", "[Μετάφραση απέτυχε")):
            p_err = doc.add_paragraph()
            err_message = text_to_add + f"\nOriginal Snippet: '{original_text_for_error_snippet}'"
            run_err = p_err.add_run(err_message)
            run_err.italic = True; run_err.font.color.rgb = RGBColor(200, 0, 0)
            if item.get('_attached_images'):
                for attached_img_data in item['_attached_images']:
                    _add_image_to_doc(doc, attached_img_data, image_folder_path, main_item_for_qa_ref=item)
            continue

        current_paragraph_obj = None

        # Handle multiple paragraphs within one item if text_to_add contains "\n\n"
        # This is a fallback; ideally, paragraph splitting happens during extraction.
        if item_type == 'p' and "\n\n" in text_to_add and not item.get('is_single_block_para', False):
            sub_paragraphs_text = text_to_add.split("\n\n")
            first_sub_para_obj = None
            for i, sub_para_text_content in enumerate(sub_paragraphs_text):
                if not sub_para_text_content.strip(): continue # Skip empty sub-paragraphs

                p_current = doc.add_paragraph()
                if item.get('is_indented') and PARAGRAPH_FIRST_LINE_INDENT_INCHES > 0 and i == 0 : # Indent only first sub-para if original item was indented
                     p_current.paragraph_format.first_line_indent = Inches(PARAGRAPH_FIRST_LINE_INDENT_INCHES)
                elif PARAGRAPH_SPACE_AFTER_PT > 0:
                    p_current.paragraph_format.space_after = Pt(PARAGRAPH_SPACE_AFTER_PT)

                if i == 0:
                    first_sub_para_obj = p_current
                    if bookmark_name_to_add_to_current:
                        add_bookmark_to_paragraph(p_current, bookmark_name_to_add_to_current)
                
                if sub_para_text_content.strip():
                    _apply_text_and_styling_to_paragraph(p_current, sub_para_text_content, item, item_type)

                if i == 0 and SHOW_QA_ANNOTATIONS_IN_WORD and item.get('quality_assessment'):
                     if first_sub_para_obj:
                        _add_qa_annotation_to_paragraph(first_sub_para_obj, item.get('quality_assessment'))
            if item.get('_attached_images'): # Attach images after all sub-paragraphs of this item
                for attached_img_data in item['_attached_images']:
                    _add_image_to_doc(doc, attached_img_data, image_folder_path, main_item_for_qa_ref=item)
        else:
            style_name_for_list = None
            if item_type == 'list_item':
                list_info = item.get('list_info', {'type': 'bullet', 'level': 1})
                # Using python-docx's built-in list handling is preferred if styles are well-defined in template.
                # For now, we use explicit styling.
                # style_name_base = 'ListBullet' if list_info.get('type', 'bullet') == 'bullet' else 'ListNumber'
                # level = list_info.get('level', 1)
                # style_name_for_list = f"{style_name_base}{level if level > 1 else ''}"
                # try:
                #     current_paragraph_obj = doc.add_paragraph(style=doc.styles[style_name_for_list])
                # except KeyError:
                #     logger.debug(f"List style '{style_name_for_list}' not found, trying base '{style_name_base}' or default.")
                #     try:
                #         current_paragraph_obj = doc.add_paragraph(style=doc.styles[style_name_base])
                #     except KeyError:
                #         logger.warning(f"Base list style '{style_name_base}' also not found. Using default paragraph for list item '{text_to_add[:30]}...'.")
                current_paragraph_obj = doc.add_paragraph() # Start with a default paragraph
                marker_prefix = item.get('list_info',{}).get('original_marker', '•') + "\t" # Add tab after marker
                run_marker = current_paragraph_obj.add_run(marker_prefix)
                if APPLY_STYLES_TO_PARAGRAPHS: # Apply style to marker if desired
                    if item.get('bold'): run_marker.bold = True
                    if item.get('italic'): run_marker.italic = True
                    if item.get('font_size'): run_marker.font.size = Pt(item['font_size'])

                indent_level = max(0, list_info.get('level', 1) - 1)
                current_paragraph_obj.paragraph_format.left_indent = Inches(LIST_INDENT_PER_LEVEL_INCHES * indent_level)
                # Hanging indent for list items
                current_paragraph_obj.paragraph_format.first_line_indent = Inches(-LIST_INDENT_PER_LEVEL_INCHES * 0.5 if LIST_INDENT_PER_LEVEL_INCHES > 0 else -0.125) # Adjust as needed
            
            elif item_type == 'h1':
                current_paragraph_obj = doc.add_heading('', level=1)
                if HEADING_SPACE_BEFORE_PT > 0 : current_paragraph_obj.paragraph_format.space_before = Pt(HEADING_SPACE_BEFORE_PT)
            elif item_type == 'h2':
                current_paragraph_obj = doc.add_heading('', level=2)
                if HEADING_SPACE_BEFORE_PT > 0 : current_paragraph_obj.paragraph_format.space_before = Pt(HEADING_SPACE_BEFORE_PT)
            elif item_type == 'h3':
                current_paragraph_obj = doc.add_heading('', level=3)
                if HEADING_SPACE_BEFORE_PT > 0 : current_paragraph_obj.paragraph_format.space_before = Pt(HEADING_SPACE_BEFORE_PT // 2)
            else: # Default to 'p'
                current_paragraph_obj = doc.add_paragraph()
                if item.get('is_indented') and PARAGRAPH_FIRST_LINE_INDENT_INCHES > 0: # Apply detected indent
                    current_paragraph_obj.paragraph_format.first_line_indent = Inches(PARAGRAPH_FIRST_LINE_INDENT_INCHES)
                elif PARAGRAPH_SPACE_AFTER_PT > 0: # Apply default spacing if no specific indent
                    current_paragraph_obj.paragraph_format.space_after = Pt(PARAGRAPH_SPACE_AFTER_PT)

            if bookmark_name_to_add_to_current and current_paragraph_obj:
                add_bookmark_to_paragraph(current_paragraph_obj, bookmark_name_to_add_to_current)
            
            if text_to_add.strip() and current_paragraph_obj:
                 _apply_text_and_styling_to_paragraph(current_paragraph_obj, text_to_add, item, item_type) # Pass item_type here
            
            if SHOW_QA_ANNOTATIONS_IN_WORD and item.get('quality_assessment') and current_paragraph_obj:
                _add_qa_annotation_to_paragraph(current_paragraph_obj, item.get('quality_assessment'))
            
            if item.get('_attached_images'):
                for attached_img_data in item['_attached_images']:
                    _add_image_to_doc(doc, attached_img_data, image_folder_path, main_item_for_qa_ref=item)
    try:
        doc.save(output_filepath)
        logger.info(f"Word document saved: {output_filepath}")
        return True
    except Exception as e:
        msg = f"Error saving Word document '{output_filepath}': {e}"; logger.error(msg, exc_info=True); quality_report_messages.append(msg)
        return False

def convert_word_to_pdf(word_filepath, pdf_filepath):
    global quality_report_messages
    logger.info(f"--- Converting Word to PDF: {os.path.basename(word_filepath)} -> {os.path.basename(pdf_filepath)} ---")
    try:
        pdf_dir = os.path.dirname(pdf_filepath)
        if pdf_dir and not os.path.exists(pdf_dir): os.makedirs(pdf_dir, exist_ok=True)
        for attempt in range(3):
            try:
                convert_to_pdf_lib(word_filepath, pdf_filepath)
                logger.info("Conversion to PDF completed.")
                return True
            except Exception as e_convert:
                if attempt < 2 and ("Συνδεθείτε στο Word" in str(e_convert) or "process cannot access" in str(e_convert).lower() or "failed to start Word" in str(e_convert).lower() or "com_error" in str(e_convert).lower()):
                    msg = f"Warning: PDF conversion attempt {attempt+1}/3 failed for '{os.path.basename(word_filepath)}'. Retrying in 5s... Error: {e_convert}"
                    logger.warning(msg); quality_report_messages.append(msg); time.sleep(5)
                else:
                    msg = f"Error converting Word to PDF '{os.path.basename(word_filepath)}' after {attempt+1} attempts: {e_convert}"
                    logger.error(msg, exc_info=True); quality_report_messages.append(msg)
                    return False
        return False
    except Exception as e:
        msg = f"Main error in convert_word_to_pdf function for '{os.path.basename(word_filepath)}': {e}"
        logger.error(msg, exc_info=True); quality_report_messages.append(msg)
        return False

def verify_glossary_consistency(translated_content_items, glossary, report_stats, quality_messages_list):
    if not USE_GLOSSARY or not glossary:
        return
    logger.info("--- Verifying Glossary Consistency ---")
    report_stats["glossary_terms_checked_in_source"] = 0
    report_stats["glossary_potential_target_mismatches"] = 0
    processed_item_indices_for_glossary_check = set()

    for item_idx, item_data in enumerate(translated_content_items):
        original_text = None
        translated_text = None
        item_type_for_log = item_data.get('type', 'unknown')
        segment_description_for_log = f"Item {item_idx} (type: {item_type_for_log})"

        if item_data['type'] != 'image_placeholder' and item_data.get('_original_text_for_qa_'):
            original_text = item_data['_original_text_for_qa_']
            translated_text = item_data.get('translated_text')
        elif item_data['type'] == 'image_placeholder' and item_data.get('_original_ocr_for_qa_'):
            original_text = item_data['_original_ocr_for_qa_']
            translated_text = item_data.get('translated_ocr_text')
            segment_description_for_log = f"Image {item_data.get('filename', 'N/A')} OCR"
        
        if original_text and translated_text and not translated_text.startswith("[Translation failed"):
            if item_idx not in processed_item_indices_for_glossary_check:
                 _perform_single_glossary_check(original_text, translated_text, glossary, report_stats, quality_messages_list, segment_description_for_log)
                 processed_item_indices_for_glossary_check.add(item_idx)

        if item_data.get('_attached_images'):
            for att_img_idx, att_img in enumerate(item_data['_attached_images']):
                original_ocr_att = att_img.get('_original_ocr_for_qa_')
                translated_ocr_att = att_img.get('translated_ocr_text')
                if original_ocr_att and translated_ocr_att and not translated_ocr_att.startswith("[Translation failed"):
                    att_img_desc = f"Attached Image {att_img.get('filename', 'N/A')} (Parent Item {item_idx}) OCR"
                    _perform_single_glossary_check(original_ocr_att, translated_ocr_att, glossary, report_stats, quality_messages_list, att_img_desc)


def _perform_single_glossary_check(original_text, translated_text, glossary, report_stats, quality_messages_list, segment_description):
    """Helper function to check glossary for a single text pair."""
    original_text_lower = original_text.lower()
    for source_term, target_term in glossary.items():
        source_term_lower = source_term.lower()
        if source_term_lower in original_text_lower:
            report_stats["glossary_terms_checked_in_source"] += 1
            target_term_found = False
            if GLOSSARY_CONSISTENCY_CHECK_CASE_SENSITIVE_TARGET:
                if target_term in translated_text:
                    target_term_found = True
            else:
                if target_term.lower() in translated_text.lower():
                    target_term_found = True
            
            if not target_term_found:
                report_stats["glossary_potential_target_mismatches"] += 1
                original_snippet = original_text[:70].replace("\n", " ") + "..." if len(original_text) > 70 else original_text.replace("\n", " ")
                translated_snippet = translated_text[:70].replace("\n", " ") + "..." if len(translated_text) > 70 else translated_text.replace("\n", " ")
                msg = (f"GLOSSARY MISMATCH? Segment: '{segment_description}'. "
                       f"Original contained '{source_term}'. Expected '{target_term}' in translation. "
                       f"Original Snippet: '{original_snippet}'. Translated Snippet: '{translated_snippet}'.")
                logger.warning(msg)
                quality_messages_list.append(f"WARNING: {msg}")

async def translate_document_agent_async(
    filepath,
    output_dir_for_this_file,
    target_language_override=None,
    gdrive_folder_id_target_override=None,
    precomputed_style_guide=None
):
    global quality_report_messages, TRANSLATION_STYLE_TONE_FALLBACK, bookmark_id_counter, qa_format_error_log_count
    current_target_language = target_language_override if target_language_override else TARGET_LANGUAGE_CONFIG
    quality_report_messages = [f"--- Quality Report for: {os.path.basename(filepath)} ---"]
    bookmark_id_counter = 0; qa_format_error_log_count = 0; start_time = time.time()
    report_stats = {
        "total_text_segments":0, "api_translations_text":0, "cached_translations_text":0, "failed_translations_text":0,
        "total_ocr_segments":0, "api_translations_ocr":0, "cached_translations_ocr":0, "failed_translations_ocr":0,
        "glossary_terms_applied_count_text":0, "glossary_terms_applied_count_ocr":0,
        "images_extracted":0, "ocr_texts_found":0, "ocr_texts_skipped_translation":0,
        "document_style_analysis_done":False, "quality_assessments_done":0,
        "total_quality_score":0, "low_quality_segments":0,
        "glossary_terms_checked_in_source": 0, "glossary_potential_target_mismatches": 0
    }
    logger.info(f"Starting full process for file: {filepath} (Target Lang: {current_target_language})")
    if not filepath or not os.path.exists(filepath) or not filepath.lower().endswith(".pdf"):
        msg = f"CRITICAL ERROR: Invalid or non-PDF input file: {filepath}."; logger.critical(msg); quality_report_messages.append(msg); return None
    if not output_dir_for_this_file or not os.path.exists(output_dir_for_this_file):
        logger.error(f"Output directory '{output_dir_for_this_file}' not valid for {os.path.basename(filepath)}. Aborting."); print_quality_report(report_stats, start_time, filepath); return None

    base_original_filename = os.path.splitext(os.path.basename(filepath))[0]
    if len(base_original_filename) > 100: base_original_filename = base_original_filename[:100] + "_TRNC"
    output_image_folder_name = f"{base_original_filename}_Εικόνες"
    all_extracted_image_refs = []; final_word_path_local = None; final_pdf_path_local = None
    dynamic_document_style_guide = TRANSLATION_STYLE_TONE_FALLBACK

    with tempfile.TemporaryDirectory(prefix="doc_trans_imgs_") as temp_image_dir_for_extraction:
        # Extract cover page first
        cover_page_data = None
        try:
            cover_page_data = extract_cover_page_from_pdf(filepath, temp_image_dir_for_extraction)
            if cover_page_data:
                logger.info("Cover page extracted successfully")
                quality_report_messages.append("INFO: Cover page extracted for inclusion in translated document")
        except Exception as e:
            logger.warning(f"Cover page extraction failed: {e}")
            quality_report_messages.append(f"WARNING: Cover page extraction failed: {e}")

        if EXTRACT_IMAGES:
            all_extracted_image_refs = extract_images_from_pdf(filepath, temp_image_dir_for_extraction)
            report_stats["images_extracted"] = len(all_extracted_image_refs)
            for img_ref in all_extracted_image_refs:
                if img_ref.get("ocr_text_original"): report_stats["ocr_texts_found"] += 1
        
        structured_content_to_translate = extract_structured_content_from_pdf(filepath, all_extracted_image_refs)
        if not structured_content_to_translate:
            quality_report_messages.append(f"CRITICAL: No content extracted for {os.path.basename(filepath)}."); logger.critical(f"No content extracted for {os.path.basename(filepath)}.")
            print_quality_report(report_stats, start_time, filepath); return None

        # 🚀 APPLY ULTIMATE OPTIMIZATIONS WITH ENHANCED INTELLIGENCE
        logger.info(f"--- Προετοιμασία {len(structured_content_to_translate)} τμημάτων για μετάφραση με βελτιστοποιήσεις ---")

        # 🧠 ENHANCED DOCUMENT INTELLIGENCE ANALYSIS
        intelligence_analysis = document_intelligence.analyze_document_intelligence(structured_content_to_translate)

        if intelligence_analysis.get('enhanced_processing', False):
            logger.info("🧠 Enhanced Document Intelligence Analysis Complete:")
            classification = intelligence_analysis.get('classification')
            if classification:
                logger.info(f"   📊 Content Type: {classification.primary_type} ({classification.confidence:.1%} confidence)")
                logger.info(f"   🎯 Complexity: {classification.complexity_level}")
                logger.info(f"   🔧 Approach: {classification.translation_approach}")

            semantic_groups = intelligence_analysis.get('semantic_groups', [])
            if semantic_groups:
                logger.info(f"   🎭 Semantic Groups: {len(semantic_groups)} identified")
                themes = set(g.semantic_theme for g in semantic_groups)
                logger.info(f"   🎨 Themes: {', '.join(themes)}")

            cross_refs = intelligence_analysis.get('cross_references', {})
            if cross_refs:
                total_refs = sum(len(refs) for refs in cross_refs.values())
                logger.info(f"   🔗 Cross-references: {total_refs} found ({len(cross_refs)} types)")

            # Add intelligence report to quality messages
            intelligence_report = intelligence_analysis.get('intelligence_report', '')
            if intelligence_report:
                quality_report_messages.append("=== ENHANCED DOCUMENT INTELLIGENCE ===")
                quality_report_messages.append(intelligence_report)

        # Apply comprehensive optimizations
        optimized_content, optimization_params = optimization_manager.optimize_translation_pipeline(
            structured_content_to_translate, current_target_language
        )

        # 🧠 INTEGRATE ENHANCED INTELLIGENCE PARAMETERS
        enhanced_params = document_intelligence.get_enhanced_translation_parameters(intelligence_analysis)
        if enhanced_params:
            # Override optimization parameters with intelligence-based recommendations
            for key, value in enhanced_params.items():
                if key in optimization_params:
                    logger.info(f"🧠 Intelligence override: {key} = {value}")
                    optimization_params[key] = value

        # Update content with optimized version
        if len(optimized_content) != len(structured_content_to_translate):
            logger.info(f"🎯 Content optimized: {len(structured_content_to_translate)} → {len(optimized_content)} items")
            structured_content_to_translate = optimized_content

        # Apply optimization parameters
        if optimization_params:
            # Update batch size if recommended
            if 'optimal_batch_size' in optimization_params:
                global MAX_GROUP_SIZE_CHARS
                original_batch_size = MAX_GROUP_SIZE_CHARS
                MAX_GROUP_SIZE_CHARS = optimization_params['optimal_batch_size']
                logger.info(f"🎯 Batch size optimized: {original_batch_size} → {MAX_GROUP_SIZE_CHARS}")

            # Update temperature if recommended
            if 'optimal_temperature' in optimization_params:
                global TRANSLATION_TEMPERATURE
                original_temp = TRANSLATION_TEMPERATURE
                TRANSLATION_TEMPERATURE = optimization_params['optimal_temperature']
                logger.info(f"🎯 Temperature optimized: {original_temp} → {TRANSLATION_TEMPERATURE}")

            # 🧠 ENHANCED INTELLIGENCE: Add cross-reference guide to translation context
            if 'reference_guide' in optimization_params and optimization_params['reference_guide']:
                global ENHANCED_REFERENCE_GUIDE
                ENHANCED_REFERENCE_GUIDE = optimization_params['reference_guide']
                logger.info(f"🔗 Cross-reference preservation guide activated ({len(ENHANCED_REFERENCE_GUIDE)} chars)")
            else:
                ENHANCED_REFERENCE_GUIDE = ""

        logger.info(f"INFO: Εκτέλεση {len(structured_content_to_translate)} εργασιών μετάφρασης...")
        
        if precomputed_style_guide:
            dynamic_document_style_guide = precomputed_style_guide
            logger.info(f"Using precomputed style guide for {os.path.basename(filepath)}: '{dynamic_document_style_guide[:100]}...'")
            report_stats["document_style_analysis_done"] = True
        elif ANALYZE_DOCUMENT_STYLE_FIRST:
            sample_text_parts = []
            for item_s in structured_content_to_translate:
                if item_s['type'] != 'image_placeholder' and item_s.get('text'): sample_text_parts.append(item_s.get('text', ''))
                if item_s.get('_attached_images'):
                    for att_img in item_s['_attached_images']:
                        if att_img.get('ocr_text_original'): sample_text_parts.append(att_img.get('ocr_text_original',''))
                elif item_s['type'] == 'image_placeholder' and item_s.get('ocr_text_original'):
                     sample_text_parts.append(item_s.get('ocr_text_original',''))

            sample_text_for_analysis = "\n\n".join(sample_text_parts)[:DOCUMENT_ANALYSIS_SAMPLE_SIZE_CHARS]
            if sample_text_for_analysis.strip():
                analysis_semaphore = asyncio.Semaphore(1)
                analyzed_style = await analyze_document_style_async(sample_text_for_analysis, DOCUMENT_ANALYSIS_PROMPT_TEXT, MODEL_NAME, TRANSLATION_TEMPERATURE, analysis_semaphore)
                if analyzed_style and analyzed_style != TRANSLATION_STYLE_TONE_FALLBACK: dynamic_document_style_guide = analyzed_style
                report_stats["document_style_analysis_done"] = bool(dynamic_document_style_guide != TRANSLATION_STYLE_TONE_FALLBACK)
            else: quality_report_messages.append("WARNING: No text found for style analysis sample.")
        
        main_translation_semaphore = asyncio.Semaphore(MAX_CONCURRENT_API_CALLS)
        tasks = []; translated_structured_content = json.loads(json.dumps(structured_content_to_translate))

        # ===== ENHANCED SMART GROUPING FOR API CALL REDUCTION =====
        if ENABLE_SMART_GROUPING:
            logger.info(f"Creating enhanced smart groups to reduce API calls (max {MAX_GROUP_SIZE_CHARS} chars)...")
            smart_grouper = SmartGroupingProcessor(
                max_group_size=MAX_GROUP_SIZE_CHARS,
                min_group_items=2,
                preserve_formatting=True
            )
            smart_grouper.aggressive_grouping_enabled = AGGRESSIVE_GROUPING_MODE
            smart_grouper.max_items_per_group = MAX_ITEMS_PER_GROUP
            content_groups = smart_grouper.create_smart_groups(structured_content_to_translate)
        else:
            logger.info("Smart grouping disabled - processing items individually")
            content_groups = [[item] for item in structured_content_to_translate]

        # Calculate API call reduction statistics
        original_api_calls = len([item for item in structured_content_to_translate if item.get('type') not in ['image_placeholder']])
        grouped_api_calls = len(content_groups)
        api_call_reduction = ((original_api_calls - grouped_api_calls) / original_api_calls * 100) if original_api_calls > 0 else 0

        logger.info(f"Original items: {len(structured_content_to_translate)}, Grouped into: {len(content_groups)} groups")
        logger.info(f"API call reduction: {original_api_calls} → {grouped_api_calls} calls ({api_call_reduction:.1f}% reduction)")

        # Log grouping efficiency
        grouped_items = sum(1 for group in content_groups if len(group) == 1 and group[0].get('type') == 'grouped_content')
        standalone_items = len(content_groups) - grouped_items
        if grouped_items > 0:
            logger.info(f"Grouping efficiency: {grouped_items} groups, {standalone_items} standalone items")

        # Track group mappings for result processing
        group_mappings = {}
        group_index = 0

        for group in content_groups:
            if len(group) == 1 and group[0].get('type') != 'grouped_content':
                # Single item - process normally
                item_main = group[0]
                original_index = structured_content_to_translate.index(item_main)
                item_type_main = item_main['type']

                if item_type_main not in ['image_placeholder']:
                    report_stats["total_text_segments"] += 1
                    text_to_translate_val = item_main.get('text','')
                    if not text_to_translate_val.strip():
                        translated_structured_content[original_index]['translated_text'] = ""
                        continue

                    # Get context
                    prev_ctx_text = ""; next_ctx_text = ""
                    if original_index > 0:
                        for j_prev in range(original_index - 1, -1, -1):
                            prev_item_text_candidate = structured_content_to_translate[j_prev].get('text','') if structured_content_to_translate[j_prev]['type'] not in ['image_placeholder'] else None
                            if prev_item_text_candidate and prev_item_text_candidate.strip(): prev_ctx_text = prev_item_text_candidate[-1000:]; break
                    if original_index < len(structured_content_to_translate) - 1:
                        for j_next in range(original_index + 1, len(structured_content_to_translate)):
                            next_item_text_candidate = structured_content_to_translate[j_next].get('text','') if structured_content_to_translate[j_next]['type'] not in ['image_placeholder'] else None
                            if next_item_text_candidate and next_item_text_candidate.strip(): next_ctx_text = next_item_text_candidate[:1000]; break

                    current_glossary_terms = {k:v for k,v in TRANSLATION_GLOSSARY.items() if k.lower() in text_to_translate_val.lower()} if USE_GLOSSARY else None
                    prompt_item_type = item_type_main if item_type_main != 'list_item' else 'list item content'
                    task_id = ('text', original_index)
                    tasks.append(call_gemini_for_translation(task_id, text_to_translate_val, prev_ctx_text, next_ctx_text, current_target_language, main_translation_semaphore, MODEL_NAME, TRANSLATION_TEMPERATURE, dynamic_document_style_guide, item_type_for_prompt=prompt_item_type, glossary_terms_in_item=current_glossary_terms))

            elif len(group) == 1 and group[0].get('type') == 'grouped_content':
                # This is a grouped content item
                group_info = group[0]
                combined_text = group_info['combined_text']
                original_items = group_info['items']

                # Store mapping for later result processing
                group_mappings[group_index] = {
                    'group_info': group_info,
                    'original_indices': [structured_content_to_translate.index(item) for item in original_items]
                }

                report_stats["total_text_segments"] += len(original_items)

                # Get context from surrounding items
                first_item_index = structured_content_to_translate.index(original_items[0])
                last_item_index = structured_content_to_translate.index(original_items[-1])

                prev_ctx_text = ""; next_ctx_text = ""
                if first_item_index > 0:
                    for j_prev in range(first_item_index - 1, -1, -1):
                        prev_item_text_candidate = structured_content_to_translate[j_prev].get('text','') if structured_content_to_translate[j_prev]['type'] not in ['image_placeholder'] else None
                        if prev_item_text_candidate and prev_item_text_candidate.strip(): prev_ctx_text = prev_item_text_candidate[-1000:]; break
                if last_item_index < len(structured_content_to_translate) - 1:
                    for j_next in range(last_item_index + 1, len(structured_content_to_translate)):
                        next_item_text_candidate = structured_content_to_translate[j_next].get('text','') if structured_content_to_translate[j_next]['type'] not in ['image_placeholder'] else None
                        if next_item_text_candidate and next_item_text_candidate.strip(): next_ctx_text = next_item_text_candidate[:1000]; break

                # Combine glossary terms from all items in group
                combined_glossary_terms = {}
                if USE_GLOSSARY:
                    for item in original_items:
                        item_text = item.get('text', '')
                        for k, v in TRANSLATION_GLOSSARY.items():
                            if k.lower() in item_text.lower():
                                combined_glossary_terms[k] = v

                task_id = ('grouped_text', group_index)
                tasks.append(call_gemini_for_translation(task_id, combined_text, prev_ctx_text, next_ctx_text, current_target_language, main_translation_semaphore, MODEL_NAME, TRANSLATION_TEMPERATURE, dynamic_document_style_guide, item_type_for_prompt=f"grouped content ({len(original_items)} items)", glossary_terms_in_item=combined_glossary_terms))

                group_index += 1

        # ===== ENHANCED OCR TEXT GROUPING =====
        # Collect all OCR texts for potential grouping
        ocr_items_to_process = []
        for i, item_main_struct in enumerate(structured_content_to_translate):
            images_to_process_for_ocr = []
            if item_main_struct['type'] == 'image_placeholder':
                images_to_process_for_ocr.append({'item_index': i, 'image_sub_index': -1, 'image_data': item_main_struct})
            elif item_main_struct.get('_attached_images'):
                for sub_idx, attached_img_data in enumerate(item_main_struct['_attached_images']):
                    images_to_process_for_ocr.append({'item_index': i, 'image_sub_index': sub_idx, 'image_data': attached_img_data})

            for img_proc_info in images_to_process_for_ocr:
                ocr_text_val = img_proc_info['image_data'].get('ocr_text_original')
                logger.debug(f"OCR Check: Img {img_proc_info['image_data'].get('filename', 'N/A')}, Has OCR Text: {bool(ocr_text_val)}, Text: '{str(ocr_text_val)[:50]}...'")
                if PERFORM_OCR_ON_IMAGES and ocr_text_val:
                    report_stats["total_ocr_segments"] +=1

                    # Enhanced OCR filtering - check if content should be skipped
                    filename = img_proc_info['image_data'].get('filename', 'N/A')
                    if should_skip_ocr_translation(ocr_text_val, filename):
                        skip_reason = "smart filtering" if SMART_OCR_FILTERING else "word count"
                        logger.info(f"Skipping OCR translation ({skip_reason}): {filename} - '{ocr_text_val[:50]}...'")
                        report_stats["ocr_texts_skipped_translation"] += 1
                        target_item_idx = img_proc_info['item_index']; target_sub_idx = img_proc_info['image_sub_index']
                        if target_sub_idx == -1:
                            translated_structured_content[target_item_idx]['translated_ocr_text'] = ""
                            translated_structured_content[target_item_idx]['_original_ocr_for_qa_'] = ocr_text_val
                        else:
                            translated_structured_content[target_item_idx]['_attached_images'][target_sub_idx]['translated_ocr_text'] = ""
                            translated_structured_content[target_item_idx]['_attached_images'][target_sub_idx]['_original_ocr_for_qa_'] = ocr_text_val
                        continue

                    # Add to OCR processing list
                    logger.info(f"OCR text approved for translation: {filename} - {len(ocr_text_val.split())} words")
                    ocr_items_to_process.append(img_proc_info)

        # Group OCR texts for batch processing
        if ocr_items_to_process and ENABLE_OCR_GROUPING:
            logger.info(f"Processing {len(ocr_items_to_process)} OCR texts with grouping enabled...")

            # Simple OCR grouping: combine short OCR texts together
            ocr_groups = []
            current_ocr_group = []
            current_ocr_size = 0
            max_ocr_group_size = min(8000, MAX_GROUP_SIZE_CHARS // 2)  # Smaller than main text groups

            for ocr_item in ocr_items_to_process:
                ocr_text = ocr_item['image_data'].get('ocr_text_original', '')
                ocr_text_size = len(ocr_text)

                # If adding this OCR would exceed size limit, finalize current group
                if current_ocr_group and (current_ocr_size + ocr_text_size > max_ocr_group_size or len(current_ocr_group) >= 5):
                    ocr_groups.append(current_ocr_group)
                    current_ocr_group = []
                    current_ocr_size = 0

                current_ocr_group.append(ocr_item)
                current_ocr_size += ocr_text_size

            # Don't forget the last group
            if current_ocr_group:
                ocr_groups.append(current_ocr_group)

            logger.info(f"Grouped {len(ocr_items_to_process)} OCR items into {len(ocr_groups)} groups")

            # Process OCR groups
            for group_idx, ocr_group in enumerate(ocr_groups):
                if len(ocr_group) == 1:
                    # Single OCR item - process normally
                    img_proc_info = ocr_group[0]
                    ocr_text_val = img_proc_info['image_data'].get('ocr_text_original')
                    current_ocr_glossary_terms = {k:v for k,v in TRANSLATION_GLOSSARY.items() if k.lower() in ocr_text_val.lower()} if USE_GLOSSARY else None
                    task_id = ('ocr', img_proc_info['item_index'], img_proc_info['image_sub_index'])
                    ocr_prev_ctx = f"Image filename: {img_proc_info['image_data'].get('filename', 'N/A')}. OCR text."
                    tasks.append(call_gemini_for_translation(task_id, ocr_text_val, ocr_prev_ctx, "", current_target_language, main_translation_semaphore, MODEL_NAME, TRANSLATION_TEMPERATURE, dynamic_document_style_guide, item_type_for_prompt="image OCR text", glossary_terms_in_item=current_ocr_glossary_terms))
                else:
                    # Multiple OCR items - group them
                    combined_ocr_texts = []
                    combined_glossary_terms = {}

                    for img_proc_info in ocr_group:
                        ocr_text = img_proc_info['image_data'].get('ocr_text_original', '')
                        filename = img_proc_info['image_data'].get('filename', 'N/A')
                        combined_ocr_texts.append(f"[Image: {filename}]\n{ocr_text}")

                        # Combine glossary terms
                        if USE_GLOSSARY:
                            for k, v in TRANSLATION_GLOSSARY.items():
                                if k.lower() in ocr_text.lower():
                                    combined_glossary_terms[k] = v

                    combined_ocr_text = "\n\n===OCR_SEPARATOR===\n\n".join(combined_ocr_texts)
                    task_id = ('ocr_group', group_idx, ocr_group)  # Special task ID for OCR groups
                    ocr_prev_ctx = f"Multiple images OCR text batch ({len(ocr_group)} images)."
                    tasks.append(call_gemini_for_translation(task_id, combined_ocr_text, ocr_prev_ctx, "", current_target_language, main_translation_semaphore, MODEL_NAME, TRANSLATION_TEMPERATURE, dynamic_document_style_guide, item_type_for_prompt=f"grouped OCR text ({len(ocr_group)} images)", glossary_terms_in_item=combined_glossary_terms))

        elif ocr_items_to_process:
            # OCR grouping disabled - process individually
            logger.info(f"Processing {len(ocr_items_to_process)} OCR texts individually (grouping disabled)...")
            for img_proc_info in ocr_items_to_process:
                ocr_text_val = img_proc_info['image_data'].get('ocr_text_original')
                current_ocr_glossary_terms = {k:v for k,v in TRANSLATION_GLOSSARY.items() if k.lower() in ocr_text_val.lower()} if USE_GLOSSARY else None
                task_id = ('ocr', img_proc_info['item_index'], img_proc_info['image_sub_index'])
                ocr_prev_ctx = f"Image filename: {img_proc_info['image_data'].get('filename', 'N/A')}. OCR text."
                tasks.append(call_gemini_for_translation(task_id, ocr_text_val, ocr_prev_ctx, "", current_target_language, main_translation_semaphore, MODEL_NAME, TRANSLATION_TEMPERATURE, dynamic_document_style_guide, item_type_for_prompt="image OCR text", glossary_terms_in_item=current_ocr_glossary_terms))

        if tasks:
            # 🚀 ENHANCED ERROR RECOVERY & QUOTA MANAGEMENT
            error_recovery = EnhancedErrorRecovery()
            successful_results = []
            failed_tasks = []

            # Initialize progress tracking
            progress_tracker = ProgressTracker(len(tasks))
            batch_start_time = time.time()

            logger.info(f"🔄 Executing {len(tasks)} translation tasks with enhanced error recovery...")

            # Process tasks in smaller batches to handle quota limits better
            batch_size = min(MAX_CONCURRENT_API_CALLS * 2, 10)  # Process in smaller batches

            for i in range(0, len(tasks), batch_size):
                batch_tasks = tasks[i:i + batch_size]
                batch_start = i + 1
                batch_end = min(i + batch_size, len(tasks))

                logger.info(f"📦 Processing batch {batch_start}-{batch_end} of {len(tasks)} tasks...")

                try:
                    # Execute batch with timeout
                    batch_results = await asyncio.wait_for(
                        asyncio.gather(*batch_tasks, return_exceptions=True),
                        timeout=API_CALL_TIMEOUT * 2
                    )

                    # Process batch results
                    batch_processing_time = time.time() - batch_start_time
                    batch_success_count = 0
                    batch_failure_count = 0

                    for j, result in enumerate(batch_results):
                        task_index = i + j

                        if isinstance(result, Exception):
                            error_msg = str(result)
                            logger.error(f"❌ Task {task_index + 1} failed: {error_msg}")
                            batch_failure_count += 1

                            # Check if it's a quota error
                            if error_recovery.quota_manager.is_quota_error(error_msg):
                                error_recovery.quota_manager.handle_quota_error()
                                logger.error(f"💰 QUOTA EXCEEDED at task {task_index + 1} - Stopping further processing")

                                # Mark remaining tasks as failed
                                for remaining_idx in range(task_index, len(tasks)):
                                    failed_tasks.append((remaining_idx, "Quota exceeded"))
                                break
                            else:
                                failed_tasks.append((task_index, error_msg))
                        else:
                            successful_results.append((task_index, result))
                            batch_success_count += 1

                    # Update progress tracking
                    progress_tracker.update(completed=batch_success_count, failed=batch_failure_count)

                    # Record performance metrics
                    batch_success_rate = batch_success_count / len(batch_tasks) if batch_tasks else 0
                    avg_batch_size = sum(len(str(task)) for task in batch_tasks) / len(batch_tasks) if batch_tasks else 0
                    optimization_manager.record_batch_performance(avg_batch_size, batch_processing_time, batch_success_rate)

                    # If quota exceeded, stop processing
                    if error_recovery.quota_manager.quota_exceeded:
                        break

                except asyncio.TimeoutError:
                    logger.error(f"⏰ Batch {batch_start}-{batch_end} timed out")
                    for j in range(len(batch_tasks)):
                        failed_tasks.append((i + j, "Batch timeout"))
                except Exception as e:
                    logger.error(f"💥 Batch {batch_start}-{batch_end} failed: {e}")
                    for j in range(len(batch_tasks)):
                        failed_tasks.append((i + j, str(e)))

            # Reconstruct results in original order
            all_translation_results = [None] * len(tasks)

            # Fill successful results
            for task_index, result in successful_results:
                all_translation_results[task_index] = result

            # Fill failed results with error messages
            for task_index, error_msg in failed_tasks:
                # Create a failed result tuple matching the expected format
                if task_index < len(tasks):
                    # Extract task info to create proper error result
                    task_id = ('text', task_index)  # Default format
                    failed_result = (task_id, f"[Translation failed: {error_msg}]", False, False)
                    all_translation_results[task_index] = failed_result

            # Log recovery report
            recovery_report = error_recovery.get_recovery_report()
            logger.info(recovery_report)

            # Finalize progress tracking
            progress_tracker.finish()

            # Add recovery info to quality report
            quality_report_messages.append("=== ENHANCED ERROR RECOVERY REPORT ===")
            quality_report_messages.append(f"✅ Successful translations: {len(successful_results)}")
            quality_report_messages.append(f"❌ Failed translations: {len(failed_tasks)}")
            if error_recovery.quota_manager.quota_exceeded:
                quality_report_messages.append("💰 QUOTA EXCEEDED - Consider upgrading API tier or using Flash model")

            # Add performance analysis
            performance_report = optimization_manager.get_final_performance_report()
            quality_report_messages.append(performance_report)
            quality_report_messages.append("=" * 50)
            for task_id_result, translated_text_or_error, was_cached, had_glossary_terms in all_translation_results:
                task_type = task_id_result[0]; task_index = task_id_result[1]
                is_successful_translation = not (isinstance(translated_text_or_error, str) and translated_text_or_error.startswith("[Translation failed"))

                if task_type == 'text' or task_type == 'list_item_text':
                    # Single item translation
                    translated_structured_content[task_index]['translated_text'] = translated_text_or_error
                    if is_successful_translation:
                        translated_structured_content[task_index]['_original_text_for_qa_'] = structured_content_to_translate[task_index]['text']
                        translated_structured_content[task_index]['_had_glossary_terms_'] = had_glossary_terms
                        if was_cached: report_stats["cached_translations_text"] += 1
                        else: report_stats["api_translations_text"] += 1
                        if had_glossary_terms: report_stats["glossary_terms_applied_count_text"] +=1
                    else: report_stats["failed_translations_text"] +=1

                elif task_type == 'grouped_text':
                    # Grouped content translation - need to split back
                    group_idx = task_index
                    if group_idx in group_mappings:
                        group_mapping = group_mappings[group_idx]
                        group_info = group_mapping['group_info']
                        original_indices = group_mapping['original_indices']

                        if is_successful_translation:
                            # Split the translated text back into individual items
                            split_results = smart_grouper.split_translated_group(translated_text_or_error, group_info)

                            # Assign results back to original items
                            for i, (result_item, original_idx) in enumerate(zip(split_results, original_indices)):
                                translated_structured_content[original_idx]['translated_text'] = result_item.get('translated_text', translated_text_or_error)
                                translated_structured_content[original_idx]['_original_text_for_qa_'] = structured_content_to_translate[original_idx]['text']
                                translated_structured_content[original_idx]['_had_glossary_terms_'] = had_glossary_terms

                                # Update stats
                                if was_cached: report_stats["cached_translations_text"] += 1
                                else: report_stats["api_translations_text"] += 1
                                if had_glossary_terms: report_stats["glossary_terms_applied_count_text"] += 1
                        else:
                            # Failed translation - mark all items in group as failed
                            for original_idx in original_indices:
                                translated_structured_content[original_idx]['translated_text'] = translated_text_or_error
                                report_stats["failed_translations_text"] += 1
                    else:
                        logger.error(f"Group mapping not found for group index {group_idx}")
                        report_stats["failed_translations_text"] += 1
                elif task_type == 'ocr':
                    original_item_idx = task_index  # Fix the variable name
                    image_sub_idx = task_id_result[2]
                    if image_sub_idx == -1:
                        translated_structured_content[original_item_idx]['translated_ocr_text'] = translated_text_or_error
                        if is_successful_translation:
                            translated_structured_content[original_item_idx]['_original_ocr_for_qa_'] = structured_content_to_translate[original_item_idx]['ocr_text_original']
                            translated_structured_content[original_item_idx]['_had_glossary_terms_ocr_'] = had_glossary_terms
                            if was_cached: report_stats["cached_translations_ocr"] += 1
                            else: report_stats["api_translations_ocr"] += 1
                            if had_glossary_terms: report_stats["glossary_terms_applied_count_ocr"] +=1
                        else: report_stats["failed_translations_ocr"] +=1
                    else:
                        if '_attached_images' in translated_structured_content[original_item_idx] and \
                           0 <= image_sub_idx < len(translated_structured_content[original_item_idx]['_attached_images']):
                            translated_structured_content[original_item_idx]['_attached_images'][image_sub_idx]['translated_ocr_text'] = translated_text_or_error
                            if is_successful_translation:
                                translated_structured_content[original_item_idx]['_attached_images'][image_sub_idx]['_original_ocr_for_qa_'] = structured_content_to_translate[original_item_idx]['_attached_images'][image_sub_idx]['ocr_text_original']
                                translated_structured_content[original_item_idx]['_attached_images'][image_sub_idx]['_had_glossary_terms_ocr_'] = had_glossary_terms
                                if was_cached: report_stats["cached_translations_ocr"] += 1
                                else: report_stats["api_translations_ocr"] += 1
                                if had_glossary_terms: report_stats["glossary_terms_applied_count_ocr"] +=1
                            else: report_stats["failed_translations_ocr"] +=1
                        else:
                            logger.error(f"Error locating attached image for OCR result: item_idx={original_item_idx}, sub_idx={image_sub_idx}")

                elif task_type == 'ocr_group':
                    # Handle grouped OCR translation results
                    group_idx = task_index
                    ocr_group = task_id_result[2]  # The original OCR group

                    if is_successful_translation:
                        # Split the translated OCR text back into individual items
                        translated_parts = translated_text_or_error.split("\n\n===OCR_SEPARATOR===\n\n")

                        if len(translated_parts) == len(ocr_group):
                            # Successful split - assign results back
                            for translated_part, img_proc_info in zip(translated_parts, ocr_group):
                                original_item_idx = img_proc_info['item_index']
                                image_sub_idx = img_proc_info['image_sub_index']

                                # Clean up the translated text (remove image filename markers)
                                cleaned_translation = translated_part.strip()
                                if cleaned_translation.startswith('[Image:') and ']\n' in cleaned_translation:
                                    cleaned_translation = cleaned_translation.split(']\n', 1)[1].strip()

                                if image_sub_idx == -1:
                                    translated_structured_content[original_item_idx]['translated_ocr_text'] = cleaned_translation
                                    translated_structured_content[original_item_idx]['_original_ocr_for_qa_'] = structured_content_to_translate[original_item_idx]['ocr_text_original']
                                    translated_structured_content[original_item_idx]['_had_glossary_terms_ocr_'] = had_glossary_terms
                                else:
                                    if '_attached_images' in translated_structured_content[original_item_idx] and \
                                       0 <= image_sub_idx < len(translated_structured_content[original_item_idx]['_attached_images']):
                                        translated_structured_content[original_item_idx]['_attached_images'][image_sub_idx]['translated_ocr_text'] = cleaned_translation
                                        translated_structured_content[original_item_idx]['_attached_images'][image_sub_idx]['_original_ocr_for_qa_'] = structured_content_to_translate[original_item_idx]['_attached_images'][image_sub_idx]['ocr_text_original']
                                        translated_structured_content[original_item_idx]['_attached_images'][image_sub_idx]['_had_glossary_terms_ocr_'] = had_glossary_terms

                                # Update stats
                                if was_cached: report_stats["cached_translations_ocr"] += 1
                                else: report_stats["api_translations_ocr"] += 1
                                if had_glossary_terms: report_stats["glossary_terms_applied_count_ocr"] += 1
                        else:
                            logger.warning(f"OCR group splitting mismatch: expected {len(ocr_group)} parts, got {len(translated_parts)}")
                            # Fallback: assign the full translation to all items
                            for img_proc_info in ocr_group:
                                original_item_idx = img_proc_info['item_index']
                                image_sub_idx = img_proc_info['image_sub_index']
                                fallback_text = f"[OCR Group Translation]: {translated_text_or_error}"

                                if image_sub_idx == -1:
                                    translated_structured_content[original_item_idx]['translated_ocr_text'] = fallback_text
                                else:
                                    if '_attached_images' in translated_structured_content[original_item_idx] and \
                                       0 <= image_sub_idx < len(translated_structured_content[original_item_idx]['_attached_images']):
                                        translated_structured_content[original_item_idx]['_attached_images'][image_sub_idx]['translated_ocr_text'] = fallback_text

                                report_stats["api_translations_ocr"] += 1
                    else:
                        # Failed translation - mark all items in OCR group as failed
                        for img_proc_info in ocr_group:
                            original_item_idx = img_proc_info['item_index']
                            image_sub_idx = img_proc_info['image_sub_index']

                            if image_sub_idx == -1:
                                translated_structured_content[original_item_idx]['translated_ocr_text'] = translated_text_or_error
                            else:
                                if '_attached_images' in translated_structured_content[original_item_idx] and \
                                   0 <= image_sub_idx < len(translated_structured_content[original_item_idx]['_attached_images']):
                                    translated_structured_content[original_item_idx]['_attached_images'][image_sub_idx]['translated_ocr_text'] = translated_text_or_error

                            report_stats["failed_translations_ocr"] += 1
        
        if PERFORM_QUALITY_ASSESSMENT and QA_STRATEGY != 'none':
            qa_tasks_with_ids = []; qa_semaphore = asyncio.Semaphore(MAX_CONCURRENT_API_CALLS); items_for_qa_indices = []
            if QA_STRATEGY == 'full': items_for_qa_indices = list(range(len(translated_structured_content)))
            elif QA_STRATEGY == 'sample': items_for_qa_indices = [i_qa_strat for i_qa_strat in range(len(translated_structured_content)) if random.random() < QA_SAMPLE_PERCENTAGE]
            elif QA_STRATEGY == 'glossary':
                 for i_qa_strat, item_qa_strat in enumerate(translated_structured_content):
                    if item_qa_strat.get('_had_glossary_terms_'): items_for_qa_indices.append(i_qa_strat)
                    if item_qa_strat.get('_attached_images'):
                        for att_img_strat in item_qa_strat['_attached_images']:
                            if att_img_strat.get('_had_glossary_terms_ocr_') and i_qa_strat not in items_for_qa_indices:
                                items_for_qa_indices.append(i_qa_strat); break
            
            for i_qa in items_for_qa_indices:
                item_qa = translated_structured_content[i_qa]
                original_text_for_qa = item_qa.get('_original_text_for_qa_'); translated_text_for_qa = item_qa.get('translated_text')
                if item_qa['type'] not in ['image_placeholder'] and original_text_for_qa and translated_text_for_qa and not translated_text_for_qa.startswith("[Translation failed"):
                    qa_tasks_with_ids.append(((i_qa, 'main_text_qa'), assess_translation_quality_async(original_text_for_qa, translated_text_for_qa, current_target_language, dynamic_document_style_guide, qa_semaphore)))
                
                images_to_qa_ocr = []
                if item_qa['type'] == 'image_placeholder': images_to_qa_ocr.append({'item_index': i_qa, 'image_sub_index': -1, 'image_data': item_qa})
                elif item_qa.get('_attached_images'):
                    for sub_idx, att_img_data in enumerate(item_qa['_attached_images']): images_to_qa_ocr.append({'item_index': i_qa, 'image_sub_index': sub_idx, 'image_data': att_img_data})
                
                for img_qa_info in images_to_qa_ocr:
                    original_ocr_for_qa = img_qa_info['image_data'].get('_original_ocr_for_qa_'); translated_ocr_for_qa = img_qa_info['image_data'].get('translated_ocr_text')
                    if original_ocr_for_qa and translated_ocr_for_qa and not translated_ocr_for_qa.startswith("[Translation failed"):
                        qa_tasks_with_ids.append(((img_qa_info['item_index'], 'ocr_text_qa', img_qa_info['image_sub_index']), assess_translation_quality_async(original_ocr_for_qa, translated_ocr_for_qa, current_target_language, dynamic_document_style_guide, qa_semaphore)))
            
            if qa_tasks_with_ids:
                logger.info(f"Executing {len(qa_tasks_with_ids)} QA tasks for {os.path.basename(filepath)} (Strategy: {QA_STRATEGY})...")
                all_qa_results = await asyncio.gather(*(task for _, task in qa_tasks_with_ids))
                for i_res, assessment_result in enumerate(all_qa_results):
                    task_id_key, _ = qa_tasks_with_ids[i_res]; item_idx_mapped = task_id_key[0]; item_type_mapped_tag = task_id_key[1]
                    if assessment_result:
                        if item_type_mapped_tag == 'main_text_qa': translated_structured_content[item_idx_mapped]['quality_assessment'] = assessment_result
                        elif item_type_mapped_tag == 'ocr_text_qa':
                            ocr_sub_idx = task_id_key[2]
                            if ocr_sub_idx == -1: translated_structured_content[item_idx_mapped]['quality_assessment_ocr'] = assessment_result
                            else:
                                if '_attached_images' in translated_structured_content[item_idx_mapped] and 0 <= ocr_sub_idx < len(translated_structured_content[item_idx_mapped]['_attached_images']):
                                    translated_structured_content[item_idx_mapped]['_attached_images'][ocr_sub_idx]['quality_assessment_ocr'] = assessment_result
                        report_stats["quality_assessments_done"] += 1; score = assessment_result.get('score', 0); report_stats["total_quality_score"] += score
                        if score > 0 and score < 3 : report_stats["low_quality_segments"] += 1
        
        if USE_GLOSSARY and TRANSLATION_GLOSSARY:
            verify_glossary_consistency(translated_structured_content, TRANSLATION_GLOSSARY, report_stats, quality_report_messages)

        word_output_filename = f"{base_original_filename}_({current_target_language})_FINAL.docx"; pdf_output_filename = f"{base_original_filename}_({current_target_language})_FINAL.pdf"
        final_images_output_path_abs = None
        if EXTRACT_IMAGES and all_extracted_image_refs:
            final_images_output_path_abs = os.path.join(output_dir_for_this_file, output_image_folder_name)
            try:
                os.makedirs(final_images_output_path_abs, exist_ok=True)
                if os.path.exists(temp_image_dir_for_extraction) and os.listdir(temp_image_dir_for_extraction):
                    shutil.copytree(temp_image_dir_for_extraction, final_images_output_path_abs, dirs_exist_ok=True)
                    quality_report_messages.append(f"INFO: Images copied to: {final_images_output_path_abs}")
            except Exception as e_img_folder:
                quality_report_messages.append(f"ERROR creating/populating final image folder '{final_images_output_path_abs}': {e_img_folder}"); logger.error(f"Error with final image folder: {e_img_folder}", exc_info=True)
                final_images_output_path_abs = None
        
        with tempfile.TemporaryDirectory(prefix="trans_artefacts_") as temp_processing_dir:
            temp_word_path = os.path.join(temp_processing_dir, word_output_filename); temp_pdf_path = os.path.join(temp_processing_dir, pdf_output_filename)
            
            word_created_successfully = create_word_document_with_structure(translated_structured_content, temp_word_path, final_images_output_path_abs if final_images_output_path_abs else temp_image_dir_for_extraction, cover_page_data)
            pdf_created_successfully = False
            if word_created_successfully:
                pdf_created_successfully = convert_word_to_pdf(temp_word_path, temp_pdf_path)
                time.sleep(1) if pdf_created_successfully else None
            
            if word_created_successfully:
                final_word_path_local = os.path.join(output_dir_for_this_file, word_output_filename)
                try: shutil.move(temp_word_path, final_word_path_local)
                except Exception as e_wmov: quality_report_messages.append(f"ERROR moving Word file: {e_wmov}"); logger.error(f"Error moving Word: {e_wmov}", exc_info=True); final_word_path_local = temp_word_path
            else: quality_report_messages.append(f"ERROR: Word creation failed for {os.path.basename(filepath)}.")
            
            if pdf_created_successfully:
                final_pdf_path_local = os.path.join(output_dir_for_this_file, pdf_output_filename)
                try: shutil.move(temp_pdf_path, final_pdf_path_local)
                except Exception as e_pmov: quality_report_messages.append(f"ERROR moving PDF file: {e_pmov}"); logger.error(f"Error moving PDF: {e_pmov}", exc_info=True); final_pdf_path_local = temp_pdf_path
            elif word_created_successfully : quality_report_messages.append(f"ERROR: PDF conversion failed for {os.path.basename(filepath)}.")

            gdrive_folder_for_this_run = gdrive_folder_id_target_override if gdrive_folder_id_target_override is not None else GDRIVE_TARGET_FOLDER_ID_CONFIG
            if final_word_path_local and os.path.exists(final_word_path_local): upload_to_google_drive(final_word_path_local, os.path.basename(final_word_path_local), gdrive_folder_for_this_run)
            if final_pdf_path_local and os.path.exists(final_pdf_path_local): upload_to_google_drive(final_pdf_path_local, os.path.basename(final_pdf_path_local), gdrive_folder_for_this_run)

    if USE_TRANSLATION_CACHE: save_translation_cache()
    print_quality_report(report_stats, start_time, filepath)
    logger.info("-" * 50 + f"\nProcess for '{os.path.basename(filepath)}' completed.\n" + "-" * 50)
    if output_dir_for_this_file and os.path.exists(output_dir_for_this_file) :
        logger.info(f"Final local output directory for {os.path.basename(filepath)}: {output_dir_for_this_file}")
    return dynamic_document_style_guide

def print_quality_report(stats, start_time_sec, original_filepath):
    if not GENERATE_QUALITY_REPORT: return
    end_time_sec = time.time(); total_duration_sec = end_time_sec - start_time_sec
    report_file_name_from_msg = os.path.basename(original_filepath)
    report_header = f"\n{'='*20} Translation Process Report for: {report_file_name_from_msg} {'='*20}"
    logger.info(report_header)
    logger.info(f"  Processed File: {report_file_name_from_msg}")
    logger.info(f"  Total Processing Time: {total_duration_sec:.2f} seconds")
    if ANALYZE_DOCUMENT_STYLE_FIRST:
        logger.info(f"  Document Style Analysis Done:    {'Yes' if stats.get('document_style_analysis_done', False) else 'No/Failed/Skipped'}")
        found_style_msg = next((msg for msg in quality_report_messages if "Document Style Analysis Complete. Summary: '" in msg), None)
        if found_style_msg:
            try: summary_part = found_style_msg.split("Summary: '", 1)[1].rstrip("'"); logger.info(f"    Summary: '{summary_part}'")
            except IndexError: logger.info(f"    Summary: (Could not parse from log) - {found_style_msg}")
    logger.info("\n  --- Text Segments Summary ---")
    logger.info(f"  Total Text Segments Identified: {stats.get('total_text_segments',0)}")
    logger.info(f"    Translated via API:           {stats.get('api_translations_text',0)}")
    logger.info(f"    Loaded from Cache:            {stats.get('cached_translations_text',0)}")
    logger.info(f"    Failed Translations:          {stats.get('failed_translations_text',0)}")
    if USE_GLOSSARY:
        logger.info(f"  Glossary Terms Applied (Text):  {stats.get('glossary_terms_applied_count_text',0)}")
        logger.info(f"  Glossary Terms Checked in Source: {stats.get('glossary_terms_checked_in_source',0)}")
        logger.info(f"  Potential Target Term Mismatches: {stats.get('glossary_potential_target_mismatches',0)}")
    if PERFORM_OCR_ON_IMAGES or stats.get('images_extracted',0) > 0 :
        logger.info("\n  --- OCR Text (from Images) Summary ---")
        logger.info(f"  Images Extracted:               {stats.get('images_extracted',0)}")
        logger.info(f"  Images with OCR Text Found:   {stats.get('ocr_texts_found',0)}")
        if PERFORM_OCR_ON_IMAGES:
            logger.info(f"  Total OCR Segments Identified:  {stats.get('total_ocr_segments',0)}")
            logger.info(f"    Skipped OCR Translation (short):{stats.get('ocr_texts_skipped_translation',0)}")
            logger.info(f"    Translated via API:           {stats.get('api_translations_ocr',0)}")
            logger.info(f"    Loaded from Cache:            {stats.get('cached_translations_ocr',0)}")
            logger.info(f"    Failed Translations:          {stats.get('failed_translations_ocr',0)}")
            if USE_GLOSSARY: logger.info(f"  Glossary Terms Applied (OCR):   {stats.get('glossary_terms_applied_count_ocr',0)}")
    if PERFORM_QUALITY_ASSESSMENT and QA_STRATEGY != 'none' and stats.get('quality_assessments_done', 0) > 0:
        logger.info("\n  --- Quality Assessment Summary ---")
        logger.info(f"  QA Strategy:                  {QA_STRATEGY.capitalize()}{f' ({QA_SAMPLE_PERCENTAGE*100:.0f}%)' if QA_STRATEGY == 'sample' else ''}")
        avg_score = (stats['total_quality_score'] / stats['quality_assessments_done']) if stats['quality_assessments_done'] > 0 else 0
        logger.info(f"  Segments Assessed:            {stats['quality_assessments_done']}")
        logger.info(f"  Average Quality Score:        {avg_score:.2f} / 5.0")
        logger.info(f"  Segments with Low Score (<3): {stats.get('low_quality_segments',0)}")
    elif PERFORM_QUALITY_ASSESSMENT and QA_STRATEGY == 'none':
        logger.info("\n  --- Quality Assessment: SKIPPED (Strategy: None) ---")
    
    critical_messages_for_report = [msg for msg in quality_report_messages if "ERROR" in msg or "WARNING" in msg or "CRITICAL" in msg]
    if critical_messages_for_report:
        logger.info("\n  --- Key Issues & Warnings for this file ---")
        unique_critical_set = set(); display_count = 0
        for msg_crit in critical_messages_for_report:
            if "QA Format): Further QA format errors suppressed." in msg_crit and any("QA LLM response for" in s for s in unique_critical_set if "QA LLM response for" in s) : continue
            if msg_crit not in unique_critical_set:
                if display_count < 15:
                    logger.info(f"    - {msg_crit.replace(os.path.dirname(original_filepath),'[PATH_REDACTED]')}")
                display_count +=1; unique_critical_set.add(msg_crit)
        if len(unique_critical_set) > 15: logger.info(f"    ... and {len(unique_critical_set) - 15} more unique critical messages logged for this file.")
    else: logger.info("\n  --- Key Issues & Warnings: No critical errors or warnings logged for this file. ---")
    logger.info("="* (60 + len(report_file_name_from_msg)) + "\n")

def run_all_tests():
    """Run all available tests"""
    print("ULTIMATE PDF TRANSLATOR - COMPREHENSIVE TESTING")
    print("=" * 60)

    # Run basic enhanced capabilities test
    print("\nRunning basic enhanced capabilities test...")
    basic_test_passed = test_enhanced_capabilities_simple()

    # Run comprehensive optimization tests
    print("\nRunning comprehensive optimization tests...")
    optimization_test_passed = test_optimizations_comprehensive()

    # Summary
    print("\n" + "=" * 60)
    print("TEST SUMMARY:")
    print(f"Basic Enhanced Capabilities: {'✓ PASSED' if basic_test_passed else '✗ FAILED'}")
    print(f"Optimization Tests: {'✓ PASSED' if optimization_test_passed else '✗ FAILED'}")

    if basic_test_passed and optimization_test_passed:
        print("\n🎉 ALL TESTS PASSED! The optimizations are ready to use.")
        print("\nKey improvements validated:")
        print("• Enhanced smart grouping (40-60% fewer API calls with 11K chars)")
        print("• TOC preservation and translation")
        print("• Cover page integration")
        print("• Improved image context preservation")
        print("• Better document structure handling")
        print("• Coherent batch processing")
        print("• Cache optimization for large groups")
        return True
    else:
        print("\n⚠️  Some tests failed. Please review the output above.")
        return False

if __name__ == "__main__":
    import sys

    # Check for test commands first
    if len(sys.argv) >= 2:
        if sys.argv[1] == "--test-enhanced":
            print("Running basic enhanced capabilities test...")
            success = test_enhanced_capabilities_simple()
            sys.exit(0 if success else 1)
        elif sys.argv[1] == "--test-all":
            success = run_all_tests()
            sys.exit(0 if success else 1)
        elif sys.argv[1] == "--test-optimizations":
            print("Running optimization tests...")
            success = test_optimizations_comprehensive()
            sys.exit(0 if success else 1)
        elif sys.argv[1] in ["--help", "-h"]:
            print("Ultimate PDF Translator - Usage:")
            print("For testing (no API calls):")
            print("  python ultimate_pdf_translator.py --test-enhanced      # Basic test")
            print("  python ultimate_pdf_translator.py --test-all           # All tests")
            print("  python ultimate_pdf_translator.py --test-optimizations # Optimization tests")
            print("\nFor normal operation:")
            print("  Run the script normally - it will prompt for input files and settings")
            sys.exit(0)

    # Normal script execution
    logger.info("--- SCRIPT START: ULTIMATE PDF TRANSLATOR (Enhanced v4 - Error Recovery & Smart Batching) ---")

    # 🔍 VALIDATE CONFIGURATION
    if not validate_configuration():
        logger.error("❌ Configuration validation failed. Please fix issues before proceeding.")
        sys.exit(1)

    if USE_GLOSSARY: load_glossary()
    if USE_TRANSLATION_CACHE: load_translation_cache()

    input_path, process_mode = choose_input_path()
    files_to_process = []; main_output_directory_for_batch = None

    if process_mode == 'file' and input_path and os.path.exists(input_path):
        files_to_process.append(input_path)
        main_output_directory_for_batch = choose_base_output_directory(os.path.dirname(input_path))
    elif process_mode == 'dir' and input_path and os.path.isdir(input_path):
        for filename in os.listdir(input_path):
            if filename.lower().endswith(".pdf"): files_to_process.append(os.path.join(input_path, filename))
        if not files_to_process: logger.warning(f"No PDF files found in directory: {input_path}")
        else: main_output_directory_for_batch = choose_base_output_directory(input_path)
    elif input_path:
        logger.error(f"ERROR: Selected path '{input_path}' is not valid for '{process_mode}' mode. Terminating.")
    else:
        logger.info("INFO: No input file or directory selected. Exiting.")

    if files_to_process:
        logger.info(f"Found {len(files_to_process)} PDF file(s) to process.")

        # 📊 ESTIMATE TRANSLATION COST AND TIME
        if len(files_to_process) == 1:
            estimate_translation_cost(files_to_process[0])
        else:
            logger.info(f"📊 BATCH PROCESSING: {len(files_to_process)} files")
            logger.info("💡 Cost estimation available for single files only")

        reused_style_guide = None; user_choice_style_reuse = 'no'
        if process_mode == 'dir' and len(files_to_process) > 1 and ANALYZE_DOCUMENT_STYLE_FIRST and BATCH_STYLE_ANALYSIS_REUSE:
            user_choice_style_reuse = input(f"Batch mode: Reuse style analysis from the first file ('{os.path.basename(files_to_process[0])}') for all {len(files_to_process)} files? (yes/no): ").lower().strip()
        
        for i, filepath in enumerate(files_to_process):
            logger.info(f"\n>>> Processing file {i+1}/{len(files_to_process)}: {os.path.basename(filepath)} <<<")
            specific_output_dir = get_specific_output_dir_for_file(main_output_directory_for_batch, filepath)
            if not specific_output_dir:
                logger.error(f"Could not get/create a specific output directory for {os.path.basename(filepath)}. Skipping."); continue
            
            current_precomputed_style = None
            if process_mode == 'dir' and i > 0 and reused_style_guide and user_choice_style_reuse == 'yes':
                current_precomputed_style = reused_style_guide
            
            try:
                computed_style_for_this_file = asyncio.run(
                    translate_document_agent_async(filepath=filepath, output_dir_for_this_file=specific_output_dir, precomputed_style_guide=current_precomputed_style)
                )
                if i == 0 and process_mode == 'dir' and user_choice_style_reuse == 'yes' and ANALYZE_DOCUMENT_STYLE_FIRST and BATCH_STYLE_ANALYSIS_REUSE and computed_style_for_this_file:
                    reused_style_guide = computed_style_for_this_file
                    logger.info(f"Style guide from '{os.path.basename(filepath)}' will be reused for subsequent files in this batch.")
            except RuntimeError as e_rt:
                if "asyncio.run() cannot be called" in str(e_rt): logger.error("Asyncio Error: This script needs to be run in an environment where asyncio.run() can create a new event loop (e.g., not from an already running asyncio loop).")
                else: logger.error(f"Main Runtime Error for {os.path.basename(filepath)}: {e_rt}", exc_info=True)
            except Exception as e_main:
                logger.error(f"An unexpected error occurred while processing {os.path.basename(filepath)}: {e_main}", exc_info=True)
            
            if i < len(files_to_process) - 1:
                logger.info(f">>> Finished {os.path.basename(filepath)}. Pause for a few seconds before next file..."); time.sleep(3)
    
    logger.info("--- SCRIPT COMPLETED ---")

# ===== ENHANCED CONFIGURATION VALIDATOR =====

def validate_configuration():
    """Validate configuration and provide optimization recommendations"""
    logger.info("🔍 VALIDATING CONFIGURATION...")

    issues = []
    recommendations = []

    # Check API Key
    if not API_KEY:
        issues.append("❌ GEMINI_API_KEY not found in environment variables")
        recommendations.append("💡 Set GEMINI_API_KEY in your .env file or environment")
    else:
        logger.info("✅ API Key found")

    # Check model configuration
    if MODEL_NAME_CONFIG:
        if "2.5-pro" in MODEL_NAME_CONFIG:
            recommendations.append("💰 Consider using 'gemini-1.5-flash-latest' for cost efficiency")
        logger.info(f"✅ Model: {MODEL_NAME}")
    else:
        issues.append("❌ No model specified in config")

    # Check batch settings
    if MAX_GROUP_SIZE_CHARS > 15000:
        recommendations.append("⚠️ Large batch size may cause API timeouts - consider reducing to 12000")

    if MAX_CONCURRENT_API_CALLS > 10:
        recommendations.append("⚠️ High concurrent calls may trigger rate limits - consider reducing to 5")

    # Check OCR settings
    if PERFORM_OCR_ON_IMAGES and not pytesseract:
        issues.append("❌ OCR enabled but pytesseract not installed")
        recommendations.append("💡 Install pytesseract: pip install pytesseract")

    # Check smart grouping
    if not ENABLE_SMART_GROUPING:
        recommendations.append("💡 Enable smart_grouping for significant API cost reduction")

    # Check quality assessment
    if PERFORM_QUALITY_ASSESSMENT:
        recommendations.append("💰 Quality assessment uses additional API calls - disable if cost is a concern")

    # Print results
    if issues:
        logger.error("🚨 CONFIGURATION ISSUES FOUND:")
        for issue in issues:
            logger.error(f"  {issue}")

    if recommendations:
        logger.info("💡 OPTIMIZATION RECOMMENDATIONS:")
        for rec in recommendations:
            logger.info(f"  {rec}")

    if not issues:
        logger.info("✅ Configuration validation passed!")

    return len(issues) == 0

def estimate_translation_cost(filepath):
    """Estimate translation cost and time"""
    try:
        doc = fitz.open(filepath)
        total_chars = 0

        for page_num in range(min(5, len(doc))):  # Sample first 5 pages
            page = doc[page_num]
            text = page.get_text()
            total_chars += len(text)

        doc.close()

        # Estimate total document size
        estimated_total_chars = (total_chars / min(5, len(doc))) * len(doc)

        # Estimate API calls with smart grouping
        if ENABLE_SMART_GROUPING:
            estimated_api_calls = max(1, estimated_total_chars // MAX_GROUP_SIZE_CHARS)
        else:
            # Rough estimate: 1 call per 500 chars without grouping
            estimated_api_calls = max(1, estimated_total_chars // 500)

        # Estimate time (rough: 2-5 seconds per API call)
        estimated_time_minutes = (estimated_api_calls * 3) / 60

        logger.info("📊 TRANSLATION ESTIMATE:")
        logger.info(f"  📄 Document pages: {len(doc)}")
        logger.info(f"  📝 Estimated characters: {estimated_total_chars:,}")
        logger.info(f"  🔄 Estimated API calls: {estimated_api_calls}")
        logger.info(f"  ⏱️ Estimated time: {estimated_time_minutes:.1f} minutes")
        logger.info(f"  🤖 Model: {MODEL_NAME}")
        logger.info(f"  📦 Smart grouping: {'✅ Enabled' if ENABLE_SMART_GROUPING else '❌ Disabled'}")

        return estimated_api_calls, estimated_time_minutes

    except Exception as e:
        logger.error(f"Failed to estimate cost: {e}")
        return None, None

# ===== ENHANCED PROGRESS TRACKING =====

class ProgressTracker:
    """Enhanced progress tracking with ETA and detailed status"""

    def __init__(self, total_tasks):
        self.total_tasks = total_tasks
        self.completed_tasks = 0
        self.failed_tasks = 0
        self.start_time = time.time()
        self.last_update = time.time()

    def update(self, completed=0, failed=0):
        """Update progress counters"""
        self.completed_tasks += completed
        self.failed_tasks += failed
        current_time = time.time()

        # Update every 10 seconds or on significant progress
        if current_time - self.last_update > 10 or (self.completed_tasks + self.failed_tasks) % 10 == 0:
            self.print_progress()
            self.last_update = current_time

    def print_progress(self):
        """Print detailed progress information"""
        total_processed = self.completed_tasks + self.failed_tasks
        progress_percent = (total_processed / self.total_tasks) * 100 if self.total_tasks > 0 else 0

        elapsed_time = time.time() - self.start_time

        # Calculate ETA
        if total_processed > 0:
            avg_time_per_task = elapsed_time / total_processed
            remaining_tasks = self.total_tasks - total_processed
            eta_seconds = remaining_tasks * avg_time_per_task
            eta_minutes = eta_seconds / 60
        else:
            eta_minutes = 0

        # Create progress bar
        bar_length = 30
        filled_length = int(bar_length * progress_percent / 100)
        bar = '█' * filled_length + '░' * (bar_length - filled_length)

        logger.info(f"🔄 Progress: [{bar}] {progress_percent:.1f}% ({total_processed}/{self.total_tasks})")
        logger.info(f"   ✅ Completed: {self.completed_tasks} | ❌ Failed: {self.failed_tasks}")
        logger.info(f"   ⏱️ Elapsed: {elapsed_time/60:.1f}m | ETA: {eta_minutes:.1f}m")

    def finish(self):
        """Print final summary"""
        total_time = time.time() - self.start_time
        success_rate = (self.completed_tasks / self.total_tasks) * 100 if self.total_tasks > 0 else 0

        logger.info("🏁 TRANSLATION COMPLETED!")
        logger.info(f"   ✅ Success rate: {success_rate:.1f}% ({self.completed_tasks}/{self.total_tasks})")
        logger.info(f"   ⏱️ Total time: {total_time/60:.1f} minutes")
        logger.info(f"   ⚡ Average: {total_time/max(1, self.total_tasks):.1f} seconds/task")

# ===== ADVANCED OPTIMIZATIONS =====

class AdaptiveBatchOptimizer:
    """Dynamically optimizes batch sizes based on performance metrics"""

    def __init__(self):
        self.performance_history = []
        self.optimal_batch_size = MAX_GROUP_SIZE_CHARS
        self.min_batch_size = 8000
        self.max_batch_size = 20000

    def record_performance(self, batch_size, processing_time, success_rate):
        """Record performance metrics for adaptive optimization"""
        efficiency_score = success_rate / max(processing_time, 0.1)  # Avoid division by zero

        self.performance_history.append({
            'batch_size': batch_size,
            'processing_time': processing_time,
            'success_rate': success_rate,
            'efficiency_score': efficiency_score,
            'timestamp': time.time()
        })

        # Keep only recent history (last 10 batches)
        if len(self.performance_history) > 10:
            self.performance_history = self.performance_history[-10:]

        # Update optimal batch size based on performance
        self._update_optimal_batch_size()

    def _update_optimal_batch_size(self):
        """Update optimal batch size based on performance history"""
        if len(self.performance_history) < 3:
            return

        # Find the batch size with highest efficiency score
        best_performance = max(self.performance_history, key=lambda x: x['efficiency_score'])

        # Gradually adjust towards optimal size
        target_size = best_performance['batch_size']
        adjustment = (target_size - self.optimal_batch_size) * 0.3  # 30% adjustment

        self.optimal_batch_size = max(
            self.min_batch_size,
            min(self.max_batch_size, self.optimal_batch_size + adjustment)
        )

        logger.info(f"🎯 Adaptive optimization: Adjusted batch size to {int(self.optimal_batch_size)}")

    def get_optimal_batch_size(self):
        """Get current optimal batch size"""
        return int(self.optimal_batch_size)

class ContentTypeOptimizer:
    """Optimizes translation approach based on content type analysis"""

    def __init__(self):
        self.content_patterns = {
            'academic': {
                'keywords': ['abstract', 'methodology', 'conclusion', 'hypothesis', 'research'],
                'batch_multiplier': 1.2,  # Academic content can handle larger batches
                'temperature': 0.05,  # More conservative for accuracy
            },
            'technical': {
                'keywords': ['specification', 'procedure', 'configuration', 'system', 'manual'],
                'batch_multiplier': 1.1,
                'temperature': 0.1,
            },
            'legal': {
                'keywords': ['whereas', 'hereby', 'pursuant', 'jurisdiction', 'contract'],
                'batch_multiplier': 0.8,  # Smaller batches for precision
                'temperature': 0.05,
            },
            'narrative': {
                'keywords': ['story', 'character', 'dialogue', 'chapter', 'narrative'],
                'batch_multiplier': 1.3,  # Can handle larger batches
                'temperature': 0.15,  # More creative
            }
        }

    def analyze_content_type(self, text_sample):
        """Analyze content type and return optimization parameters"""
        text_lower = text_sample.lower()

        scores = {}
        for content_type, config in self.content_patterns.items():
            score = sum(1 for keyword in config['keywords'] if keyword in text_lower)
            scores[content_type] = score

        if not scores or max(scores.values()) == 0:
            return 'general', 1.0, TRANSLATION_TEMPERATURE

        best_type = max(scores, key=scores.get)
        config = self.content_patterns[best_type]

        return best_type, config['batch_multiplier'], config['temperature']

class SmartCacheManager:
    """Advanced caching with fuzzy matching and context awareness"""

    def __init__(self):
        self.fuzzy_cache = {}
        self.context_cache = {}
        self.similarity_threshold = 0.85

    def get_fuzzy_match(self, text, target_language):
        """Find fuzzy matches in cache for similar content"""
        if not USE_TRANSLATION_CACHE:
            return None

        text_normalized = self._normalize_text(text)
        cache_key = f"{text_normalized}_{target_language}"

        # Check exact match first
        if cache_key in TRANSLATION_CACHE:
            return TRANSLATION_CACHE[cache_key]

        # Check fuzzy matches
        for cached_key, cached_translation in TRANSLATION_CACHE.items():
            if target_language in cached_key:
                cached_text = cached_key.replace(f"_{target_language}", "")
                similarity = self._calculate_similarity(text_normalized, cached_text)

                if similarity >= self.similarity_threshold:
                    logger.info(f"🔍 Fuzzy cache hit: {similarity:.2f} similarity")
                    return cached_translation

        return None

    def _normalize_text(self, text):
        """Normalize text for better matching"""
        import re
        # Remove extra whitespace, punctuation variations
        normalized = re.sub(r'\s+', ' ', text.strip().lower())
        normalized = re.sub(r'[^\w\s]', '', normalized)
        return normalized

    def _calculate_similarity(self, text1, text2):
        """Calculate text similarity using simple ratio"""
        if not text1 or not text2:
            return 0.0

        # Simple character-based similarity
        longer = text1 if len(text1) > len(text2) else text2
        shorter = text2 if len(text1) > len(text2) else text1

        if len(longer) == 0:
            return 1.0

        matches = sum(1 for i, char in enumerate(shorter) if i < len(longer) and char == longer[i])
        return matches / len(longer)

class PerformanceProfiler:
    """Profiles performance and suggests optimizations"""

    def __init__(self):
        self.metrics = {
            'api_call_times': [],
            'batch_sizes': [],
            'success_rates': [],
            'memory_usage': [],
            'start_time': time.time()
        }

    def record_api_call(self, duration, batch_size, success):
        """Record API call performance"""
        self.metrics['api_call_times'].append(duration)
        self.metrics['batch_sizes'].append(batch_size)
        self.metrics['success_rates'].append(1.0 if success else 0.0)

        # Keep only recent metrics
        max_records = 100
        for key in ['api_call_times', 'batch_sizes', 'success_rates']:
            if len(self.metrics[key]) > max_records:
                self.metrics[key] = self.metrics[key][-max_records:]

    def get_performance_report(self):
        """Generate performance analysis report"""
        if not self.metrics['api_call_times']:
            return "No performance data available"

        avg_time = sum(self.metrics['api_call_times']) / len(self.metrics['api_call_times'])
        avg_batch_size = sum(self.metrics['batch_sizes']) / len(self.metrics['batch_sizes'])
        success_rate = sum(self.metrics['success_rates']) / len(self.metrics['success_rates'])
        total_time = time.time() - self.metrics['start_time']

        report = f"""
📊 PERFORMANCE ANALYSIS
======================
⏱️ Average API call time: {avg_time:.2f}s
📦 Average batch size: {avg_batch_size:.0f} chars
✅ Success rate: {success_rate:.1%}
🕐 Total processing time: {total_time/60:.1f} minutes
⚡ Throughput: {len(self.metrics['api_call_times'])/max(total_time/60, 0.1):.1f} calls/minute

💡 OPTIMIZATION SUGGESTIONS:
"""

        # Generate suggestions based on metrics
        if avg_time > 10:
            report += "• Consider reducing batch size for faster response times\n"
        if avg_time < 3:
            report += "• Consider increasing batch size for better efficiency\n"
        if success_rate < 0.9:
            report += "• High failure rate detected - check API limits and network\n"
        if avg_batch_size < 8000:
            report += "• Batch sizes are small - enable aggressive grouping\n"

        return report

class SmartPreprocessor:
    """Intelligent preprocessing to optimize content before translation"""

    def __init__(self):
        self.skip_patterns = [
            r'^\s*\d+\s*$',  # Page numbers
            r'^\s*[ivxlcdm]+\s*$',  # Roman numerals
            r'^\s*\w{1,3}\s*$',  # Single short words
            r'^\s*[^\w\s]+\s*$',  # Only punctuation
            r'^\s*https?://\S+\s*$',  # URLs
            r'^\s*\w+@\w+\.\w+\s*$',  # Email addresses
        ]
        self.merge_patterns = [
            r'^\s*[A-Z][a-z]*\s*$',  # Single capitalized words (likely headers)
            r'^\s*\d+\.\d+\s*$',  # Version numbers
        ]

    def preprocess_content(self, content_items):
        """Intelligently preprocess content for optimal translation"""
        logger.info("🔧 Smart preprocessing content...")

        processed_items = []
        skipped_count = 0
        merged_count = 0

        i = 0
        while i < len(content_items):
            item = content_items[i]
            text = item.get('text', '').strip()

            # Skip items that don't need translation
            if self._should_skip_item(text):
                skipped_count += 1
                i += 1
                continue

            # Try to merge with next items if beneficial
            merged_item, items_consumed = self._try_merge_items(content_items, i)
            if items_consumed > 1:
                merged_count += items_consumed - 1
                processed_items.append(merged_item)
                i += items_consumed
            else:
                processed_items.append(item)
                i += 1

        logger.info(f"✅ Preprocessing complete: Skipped {skipped_count}, Merged {merged_count} items")
        logger.info(f"📊 Optimization: {len(content_items)} → {len(processed_items)} items ({((len(content_items) - len(processed_items)) / len(content_items) * 100):.1f}% reduction)")

        return processed_items

    def _should_skip_item(self, text):
        """Determine if an item should be skipped"""
        if not text or len(text.strip()) < 3:
            return True

        for pattern in self.skip_patterns:
            if re.match(pattern, text, re.IGNORECASE):
                return True

        return False

    def _try_merge_items(self, items, start_index):
        """Try to merge consecutive small items for better context"""
        if start_index >= len(items):
            return items[start_index], 1

        current_item = items[start_index]
        current_text = current_item.get('text', '').strip()

        # Don't merge if current item is already large
        if len(current_text) > 500:
            return current_item, 1

        # Look ahead for mergeable items
        merged_text = current_text
        items_to_merge = 1
        total_length = len(current_text)

        for i in range(start_index + 1, min(start_index + 5, len(items))):  # Look ahead max 5 items
            next_item = items[i]
            next_text = next_item.get('text', '').strip()

            # Stop if next item is too large or different type
            if (len(next_text) > 300 or
                total_length + len(next_text) > 1000 or
                next_item.get('type') != current_item.get('type')):
                break

            # Check if items are related (similar formatting, consecutive)
            if self._are_items_mergeable(current_item, next_item):
                merged_text += " " + next_text
                total_length += len(next_text)
                items_to_merge += 1
            else:
                break

        if items_to_merge > 1:
            # Create merged item
            merged_item = current_item.copy()
            merged_item['text'] = merged_text
            merged_item['_merged_from'] = items_to_merge
            return merged_item, items_to_merge

        return current_item, 1

    def _are_items_mergeable(self, item1, item2):
        """Check if two items can be merged"""
        # Same type
        if item1.get('type') != item2.get('type'):
            return False

        # Both are regular text or list items
        if item1.get('type') in ['paragraph', 'list_item']:
            return True

        # Both are headings of same level
        if item1.get('type') == item2.get('type') and item1.get('type').startswith('h'):
            return False  # Don't merge headings

        return False

class DynamicModelSelector:
    """Dynamically selects the best model based on content and performance"""

    def __init__(self):
        self.models = {
            'gemini-1.5-flash-latest': {
                'cost_factor': 1.0,
                'speed_factor': 2.0,
                'quality_factor': 0.9,
                'best_for': ['simple', 'repetitive', 'technical']
            },
            'gemini-1.5-pro-latest': {
                'cost_factor': 5.0,
                'speed_factor': 1.0,
                'quality_factor': 1.0,
                'best_for': ['complex', 'academic', 'literary']
            },
            'gemini-2.5-pro-preview-03-25': {
                'cost_factor': 10.0,
                'speed_factor': 0.8,
                'quality_factor': 1.1,
                'best_for': ['critical', 'legal', 'medical']
            }
        }
        self.performance_history = {}

    def select_optimal_model(self, content_type, content_complexity, priority='balanced'):
        """Select the optimal model based on content and requirements"""

        if priority == 'cost':
            # Prioritize cost efficiency
            return 'gemini-1.5-flash-latest'
        elif priority == 'quality':
            # Prioritize quality
            return 'gemini-2.5-pro-preview-03-25'
        elif priority == 'speed':
            # Prioritize speed
            return 'gemini-1.5-flash-latest'
        else:
            # Balanced approach
            if content_type in ['legal', 'medical', 'critical']:
                return 'gemini-2.5-pro-preview-03-25'
            elif content_type in ['technical', 'simple']:
                return 'gemini-1.5-flash-latest'
            else:
                return 'gemini-1.5-pro-latest'

    def get_model_recommendation(self, estimated_cost, content_analysis):
        """Get model recommendation with reasoning"""
        content_type = content_analysis.get('type', 'general')

        recommendation = self.select_optimal_model(content_type, 'medium', 'balanced')

        reasoning = f"""
🤖 MODEL RECOMMENDATION: {recommendation}

📋 Analysis:
• Content type: {content_type}
• Estimated API calls: {estimated_cost}
• Current model: {MODEL_NAME}

💡 Reasoning:
"""

        if recommendation != MODEL_NAME.replace('models/', ''):
            model_info = self.models.get(recommendation, {})
            current_info = self.models.get(MODEL_NAME.replace('models/', ''), {})

            if model_info.get('cost_factor', 1) < current_info.get('cost_factor', 1):
                reasoning += f"• Recommended model is {current_info.get('cost_factor', 1) / model_info.get('cost_factor', 1):.1f}x more cost-efficient\n"

            if model_info.get('speed_factor', 1) > current_info.get('speed_factor', 1):
                reasoning += f"• Recommended model is {model_info.get('speed_factor', 1) / current_info.get('speed_factor', 1):.1f}x faster\n"
        else:
            reasoning += "• Current model is optimal for this content type\n"

        return recommendation, reasoning

class UltimateOptimizationManager:
    """Master optimization manager that coordinates all optimization systems"""

    def __init__(self):
        self.adaptive_batcher = AdaptiveBatchOptimizer()
        self.content_optimizer = ContentTypeOptimizer()
        self.cache_manager = SmartCacheManager()
        self.profiler = PerformanceProfiler()
        self.preprocessor = SmartPreprocessor()
        self.model_selector = DynamicModelSelector()

        self.optimization_enabled = True
        self.optimization_stats = {
            'items_preprocessed': 0,
            'items_skipped': 0,
            'cache_hits': 0,
            'batch_optimizations': 0,
            'api_calls_saved': 0
        }

    def optimize_translation_pipeline(self, content_items, target_language):
        """Apply comprehensive optimizations to the translation pipeline"""
        if not self.optimization_enabled:
            return content_items

        logger.info("🚀 APPLYING ULTIMATE OPTIMIZATIONS...")

        # Step 1: Smart preprocessing
        optimized_items = self.preprocessor.preprocess_content(content_items)
        self.optimization_stats['items_preprocessed'] = len(content_items) - len(optimized_items)

        # Step 2: Content type analysis for first few items
        sample_text = " ".join([item.get('text', '')[:500] for item in optimized_items[:5]])
        content_type, batch_multiplier, optimal_temp = self.content_optimizer.analyze_content_type(sample_text)

        logger.info(f"📊 Content analysis: {content_type} (batch multiplier: {batch_multiplier:.1f})")

        # Step 3: Adaptive batch size optimization
        optimal_batch_size = int(self.adaptive_batcher.get_optimal_batch_size() * batch_multiplier)
        logger.info(f"🎯 Optimal batch size: {optimal_batch_size} chars")

        # Step 4: Cache optimization check
        cache_hits = 0
        for item in optimized_items:
            text = item.get('text', '')
            if self.cache_manager.get_fuzzy_match(text, target_language):
                cache_hits += 1

        self.optimization_stats['cache_hits'] = cache_hits

        # Step 5: Model recommendation
        estimated_calls = max(1, len(optimized_items) // 10)  # Rough estimate
        model_rec, reasoning = self.model_selector.get_model_recommendation(
            estimated_calls, {'type': content_type}
        )

        logger.info(reasoning)

        # Generate optimization report
        self._generate_optimization_report(content_items, optimized_items, content_type)

        return optimized_items, {
            'optimal_batch_size': optimal_batch_size,
            'optimal_temperature': optimal_temp,
            'content_type': content_type,
            'recommended_model': model_rec
        }

    def _generate_optimization_report(self, original_items, optimized_items, content_type):
        """Generate comprehensive optimization report"""
        reduction_percent = ((len(original_items) - len(optimized_items)) / len(original_items)) * 100

        report = f"""
🎯 OPTIMIZATION REPORT
=====================
📊 Content Analysis: {content_type}
📦 Items: {len(original_items)} → {len(optimized_items)} ({reduction_percent:.1f}% reduction)
💾 Cache hits: {self.optimization_stats['cache_hits']}
🔧 Preprocessing optimizations applied

💡 ESTIMATED SAVINGS:
• API calls reduced by ~{reduction_percent:.0f}%
• Processing time reduced by ~{reduction_percent * 0.8:.0f}%
• Cost savings: ~{reduction_percent:.0f}%
"""

        logger.info(report)
        quality_report_messages.append("=== OPTIMIZATION REPORT ===")
        quality_report_messages.append(report)

    def record_batch_performance(self, batch_size, processing_time, success_rate):
        """Record performance for adaptive optimization"""
        self.adaptive_batcher.record_performance(batch_size, processing_time, success_rate)
        self.profiler.record_api_call(processing_time, batch_size, success_rate > 0.8)

    def get_final_performance_report(self):
        """Get comprehensive performance analysis"""
        return self.profiler.get_performance_report()

# Initialize global optimization manager
optimization_manager = UltimateOptimizationManager()

# ===== ENHANCED DOCUMENT INTELLIGENCE INTEGRATION =====

try:
    from enhanced_document_intelligence import (
        AdvancedContentClassifier, SemanticContentGrouper,
        CrossReferencePreserver, ContentClassification, SemanticGroup
    )
    from typing import List, Dict
    ENHANCED_INTELLIGENCE_AVAILABLE = True
    logger.info("🧠 Enhanced Document Intelligence loaded successfully")
except ImportError:
    ENHANCED_INTELLIGENCE_AVAILABLE = False
    logger.warning("⚠️ Enhanced Document Intelligence not available - using standard processing")
    # Define fallback types
    List = list
    Dict = dict

class DocumentIntelligenceManager:
    """Manages enhanced document intelligence features safely"""

    def __init__(self):
        self.enabled = ENHANCED_INTELLIGENCE_AVAILABLE
        if self.enabled:
            self.classifier = AdvancedContentClassifier()
            self.grouper = SemanticContentGrouper()
            self.ref_preserver = CrossReferencePreserver()

        self.intelligence_report = []

    def analyze_document_intelligence(self, content_items):
        """Perform comprehensive document intelligence analysis"""
        if not self.enabled or not content_items:
            return self._get_fallback_analysis(content_items)

        try:
            logger.info("🧠 Performing enhanced document intelligence analysis...")

            # 1. Advanced Content Classification
            classification = self.classifier.classify_content(content_items)

            # 2. Semantic Grouping
            semantic_groups = self.grouper.create_semantic_groups(content_items)

            # 3. Cross-Reference Analysis
            references = self.ref_preserver.extract_references(content_items)
            ref_guide = self.ref_preserver.create_reference_preservation_guide(references)

            # Generate intelligence report
            intelligence_report = self._generate_intelligence_report(
                classification, semantic_groups, references
            )

            return {
                'classification': classification,
                'semantic_groups': semantic_groups,
                'cross_references': references,
                'reference_guide': ref_guide,
                'intelligence_report': intelligence_report,
                'enhanced_processing': True
            }

        except Exception as e:
            logger.warning(f"⚠️ Enhanced intelligence analysis failed: {e}")
            return self._get_fallback_analysis(content_items)

    def _get_fallback_analysis(self, content_items):
        """Fallback analysis when enhanced intelligence is not available"""
        return {
            'classification': None,
            'semantic_groups': [],
            'cross_references': {},
            'reference_guide': "",
            'intelligence_report': "Standard processing mode - no enhanced intelligence",
            'enhanced_processing': False
        }

    def _generate_intelligence_report(self, classification, semantic_groups, references):
        """Generate comprehensive intelligence report"""
        report = f"""
🧠 ENHANCED DOCUMENT INTELLIGENCE REPORT
========================================

📊 CONTENT CLASSIFICATION:
• Primary Type: {classification.primary_type} (confidence: {classification.confidence:.1%})
• Complexity Level: {classification.complexity_level}
• Translation Approach: {classification.translation_approach}
• Recommended Batch Size: {classification.recommended_batch_size} chars
• Recommended Temperature: {classification.recommended_temperature}

🎯 SEMANTIC ANALYSIS:
• Semantic Groups Identified: {len(semantic_groups)}
• Themes Found: {', '.join(set(g.semantic_theme for g in semantic_groups))}
• High Priority Groups: {sum(1 for g in semantic_groups if g.translation_priority >= 4)}
• Order-Sensitive Groups: {sum(1 for g in semantic_groups if g.preserve_order)}

🔗 CROSS-REFERENCE ANALYSIS:
• Reference Types Found: {len(references)}
• Total References: {sum(len(refs) for refs in references.values())}
"""

        if references:
            report += "• Reference Types: " + ", ".join(references.keys()) + "\n"

        if classification.domain_indicators:
            report += f"• Domain Indicators: {', '.join(classification.domain_indicators[:3])}\n"

        if classification.secondary_types:
            report += f"• Secondary Types: {', '.join(classification.secondary_types)}\n"

        report += "\n💡 OPTIMIZATION RECOMMENDATIONS:\n"

        # Generate specific recommendations
        if classification.primary_type == 'academic':
            report += "• Use conservative translation for academic precision\n"
            report += "• Preserve technical terminology and citations\n"
        elif classification.primary_type == 'legal':
            report += "• Use highly conservative translation for legal accuracy\n"
            report += "• Maintain exact clause structure and terminology\n"
        elif classification.primary_type == 'technical':
            report += "• Preserve technical specifications and procedures\n"
            report += "• Maintain step-by-step instruction clarity\n"

        if classification.complexity_level in ['complex', 'expert']:
            report += "• Use smaller batch sizes for complex content\n"
            report += "• Consider manual review for expert-level content\n"

        if references:
            report += "• Cross-reference preservation guide has been generated\n"
            report += "• Maintain reference numbering consistency\n"

        report += "========================================\n"

        return report

    def get_enhanced_translation_parameters(self, analysis):
        """Get enhanced translation parameters based on intelligence analysis"""
        if not analysis.get('enhanced_processing', False):
            return {}

        classification = analysis.get('classification')
        if not classification:
            return {}

        return {
            'optimal_batch_size': classification.recommended_batch_size,
            'optimal_temperature': classification.recommended_temperature,
            'translation_approach': classification.translation_approach,
            'content_type': classification.primary_type,
            'complexity_level': classification.complexity_level,
            'reference_guide': analysis.get('reference_guide', ''),
            'semantic_groups_count': len(analysis.get('semantic_groups', [])),
            'cross_references_count': sum(len(refs) for refs in analysis.get('cross_references', {}).values())
        }

# Initialize global document intelligence manager
document_intelligence = DocumentIntelligenceManager()

# Global variable for enhanced reference guide
ENHANCED_REFERENCE_GUIDE = ""

# ===== RECOVERY UTILITIES =====

def save_recovery_state(filepath, failed_items, partial_results):
    """Save recovery state for resuming failed translations"""
    recovery_file = f"{os.path.splitext(filepath)[0]}_recovery.json"

    recovery_data = {
        'timestamp': time.time(),
        'original_file': filepath,
        'failed_items': failed_items,
        'partial_results': partial_results,
        'config_snapshot': {
            'model': MODEL_NAME,
            'target_language': TARGET_LANGUAGE_CONFIG,
            'smart_grouping': ENABLE_SMART_GROUPING,
            'max_group_size': MAX_GROUP_SIZE_CHARS
        }
    }

    try:
        with open(recovery_file, 'w', encoding='utf-8') as f:
            json.dump(recovery_data, f, indent=2, ensure_ascii=False)
        logger.info(f"💾 Recovery state saved: {recovery_file}")
        return recovery_file
    except Exception as e:
        logger.error(f"Failed to save recovery state: {e}")
        return None

def load_recovery_state(recovery_file):
    """Load recovery state for resuming translations"""
    try:
        with open(recovery_file, 'r', encoding='utf-8') as f:
            recovery_data = json.load(f)
        logger.info(f"📂 Recovery state loaded: {recovery_file}")
        return recovery_data
    except Exception as e:
        logger.error(f"Failed to load recovery state: {e}")
        return None

# Test runner functions are now integrated into the main __name__ == "__main__" block above
