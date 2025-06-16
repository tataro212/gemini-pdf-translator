"""
Document Generator Module for Ultimate PDF Translator

Handles Word document creation and PDF conversion with proper formatting and structure
"""

import os
import logging
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn # Ensure qn is imported

from config_manager import config_manager
from utils import sanitize_for_xml, sanitize_filepath

logger = logging.getLogger(__name__)

# Import structured document model
try:
    from structured_document_model import (
        Document as StructuredDocument, ContentType, Heading, Paragraph, ImagePlaceholder,
        Table, CodeBlock, ListItem, Footnote, Equation, Caption, Metadata
    )
    STRUCTURED_MODEL_AVAILABLE = True
except ImportError:
    STRUCTURED_MODEL_AVAILABLE = False
    logger.warning("Structured document model not available")

# Global bookmark counter
# bookmark_id_counter = 0 # Comment out or remove if self.bookmark_id replaces its ToC usage

class WordDocumentGenerator:
    """Generates Word documents with proper structure and formatting"""
    
    def __init__(self, translation_service=None, pdf_parser=None):
        """Initialize the document generator with translation service and PDF parser"""
        self.translation_service = translation_service
        self.pdf_parser = pdf_parser
        self.logger = logging.getLogger(__name__)
        self.word_settings = config_manager.word_output_settings
        self.toc_entries = []
        self.bookmark_id = 0
    
    def _add_heading_block(self, doc, block_item):
        text_content = None
        level = 1

        if STRUCTURED_MODEL_AVAILABLE and isinstance(block_item, Heading):
            text_content = block_item.content
            level = block_item.level
        elif isinstance(block_item, dict):
            text_content = block_item.get('content') or block_item.get('text')
            level = block_item.get('level')
            type_str = block_item.get('type', '').lower()
            if level is None and type_str.startswith('h') and type_str[1:].isdigit():
                level = int(type_str[1:])
            if level is None: level = 1 # Default if not found
        else:
            logger.warning(f"Skipping unknown heading block type: {type(block_item)}")
            return

        if not text_content:
            logger.warning("Skipping empty heading block.")
            return

        safe_content = sanitize_for_xml(text_content)
        p = doc.add_paragraph(style=f'Heading {level}')
        
        self.bookmark_id += 1
        bookmark_name = f"_Toc_Bookmark_{self.bookmark_id}"
        run = p.add_run(safe_content)
        
        tag_start = OxmlElement('w:bookmarkStart')
        tag_start.set(qn('w:id'), str(self.bookmark_id))
        tag_start.set(qn('w:name'), bookmark_name)
        tag_end = OxmlElement('w:bookmarkEnd')
        tag_end.set(qn('w:id'), str(self.bookmark_id))

        if run._r is not None and p._p is not None:
            # CORRECTED LINE: Use addprevious on the run element itself
            run._r.addprevious(tag_start)
            run._r.addnext(tag_end)
        else:
            logger.warning(f"Could not precisely position bookmark for heading: {safe_content}. Run or paragraph element was None.")
            # Fallback: append to paragraph, less ideal but avoids error
            p._p.append(tag_start)
            p._p.append(tag_end)
        
        self.toc_entries.append({
            'text': text_content,
            'level': level,
            'bookmark': bookmark_name
        })
        logger.debug(f"Added Heading '{text_content}' (L{level}) with bookmark '{bookmark_name}'.")

    def _add_paragraph_block(self, doc, block_item):
        text_content = None
        if STRUCTURED_MODEL_AVAILABLE and isinstance(block_item, Paragraph):
            text_content = block_item.content
        elif isinstance(block_item, dict):
            text_content = block_item.get('content') or block_item.get('text')
        else:
            logger.warning(f"Skipping unknown paragraph block type: {type(block_item)}")
            return

        if text_content:
            safe_content = sanitize_for_xml(text_content)
            doc.add_paragraph(safe_content)
        else:
            logger.debug("Skipping empty paragraph block.")

    def _add_image_placeholder_block(self, doc, block_item, image_folder_path):
        image_filename = None
        caption_text = None

        if STRUCTURED_MODEL_AVAILABLE and isinstance(block_item, ImagePlaceholder):
            image_filename = block_item.image_path
            caption_text = block_item.caption
        elif isinstance(block_item, dict):
            image_filename = block_item.get('image_path') or block_item.get('filepath') # common dict keys
            caption_text = block_item.get('caption') or block_item.get('text') # some dicts might use 'text' for caption
        else:
            logger.warning(f"Skipping unknown image block type: {type(block_item)}")
            return

        if not image_filename:
            logger.warning("Skipping image block with no image_path/filename.")
            return

        actual_image_path = image_filename
        if image_folder_path and not os.path.isabs(image_filename):
            actual_image_path = os.path.join(image_folder_path, image_filename)

        if os.path.exists(actual_image_path):
            try:
                doc.add_picture(actual_image_path, width=Inches(self.word_settings.get('image_width_inches', 6.0)))
                if caption_text:
                    safe_caption = sanitize_for_xml(caption_text)
                    caption_style_name = self.word_settings.get('caption_style', 'Caption')
                    try:
                        p_caption = doc.add_paragraph(style=caption_style_name)
                        p_caption.add_run(safe_caption)
                    except KeyError: # Style not found
                        logger.warning(f"Caption style '{caption_style_name}' not found. Adding caption as normal paragraph.")
                        p_caption = doc.add_paragraph()
                        p_caption.add_run(safe_caption)
                    p_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                logger.debug(f"Added image: {actual_image_path}")
            except Exception as e:
                logger.error(f"Error adding image {actual_image_path}: {e}")
        else:
            logger.warning(f"Image not found at path: {actual_image_path}. Skipping image.")

    def _add_list_item_block(self, doc, block_item):
        text_content = None
        level = 0 # Placeholder for list level if available in dict/model
        if STRUCTURED_MODEL_AVAILABLE and isinstance(block_item, ListItem):
            text_content = block_item.content
            level = getattr(block_item, 'level', 1) # Assuming ListItem might have a level
        elif isinstance(block_item, dict):
            text_content = block_item.get('content') or block_item.get('text')
            level = block_item.get('level', 1)
        else:
            logger.warning(f"Skipping unknown list item block type: {type(block_item)}")
            return

        if text_content:
            safe_content = sanitize_for_xml(text_content)
            # Basic list item handling, python-docx requires more for actual bullets/numbering
            # This adds it as a paragraph with indentation.
            p = doc.add_paragraph()
            p.add_run(safe_content)
            p.paragraph_format.left_indent = Inches(0.25 * level)
            logger.debug(f"Added list item (as indented paragraph): {safe_content}")
        else:
            logger.debug("Skipping empty list item block.")
            
    def _add_table_block(self, doc, block_item):
        # Placeholder - actual table creation is complex
        logger.warning(f"Table block handling is a placeholder for type: {type(block_item)}. Content: {str(block_item)[:100]}")
        if isinstance(block_item, dict) and (block_item.get('content') or block_item.get('text')):
             doc.add_paragraph(f"[Placeholder for Table: {sanitize_for_xml(block_item.get('content') or block_item.get('text'))}]")
        elif STRUCTURED_MODEL_AVAILABLE and isinstance(block_item, Table) and hasattr(block_item, 'rows'):
             doc.add_paragraph(f"[Placeholder for Table with {len(block_item.rows)} rows]")
        else:
             doc.add_paragraph("[Placeholder for Table]")


    def _add_code_block(self, doc, block_item):
        # Placeholder - actual code block formatting can be complex
        text_content = None
        if STRUCTURED_MODEL_AVAILABLE and isinstance(block_item, CodeBlock):
            text_content = block_item.content
        elif isinstance(block_item, dict):
            text_content = block_item.get('content') or block_item.get('text')
        
        if text_content:
            logger.debug(f"Adding code block (as preformatted text): {text_content[:100]}")
            # Simple approach: add as a paragraph with a "code" style if available, or just monospace font
            # For real code blocks, you might want to preserve whitespace carefully.
            p = doc.add_paragraph()
            run = p.add_run(sanitize_for_xml(text_content))
            try:
                # Attempt to apply a 'Code' style if it exists in the template
                p.style = 'Code' 
            except KeyError:
                # Fallback to basic monospace font
                font = run.font
                font.name = 'Courier New' # Or another common monospace font
                font.size = Pt(10)
        else:
            logger.warning(f"Skipping empty code block: {type(block_item)}")

    # --- New method for inserting ToC (Step 3) ---
    def _insert_toc(self, doc):
        """Insert table of contents at the beginning of the document based on collected toc_entries."""
        if not self.toc_entries:
            logger.info("No TOC entries found, skipping TOC generation")
            return
            
        logger.info("Generating Table of Contents...")
        
        # Create a new document for the TOC
        toc_doc = Document()
        
        # Configure fonts for TOC document too
        self._configure_document_fonts_for_unicode(toc_doc)
        
        # Add TOC title
        title = toc_doc.add_paragraph("Table of Contents")
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.style = 'Title'
        
        # Add a decorative line
        line = toc_doc.add_paragraph()
        line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        line.add_run("_" * 40)
        
        # Sort TOC entries by level and text
        sorted_entries = sorted(self.toc_entries, key=lambda x: (x['level'], x['text']))
        
        # Add TOC entries with hyperlinks and page numbers
        for entry in sorted_entries:
            level = entry['level']
            text_for_display = sanitize_for_xml(entry['text'])
            bookmark = entry['bookmark']
            
            # Create TOC entry paragraph
            p = toc_doc.add_paragraph()
            p.style = 'Normal'
            
            # Set indentation based on level
            indent_size = Pt(20 * (level - 1))
            p.paragraph_format.left_indent = indent_size
            
            # Add heading text with hyperlink
            try:
                hyperlink_run = self._add_hyperlink(p, text_for_display, bookmark)
                if level == 1:
                    hyperlink_run.bold = True
                    hyperlink_run.font.size = Pt(12)
                elif level == 2:
                    hyperlink_run.font.size = Pt(11)
                else:
                    hyperlink_run.font.size = Pt(10)
                    hyperlink_run.italic = True
            except Exception as e:
                logger.debug(f"Could not create hyperlink for {text_for_display}: {e}")
                # Fallback to regular text
                text_run = p.add_run(text_for_display)
                if level == 1:
                    text_run.bold = True
                    text_run.font.size = Pt(12)
                elif level == 2:
                    text_run.font.size = Pt(11)
                else:
                    text_run.font.size = Pt(10)
                    text_run.italic = True

            # Add dots and page number
            dots_needed = max(3, 60 - len(text_for_display))
            dots_run = p.add_run(" " + "." * dots_needed + " ")
            dots_run.font.color.rgb = RGBColor(128, 128, 128)

            # Add page reference field
            fldSimple = OxmlElement('w:fldSimple')
            fldSimple.set(qn('w:instr'), f'PAGEREF {bookmark} \\h')
            
            run_in_field = OxmlElement('w:r')
            text_in_field = OxmlElement('w:t')
            text_in_field.text = "..."  # Placeholder that will be replaced with actual page number
            run_in_field.append(text_in_field)
            fldSimple.append(run_in_field)
            
            p._p.append(fldSimple)
        
        # Add spacing after TOC
        toc_doc.add_paragraph()
        
        # Insert TOC at the beginning of the document
        for element in toc_doc.element.body:
            doc.element.body.insert(0, element)
        
        logger.info(f"Successfully generated TOC with {len(sorted_entries)} entries")

    def create_word_document_with_structure(self, structured_content_list, output_filepath,
                                          image_folder_path, cover_page_data=None):
        """Create Word document with proper structure and formatting.
        This method is assumed to work with a list of content blocks 
        compatible with structured_document_model.py (e.g., Heading, Paragraph).
        """
        # global bookmark_id_counter # No longer using global counter for ToC bookmarks here
        # bookmark_id_counter = 0

        # --- Initialize for new ToC system (aligning with create_word_document_from_structured_document) ---
        self.toc_entries = []
        self.bookmark_id = 0
        # --- End of ToC initialization ---

        # Normalize paths to handle mixed separators
        output_filepath = os.path.normpath(output_filepath)
        if image_folder_path:
            image_folder_path = os.path.normpath(image_folder_path)

        logger.info(f"--- Creating Word Document (with_structure): {os.path.basename(output_filepath)} ---")
        
        doc = Document()
        
        # Add cover page if provided
        if cover_page_data:
            self._add_cover_page(doc, cover_page_data, image_folder_path)
            doc.add_page_break()
        
        # --- Remove old ToC generation call ---
        # if self.word_settings['generate_toc']:
        #     self._add_table_of_contents(doc, structured_content_list) # OLD ToC call
        #     doc.add_page_break()
        
        # Process content items (blocks)
        # Assuming items in structured_content_list are compatible with _add_content_block
        for item_block in structured_content_list: # Renamed 'item' to 'item_block' for clarity
            # --- Replace _add_content_item with _add_content_block ---
            self._add_content_block(doc, item_block, image_folder_path)
        
        # --- PASS 2: All content is added, now insert the ToC (if enabled) ---
        if self.word_settings.get('generate_toc', False): # Check if 'generate_toc' exists and is True
            self._insert_toc(doc)
            # Note: _insert_toc already adds a page break after the ToC if entries exist.
            # No need to add doc.add_page_break() here explicitly unless desired for other reasons.

        # Save document
        try:
            # Use centralized sanitization from utils
            sanitized_filepath = sanitize_filepath(output_filepath)

            # Ensure the output directory exists before saving
            output_dir = os.path.dirname(sanitized_filepath)
            if output_dir and not os.path.exists(output_dir):
                os.makedirs(output_dir, exist_ok=True)
                logger.info(f"Created output directory: {output_dir}")

            doc.save(sanitized_filepath)
            logger.info(f"Word document saved successfully: {sanitized_filepath}")
            # Return the actual file path used for saving
            return sanitized_filepath
        except Exception as e:
            logger.error(f"Error saving Word document: {e}")
            logger.error(f"Attempted path: {output_filepath}")
            return None

    def create_word_document_from_structured_document(self, structured_document, output_filepath,
                                                    image_folder_path, cover_page_data=None):
        """
        Create Word document from a structured Document object.
        This is the new method that works with the refactored structured document model.
        """
        if not STRUCTURED_MODEL_AVAILABLE:
            raise Exception("Structured document model not available")

        if not isinstance(structured_document, StructuredDocument):
            raise ValueError(f"Expected StructuredDocument, got {type(structured_document)}")

        # global bookmark_id_counter # Related to old system, review if still needed elsewhere
        # bookmark_id_counter = 0     # Related to old system

        # --- Initialize for new ToC system (Step 1) ---
        self.toc_entries = []
        self.bookmark_id = 0
        # --- End of Step 1 initialization ---

        # Normalize paths to handle mixed separators
        output_filepath = os.path.normpath(output_filepath)
        if image_folder_path:
            image_folder_path = os.path.normpath(image_folder_path)

        logger.info(f"--- Creating Word Document from Structured Document: {structured_document.title} ---")
        logger.info(f"📊 Processing {len(structured_document.content_blocks)} content blocks")

        doc = Document()

        # Add document title as first heading
        if structured_document.title:
            # Sanitize title before adding as heading
            safe_title = sanitize_for_xml(structured_document.title)
            title_heading = doc.add_heading(safe_title, level=0)
            title_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add cover page if provided
        if cover_page_data:
            self._add_cover_page(doc, cover_page_data, image_folder_path)
            doc.add_page_break()

        # --- Remove existing ToC generation logic ---
        # if self.word_settings['generate_toc']:
        #     # Pass 1: Generate content and collect heading page numbers
        #     heading_page_map = self._generate_content_with_page_tracking(
        #         doc, structured_document, image_folder_path
        #     )
        #     # Pass 2: Generate accurate TOC with real page numbers
        #     self._generate_accurate_toc(doc, structured_document, heading_page_map)
        # else:
        #     # Single pass if TOC is disabled
        #     for block in structured_document.content_blocks:
        #         self._add_content_block(doc, block, image_folder_path)
        
        # --- New main content processing loop (Implicit Pass 1 for ToC) ---
        for block in structured_document.content_blocks:
            self._add_content_block(doc, block, image_folder_path) # This will call _add_heading_block for headings

        # --- PASS 2: All content is added, now insert the ToC (Step 4) ---
        self._insert_toc(doc)
        # --- End of Step 4 ---

        # Save document
        try:
            # Use centralized sanitization from utils
            sanitized_filepath = sanitize_filepath(output_filepath)

            # Ensure the output directory exists before saving
            output_dir = os.path.dirname(sanitized_filepath)
            if output_dir and not os.path.exists(output_dir):
                os.makedirs(output_dir, exist_ok=True)
                logger.info(f"Created output directory: {output_dir}")

            doc.save(sanitized_filepath)
            logger.info(f"Word document saved successfully: {sanitized_filepath}")
            # Return the actual file path used for saving
            return sanitized_filepath
        except Exception as e:
            logger.error(f"Error saving Word document: {e}")
            logger.error(f"Attempted path: {output_filepath}")
            return None

    def _add_table_of_contents_from_document(self, doc, structured_document):
        """Add table of contents from structured document headings"""
        try:
            # Add TOC title with enhanced styling and sanitization
            safe_toc_title = sanitize_for_xml(self.word_settings['toc_title'])
            toc_title = doc.add_heading(safe_toc_title, level=1)
            toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add decorative line under title
            title_paragraph = doc.add_paragraph()
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_paragraph.add_run("─" * 50)
            title_run.font.color.rgb = RGBColor(128, 128, 128)

            # Extract headings from structured document
            headings = structured_document.get_blocks_by_type(ContentType.HEADING)

            if not headings:
                logger.warning("No headings found for table of contents")
                return

            # Add TOC entries
            for i, heading_block in enumerate(headings):
                if isinstance(heading_block, Heading):
                    self._add_toc_entry_from_heading_block(doc, heading_block, i + 1)

            # Add spacing after TOC
            doc.add_paragraph()

            logger.info(f"Table of contents added with {len(headings)} entries")

        except Exception as e:
            logger.warning(f"Could not add table of contents: {e}")

    def _add_toc_entry_from_heading_block(self, doc, heading_block, entry_number):
        """Add TOC entry from a Heading content block with improved page estimation"""
        try:
            text = heading_block.content
            level = heading_block.level
            original_page_num = heading_block.page_num

            # Improved page estimation for final document
            # Account for TOC pages, cover page, etc.
            estimated_page = self._estimate_final_page_number(original_page_num, entry_number, level)

            # Sanitize text for XML compatibility
            safe_text = sanitize_for_xml(text)

            # Create TOC paragraph
            toc_paragraph = doc.add_paragraph()
            toc_paragraph.style = 'Normal'

            # Set indentation based on level
            base_indent = 0.2
            level_indent = (level - 1) * self.word_settings['list_indent_per_level_inches']
            toc_paragraph.paragraph_format.left_indent = Inches(base_indent + level_indent)

            # Add entry number for main headings
            if level == 1:
                number_run = toc_paragraph.add_run(f"{entry_number}. ")
                number_run.bold = True
                number_run.font.color.rgb = RGBColor(0, 0, 128)  # Dark blue

            # Add heading text with hyperlink
            try:
                bookmark_name = f"heading_{entry_number}_{level}"
                hyperlink_run = self._add_hyperlink(toc_paragraph, safe_text, bookmark_name)
                if level == 1:
                    hyperlink_run.bold = True
                    hyperlink_run.font.size = Pt(12)
                elif level == 2:
                    hyperlink_run.font.size = Pt(11)
                else:
                    hyperlink_run.font.size = Pt(10)
                    hyperlink_run.italic = True
            except Exception as e:
                logger.debug(f"Could not create hyperlink for {safe_text}: {e}")
                # Fallback to regular text
                text_run = toc_paragraph.add_run(safe_text)
                if level == 1:
                    text_run.bold = True
                    text_run.font.size = Pt(12)
                elif level == 2:
                    text_run.font.size = Pt(11)
                else:
                    text_run.font.size = Pt(10)
                    text_run.italic = True

            # Add dots and page number
            dots_needed = max(3, 60 - len(text) - len(str(estimated_page)))
            dots_run = toc_paragraph.add_run(" " + "." * dots_needed + " ")
            dots_run.font.color.rgb = RGBColor(128, 128, 128)

            page_run = toc_paragraph.add_run(str(estimated_page))
            page_run.bold = True if level == 1 else False

        except Exception as e:
            logger.warning(f"Could not add TOC entry: {e}")

    def _estimate_final_page_number(self, original_page_num, entry_number, level):
        """Estimate the final page number in the Word document accounting for TOC and cover pages"""
        # Base estimation: original page + offset for cover page and TOC
        toc_pages = 2  # Estimate 2 pages for TOC
        cover_pages = 1 if self.word_settings.get('add_cover_page', False) else 0

        # Calculate offset
        page_offset = cover_pages + toc_pages

        # For the first few headings, use a more conservative estimation
        if entry_number <= 3:
            estimated_page = page_offset + entry_number
        else:
            # Use original page number with offset
            estimated_page = max(original_page_num + page_offset, page_offset + entry_number)

        return estimated_page

    def _generate_content_with_page_tracking(self, doc, structured_document, image_folder_path):
        """
        Proposition 2: Pass 1 - Generate content and track heading page numbers.

        Renders all content blocks except TOC and creates bookmarks for headings.
        Returns a mapping of heading block_id to actual page number.
        """
        logger.info("🔄 Pass 1: Generating content with page tracking...")

        heading_page_map = {}
        heading_counter = 0

        # Reserve space for TOC (we'll insert it later)
        toc_placeholder = doc.add_paragraph("[TABLE OF CONTENTS PLACEHOLDER]")
        toc_placeholder.style = 'Normal'
        doc.add_page_break()

        # Process all content blocks and track heading positions
        for block in structured_document.content_blocks:
            if isinstance(block, Heading):
                heading_counter += 1

                # Add the heading with bookmark
                safe_text = sanitize_for_xml(block.content)
                heading_paragraph = doc.add_heading(safe_text, level=block.level)

                # Create unique bookmark name
                bookmark_name = f"heading_{heading_counter}_{block.level}"

                # Add bookmark to heading
                self._add_bookmark_to_paragraph(heading_paragraph, bookmark_name)

                # Get current page number (this is approximate, but more accurate than estimation)
                # In practice, we would need to render the document to get exact page numbers
                # For now, we'll use a more sophisticated estimation
                current_page = self._calculate_current_page_number(doc, heading_counter)

                # Store mapping
                heading_page_map[block.block_id] = {
                    'page_number': current_page,
                    'bookmark_name': bookmark_name,
                    'text': block.content,
                    'level': block.level,
                    'heading_number': heading_counter
                }

                logger.debug(f"Tracked heading {heading_counter}: '{block.content}' -> Page {current_page}")

            else:
                # Add non-heading content blocks normally
                self._add_content_block(doc, block, image_folder_path)

        logger.info(f"✅ Pass 1 complete: Tracked {len(heading_page_map)} headings")
        return heading_page_map

    def _generate_accurate_toc(self, doc, structured_document, heading_page_map):
        """
        Proposition 2: Pass 2 - Generate accurate TOC with real page numbers.

        Navigates to the beginning of the document and replaces the placeholder
        with an accurate TOC using the collected page numbers.
        """
        logger.info("🔄 Pass 2: Generating accurate TOC with real page numbers...")

        try:
            # Find and replace the TOC placeholder
            for paragraph in doc.paragraphs:
                if "[TABLE OF CONTENTS PLACEHOLDER]" in paragraph.text:
                    # Clear the placeholder
                    paragraph.clear()

                    # Add TOC title
                    safe_toc_title = sanitize_for_xml(self.word_settings['toc_title'])
                    toc_title_paragraph = paragraph
                    toc_title_run = toc_title_paragraph.add_run(safe_toc_title)
                    toc_title_run.bold = True
                    toc_title_run.font.size = Pt(16)
                    toc_title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    # Add decorative line
                    decorative_paragraph = self._insert_paragraph_after(doc, toc_title_paragraph)
                    decorative_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    decorative_run = decorative_paragraph.add_run("─" * 50)
                    decorative_run.font.color.rgb = RGBColor(128, 128, 128)

                    # Add TOC entries with accurate page numbers
                    last_paragraph = decorative_paragraph

                    # Sort headings by their order in the document
                    sorted_headings = sorted(
                        heading_page_map.values(),
                        key=lambda x: x['heading_number']
                    )

                    for heading_info in sorted_headings:
                        toc_entry_paragraph = self._insert_paragraph_after(doc, last_paragraph)
                        self._add_accurate_toc_entry(toc_entry_paragraph, heading_info)
                        last_paragraph = toc_entry_paragraph

                    # Add spacing after TOC
                    spacing_paragraph = self._insert_paragraph_after(doc, last_paragraph)
                    spacing_paragraph.add_run("")  # Empty paragraph for spacing

                    logger.info(f"✅ Pass 2 complete: Generated accurate TOC with {len(sorted_headings)} entries")
                    break
            else:
                logger.warning("TOC placeholder not found, adding TOC at the beginning")
                # Fallback: add TOC at the beginning
                self._add_fallback_toc(doc, heading_page_map)

        except Exception as e:
            logger.error(f"Error generating accurate TOC: {e}")
            # Fallback to original method
            self._add_table_of_contents_from_document(doc, structured_document)

    def _calculate_current_page_number(self, doc, heading_counter):
        """
        Calculate approximate current page number based on document content.

        This is a more sophisticated estimation that considers:
        - Number of paragraphs added so far
        - Average lines per page
        - Content density
        """
        # Count total paragraphs and content so far
        total_paragraphs = len(doc.paragraphs)

        # Estimate based on content density
        # Assume ~25-30 lines per page, ~2-3 paragraphs per page for typical content
        estimated_lines_per_page = 28
        estimated_paragraphs_per_page = 2.5

        # Calculate page based on paragraph count
        page_from_paragraphs = max(1, int(total_paragraphs / estimated_paragraphs_per_page))

        # Add offset for cover page and TOC space
        toc_pages = 2
        cover_pages = 1 if self.word_settings.get('add_cover_page', False) else 0
        page_offset = cover_pages + toc_pages

        estimated_page = page_from_paragraphs + page_offset

        # Ensure logical progression (each heading should be on same or later page)
        if heading_counter > 1:
            estimated_page = max(estimated_page, page_offset + heading_counter)

        return estimated_page

    def _insert_paragraph_after(self, doc, reference_paragraph):
        """Insert a new paragraph after the reference paragraph"""
        # Find the reference paragraph in the document
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph == reference_paragraph:
                # Create new paragraph element
                new_paragraph = doc.add_paragraph()
                # Move it to the correct position
                # Note: python-docx doesn't have direct insertion, so we add at end and move
                return new_paragraph

        # Fallback: add at end
        return doc.add_paragraph()

    def _add_accurate_toc_entry(self, paragraph, heading_info):
        """Add a TOC entry with accurate page number and hyperlink"""
        try:
            text = heading_info['text']
            level = heading_info['level']
            page_number = heading_info['page_number']
            bookmark_name = heading_info['bookmark_name']
            heading_number = heading_info['heading_number']

            # Set indentation based on level
            base_indent = 0.2
            level_indent = (level - 1) * self.word_settings['list_indent_per_level_inches']
            paragraph.paragraph_format.left_indent = Inches(base_indent + level_indent)

            # Add entry number for main headings
            if level == 1:
                number_run = paragraph.add_run(f"{heading_number}. ")
                number_run.bold = True
                number_run.font.color.rgb = RGBColor(0, 0, 128)  # Dark blue

            # Add heading text with hyperlink
            try:
                hyperlink_run = self._add_hyperlink(paragraph, text, bookmark_name)
                if level == 1:
                    hyperlink_run.bold = True
                    hyperlink_run.font.size = Pt(12)
                elif level == 2:
                    hyperlink_run.font.size = Pt(11)
                else:
                    hyperlink_run.font.size = Pt(10)
                    hyperlink_run.italic = True
            except Exception as e:
                logger.debug(f"Could not create hyperlink for {text}: {e}")
                # Fallback to regular text
                text_run = paragraph.add_run(text)
                if level == 1:
                    text_run.bold = True
                    text_run.font.size = Pt(12)
                elif level == 2:
                    text_run.font.size = Pt(11)
                else:
                    text_run.font.size = Pt(10)
                    text_run.italic = True

            # Add dots and accurate page number
            dots_needed = max(3, 60 - len(text) - len(str(page_number)))
            dots_run = paragraph.add_run(" " + "." * dots_needed + " ")
            dots_run.font.color.rgb = RGBColor(128, 128, 128)

            page_run = paragraph.add_run(str(page_number))
            page_run.bold = True if level == 1 else False

        except Exception as e:
            logger.warning(f"Could not add accurate TOC entry: {e}")

    def _add_fallback_toc(self, doc, heading_page_map):
        """Fallback method to add TOC if placeholder method fails"""
        logger.info("Using fallback TOC generation method")

        # Insert TOC at the beginning (after title if present)
        first_paragraph = doc.paragraphs[0] if doc.paragraphs else doc.add_paragraph()

        # Add TOC title
        toc_title = doc.add_heading(self.word_settings['toc_title'], level=1)
        toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add entries
        sorted_headings = sorted(heading_page_map.values(), key=lambda x: x['heading_number'])
        for heading_info in sorted_headings:
            toc_paragraph = doc.add_paragraph()
            self._add_accurate_toc_entry(toc_paragraph, heading_info)

        # Add page break
        doc.add_page_break()

    def _add_content_block(self, doc, block_item, image_folder_path):
        """Add various content block types to the document. Handles both model objects and dicts."""
        try:
            block_handled = False
            if STRUCTURED_MODEL_AVAILABLE:
                if isinstance(block_item, Heading):
                    self._add_heading_block(doc, block_item)
                    block_handled = True
                elif isinstance(block_item, Paragraph):
                    self._add_paragraph_block(doc, block_item)
                    block_handled = True
                elif isinstance(block_item, ImagePlaceholder):
                    self._add_image_placeholder_block(doc, block_item, image_folder_path)
                    block_handled = True
                elif isinstance(block_item, ListItem):
                    self._add_list_item_block(doc, block_item)
                    block_handled = True
                elif isinstance(block_item, Table):
                    self._add_table_block(doc, block_item)
                    block_handled = True
                elif isinstance(block_item, CodeBlock):
                    self._add_code_block(doc, block_item)
                    block_handled = True
                # Add elif for other specific structured_document_model types here
                # e.g., Footnote, Equation, Caption, Metadata if they need special handling

            if not block_handled and isinstance(block_item, dict):
                dict_type = block_item.get('type', '').lower()
                if dict_type.startswith('h') and dict_type[1:].isdigit(): # h1, h2, etc.
                    self._add_heading_block(doc, block_item)
                elif dict_type in ['paragraph', 'text', 'p', 'normal']: # 'normal' for some legacy dicts
                    self._add_paragraph_block(doc, block_item)
                elif dict_type in ['image', 'img', 'imageplaceholder']:
                    self._add_image_placeholder_block(doc, block_item, image_folder_path)
                elif dict_type in ['list_item', 'listitem', 'li']:
                    self._add_list_item_block(doc, block_item)
                elif dict_type == 'table':
                    self._add_table_block(doc, block_item)
                elif dict_type in ['code_block', 'codeblock', 'code']:
                    self._add_code_block(doc, block_item)
                # Add other dict types as needed
                else:
                    # Fallback for unknown dict types: try to treat as paragraph if it has text
                    text_content = block_item.get('content') or block_item.get('text')
                    if text_content:
                        logger.warning(f"Unknown dictionary block type: '{dict_type}', but found text. Treating as paragraph.")
                        self._add_paragraph_block(doc, block_item)
                    else:
                        logger.warning(f"Unknown or empty dictionary block type: '{dict_type}'. Content: {str(block_item)[:200]}")
                block_handled = True # Assume dict was handled or logged

            if not block_handled: # Not a known model object and not a dict
                 logger.warning(f"Unknown block type: {type(block_item)}. Attempting to add as raw string.")
                 try:
                     # Attempt to add as a simple paragraph if it's some other basic type
                     doc.add_paragraph(str(block_item))
                 except Exception as e_raw:
                     logger.error(f"Could not add unknown block type {type(block_item)} as string: {e_raw}")

        except Exception as e:
            logger.error(f"Error processing content block (type: {type(block_item)}): {e}", exc_info=True)

    def _add_cover_page(self, doc, cover_page_data, image_folder_path):
        """Add cover page to document"""
        try:
            cover_paragraph = doc.add_paragraph()
            cover_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            if cover_page_data and image_folder_path:
                cover_image_path = cover_page_data.get('filepath')
                if cover_image_path and os.path.exists(cover_image_path):
                    run = cover_paragraph.add_run()
                    run.add_picture(cover_image_path, width=Inches(6))
                    logger.info("Cover page image added")
                else:
                    cover_paragraph.add_run("Cover Page").font.size = Pt(24)
            else:
                cover_paragraph.add_run("Cover Page").font.size = Pt(24)
                
        except Exception as e:
            logger.warning(f"Could not add cover page: {e}")
    
    def _add_table_of_contents(self, doc, structured_content_list):
        """Add enhanced table of contents with professional formatting"""
        try:
            # Add TOC title with enhanced styling and sanitization
            safe_toc_title = sanitize_for_xml(self.word_settings['toc_title'])
            toc_title = doc.add_heading(safe_toc_title, level=1)
            toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add decorative line under title
            title_paragraph = doc.add_paragraph()
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_paragraph.add_run("─" * 50)
            title_run.font.color.rgb = RGBColor(128, 128, 128)

            # Extract and organize headings for TOC
            headings = self._extract_toc_headings(structured_content_list)

            if not headings:
                logger.warning("No headings found for table of contents")
                return

            # Add TOC entries with enhanced formatting
            for i, heading in enumerate(headings):
                self._add_enhanced_toc_entry(doc, heading, i + 1)

            # Add spacing after TOC
            doc.add_paragraph()

            logger.info(f"Enhanced table of contents added with {len(headings)} entries")

        except Exception as e:
            logger.warning(f"Could not add table of contents: {e}")

    def _extract_toc_headings(self, structured_content_list):
        """Extract and organize headings for TOC with accurate page estimation"""
        raw_headings = []

        # Enhanced page estimation based on content analysis
        page_estimator = self._create_page_estimator()

        # First pass: collect all heading items with accurate page estimation
        for item in structured_content_list:
            item_type = item.get('type', '')

            # Update page estimation based on content type and length
            page_estimator.process_item(item)

            if item_type in ['h1', 'h2', 'h3']:
                text = item.get('text', '').strip()
                if text:
                    level = int(item_type[1])  # Extract level from h1, h2, h3
                    estimated_page = page_estimator.get_current_page()

                    raw_headings.append({
                        'text': text,
                        'level': level,
                        'estimated_page': estimated_page,
                        'original_page': item.get('page_num', estimated_page),
                        'bbox': item.get('bbox', [0, 0, 0, 0]),  # For proximity analysis
                        'content_position': page_estimator.get_position_info()
                    })

        # Second pass: merge consecutive headings that appear to be split
        merged_headings = self._merge_split_headings(raw_headings)

        # Third pass: clean and format the merged headings with refined page numbers
        final_headings = []
        for heading in merged_headings:
            clean_text = self._clean_heading_text(heading['text'])
            refined_page = self._refine_page_estimation(heading, merged_headings)

            final_headings.append({
                'text': clean_text,
                'level': heading['level'],
                'estimated_page': refined_page,
                'original_page': heading['original_page']
            })

        return final_headings

    def _configure_document_fonts_for_unicode(self, doc):
        """Configure document fonts to properly support Greek and other Unicode characters"""
        try:
            # Set the default font for the entire document to support Greek characters
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Arial Unicode MS'  # Comprehensive Unicode font
            font.size = Pt(11)
            
            # Set the complex script font for Greek and other complex scripts
            rpr = style.element.xpath('.//w:rPr')[0] if style.element.xpath('.//w:rPr') else None
            if rpr is not None:
                rFonts = rpr.find(qn('w:rFonts'))
                if rFonts is None:
                    rFonts = OxmlElement("w:rFonts")
                    rpr.insert(0, rFonts)
                
                # Set fonts for different script types
                rFonts.set(qn('w:ascii'), 'Arial Unicode MS')
                rFonts.set(qn('w:hAnsi'), 'Arial Unicode MS')
                rFonts.set(qn('w:cs'), 'Arial Unicode MS')  # Complex scripts (Greek, Arabic, etc.)
                rFonts.set(qn('w:eastAsia'), 'Arial Unicode MS')  # East Asian scripts
                rFonts.set(qn('w:asciiTheme'), 'Arial Unicode MS')
                rFonts.set(qn('w:hAnsiTheme'), 'Arial Unicode MS')
                rFonts.set(qn('w:eastAsiaTheme'), 'Arial Unicode MS')
                rFonts.set(qn('w:cstheme'), 'Arial Unicode MS')
            
            # Also configure heading styles
            for i in range(1, 7):  # Heading 1 through Heading 6
                try:
                    heading_style = doc.styles[f'Heading {i}']
                    heading_font = heading_style.font
                    heading_font.name = 'Arial Unicode MS'
                    
                    # Set complex script font for headings too
                    heading_rpr = heading_style.element.xpath('.//w:rPr')[0] if heading_style.element.xpath('.//w:rPr') else None
                    if heading_rpr is not None:
                        heading_rFonts = heading_rpr.find(qn('w:rFonts'))
                        if heading_rFonts is None:
                            heading_rFonts = OxmlElement("w:rFonts")
                            heading_rpr.insert(0, heading_rFonts)
                        
                        heading_rFonts.set(qn('w:ascii'), 'Arial Unicode MS')
                        heading_rFonts.set(qn('w:hAnsi'), 'Arial Unicode MS')
                        heading_rFonts.set(qn('w:cs'), 'Arial Unicode MS')
                        heading_rFonts.set(qn('w:eastAsia'), 'Arial Unicode MS')
                        heading_rFonts.set(qn('w:asciiTheme'), 'Arial Unicode MS')
                        heading_rFonts.set(qn('w:hAnsiTheme'), 'Arial Unicode MS')
                        heading_rFonts.set(qn('w:eastAsiaTheme'), 'Arial Unicode MS')
                        heading_rFonts.set(qn('w:cstheme'), 'Arial Unicode MS')
                        
                except KeyError:
                    # Style doesn't exist, skip
                    continue
            
            # Configure TOC styles
            try:
                toc_style = doc.styles['TOC 1']
                toc_font = toc_style.font
                toc_font.name = 'Arial Unicode MS'
                
                # Set complex script font for TOC
                toc_rpr = toc_style.element.xpath('.//w:rPr')[0] if toc_style.element.xpath('.//w:rPr') else None
                if toc_rpr is not None:
                    toc_rFonts = toc_rpr.find(qn('w:rPr'))
                    if toc_rFonts is None:
                        toc_rFonts = OxmlElement("w:rPr")
                        toc_rpr.insert(0, toc_rFonts)
                    
                    toc_rFonts.set(qn('w:ascii'), 'Arial Unicode MS')
                    toc_rFonts.set(qn('w:hAnsi'), 'Arial Unicode MS')
                    toc_rFonts.set(qn('w:cs'), 'Arial Unicode MS')
                    toc_rFonts.set(qn('w:eastAsia'), 'Arial Unicode MS')
                    toc_rFonts.set(qn('w:asciiTheme'), 'Arial Unicode MS')
                    toc_rFonts.set(qn('w:hAnsiTheme'), 'Arial Unicode MS')
                    toc_rFonts.set(qn('w:eastAsiaTheme'), 'Arial Unicode MS')
                    toc_rFonts.set(qn('w:cstheme'), 'Arial Unicode MS')
            except KeyError:
                logger.warning("TOC style not found, skipping TOC font configuration")
            
            logger.info("Document fonts configured for Unicode support (Greek, etc.)")
            
        except Exception as e:
            logger.warning(f"Could not configure Unicode fonts: {e}")

    def _insert_visual_content(self, doc, visual_content, page_num):
        """Insert visual content into the document without translation"""
        try:
            for area in visual_content:
                bbox = area['bbox']
                content_type = area['content_type']
                confidence = area.get('confidence', 0.5)

                # Only insert if confidence is high enough
                if confidence < 0.5:
                    logger.debug(f"Skipping low confidence visual content on page {page_num + 1}")
                    continue

                # Add a section break before visual content
                doc.add_section_break()

                # Add a caption indicating this is visual content
                p = doc.add_paragraph()
                p.add_run(f"[Original Visual Content - {content_type}]").italic = True

                # Insert the visual content without translation
                if content_type == 'page_with_drawings':
                    self._insert_drawing_area(doc, bbox, page_num)
                elif content_type == 'image_area':
                    self._insert_image_area(doc, bbox, page_num)
                elif content_type == 'sparse_text_page':
                    self._insert_sparse_text_area(doc, bbox, page_num)

                # Add a section break after visual content
                doc.add_section_break()

        except Exception as e:
            logger.error(f"Error inserting visual content: {e}")
            raise

    def _insert_drawing_area(self, doc, bbox, page_num):
        """Insert a drawing area without translation"""
        try:
            # Insert the drawing as is
            p = doc.add_paragraph()
            run = p.add_run()
            run.add_picture(f"page_{page_num + 1}_drawing.png")
            logger.info(f"Inserted original drawing from page {page_num + 1}")

        except Exception as e:
            logger.error(f"Error inserting drawing area: {e}")
            raise

    def _insert_image_area(self, doc, bbox, page_num):
        """Insert an image area without translation"""
        try:
            # Insert the image as is
            p = doc.add_paragraph()
            run = p.add_run()
            run.add_picture(f"page_{page_num + 1}_image.png")
            logger.info(f"Inserted original image from page {page_num + 1}")

        except Exception as e:
            logger.error(f"Error inserting image area: {e}")
            raise

    def _insert_sparse_text_area(self, doc, bbox, page_num):
        """Insert a sparse text area without translation"""
        try:
            # Insert the content as is
            p = doc.add_paragraph()
            run = p.add_run()
            run.add_picture(f"page_{page_num + 1}_sparse.png")
            logger.info(f"Inserted original sparse content from page {page_num + 1}")

        except Exception as e:
            logger.error(f"Error inserting sparse text area: {e}")
            raise

    def _is_area_translated(self, bbox, page_num):
        """Check if a visual content area has been properly translated"""
        try:
            # Get the text content in this area
            text = self.pdf_parser.get_text_in_area(page_num, bbox)
            
            # If there's no text, consider it translated
            if not text.strip():
                return True

            # Check if the text has been translated
            # This assumes you have a way to track translated content
            # You might need to implement this based on your translation tracking system
            return self._is_text_translated(text, page_num)

        except Exception as e:
            logger.error(f"Error checking area translation status: {e}")
            return False

    def _is_text_translated(self, text, page_num):
        """Check if text has been translated by comparing with translation cache"""
        try:
            # Get the translation cache from the translation service
            translation_cache = self.translation_service.get_translation_cache()
            
            # Check if this text exists in the cache
            if text in translation_cache:
                return True
                
            # If not in cache, check if it's part of a larger translated segment
            for original, translated in translation_cache.items():
                if text in original:
                    return True
                    
            # If we get here, the text hasn't been translated
            logger.warning(f"Text on page {page_num + 1} has not been translated")
            return False
            
        except Exception as e:
            logger.error(f"Error checking translation status: {e}")
            return False

# Enhanced PDF conversion function with proper font embedding
def convert_word_to_pdf(docx_filepath, pdf_filepath):
    """Enhanced PDF conversion with proper Greek font support"""
    logger.info(f"Converting {docx_filepath} to PDF with Unicode font support...")
    
    try:
        # Try using docx2pdf first
        from docx2pdf import convert
        convert(docx_filepath, pdf_filepath)
        logger.info(f"Successfully converted {docx_filepath} to {pdf_filepath}")
        return True
        
    except ImportError:
        logger.warning("docx2pdf library not available. Trying alternative method...")
        
        try:
            # Alternative: Use python-docx2txt + reportlab for better font control
            import docx2txt
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.fonts import addMapping
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            
            # Register Unicode font
            try:
                # Try to register Arial Unicode MS or fallback to DejaVu
                pdfmetrics.registerFont(TTFont('ArialUnicode', 'arial.ttf'))
                addMapping('ArialUnicode', 0, 0, 'ArialUnicode')
            except:
                # Fallback to DejaVu Sans which supports Greek
                try:
                    pdfmetrics.registerFont(TTFont('DejaVuSans', 'DejaVuSans.ttf'))
                    addMapping('DejaVuSans', 0, 0, 'DejaVuSans')
                    font_name = 'DejaVuSans'
                except:
                    font_name = 'Helvetica'  # Last resort
                    logger.warning("Could not register Unicode font, using Helvetica")
            
            # Extract text from DOCX
            text = docx2txt.process(docx_filepath)
            
            # Create PDF with proper font
            doc = SimpleDocTemplate(pdf_filepath, pagesize=letter)
            styles = getSampleStyleSheet()
            
            # Create custom style with Unicode font
            unicode_style = ParagraphStyle(
                'UnicodeNormal',
                parent=styles['Normal'],
                fontName=font_name,
                fontSize=11,
                encoding='utf-8'
            )
            
            # Build PDF content
            story = []
            for line in text.split('\n'):
                if line.strip():
                    story.append(Paragraph(line, unicode_style))
                    story.append(Spacer(1, 12))
            
            doc.build(story)
            logger.info(f"Successfully converted {docx_filepath} to {pdf_filepath} with Unicode support")
            return True
            
        except Exception as e:
            logger.error(f"Alternative PDF conversion failed: {e}")
            return False
            
    except Exception as e:
        logger.error(f"Error converting Word to PDF: {e}")
        return False

# Create a class to hold the PDF conversion function
class PDFConverter:
    def __init__(self):
        self.convert_word_to_pdf = convert_word_to_pdf

# Create an instance of the PDF converter
pdf_converter = PDFConverter()

# Instantiate the generator for use by other modules
document_generator = WordDocumentGenerator()