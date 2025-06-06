import os
import asyncio
import re
from dotenv import load_dotenv
import google.generativeai as genai
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx2pdf import convert as convert_to_pdf_lib
from pydrive2.auth import GoogleAuth
from pydrive2.drive import GoogleDrive
from collections import Counter
import tkinter as tk
from tkinter import filedialog
import shutil
import tempfile

# --- ΣΤΑΘΕΡΕΣ ΚΑΙ ΒΟΗΘΗΤΙΚΕΣ ΣΥΝΑΡΤΗΣΕΙΣ ---
load_dotenv()
API_KEY = os.getenv('GEMINI_API_KEY')
if not API_KEY:
    raise ValueError("Το GEMINI_API_KEY δεν βρέθηκε. Παρακαλώ ρυθμίστε το στο .env αρχείο σας.")
genai.configure(api_key=API_KEY)

# ΟΡΙΣΤΙΚΟ ΟΝΟΜΑ ΜΟΝΤΕΛΟΥ ΒΑΣΕΙ ΤΩΝ ΠΛΗΡΟΦΟΡΙΩΝ ΣΑΣ ΑΠΟ ΤΟ AI STUDIO
# Βεβαιωθείτε ότι αυτό είναι το ακριβές API ID που σας δίνει το "Get Code" στο AI Studio.
# Αν το "Get Code" στο AI Studio δείχνει π.χ. client = GenerativeModel('gemini-2.5-pro-preview-0506')
# τότε η παρακάτω τιμή είναι σωστή. Αν δείχνει κάτι με "models/" μπροστά, προσαρμόστε το.
MODEL_NAME = "models/gemini-2.5-pro-preview-03-25"
print(f"INFO: Χρησιμοποιείται το μοντέλο API: '{MODEL_NAME}'")

FOOTNOTE_MARKER_REGEX = re.compile(r"[\u00B9\u00B2\u00B3\u2070\u2074-\u2079]+|\[\s*\d+\s*\]")
BIBLIOGRAPHY_KEYWORDS_LOWER = [
    "bibliography", "references", "sources", "literature cited",
    "πηγές", "βιβλιογραφία", "αναφορές", "εργα που αναφερονται", "works cited"
]

def clean_text_of_markers(text):
    if not text: return ""
    return FOOTNOTE_MARKER_REGEX.sub("", text)

def get_desktop_path():
    desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
    if os.path.exists(desktop_path): return desktop_path
    desktop_path_el = os.path.join(os.path.expanduser("~"), "Επιφάνεια εργασίας")
    if os.path.exists(desktop_path_el): return desktop_path_el
    print("Προσοχή: Δεν εντοπίστηκε φάκελος Desktop/Επιφάνεια εργασίας. Χρήση αρχικού φακέλου χρήστη για fallback.")
    return os.path.expanduser("~")

def choose_input_pdf_file():
    print("Παρακαλώ επιλέξτε το αρχείο PDF για μετάφραση από το παράθυρο διαλόγου...")
    root = tk.Tk()
    root.withdraw()
    file_path = filedialog.askopenfilename(
        parent=root, title="Επιλέξτε αρχείο PDF",
        filetypes=(("PDF files", "*.pdf"), ("All files", "*.*"))
    )
    root.destroy()
    if file_path: print(f"Επιλέχθηκε το αρχείο: {file_path}")
    else: print("Δεν επιλέχθηκε αρχείο.")
    return file_path

def choose_output_directory(initial_dir=None):
    print("\nΠαρακαλώ επιλέξτε τον φάκελο όπου θα δημιουργηθεί ο υποφάκελος για τα μεταφρασμένα αρχεία...")
    root = tk.Tk()
    root.withdraw()
    if initial_dir is None: initial_dir = os.path.expanduser("~")
    chosen_directory = filedialog.askdirectory(
        parent=root, title="Επιλέξτε βασικό φάκελο αποθήκευσης μεταφρασμένων αρχείων",
        initialdir=initial_dir
    )
    root.destroy()
    if chosen_directory: print(f"Επιλεγμένος βασικός φάκελος αποθήκευσης: {chosen_directory}")
    else: print("Δεν επιλέχθηκε προσαρμοσμένος φάκελος αποθήκευσης.")
    return chosen_directory

def extract_structured_content_from_pdf(filepath):
    print(f"\n--- Εξαγωγή, Ανάλυση Δομής & Φιλτράρισμα Περιεχομένου από PDF (απλοποιημένη ανίχνευση bold): {filepath} ---")
    doc = fitz.open(filepath) # Άνοιγμα του PDF
    raw_extracted_items = []

    for page_num in range(len(doc)): # Για κάθε σελίδα του PDF
        page = doc.load_page(page_num)
        # Εξαγωγή κειμένου ως λεξικό, χωρίς το TEXTFLAGS_FONTS
        page_dict = page.get_text("dict", flags=fitz.TEXT_PRESERVE_LIGATURES | fitz.TEXT_PRESERVE_WHITESPACE)
        for block in page_dict["blocks"]: # Για κάθε μπλοκ κειμένου στη σελίδα
            if block['type'] == 0: # Αν το μπλοκ περιέχει κείμενο
                for line in block["lines"]: # Για κάθε γραμμή στο μπλοκ
                    line_text_parts = []
                    first_span = line["spans"][0] if line["spans"] else None
                    if not first_span: continue # Παράλειψη κενών γραμμών (χωρίς spans)
                    
                    current_line_font_size = round(first_span["size"], 1)
                    font_name_lower = first_span["font"].lower()
                    current_line_is_bold = "bold" in font_name_lower or "heavy" in font_name_lower or "black" in font_name_lower

                    line_text = "".join([s["text"] for s in line["spans"]]) # Συνένωση όλων των spans της γραμμής
                    if not line_text.strip(): continue # Παράλειψη αν η γραμμή είναι κενή μετά το strip
                        
                    raw_extracted_items.append({
                        'original_text': line_text,
                        'font_size': current_line_font_size,
                        'bold': current_line_is_bold,
                        'page_num': page_num
                    })
    
    doc.close() # Κλείσιμο του PDF
    print(f"Αρχική εξαγωγή: {len(raw_extracted_items)} γραμμές κειμένου.")
    if not raw_extracted_items: return []

    # Αυτόματος υπολογισμός ορίων γραμματοσειράς (όπως πριν)
    font_sizes_in_doc = [item['font_size'] for item in raw_extracted_items if item['font_size'] > 0]
    auto_min_font_for_inclusion = 7.0 
    dominant_font_size = 10.0 
    auto_pdf_h1_min_font = 14.0 
    auto_pdf_h2_min_font = 12.0 

    if font_sizes_in_doc:
        plausible_body_text_sizes = [s for s in font_sizes_in_doc if 7.0 <= s <= 18.0]
        if plausible_body_text_sizes:
            font_counts = Counter(plausible_body_text_sizes)
            dominant_font_size = font_counts.most_common(1)[0][0]
            auto_min_font_for_inclusion = max(7.0, dominant_font_size - 2.5) 
            print(f"INFO: Αυτόματος εντοπισμός: Κυρίαρχο μέγεθος (σώμα): {dominant_font_size}pt. Ελάχιστο για συμπερίληψη: {auto_min_font_for_inclusion:.1f}pt.")
        else: print(f"INFO: Δεν βρέθηκαν κυρίαρχα μεγέθη γραμματοσειράς. Χρήση προεπιλογών για όρια φιλτραρίσματος ({auto_min_font_for_inclusion:.1f}pt).")
        
        unique_larger_sizes = sorted(list(set(s for s in font_sizes_in_doc if s > dominant_font_size + 0.5 and s >= auto_min_font_for_inclusion)), reverse=True)
        if len(unique_larger_sizes) >= 2:
            auto_pdf_h1_min_font = unique_larger_sizes[0]
            auto_pdf_h2_min_font = unique_larger_sizes[1]
            if auto_pdf_h2_min_font >= auto_pdf_h1_min_font: auto_pdf_h2_min_font = max(dominant_font_size + 0.5, auto_pdf_h1_min_font - 0.1)
        elif len(unique_larger_sizes) == 1:
            auto_pdf_h1_min_font = unique_larger_sizes[0]
            auto_pdf_h2_min_font = max(dominant_font_size + 0.5, auto_pdf_h1_min_font - 2.0) # Προσπάθεια για λογικό H2
            if auto_pdf_h2_min_font >= auto_pdf_h1_min_font : auto_pdf_h2_min_font = auto_pdf_h1_min_font # Ή ίδιο αν μόνο ένα επίπεδο
        else:
            auto_pdf_h1_min_font = dominant_font_size + 3.0
            auto_pdf_h2_min_font = dominant_font_size + 1.5
            if auto_pdf_h2_min_font >= auto_pdf_h1_min_font : auto_pdf_h2_min_font = auto_pdf_h1_min_font - 0.1
        print(f"INFO: Αυτόματος εντοπισμός μεγεθών τίτλων: H1 >= {auto_pdf_h1_min_font:.1f}pt, H2 >= {auto_pdf_h2_min_font:.1f}pt.")
    else: print(f"INFO: Δεν βρέθηκαν πληροφορίες γραμματοσειράς. Χρήση προεπιλεγμένων τιμών για H1, H2 ({auto_pdf_h1_min_font}pt, {auto_pdf_h2_min_font}pt) και φιλτράρισμα ({auto_min_font_for_inclusion}pt).")

    processed_content = []
    bibliography_started = False
    for item in raw_extracted_items:
        cleaned_text_line = clean_text_of_markers(item['original_text'])
        cleaned_text_line_stripped = cleaned_text_line.strip()
        if not cleaned_text_line_stripped: continue
        if item['font_size'] < auto_min_font_for_inclusion: continue
            
        current_text_lower_stripped = cleaned_text_line_stripped.lower()
        is_potential_bib_heading = any(keyword == current_text_lower_stripped for keyword in BIBLIOGRAPHY_KEYWORDS_LOWER) or \
                                   any(current_text_lower_stripped.startswith(keyword + ":") for keyword in BIBLIOGRAPHY_KEYWORDS_LOWER) or \
                                   any(current_text_lower_stripped.startswith(keyword + " ") for keyword in BIBLIOGRAPHY_KEYWORDS_LOWER)
        
        # Το item['bold'] εδώ προέρχεται από την ευρετική ανίχνευση στο όνομα της γραμματοσειράς
        is_styled_as_heading_for_bib_check = item['font_size'] >= auto_pdf_h2_min_font -1 or \
                                            (item['bold'] and item['font_size'] >= auto_pdf_h2_min_font - 2) 
        if not bibliography_started and is_potential_bib_heading and is_styled_as_heading_for_bib_check:
            if len(cleaned_text_line_stripped.split()) < 10: # Ευρετικός κανόνας για μήκος επικεφαλίδας βιβλιογραφίας
                bibliography_started = True
                print(f"INFO: Εντοπίστηκε ενότητα βιβλιογραφίας: '{cleaned_text_line_stripped}'. Εξαιρείται από εδώ και πέρα.")
                break 
        
        item_type = 'p'
        if item['font_size'] >= auto_pdf_h1_min_font or \
           (item['bold'] and item['font_size'] >= auto_pdf_h1_min_font - 1.5): 
            item_type = 'h1'
        elif item['font_size'] >= auto_pdf_h2_min_font or \
             (item['bold'] and item['font_size'] >= auto_pdf_h2_min_font - 1.5): 
            item_type = 'h2'
        
        if processed_content and item_type == 'p' and processed_content[-1]['type'] == 'p' and \
           abs(item['font_size'] - processed_content[-1]['font_size']) < 1 and item['bold'] == processed_content[-1]['bold']:
            processed_content[-1]['text'] += " " + cleaned_text_line_stripped # Συνένωση παραγράφων
        else:
            processed_content.append({
                'type': item_type, 'text': cleaned_text_line_stripped,
                'font_size': item['font_size'], 'bold': item['bold']
            })
    print(f"Επεξεργασμένα δομικά στοιχεία (μετά από φίλτρα): {len(processed_content)}")
    return processed_content

async def call_gemini_for_translation(item_index, text_to_translate, prev_ctx_text, next_ctx_text,
                                      target_language, semaphore, model_name_to_use, temperature):
    # (Η συνάρτηση παραμένει αμετάβλητη)
    async with semaphore:
        model = genai.GenerativeModel(model_name_to_use)
        prompt = f"""You are an expert translator. Translate the "MAIN TEXT TO TRANSLATE" into {target_language} with precision.
PREVIOUS CONTEXT (original language, for context only, do not translate this part):
---
{prev_ctx_text if prev_ctx_text else "N/A"}
---
MAIN TEXT TO TRANSLATE:
---
{text_to_translate}
---
NEXT CONTEXT (original language, for context only, do not translate this part):
---
{next_ctx_text if next_ctx_text else "N/A"}
---
Translation of "MAIN TEXT TO TRANSLATE" into {target_language} (provide only the translation):
"""
        try:
            response = await model.generate_content_async(
                prompt,
                generation_config=genai.types.GenerationConfig(temperature=temperature)
            )
            if response.parts:
                return item_index, response.text.strip()
            else:
                err_msg = f"[Μετάφραση απέτυχε για στοιχείο {item_index + 1}: Κενή ή μπλοκαρισμένη απάντηση."
                if hasattr(response, 'prompt_feedback') and response.prompt_feedback: err_msg += f" Feedback: {response.prompt_feedback}"
                print(err_msg)
                return item_index, err_msg
        except Exception as e:
            print(f"Σφάλμα στη μετάφραση του στοιχείου {item_index + 1}: {e}")
            return item_index, f"[Μετάφραση απέτυχε για στοιχείο {item_index + 1} λόγω εξαίρεσης: {e}]"

async def perform_language_check(full_translated_text, target_language, model_name_to_use, temperature=0.2):
    # (Η συνάρτηση παραμένει αμετάβλητη)
    if not full_translated_text.strip(): return "Δεν παρασχέθηκε κείμενο για έλεγχο."
    print("\n--- Εκτέλεση AI Γλωσσικού Ελέγχου ---")
    model = genai.GenerativeModel(model_name_to_use)
    max_chars_for_review = 700000 
    text_for_review = full_translated_text
    if len(text_for_review) > max_chars_for_review:
        print(f"INFO: Γλωσσικός έλεγχος: Το κείμενο είναι πολύ μεγάλο ({len(text_for_review)} chars). Έλεγχος τμήματος ({max_chars_for_review} chars).")
        text_for_review = full_translated_text[:max_chars_for_review]
    prompt = f"""You are an expert linguistic reviewer. The following text is a translation into {target_language}.
Review for: Clarity, Coherence, Thematic Language/Terminology, Grammar/Orthography.
Provide a concise report with findings and specific suggestions.
Translated Text to Review:
---
{text_for_review}
---
Your Review and Suggestions:"""
    try:
        print("INFO: Γλωσσικός έλεγχος: Αποστολή αιτήματος...")
        response = await model.generate_content_async(prompt, generation_config=genai.types.GenerationConfig(temperature=temperature))
        if response.parts:
            review_text = response.text.strip()
            print("INFO: Γλωσσικός έλεγχος: Η αναφορά ελήφθη.")
            return review_text
        else:
            error_message = "Σφάλμα Γλωσσικού ελέγχου: Κενή ή μπλοκαρισμένη απάντηση."
            if hasattr(response, 'prompt_feedback') and response.prompt_feedback: error_message += f" Prompt feedback: {response.prompt_feedback}"
            print(error_message)
            return error_message
    except Exception as e:
        print(f"Σφάλμα Γλωσσικού ελέγχου: {e}")
        return f"Σφάλμα κατά τον γλωσσικό έλεγχο: {e}"

def add_bookmark_to_paragraph(paragraph, bookmark_name):
    # (Η συνάρτηση παραμένει αμετάβλητη)
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    tag = run._r.get_or_add_CT_Bookmark()
    if not tag.name: tag.name = bookmark_name
    tag.id = str(abs(hash(bookmark_name)) % (10**8)) 

def add_hyperlink_to_paragraph(paragraph, text_to_display, anchor_bookmark_name):
    # (Η συνάρτηση παραμένει αμετάβλητη)
    part = paragraph.part
    r_id = part.relate_to(f"#{anchor_bookmark_name}", "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=False)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('w:anchor'), anchor_bookmark_name)
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    new_run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text_to_display
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

def create_word_document_with_structure(structured_content_list, output_filepath):
    # (Η συνάρτηση παραμένει αμετάβλητη)
    print(f"\n--- Δημιουργία Δομημένου Εγγράφου Word: {output_filepath} ---")
    doc = Document()
    toc_items = []
    heading_bookmark_counter = 0
    for item in structured_content_list:
        text_to_add = item.get('translated_text', item.get('text', ''))
        if item['type'] in ['h1', 'h2'] and text_to_add.strip() and not text_to_add.startswith("[Μετάφραση απέτυχε"):
            heading_bookmark_counter += 1
            bookmark = f"_Toc_Agent_Final_{heading_bookmark_counter}"
            toc_items.append({'text': text_to_add, 'level': 1 if item['type'] == 'h1' else 2, 'bookmark': bookmark})

    if toc_items:
        toc_p_title = doc.add_paragraph()
        toc_p_title.add_run("Πίνακας Περιεχομένων").bold = True
        for toc_item in toc_items:
            p_toc_entry = doc.add_paragraph()
            if toc_item['level'] == 2: p_toc_entry.paragraph_format.left_indent = Inches(0.25)
            add_hyperlink_to_paragraph(p_toc_entry, toc_item['text'], toc_item['bookmark'])
        doc.add_page_break()

    heading_bookmark_counter = 0 
    for item in structured_content_list:
        text_to_add = item.get('translated_text', item.get('text', ''))
        if not text_to_add.strip():
            if item.get('original_text', '').strip(): doc.add_paragraph("")
            continue
        if text_to_add.startswith("[Μετάφραση απέτυχε"):
            doc.add_paragraph().add_run(text_to_add).italic = True; continue
        current_paragraph, current_bookmark_name = None, None
        if item['type'] == 'h1':
            current_paragraph = doc.add_heading(text_to_add, level=1)
            heading_bookmark_counter += 1; current_bookmark_name = f"_Toc_Agent_Final_{heading_bookmark_counter}"
        elif item['type'] == 'h2':
            current_paragraph = doc.add_heading(text_to_add, level=2)
            heading_bookmark_counter += 1; current_bookmark_name = f"_Toc_Agent_Final_{heading_bookmark_counter}"
        else: current_paragraph = doc.add_paragraph(text_to_add)
        if current_paragraph and current_bookmark_name:
            if not current_paragraph.runs: current_paragraph.add_run()
            add_bookmark_to_paragraph(current_paragraph, current_bookmark_name)
    try:
        doc.save(output_filepath)
        print(f"INFO: Το δομημένο Word αποθηκεύτηκε: {output_filepath}"); return True
    except Exception as e: print(f"Σφάλμα αποθήκευσης Word: {e}"); return False

def convert_word_to_pdf(word_filepath, pdf_filepath):
    # (Η συνάρτηση παραμένει αμετάβλητη)
    print(f"\n--- Μετατροπή Word σε PDF: {word_filepath} -> {pdf_filepath} ---")
    try:
        convert_to_pdf_lib(word_filepath, pdf_filepath)
        print("INFO: Η μετατροπή σε PDF ολοκληρώθηκε."); return True
    except Exception as e: print(f"Σφάλμα μετατροπής σε PDF: {e}"); return False

def upload_to_google_drive(filepath_to_upload, filename_on_drive, gdrive_parent_folder_id=None):
    # (Η συνάρτηση παραμένει αμετάβλητη)
    print(f"\n--- Ανέβασμα '{filename_on_drive}' στο Google Drive ---")
    try:
        gauth = GoogleAuth(); gauth.settings['client_config_file'] = "client_secrets.json"
        gauth.LoadCredentialsFile("credentials.json")
        if gauth.credentials is None: gauth.LocalWebserverAuth()
        elif gauth.access_token_expired: gauth.Refresh()
        else: gauth.Authorize()
        gauth.SaveCredentialsFile("credentials.json"); drive = GoogleDrive(gauth)
        file_metadata = {'title': filename_on_drive}
        if gdrive_parent_folder_id: file_metadata['parents'] = [{'id': gdrive_parent_folder_id}]
        gfile = drive.CreateFile(file_metadata); gfile.SetContentFile(filepath_to_upload)
        gfile.Upload()
        print(f"INFO: Το αρχείο '{filename_on_drive}' ανέβηκε. Link: {gfile.get('alternateLink', 'N/A')}")
        return gfile.get('id')
    except Exception as e: print(f"Σφάλμα ανεβάσματος στο Google Drive: {e}"); return None

async def translate_document_agent_async(
    filepath,
    translation_temperature,
    max_concurrent_requests,
    gdrive_folder_id_target=None,
    target_language="Ελληνικά"
):
    # (Η συνάρτηση παραμένει αμετάβλητη, όπως στην τελευταία πλήρη έκδοση που σας έδωσα,
    #  με τη διαφορά ότι η extract_structured_content_from_pdf καλείται χωρίς παραμέτρους μεγέθους γραμματοσειράς)
    print(f"Ξεκινά η πλήρης διαδικασία για το αρχείο: {filepath}...")
    print(f"Γλώσσα στόχος: {target_language}")
    print(f"Ρυθμίσεις Μετάφρασης: Temp={translation_temperature}, Max Concurrent Req={max_concurrent_requests}")

    if not filepath.lower().endswith(".pdf"):
        print("Σφάλμα: Υποστηρίζονται μόνο αρχεία PDF."); return

    structured_content_to_translate = extract_structured_content_from_pdf(filepath)
    if not structured_content_to_translate:
        print("Δεν εξήχθη περιεχόμενο ή όλο φιλτραρίστηκε. Τερματισμός."); return

    semaphore = asyncio.Semaphore(max_concurrent_requests)
    tasks = []
    translated_structured_content = [item.copy() for item in structured_content_to_translate]
    print(f"\n--- Προετοιμασία {len(structured_content_to_translate)} τμημάτων για μετάφραση ---")
    for i, item in enumerate(structured_content_to_translate):
        text_to_translate = item['text']
        if not text_to_translate.strip():
            translated_structured_content[i]['translated_text'] = ""; continue
        # Ο έλεγχος για πολύ μεγάλα τμήματα έχει αφαιρεθεί (υποθέτοντας μεγάλο context window)
        prev_ctx_text = structured_content_to_translate[i-1]['text'][-500:] if i > 0 and structured_content_to_translate[i-1]['text'].strip() else ""
        next_ctx_text = structured_content_to_translate[i+1]['text'][:500] if i < len(structured_content_to_translate) - 1 and structured_content_to_translate[i+1]['text'].strip() else ""
        tasks.append(call_gemini_for_translation(i, text_to_translate, prev_ctx_text, next_ctx_text, target_language, semaphore, MODEL_NAME, translation_temperature))
    
    if not tasks: print("Δεν υπάρχουν τμήματα προς μετάφραση."); return
    print(f"INFO: Εκτέλεση {len(tasks)} εργασιών μετάφρασης...")
    translation_results = await asyncio.gather(*tasks)
    for original_index, translated_text_or_error in translation_results:
        translated_structured_content[original_index]['translated_text'] = translated_text_or_error
    
    full_plain_translated_text = "\n".join(item.get('translated_text', '') for item in translated_structured_content if item.get('translated_text', '').strip() and not item['translated_text'].startswith("[Μετάφραση απέτυχε"))

    if full_plain_translated_text.strip():
        ai_review_feedback = await perform_language_check(full_plain_translated_text, target_language, MODEL_NAME)
        print("\n--- Αναφορά Γλωσσικού Ελέγχου από AI ---\n" + ai_review_feedback + "\n--- Τέλος Αναφοράς ---")
    else: print("\nΠαράλειψη γλωσσικού ελέγχου.")

    base_original_filename = os.path.splitext(os.path.basename(filepath))[0]
    word_output_filename = f"{base_original_filename}_({target_language})_FINAL.docx"
    pdf_output_filename = f"{base_original_filename}_({target_language})_FINAL.pdf"
    word_created_successfully, pdf_created_successfully = False, False
    final_word_path_local, final_pdf_path_local = None, None

    with tempfile.TemporaryDirectory() as temp_dir:
        print(f"\nINFO: Χρήση προσωρινού φακέλου: {temp_dir}")
        temp_word_path = os.path.join(temp_dir, word_output_filename)
        temp_pdf_path = os.path.join(temp_dir, pdf_output_filename)

        word_created_successfully = create_word_document_with_structure(translated_structured_content, temp_word_path)
        if word_created_successfully:
            pdf_created_successfully = convert_word_to_pdf(temp_word_path, temp_pdf_path)

        user_chosen_base_dir = choose_output_directory(initial_dir=os.path.dirname(filepath))
        target_dir_for_job_files = None
        if user_chosen_base_dir:
            job_specific_subfolder_name = f"{base_original_filename}_Μεταφρασμένο"
            target_dir_for_job_files = os.path.join(user_chosen_base_dir, job_specific_subfolder_name)
        else:
            print("\nINFO: Δεν επιλέχθηκε προσαρμοσμένος φάκελος. Αποθήκευση δίπλα στο αρχικό PDF.")
            target_dir_for_job_files = os.path.dirname(filepath) if os.path.dirname(filepath) else "."
        
        try:
            os.makedirs(target_dir_for_job_files, exist_ok=True)
            print(f"INFO: Τα τελικά αρχεία θα αποθηκευτούν/μετακινηθούν στον φάκελο: {target_dir_for_job_files}")
        except Exception as e_mkdir:
            print(f"ΠΡΟΣΟΧΗ: Σφάλμα δημιουργίας φακέλου '{target_dir_for_job_files}': {e_mkdir}. Χρήση προσωρινού για ανέβασμα.")
            target_dir_for_job_files = temp_dir 

        if word_created_successfully:
            final_word_path_local = os.path.join(target_dir_for_job_files, word_output_filename)
            try:
                if os.path.abspath(temp_word_path) != os.path.abspath(final_word_path_local): shutil.move(temp_word_path, final_word_path_local)
                print(f"INFO: Το Word είναι διαθέσιμο στο: {final_word_path_local}")
            except Exception as e_move_w:
                print(f"Σφάλμα μετακίνησης Word: {e_move_w}. Αρχείο παραμένει στο: {temp_word_path}")
                final_word_path_local = temp_word_path 
        
        if pdf_created_successfully:
            final_pdf_path_local = os.path.join(target_dir_for_job_files, pdf_output_filename)
            try:
                if os.path.abspath(temp_pdf_path) != os.path.abspath(final_pdf_path_local): shutil.move(temp_pdf_path, final_pdf_path_local)
                print(f"INFO: Το PDF είναι διαθέσιμο στο: {final_pdf_path_local}")
            except Exception as e_move_p:
                print(f"Σφάλμα μετακίνησης PDF: {e_move_p}. Αρχείο παραμένει στο: {temp_pdf_path}")
                final_pdf_path_local = temp_pdf_path
        
        if word_created_successfully and final_word_path_local and os.path.exists(final_word_path_local):
            upload_to_google_drive(final_word_path_local, os.path.basename(final_word_path_local), gdrive_folder_id_target)
        if pdf_created_successfully and final_pdf_path_local and os.path.exists(final_pdf_path_local):
            upload_to_google_drive(final_pdf_path_local, os.path.basename(final_pdf_path_local), gdrive_folder_id_target)
            
    print("-" * 50 + f"\nΗ διαδικασία για το '{filepath}' ολοκληρώθηκε.\n" + "-" * 50)
    if final_word_path_local and os.path.exists(final_word_path_local):
        print(f"Τελικός φάκελος αποθήκευσης τοπικά: {os.path.dirname(final_word_path_local)}")
    else: print("ΠΡΟΣΟΧΗ: Δεν φαίνεται να ορίστηκε ή να βρέθηκε η τελική τοπική θέση αποθήκευσης των αρχείων.")

# --- ΚΥΡΙΑ ΕΚΤΕΛΕΣΗ ---
if __name__ == "__main__":
    INPUT_FILE_PATH = choose_input_pdf_file()
    if not INPUT_FILE_PATH:
        print("Τερματισμός: Δεν επιλέχθηκε αρχείο εισόδου.")
    else:
        TRANSLATION_TEMPERATURE = 0.1
        MAX_CONCURRENT_API_CALLS = 5
        GDRIVE_TARGET_FOLDER_ID = None 

        if not os.path.exists(INPUT_FILE_PATH):
            print(f"Σφάλμα: Το αρχείο '{INPUT_FILE_PATH}' δεν βρέθηκε. Δημιουργία δοκιμαστικού.")
            try:
                doc = fitz.open() 
                page = doc.new_page(width=595, height=842) 
                page.insert_text((72, 72), "Κεφάλαιο 1: Η Αρχή", fontsize=18, fontname="helvb")
                page.insert_text((72, 108), "Αυτή είναι η πρώτη παράγραφος.", fontsize=11, fontname="helv")
                page.insert_text((72, 126), "Μια γραμμή με δείκτη υποσημείωσης1 και άλλη μία [2].", fontsize=11, fontname="helv")
                page.insert_text((72, 162), "1.1 Μια Υποενότητα", fontsize=14, fontname="helvb")
                page.insert_text((72, 180), "Περιεχόμενο για την υποενότητα.", fontsize=11, fontname="helv")
                page.insert_text((72, page.rect.height - 72), "1 Αυτό είναι κείμενο υποσημείωσης.", fontsize=7, fontname="helv")
                page2 = doc.new_page(width=595, height=842)
                page2.insert_text((72, 72), "Βιβλιογραφία", fontsize=16, fontname="helvb")
                page2.insert_text((72, 108), "Συγγραφέας Α. (2024). Τίτλος Βιβλίου. Εκδόσεις.", fontsize=10, fontname="helv")
                doc.save(INPUT_FILE_PATH)
                doc.close()
                print(f"INFO: Δημιουργήθηκε δοκιμαστικό PDF: '{INPUT_FILE_PATH}'")
            except Exception as e_pdf:
                print(f"Σφάλμα δημιουργίας δοκιμαστικού PDF: {e_pdf}. Παρακαλώ δώστε ένα υπάρχον PDF. Τερματισμός."); exit()
        
        try:
            asyncio.run(
                translate_document_agent_async(
                    filepath=INPUT_FILE_PATH,
                    translation_temperature=TRANSLATION_TEMPERATURE,
                    max_concurrent_requests=MAX_CONCURRENT_API_CALLS,
                    gdrive_folder_id_target=GDRIVE_TARGET_FOLDER_ID
                )
            )
        except RuntimeError as e:
            if str(e).startswith("asyncio.run() cannot be called"): print("Σφάλμα asyncio: Καλέστε με 'await' σε Jupyter.")
            else: raise
        except Exception as e_main:
            print(f"Κύριο σφάλμα εκτέλεσης: {e_main}")