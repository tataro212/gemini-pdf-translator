#!/usr/bin/env python3
"""
Verify Images in Translated Document

This script checks if images are actually present in the translated Word document.

Usage: python verify_images_in_translated_document.py
"""

import os
import sys
import logging
from docx import Document
import zipfile
import tempfile

# Setup logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

def check_images_in_docx(docx_path):
    """Check if a Word document contains images"""
    
    if not os.path.exists(docx_path):
        logger.error(f"❌ Document not found: {docx_path}")
        return False
    
    logger.info(f"🔍 Checking images in: {os.path.basename(docx_path)}")
    
    try:
        # Method 1: Check using python-docx
        doc = Document(docx_path)
        
        image_count = 0
        image_placeholders = 0
        
        for paragraph in doc.paragraphs:
            # Check for image placeholders in text
            if '[Image' in paragraph.text or '[image' in paragraph.text:
                image_placeholders += 1
                logger.info(f"📝 Found image placeholder: {paragraph.text[:100]}...")
            
            # Check for actual images in runs
            for run in paragraph.runs:
                if run._element.xpath('.//pic:pic', namespaces={'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture'}):
                    image_count += 1
                    logger.info(f"🖼️ Found embedded image in paragraph")
        
        logger.info(f"📊 python-docx analysis:")
        logger.info(f"   Embedded images: {image_count}")
        logger.info(f"   Image placeholders: {image_placeholders}")
        
        # Method 2: Check the ZIP structure (Word docs are ZIP files)
        logger.info(f"📦 Checking ZIP structure...")
        
        with zipfile.ZipFile(docx_path, 'r') as zip_file:
            file_list = zip_file.namelist()
            
            # Look for media files
            media_files = [f for f in file_list if f.startswith('word/media/')]
            
            logger.info(f"📁 Media files in document: {len(media_files)}")
            for media_file in media_files:
                file_info = zip_file.getinfo(media_file)
                logger.info(f"   📷 {media_file} ({file_info.file_size} bytes)")
            
            # Check document.xml for image references
            if 'word/document.xml' in file_list:
                with zip_file.open('word/document.xml') as doc_xml:
                    content = doc_xml.read().decode('utf-8')
                    
                    # Count image references
                    pic_count = content.count('<pic:pic')
                    drawing_count = content.count('<w:drawing>')
                    
                    logger.info(f"📄 XML analysis:")
                    logger.info(f"   <pic:pic> tags: {pic_count}")
                    logger.info(f"   <w:drawing> tags: {drawing_count}")
        
        # Summary
        total_images = max(image_count, len(media_files), pic_count)
        
        if total_images > 0:
            logger.info(f"✅ Document contains {total_images} images")
            return True
        elif image_placeholders > 0:
            logger.warning(f"⚠️ Document contains {image_placeholders} image placeholders but no actual images")
            return False
        else:
            logger.warning("❌ No images or image placeholders found")
            return False
            
    except Exception as e:
        logger.error(f"❌ Error checking document: {e}")
        return False

def check_recent_translations():
    """Check recent translation outputs for images"""
    
    logger.info("🔍 Checking recent translation outputs...")
    
    # Possible locations for translated documents
    possible_paths = [
        "C:\\Users\\30694\\Downloads\\sickdays\\A World Beyond Physics _ The Emergence and Evolution of Life\\A World Beyond Physics _ The Emergence and Evolution of Life _translated.docx",
        "A World Beyond Physics _ The Emergence and Evolution of Life _translated.docx",
        "translated_documents\\A World Beyond Physics _ The Emergence and Evolution of Life _translated.docx"
    ]
    
    found_documents = []
    
    for path in possible_paths:
        if os.path.exists(path):
            found_documents.append(path)
            logger.info(f"📄 Found document: {path}")
    
    if not found_documents:
        logger.warning("⚠️ No translated documents found")
        return False
    
    # Check each found document
    results = []
    for doc_path in found_documents:
        logger.info(f"\n{'='*60}")
        logger.info(f"📋 ANALYZING: {os.path.basename(doc_path)}")
        logger.info(f"{'='*60}")
        
        file_size = os.path.getsize(doc_path)
        logger.info(f"📊 File size: {file_size:,} bytes ({file_size/1024:.1f} KB)")
        
        has_images = check_images_in_docx(doc_path)
        results.append((doc_path, has_images, file_size))
    
    # Summary
    logger.info(f"\n{'='*60}")
    logger.info("📊 SUMMARY")
    logger.info(f"{'='*60}")
    
    for doc_path, has_images, file_size in results:
        status = "✅ HAS IMAGES" if has_images else "❌ NO IMAGES"
        logger.info(f"{status}: {os.path.basename(doc_path)} ({file_size:,} bytes)")
    
    return any(has_images for _, has_images, _ in results)

def create_test_document_with_images():
    """Create a test document with images to verify the process works"""
    
    logger.info("🧪 Creating test document with images...")
    
    try:
        from document_generator import WordDocumentGenerator
        from structured_document_model import Document as StructuredDocument, ImagePlaceholder, Paragraph, ContentType
        
        # Create a simple test document with image placeholders
        test_blocks = []
        
        # Add a paragraph
        test_blocks.append(Paragraph(
            block_type=ContentType.PARAGRAPH,
            original_text="This is a test document with images.",
            page_num=1,
            bbox=(0, 0, 100, 20),
            content="This is a test document with images."
        ))
        
        # Add an image placeholder (using a known image if available)
        test_image_paths = [
            "page_1_img_1.png",
            "test_image.png",
            "sample.png"
        ]
        
        image_path = None
        for path in test_image_paths:
            if os.path.exists(path):
                image_path = path
                break
        
        if image_path:
            test_blocks.append(ImagePlaceholder(
                block_type=ContentType.IMAGE_PLACEHOLDER,
                original_text="",
                page_num=1,
                bbox=(0, 30, 100, 130),
                image_path=image_path,
                width=100,
                height=100
            ))
            logger.info(f"📷 Using test image: {image_path}")
        else:
            logger.warning("⚠️ No test images found")
        
        # Create structured document
        test_document = StructuredDocument(
            title="Test Document with Images",
            content_blocks=test_blocks,
            source_filepath="test.pdf",
            total_pages=1,
            metadata={'test': True}
        )
        
        # Generate Word document
        generator = WordDocumentGenerator()
        output_path = "test_image_verification.docx"
        
        result_path = generator.create_word_document_from_structured_document(
            structured_document=test_document,
            output_filepath=output_path,
            image_folder_path="."
        )
        
        if result_path and os.path.exists(result_path):
            logger.info(f"✅ Test document created: {result_path}")
            
            # Check the test document
            has_images = check_images_in_docx(result_path)
            
            if has_images:
                logger.info("✅ Test document generation with images works correctly!")
                return True
            else:
                logger.warning("⚠️ Test document created but no images detected")
                return False
        else:
            logger.error("❌ Failed to create test document")
            return False
            
    except Exception as e:
        logger.error(f"❌ Test document creation failed: {e}")
        return False

def main():
    """Main verification function"""
    
    logger.info("🔍 VERIFYING IMAGES IN TRANSLATED DOCUMENTS")
    logger.info("="*60)
    
    # Test 1: Check recent translations
    logger.info("\n🧪 Test 1: Checking recent translation outputs")
    recent_has_images = check_recent_translations()
    
    # Test 2: Create test document
    logger.info("\n🧪 Test 2: Creating test document with images")
    test_works = create_test_document_with_images()
    
    # Final conclusion
    logger.info("\n" + "="*60)
    logger.info("🎯 FINAL CONCLUSION")
    logger.info("="*60)
    
    if recent_has_images:
        logger.info("✅ IMAGES ARE PRESENT in translated documents!")
        logger.info("💡 If you're not seeing them, check your document viewer settings")
    elif test_works:
        logger.info("⚠️ Image processing works, but recent translations may have issues")
        logger.info("💡 The translation process might be losing images during processing")
    else:
        logger.info("❌ Image processing has issues that need investigation")
        logger.info("💡 Check image paths and document generation logic")
    
    return recent_has_images or test_works

if __name__ == "__main__":
    success = main()
    sys.exit(0 if success else 1)
