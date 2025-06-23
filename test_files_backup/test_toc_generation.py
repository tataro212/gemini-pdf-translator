import logging
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def create_test_document():
    """Create a test document with various heading levels and content"""
    doc = Document()
    
    # Add title
    title = doc.add_heading('Advanced Document Processing System', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add introduction
    doc.add_heading('Introduction', level=1)
    doc.add_paragraph('This document demonstrates the capabilities of our advanced document processing system, including sophisticated table of contents generation and content organization.')
    
    # Add main sections
    doc.add_heading('System Architecture', level=1)
    doc.add_paragraph('The system is built on a modular architecture that allows for flexible processing of various document types.')
    
    # Add subsections
    doc.add_heading('Core Components', level=2)
    doc.add_paragraph('The core components form the foundation of the system.')
    
    doc.add_heading('Document Parser', level=3)
    doc.add_paragraph('Responsible for initial document analysis and structure extraction.')
    
    doc.add_heading('Content Processor', level=3)
    doc.add_paragraph('Handles the processing and transformation of document content.')
    
    doc.add_heading('Output Generator', level=3)
    doc.add_paragraph('Creates the final output in various formats with proper formatting.')
    
    # Add another main section
    doc.add_heading('Implementation Details', level=1)
    doc.add_paragraph('Detailed information about the system implementation.')
    
    # Add subsections with more content
    doc.add_heading('Data Structures', level=2)
    doc.add_paragraph('The system uses sophisticated data structures to maintain document integrity.')
    
    doc.add_heading('Content Blocks', level=3)
    doc.add_paragraph('Content blocks are the fundamental units of document structure.')
    
    doc.add_heading('Formatting Rules', level=3)
    doc.add_paragraph('Rules that govern how content is formatted in the output.')
    
    # Add another main section
    doc.add_heading('Performance Optimization', level=1)
    doc.add_paragraph('Techniques used to optimize system performance.')
    
    # Add subsections
    doc.add_heading('Caching Strategy', level=2)
    doc.add_paragraph('How the system caches processed content for improved performance.')
    
    doc.add_heading('Memory Management', level=2)
    doc.add_paragraph('Strategies for efficient memory usage during document processing.')
    
    # Add conclusion
    doc.add_heading('Conclusion', level=1)
    doc.add_paragraph('The system provides a robust solution for document processing with advanced features.')
    
    return doc

def add_toc(doc):
    """Add a professionally formatted table of contents"""
    # Add TOC title
    toc_title = doc.add_paragraph('Table of Contents', style='Title')
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add decorative line
    decorative = doc.add_paragraph()
    decorative.alignment = WD_ALIGN_PARAGRAPH.CENTER
    decorative_run = decorative.add_run("─" * 50)
    decorative_run.font.color.rgb = RGBColor(128, 128, 128)
    
    # Collect headings
    headings = []
    for para in doc.paragraphs:
        if para.style.name.startswith('Heading'):
            level = int(para.style.name[-1]) if para.style.name[-1].isdigit() else 1
            headings.append({
                'text': para.text,
                'level': level
            })
    
    # Add TOC entries
    for heading in headings:
        p = doc.add_paragraph()
        p.style = 'Normal'
        
        # Set indentation based on level
        indent_size = Pt(20 * (heading['level'] - 1))
        p.paragraph_format.left_indent = indent_size
        
        # Add heading text
        text_run = p.add_run(heading['text'])
        if heading['level'] == 1:
            text_run.bold = True
            text_run.font.size = Pt(12)
        elif heading['level'] == 2:
            text_run.font.size = Pt(11)
        else:
            text_run.font.size = Pt(10)
            text_run.italic = True
        
        # Add dots
        dots_needed = max(3, 60 - len(heading['text']))
        dots_run = p.add_run(" " + "." * dots_needed + " ")
        dots_run.font.color.rgb = RGBColor(128, 128, 128)
        
        # Add page number (placeholder)
        page_run = p.add_run("1")
        if heading['level'] == 1:
            page_run.bold = True
    
    # Add page break after TOC
    doc.add_page_break()

def main():
    """Main function to test TOC generation"""
    logger.info("Creating test document...")
    doc = create_test_document()
    
    logger.info("Adding table of contents...")
    add_toc(doc)
    
    # Save the document
    output_path = "test_toc_document.docx"
    doc.save(output_path)
    logger.info(f"Document saved as: {output_path}")
    
    # Print document statistics and structure
    logger.info("\nDocument Statistics:")
    logger.info(f"Total paragraphs: {len(doc.paragraphs)}")
    logger.info(f"Total headings: {sum(1 for p in doc.paragraphs if p.style.name.startswith('Heading'))}")
    logger.info(f"File size: {os.path.getsize(output_path)} bytes")
    
    # Print document structure
    logger.info("\nDocument Structure:")
    logger.info("=" * 50)
    for para in doc.paragraphs:
        if para.style.name.startswith('Heading'):
            level = int(para.style.name[-1]) if para.style.name[-1].isdigit() else 1
            indent = "  " * (level - 1)
            logger.info(f"{indent}{'#' * level} {para.text}")
        elif para.text.strip():
            logger.info(f"  {para.text[:100]}...")
    logger.info("=" * 50)

if __name__ == "__main__":
    main() 