#!/usr/bin/env python3
"""
Test script to diagnose PDF conversion issues
"""

import os
import sys
import logging
from pathlib import Path
from document_generator import PDFConverter

# Set up logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

def test_docx2pdf_import():
    """Test if docx2pdf can be imported and used"""
    logger.info("🧪 Testing docx2pdf import...")
    
    try:
        from docx2pdf import convert
        logger.info("✅ docx2pdf imported successfully")
        return True
    except ImportError as e:
        logger.error(f"❌ Failed to import docx2pdf: {e}")
        return False
    except Exception as e:
        logger.error(f"❌ Unexpected error importing docx2pdf: {e}")
        return False

def test_word_installation():
    """Test if Microsoft Word is available (required for docx2pdf on Windows)"""
    logger.info("🧪 Testing Microsoft Word availability...")
    
    try:
        import win32com.client
        word_app = win32com.client.Dispatch("Word.Application")
        word_app.Visible = False
        word_app.Quit()
        logger.info("✅ Microsoft Word is available")
        return True
    except Exception as e:
        logger.error(f"❌ Microsoft Word not available: {e}")
        logger.info("💡 docx2pdf requires Microsoft Word to be installed on Windows")
        return False

def create_test_word_document():
    """Create a simple test Word document"""
    logger.info("📝 Creating test Word document...")
    
    try:
        from docx import Document
        
        doc = Document()
        doc.add_heading('Test Document', 0)
        doc.add_paragraph('This is a test paragraph for PDF conversion.')
        doc.add_heading('Section 1', level=1)
        doc.add_paragraph('This is content for section 1.')
        
        test_docx_path = "test_conversion.docx"
        doc.save(test_docx_path)
        
        if os.path.exists(test_docx_path):
            logger.info(f"✅ Test Word document created: {test_docx_path}")
            return test_docx_path
        else:
            logger.error("❌ Failed to create test Word document")
            return None
            
    except Exception as e:
        logger.error(f"❌ Error creating test Word document: {e}")
        return None

def test_pdf_conversion_direct():
    """Test PDF conversion using docx2pdf directly"""
    logger.info("🧪 Testing direct docx2pdf conversion...")
    
    test_docx = create_test_word_document()
    if not test_docx:
        return False
    
    try:
        from docx2pdf import convert
        test_pdf_path = "test_conversion.pdf"
        
        logger.info(f"Converting {test_docx} to {test_pdf_path}...")
        convert(test_docx, test_pdf_path)
        
        if os.path.exists(test_pdf_path):
            file_size = os.path.getsize(test_pdf_path)
            logger.info(f"✅ PDF conversion successful: {test_pdf_path} ({file_size} bytes)")
            
            # Clean up
            os.remove(test_docx)
            os.remove(test_pdf_path)
            return True
        else:
            logger.error("❌ PDF file was not created")
            return False
            
    except Exception as e:
        logger.error(f"❌ PDF conversion failed: {e}")
        return False

def test_pdf_converter_class():
    """Test the PDFConverter class"""
    logger.info("🧪 Testing PDFConverter class...")
    
    test_docx = create_test_word_document()
    if not test_docx:
        return False
    
    try:
        converter = PDFConverter()
        test_pdf_path = "test_conversion_class.pdf"
        
        success = converter.convert_word_to_pdf(test_docx, test_pdf_path)
        
        if success and os.path.exists(test_pdf_path):
            file_size = os.path.getsize(test_pdf_path)
            logger.info(f"✅ PDFConverter class works: {test_pdf_path} ({file_size} bytes)")
            
            # Clean up
            os.remove(test_docx)
            os.remove(test_pdf_path)
            return True
        else:
            logger.error("❌ PDFConverter class failed")
            return False
            
    except Exception as e:
        logger.error(f"❌ PDFConverter class error: {e}")
        return False

def check_existing_translated_files():
    """Check for existing translated files to see if PDF conversion failed"""
    logger.info("🔍 Checking for existing translated files...")
    
    # Look for translated Word documents
    word_files = list(Path('.').glob('**/*_translated.docx'))
    pdf_files = list(Path('.').glob('**/*_translated.pdf'))
    
    logger.info(f"Found {len(word_files)} translated Word documents")
    logger.info(f"Found {len(pdf_files)} translated PDF documents")
    
    if word_files and not pdf_files:
        logger.warning("⚠️ Found Word documents but no PDFs - PDF conversion is failing")
        for word_file in word_files[:3]:  # Show first 3
            logger.info(f"   Word file: {word_file}")
    elif word_files and pdf_files:
        logger.info("✅ Both Word and PDF files found")
    else:
        logger.info("ℹ️ No translated files found")
    
    return len(word_files), len(pdf_files)

def suggest_solutions():
    """Suggest solutions for PDF conversion issues"""
    logger.info("\n💡 SUGGESTED SOLUTIONS:")
    logger.info("=" * 50)
    
    logger.info("1. Ensure Microsoft Word is installed and properly licensed")
    logger.info("2. Try alternative PDF conversion methods:")
    logger.info("   - LibreOffice (free alternative)")
    logger.info("   - python-docx2txt + reportlab")
    logger.info("   - Online conversion services")
    
    logger.info("\n3. Alternative packages to try:")
    logger.info("   pip install python-docx2txt")
    logger.info("   pip install reportlab")
    logger.info("   pip install comtypes")
    
    logger.info("\n4. Check Windows permissions and antivirus settings")
    logger.info("5. Run as administrator if needed")

def main():
    """Run all PDF conversion tests"""
    logger.info("🚀 PDF CONVERSION DIAGNOSTIC TOOL")
    logger.info("=" * 50)
    
    tests = [
        ("docx2pdf Import", test_docx2pdf_import),
        ("Microsoft Word", test_word_installation),
        ("Direct Conversion", test_pdf_conversion_direct),
        ("PDFConverter Class", test_pdf_converter_class),
    ]
    
    results = {}
    
    # Check existing files first
    word_count, pdf_count = check_existing_translated_files()
    
    # Run tests
    for test_name, test_func in tests:
        logger.info(f"\n--- {test_name} Test ---")
        try:
            result = test_func()
            results[test_name] = result
            status = "✅ PASSED" if result else "❌ FAILED"
            logger.info(f"{test_name}: {status}")
        except Exception as e:
            logger.error(f"{test_name} test failed with error: {e}")
            results[test_name] = False
    
    # Summary
    logger.info(f"\n🎯 DIAGNOSTIC SUMMARY")
    logger.info("=" * 50)
    
    passed = sum(1 for result in results.values() if result)
    total = len(results)
    
    for test_name, result in results.items():
        status = "✅ PASSED" if result else "❌ FAILED"
        logger.info(f"{test_name}: {status}")
    
    logger.info(f"\nOverall: {passed}/{total} tests passed")
    
    if passed < total:
        suggest_solutions()
    else:
        logger.info("🎉 All tests passed! PDF conversion should work.")
    
    return passed == total

if __name__ == "__main__":
    success = main()
    sys.exit(0 if success else 1)
