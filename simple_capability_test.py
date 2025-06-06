#!/usr/bin/env python3
"""
Simple capability test for Ultimate PDF Translator
Tests basic functionality step by step
"""

import sys
import os
import traceback

def test_basic_imports():
    """Test basic imports"""
    print("🔍 Testing basic imports...")
    
    try:
        import fitz
        print("✅ PyMuPDF (fitz) imported successfully")
    except ImportError as e:
        print(f"❌ PyMuPDF import failed: {e}")
        return False
    
    try:
        from docx import Document
        print("✅ python-docx imported successfully")
    except ImportError as e:
        print(f"❌ python-docx import failed: {e}")
        return False
    
    try:
        import google.generativeai as genai
        print("✅ Google Generative AI imported successfully")
    except ImportError as e:
        print(f"❌ Google Generative AI import failed: {e}")
        return False
    
    return True

def test_config_loading():
    """Test configuration loading"""
    print("\n🔧 Testing configuration loading...")
    
    try:
        # Check if config.ini exists
        if os.path.exists('config.ini'):
            print("✅ config.ini file found")
        else:
            print("❌ config.ini file not found")
            return False
        
        # Try to read config
        import configparser
        config = configparser.ConfigParser()
        config.read('config.ini', encoding='utf-8')
        
        sections = config.sections()
        print(f"✅ Config sections loaded: {sections}")
        
        return True
        
    except Exception as e:
        print(f"❌ Config loading failed: {e}")
        return False

def test_ultimate_translator_import():
    """Test importing the main translator module"""
    print("\n📦 Testing Ultimate PDF Translator import...")
    
    try:
        # Add current directory to path
        current_dir = os.path.dirname(os.path.abspath(__file__))
        if current_dir not in sys.path:
            sys.path.insert(0, current_dir)
        
        # Try importing main module
        import ultimate_pdf_translator
        print("✅ ultimate_pdf_translator module imported successfully")
        
        # Check key variables
        model_name = getattr(ultimate_pdf_translator, 'MODEL_NAME', 'Not found')
        target_lang = getattr(ultimate_pdf_translator, 'TARGET_LANGUAGE_CONFIG', 'Not found')
        
        print(f"✅ Model: {model_name}")
        print(f"✅ Target Language: {target_lang}")
        
        return True
        
    except Exception as e:
        print(f"❌ Ultimate PDF Translator import failed: {e}")
        traceback.print_exc()
        return False

def test_pdf_creation():
    """Test creating a simple PDF for testing"""
    print("\n📄 Testing PDF creation...")
    
    try:
        import fitz
        
        # Create a simple test PDF
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 100), "Test Document", fontsize=16)
        page.insert_text((72, 150), "This is a test paragraph for translation testing.", fontsize=12)
        
        test_pdf_path = "simple_test.pdf"
        doc.save(test_pdf_path)
        doc.close()
        
        # Verify file was created
        if os.path.exists(test_pdf_path):
            file_size = os.path.getsize(test_pdf_path)
            print(f"✅ Test PDF created: {test_pdf_path} ({file_size} bytes)")
            return test_pdf_path
        else:
            print("❌ Test PDF creation failed")
            return None
            
    except Exception as e:
        print(f"❌ PDF creation failed: {e}")
        traceback.print_exc()
        return None

def test_pdf_reading(pdf_path):
    """Test reading PDF content"""
    print("\n📖 Testing PDF reading...")
    
    try:
        import fitz
        
        doc = fitz.open(pdf_path)
        print(f"✅ PDF opened: {len(doc)} pages")
        
        # Extract text from first page
        if len(doc) > 0:
            page_text = doc[0].get_text()
            print(f"✅ Text extracted: {len(page_text)} characters")
            print(f"✅ Sample text: {page_text[:100]}...")
        
        doc.close()
        return True
        
    except Exception as e:
        print(f"❌ PDF reading failed: {e}")
        traceback.print_exc()
        return False

def test_word_document_creation():
    """Test creating a Word document"""
    print("\n📝 Testing Word document creation...")
    
    try:
        from docx import Document
        
        # Create a simple Word document
        doc = Document()
        doc.add_heading('Test Document', 0)
        doc.add_paragraph('This is a test paragraph.')
        doc.add_paragraph('This is another test paragraph.')
        
        test_docx_path = "simple_test.docx"
        doc.save(test_docx_path)
        
        # Verify file was created
        if os.path.exists(test_docx_path):
            file_size = os.path.getsize(test_docx_path)
            print(f"✅ Word document created: {test_docx_path} ({file_size} bytes)")
            return test_docx_path
        else:
            print("❌ Word document creation failed")
            return None
            
    except Exception as e:
        print(f"❌ Word document creation failed: {e}")
        traceback.print_exc()
        return None

def test_pdf_conversion(docx_path):
    """Test converting Word to PDF"""
    print("\n🔄 Testing Word to PDF conversion...")
    
    try:
        from docx2pdf import convert
        
        pdf_output_path = "converted_test.pdf"
        convert(docx_path, pdf_output_path)
        
        # Verify conversion
        if os.path.exists(pdf_output_path):
            file_size = os.path.getsize(pdf_output_path)
            print(f"✅ PDF conversion successful: {pdf_output_path} ({file_size} bytes)")
            return pdf_output_path
        else:
            print("❌ PDF conversion failed - output file not created")
            return None
            
    except Exception as e:
        print(f"❌ PDF conversion failed: {e}")
        print("💡 Note: This requires Microsoft Word or LibreOffice to be installed")
        return None

def cleanup_test_files():
    """Clean up test files"""
    print("\n🧹 Cleaning up test files...")
    
    test_files = ["simple_test.pdf", "simple_test.docx", "converted_test.pdf"]
    
    for file_path in test_files:
        if os.path.exists(file_path):
            try:
                os.remove(file_path)
                print(f"✅ Removed: {file_path}")
            except Exception as e:
                print(f"⚠️ Could not remove {file_path}: {e}")

def main():
    """Run simple capability tests"""
    print("🔍 ULTIMATE PDF TRANSLATOR - SIMPLE CAPABILITY TEST")
    print("=" * 60)
    
    test_results = []
    
    # Test 1: Basic imports
    test_results.append(("Basic Imports", test_basic_imports()))
    
    # Test 2: Configuration loading
    test_results.append(("Configuration Loading", test_config_loading()))
    
    # Test 3: Main module import
    test_results.append(("Main Module Import", test_ultimate_translator_import()))
    
    # Test 4: PDF creation
    pdf_path = test_pdf_creation()
    test_results.append(("PDF Creation", pdf_path is not None))
    
    # Test 5: PDF reading (if PDF was created)
    if pdf_path:
        test_results.append(("PDF Reading", test_pdf_reading(pdf_path)))
    
    # Test 6: Word document creation
    docx_path = test_word_document_creation()
    test_results.append(("Word Document Creation", docx_path is not None))
    
    # Test 7: PDF conversion (if Word doc was created)
    if docx_path:
        converted_pdf = test_pdf_conversion(docx_path)
        test_results.append(("PDF Conversion", converted_pdf is not None))
    
    # Clean up
    cleanup_test_files()
    
    # Results summary
    print("\n" + "=" * 60)
    print("📊 TEST RESULTS SUMMARY")
    print("=" * 60)
    
    passed = 0
    total = len(test_results)
    
    for test_name, result in test_results:
        status = "✅ PASS" if result else "❌ FAIL"
        print(f"{status} {test_name}")
        if result:
            passed += 1
    
    success_rate = (passed / total) * 100 if total > 0 else 0
    print(f"\n📈 OVERALL: {passed}/{total} tests passed ({success_rate:.1f}%)")
    
    if success_rate >= 80:
        print("🎉 EXCELLENT! Core capabilities are working.")
        print("💡 Ready for advanced testing and optimization.")
    elif success_rate >= 60:
        print("⚠️ GOOD! Some issues need attention.")
        print("💡 Check failed tests and dependencies.")
    else:
        print("🚨 ATTENTION NEEDED! Multiple core issues.")
        print("💡 Review installation and dependencies.")
    
    return success_rate >= 60

if __name__ == "__main__":
    try:
        success = main()
        print(f"\n🏁 Test completed. Exit code: {0 if success else 1}")
        sys.exit(0 if success else 1)
    except Exception as e:
        print(f"\n💥 Test suite crashed: {e}")
        traceback.print_exc()
        sys.exit(1)
