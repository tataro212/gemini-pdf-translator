#!/usr/bin/env python3
"""
Test script to debug the document save issue
"""

import os
import logging
from docx import Document
from utils import sanitize_filepath

# Set up logging
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

def test_document_save_issue():
    """Test the exact document save process that's failing"""
    
    print("🔧 Testing document save issue...")
    
    # Test the exact path from the logs
    original_path = r"C:\Users\30694\Downloads\words-to-fire-press-betrayal\fel-a-militant-assessment-of-the-experience-of-the-libertarian-student-fed\fel-a-militant-assessment-of-the-experience-of-the-libertarian-student-fed_translated.docx"
    
    print(f"📄 Original path: {original_path}")
    print(f"📏 Path length: {len(original_path)} characters")
    
    # Test sanitization
    sanitized_path = sanitize_filepath(original_path)
    print(f"🧹 Sanitized path: {sanitized_path}")
    print(f"📏 Sanitized length: {len(sanitized_path)} characters")
    print(f"🔄 Path changed: {original_path != sanitized_path}")
    
    # Check if directory exists
    output_dir = os.path.dirname(sanitized_path)
    print(f"📁 Output directory: {output_dir}")
    print(f"📂 Directory exists: {os.path.exists(output_dir)}")
    
    # Try to create directory
    try:
        if not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)
            print(f"✅ Created directory: {output_dir}")
        else:
            print(f"✅ Directory already exists")
    except Exception as e:
        print(f"❌ Failed to create directory: {e}")
        return
    
    # Test document creation and save
    try:
        print("🔄 Creating test document...")
        doc = Document()
        doc.add_heading("Test Document", level=1)
        doc.add_paragraph("This is a test paragraph to verify document creation works.")
        
        print(f"🔄 Attempting to save to: {sanitized_path}")
        doc.save(sanitized_path)
        
        # Check if file was created
        if os.path.exists(sanitized_path):
            size = os.path.getsize(sanitized_path)
            print(f"✅ Document saved successfully!")
            print(f"   📄 File: {sanitized_path}")
            print(f"   📏 Size: {size} bytes")
            
            # List directory contents
            files = os.listdir(output_dir)
            print(f"📂 Directory contents: {files}")
            
            # Clean up test file
            os.remove(sanitized_path)
            print("🧹 Test file cleaned up")
            
        else:
            print("❌ Document save failed - file not found after save!")
            
    except Exception as e:
        print(f"❌ Document save failed with error: {e}")
        import traceback
        traceback.print_exc()

def test_path_length_limits():
    """Test Windows path length limits"""
    print(f"\n🔧 Testing path length limits...")
    
    # Windows has a 260 character limit for full paths (unless long path support is enabled)
    test_path = r"C:\Users\30694\Downloads\words-to-fire-press-betrayal\fel-a-militant-assessment-of-the-experience-of-the-libertarian-student-fed\fel-a-militant-assessment-of-the-experience-of-the-libertarian-student-fed_translated.docx"
    
    print(f"📏 Test path length: {len(test_path)} characters")
    print(f"⚠️ Windows limit: 260 characters")
    print(f"🚨 Exceeds limit: {len(test_path) > 260}")
    
    if len(test_path) > 260:
        print("❌ PATH TOO LONG! This is likely the issue!")
        print("💡 Solutions:")
        print("   1. Enable Windows long path support")
        print("   2. Use shorter output directory names")
        print("   3. Use shorter filenames")
        
        # Test with shorter path
        shorter_dir = r"C:\Users\30694\Downloads\test_output"
        shorter_path = os.path.join(shorter_dir, "test_translated.docx")
        print(f"\n🔧 Testing with shorter path: {shorter_path}")
        print(f"📏 Shorter path length: {len(shorter_path)} characters")
        
        try:
            os.makedirs(shorter_dir, exist_ok=True)
            doc = Document()
            doc.add_heading("Test Document", level=1)
            doc.add_paragraph("This is a test with shorter path.")
            doc.save(shorter_path)
            
            if os.path.exists(shorter_path):
                size = os.path.getsize(shorter_path)
                print(f"✅ Shorter path works! File size: {size} bytes")
                os.remove(shorter_path)
                os.rmdir(shorter_dir)
            else:
                print("❌ Even shorter path failed!")
                
        except Exception as e:
            print(f"❌ Shorter path test failed: {e}")

if __name__ == "__main__":
    test_document_save_issue()
    test_path_length_limits()
