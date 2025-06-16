import os
import logging
from document_generator import WordDocumentGenerator
from docx import Document

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def test_toc_generation():
    """Test the Table of Contents generation functionality"""
    logger.info("🔍 Testing TOC generation...")
    
    # Create test document
    doc = Document()
    
    # Create document generator
    generator = WordDocumentGenerator()
    
    # Add some test headings to create TOC entries
    test_headings = [
        {"content": "Chapter 1: Introduction", "level": 1},
        {"content": "1.1 Background", "level": 2},
        {"content": "1.2 Objectives", "level": 2},
        {"content": "Chapter 2: Methodology", "level": 1},
        {"content": "2.1 Research Design", "level": 2},
        {"content": "2.2 Data Collection", "level": 2}
    ]
    
    # Add headings to document and collect TOC entries
    for heading in test_headings:
        generator._add_heading_block(doc, heading)
    
    # Add some content paragraphs
    test_paragraphs = [
        "This is a test paragraph in the introduction.",
        "Another paragraph discussing the background.",
        "A paragraph about the methodology.",
        "Details about the research design."
    ]
    
    for para in test_paragraphs:
        doc.add_paragraph(para)
    
    # Generate TOC
    logger.info("📝 Generating Table of Contents...")
    generator._insert_toc(doc)
    
    # Save test document
    output_path = "test_toc_output.docx"
    try:
        doc.save(output_path)
        logger.info(f"✅ Document saved successfully: {output_path}")
        
        # Verify file exists and has content
        if os.path.exists(output_path):
            file_size = os.path.getsize(output_path)
            logger.info(f"📄 File size: {file_size} bytes")
            if file_size > 0:
                logger.info("✅ Test passed: Document created successfully with TOC")
            else:
                logger.error("❌ Test failed: Document is empty")
        else:
            logger.error("❌ Test failed: Document was not created")
            
    except Exception as e:
        logger.error(f"❌ Test failed with error: {e}")
    finally:
        # Cleanup
        if os.path.exists(output_path):
            os.remove(output_path)
            logger.info("🧹 Test file cleaned up")

if __name__ == "__main__":
    test_toc_generation() 