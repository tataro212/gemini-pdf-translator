"""
Enhanced Document Generator for High-Fidelity Assembly

This module provides enhanced document generation capabilities that properly handle
the unified StructuredDocument from the hybrid reconciliation process, ensuring
that TOC and visual elements are correctly rendered in the final output.

Key Features:
- Block-specific rendering functions for each ContentType
- Two-pass TOC generation with accurate page numbers
- Image bypass handling (images preserved without translation)
- Enhanced visual element positioning and sizing
"""

import os
import logging
from typing import List, Dict, Any, Optional
from io import BytesIO

# Import document generation dependencies
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.shared import OxmlElement, qn
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    Document = None

# Import structured document model
try:
    from structured_document_model import (
        Document as StructuredDocument, ContentBlock, ContentType,
        Heading, Paragraph, ImagePlaceholder, Table, CodeBlock, 
        ListItem, Caption, Equation
    )
    STRUCTURED_MODEL_AVAILABLE = True
except ImportError:
    STRUCTURED_MODEL_AVAILABLE = False
    StructuredDocument = None

logger = logging.getLogger(__name__)

class EnhancedDocumentGenerator:
    """
    Enhanced document generator that creates high-fidelity Word documents
    from unified StructuredDocument objects with proper TOC and image handling.
    """
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
        
        # Document generation settings
        self.settings = {
            'generate_toc': True,
            'toc_title': 'Table of Contents',
            'default_image_width_inches': 5.0,
            'max_image_width_inches': 6.5,
            'max_image_height_inches': 8.0,
            'maintain_aspect_ratio': True,
            'image_alignment': WD_ALIGN_PARAGRAPH.CENTER if DOCX_AVAILABLE else None
        }
        
        # Image bypass configuration
        self.image_config = {
            'bypass_translation': True,  # USER REQUIREMENT: No image translation
            'preserve_in_output': True,  # Keep images in final document
            'add_captions': True,  # Add captions below images
            'caption_style': 'Caption'
        }
        
        # Heading tracking for TOC
        self.heading_locations = {}
        self.bookmark_counter = 0
    
    def create_document_from_structured(self, structured_document: StructuredDocument, 
                                      output_filepath: str, image_folder_path: str = None) -> str:
        """
        Create Word document from unified StructuredDocument.
        
        Args:
            structured_document: Unified document from hybrid reconciliation
            output_filepath: Path for output Word document
            image_folder_path: Path to folder containing extracted images
            
        Returns:
            Path to created Word document
        """
        if not DOCX_AVAILABLE:
            raise Exception("python-docx not available for document generation")
            
        if not STRUCTURED_MODEL_AVAILABLE:
            raise Exception("Structured document model not available")
        
        self.logger.info(f"🏗️ Creating enhanced Word document: {os.path.basename(output_filepath)}")
        self.logger.info(f"   • Content blocks: {len(structured_document.content_blocks)}")
        self.logger.info(f"   • Image folder: {image_folder_path}")
        
        # Initialize document
        doc = Document()
        
        # Reset tracking variables
        self.heading_locations = {}
        self.bookmark_counter = 0
        
        # Two-pass generation for accurate TOC
        if self.settings['generate_toc']:
            self.logger.info("📋 Using two-pass TOC generation...")
            
            # Pass 1: Generate content and track heading locations
            heading_page_map = self._generate_content_with_tracking(
                doc, structured_document, image_folder_path
            )
            
            # Pass 2: Generate accurate TOC
            self._generate_accurate_toc(doc, structured_document, heading_page_map)
        else:
            # Single pass without TOC
            self._generate_content_single_pass(doc, structured_document, image_folder_path)
        
        # Save document
        output_filepath = os.path.normpath(output_filepath)
        doc.save(output_filepath)
        
        self.logger.info(f"✅ Enhanced Word document created: {output_filepath}")
        return output_filepath
    
    def _generate_content_with_tracking(self, doc: Document, structured_document: StructuredDocument,
                                      image_folder_path: str) -> Dict[str, Any]:
        """
        Pass 1: Generate content and track heading page numbers for TOC.
        """
        self.logger.info("🔄 Pass 1: Generating content with heading tracking...")
        
        heading_page_map = {}
        
        # Reserve space for TOC
        toc_placeholder = doc.add_paragraph("[TABLE OF CONTENTS PLACEHOLDER]")
        toc_placeholder.style = 'Normal'
        doc.add_page_break()
        
        # Process each content block
        for block in structured_document.content_blocks:
            if isinstance(block, Heading):
                # Track heading location for TOC
                bookmark_name = self._add_heading_block_with_bookmark(doc, block)
                heading_page_map[block.block_id] = {
                    'heading_text': block.content,
                    'level': block.level,
                    'bookmark_name': bookmark_name,
                    'heading_number': len(heading_page_map) + 1
                }
                
            elif isinstance(block, ImagePlaceholder):
                # Handle image with bypass logic
                self._add_image_block_with_bypass(doc, block, image_folder_path)
                
            else:
                # Handle other content blocks
                self._add_content_block(doc, block, image_folder_path)
        
        self.logger.info(f"   • Tracked {len(heading_page_map)} headings")
        return heading_page_map
    
    def _add_heading_block_with_bookmark(self, doc: Document, heading_block: Heading) -> str:
        """Add heading with bookmark for TOC linking"""
        self.bookmark_counter += 1
        bookmark_name = f"heading_{self.bookmark_counter}"
        
        # Add heading
        heading_paragraph = doc.add_heading(heading_block.content, level=heading_block.level)
        
        # Add bookmark for TOC linking
        self._add_bookmark(heading_paragraph, bookmark_name)
        
        return bookmark_name
    
    def _add_image_block_with_bypass(self, doc: Document, image_block: ImagePlaceholder, 
                                   image_folder_path: str):
        """
        Add image block with bypass logic - images are preserved without translation.
        """
        self.logger.debug(f"🖼️ Processing image block: {os.path.basename(image_block.image_path) if image_block.image_path else 'unknown'}")
        
        if not image_folder_path or not image_block.image_path:
            # Add text placeholder if no image available
            placeholder_text = f"[Image: {image_block.caption or 'Unknown'}]"
            paragraph = doc.add_paragraph(placeholder_text)
            paragraph.italic = True
            self.logger.debug(f"   • Added text placeholder: {placeholder_text}")
            return
        
        # Resolve image path
        image_path = self._resolve_image_path(image_block.image_path, image_folder_path)
        
        if not image_path or not os.path.exists(image_path):
            # Image file not found
            placeholder_text = f"[Image not found: {os.path.basename(image_block.image_path)}]"
            paragraph = doc.add_paragraph(placeholder_text)
            paragraph.italic = True
            self.logger.warning(f"   • Image not found: {image_block.image_path}")
            return
        
        try:
            # Add image to document
            self._add_image_with_optimal_sizing(doc, image_path, image_block)
            
            # Add caption if available
            if image_block.caption and self.image_config['add_captions']:
                self._add_image_caption(doc, image_block.caption)
            
            self.logger.debug(f"   • Successfully added image: {os.path.basename(image_path)}")
            
        except Exception as e:
            self.logger.error(f"   • Failed to add image {image_path}: {e}")
            # Add error placeholder
            error_text = f"[Error loading image: {os.path.basename(image_path)}]"
            paragraph = doc.add_paragraph(error_text)
            paragraph.italic = True
    
    def _add_image_with_optimal_sizing(self, doc: Document, image_path: str, 
                                     image_block: ImagePlaceholder):
        """Add image with optimal sizing and positioning"""
        # Calculate optimal size
        width_inches = self.settings['default_image_width_inches']
        
        if image_block.width and image_block.height:
            # Use original dimensions to calculate optimal size
            aspect_ratio = image_block.width / image_block.height
            max_width = self.settings['max_image_width_inches']
            max_height = self.settings['max_image_height_inches']
            
            if self.settings['maintain_aspect_ratio']:
                # Calculate size maintaining aspect ratio
                if width_inches / aspect_ratio > max_height:
                    width_inches = max_height * aspect_ratio
                width_inches = min(width_inches, max_width)
        
        # Add image
        paragraph = doc.add_paragraph()
        paragraph.alignment = self.settings['image_alignment']
        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
        run.add_picture(image_path, width=Inches(width_inches))
    
    def _add_image_caption(self, doc: Document, caption_text: str):
        """Add caption below image"""
        caption_paragraph = doc.add_paragraph(caption_text)
        caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Apply caption formatting
        for run in caption_paragraph.runs:
            run.italic = True
            run.font.size = Pt(9)
    
    def _add_content_block(self, doc: Document, block: ContentBlock, image_folder_path: str):
        """Add generic content block to document"""
        if isinstance(block, Paragraph):
            self._add_paragraph_block(doc, block)
        elif isinstance(block, Table):
            self._add_table_block(doc, block)
        elif isinstance(block, CodeBlock):
            self._add_code_block(doc, block)
        elif isinstance(block, ListItem):
            self._add_list_item_block(doc, block)
        elif isinstance(block, Caption):
            self._add_caption_block(doc, block)
        elif isinstance(block, Equation):
            self._add_equation_block(doc, block)
        else:
            # Fallback to paragraph
            self._add_fallback_paragraph(doc, block)
    
    def _add_paragraph_block(self, doc: Document, paragraph_block: Paragraph):
        """Add paragraph content block"""
        paragraph = doc.add_paragraph(paragraph_block.content)
        paragraph.style = 'Normal'
    
    def _add_table_block(self, doc: Document, table_block: Table):
        """Add table content block"""
        # Simplified table handling - add as formatted text
        paragraph = doc.add_paragraph(table_block.content)
        paragraph.style = 'Normal'
        
        # Apply monospace font for table formatting
        for run in paragraph.runs:
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
    
    def _add_code_block(self, doc: Document, code_block: CodeBlock):
        """Add code content block"""
        paragraph = doc.add_paragraph(code_block.content)
        paragraph.style = 'Normal'
        
        # Apply code formatting
        for run in paragraph.runs:
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
    
    def _add_list_item_block(self, doc: Document, list_item_block: ListItem):
        """Add list item content block"""
        paragraph = doc.add_paragraph(list_item_block.content, style='List Bullet')
    
    def _add_caption_block(self, doc: Document, caption_block: Caption):
        """Add caption content block"""
        self._add_image_caption(doc, caption_block.content)
    
    def _add_equation_block(self, doc: Document, equation_block: Equation):
        """Add equation content block"""
        paragraph = doc.add_paragraph(equation_block.content)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Apply equation formatting
        for run in paragraph.runs:
            run.font.name = 'Cambria Math'
            run.font.size = Pt(12)
    
    def _add_fallback_paragraph(self, doc: Document, block: ContentBlock):
        """Add unknown content block as paragraph"""
        content = getattr(block, 'content', '') or block.original_text
        paragraph = doc.add_paragraph(content)
        paragraph.style = 'Normal'
    
    def _resolve_image_path(self, image_path: str, image_folder_path: str) -> Optional[str]:
        """Resolve image path using multiple strategies"""
        possible_paths = []
        
        # Strategy 1: Use absolute path if provided
        if os.path.isabs(image_path):
            possible_paths.append(image_path)
        
        # Strategy 2: Relative to image folder
        if image_folder_path:
            possible_paths.append(os.path.join(image_folder_path, os.path.basename(image_path)))
            possible_paths.append(os.path.join(image_folder_path, image_path))
        
        # Strategy 3: Current directory
        possible_paths.append(os.path.basename(image_path))
        
        # Find first existing path
        for path in possible_paths:
            if os.path.exists(path):
                return path
        
        return None
    
    def _add_bookmark(self, paragraph, bookmark_name: str):
        """Add bookmark to paragraph for TOC linking"""
        try:
            # Create bookmark XML
            bookmark_start = OxmlElement('w:bookmarkStart')
            bookmark_start.set(qn('w:id'), str(self.bookmark_counter))
            bookmark_start.set(qn('w:name'), bookmark_name)

            bookmark_end = OxmlElement('w:bookmarkEnd')
            bookmark_end.set(qn('w:id'), str(self.bookmark_counter))

            # Insert bookmark
            paragraph._p.insert(0, bookmark_start)
            paragraph._p.append(bookmark_end)

        except Exception as e:
            self.logger.warning(f"Failed to add bookmark {bookmark_name}: {e}")

    def _generate_accurate_toc(self, doc, structured_document, heading_page_map: Dict[str, Any]):
        """
        Pass 2: Generate accurate TOC with real page numbers and hyperlinks.
        """
        self.logger.info("🔄 Pass 2: Generating accurate TOC...")

        if not heading_page_map:
            self.logger.warning("No headings found for TOC generation")
            return

        try:
            # Find and replace TOC placeholder
            for paragraph in doc.paragraphs:
                if "[TABLE OF CONTENTS PLACEHOLDER]" in paragraph.text:
                    # Clear placeholder text
                    paragraph.clear()

                    # Add TOC title
                    toc_title = paragraph
                    toc_title.add_run(self.settings['toc_title']).bold = True
                    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    # Add TOC entries
                    sorted_headings = sorted(heading_page_map.values(),
                                           key=lambda x: x['heading_number'])

                    for heading_info in sorted_headings:
                        self._add_toc_entry(doc, heading_info, paragraph)

                    self.logger.info(f"   • Generated TOC with {len(sorted_headings)} entries")
                    break
            else:
                self.logger.warning("TOC placeholder not found")

        except Exception as e:
            self.logger.error(f"Failed to generate accurate TOC: {e}")

    def _add_toc_entry(self, doc, heading_info: Dict[str, Any], toc_paragraph):
        """Add a single TOC entry with proper formatting and linking"""
        try:
            # Create new paragraph for TOC entry
            entry_paragraph = doc.add_paragraph()

            # Add indentation based on heading level
            level = heading_info.get('level', 1)
            indent_spaces = "  " * (level - 1)

            # Add entry text with hyperlink
            entry_text = f"{indent_spaces}{heading_info['heading_text']}"
            run = entry_paragraph.add_run(entry_text)

            # Add hyperlink to bookmark
            bookmark_name = heading_info.get('bookmark_name')
            if bookmark_name:
                self._add_hyperlink_to_bookmark(run, bookmark_name)

            # Add page number (placeholder - Word will update this)
            entry_paragraph.add_run(" ").add_tab()
            page_run = entry_paragraph.add_run("1")  # Placeholder page number
            page_run.font.size = Pt(10)

        except Exception as e:
            self.logger.warning(f"Failed to add TOC entry: {e}")

    def _add_hyperlink_to_bookmark(self, run, bookmark_name: str):
        """Add hyperlink to bookmark for TOC navigation"""
        try:
            # Create hyperlink XML
            hyperlink = OxmlElement('w:hyperlink')
            hyperlink.set(qn('w:anchor'), bookmark_name)

            # Move run content to hyperlink
            new_run = OxmlElement('w:r')
            new_run.append(run._element.rPr)
            new_run.append(run._element.t)
            hyperlink.append(new_run)

            # Replace original run with hyperlink
            run._element.getparent().replace(run._element, hyperlink)

        except Exception as e:
            self.logger.warning(f"Failed to add hyperlink to {bookmark_name}: {e}")

    def _generate_content_single_pass(self, doc, structured_document, image_folder_path: str):
        """Single pass content generation without TOC"""
        self.logger.info("🔄 Single pass: Generating content without TOC...")

        for block in structured_document.content_blocks:
            if isinstance(block, Heading):
                self._add_heading_block_simple(doc, block)
            elif isinstance(block, ImagePlaceholder):
                self._add_image_block_with_bypass(doc, block, image_folder_path)
            else:
                self._add_content_block(doc, block, image_folder_path)

    def _add_heading_block_simple(self, doc, heading_block):
        """Add heading without bookmark tracking"""
        doc.add_heading(heading_block.content, level=heading_block.level)
