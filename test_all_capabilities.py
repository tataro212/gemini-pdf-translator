#!/usr/bin/env python3
"""
Comprehensive capability testing for Ultimate PDF Translator
Tests every feature from PDF input to final output without requiring API calls
"""

import sys
import os
import tempfile
import shutil
import json
from pathlib import Path
import fitz  # PyMuPDF
from docx import Document
import time

# Add the current directory to path
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

def create_test_pdf():
    """Create a test PDF with various content types for testing"""
    print("📄 Creating test PDF with diverse content...")
    
    # Create a simple test PDF using PyMuPDF
    doc = fitz.open()
    
    # Page 1: Title page
    page1 = doc.new_page()
    page1.insert_text((72, 100), "Test Document for Translation", fontsize=20)
    page1.insert_text((72, 150), "Ultimate PDF Translator Test", fontsize=16)
    page1.insert_text((72, 200), "Author: Test Suite", fontsize=12)
    
    # Page 2: Table of Contents
    page2 = doc.new_page()
    page2.insert_text((72, 100), "Table of Contents", fontsize=18)
    page2.insert_text((72, 150), "Chapter 1: Introduction .................. 3", fontsize=12)
    page2.insert_text((72, 170), "Chapter 2: Methodology ................. 4", fontsize=12)
    page2.insert_text((72, 190), "Chapter 3: Results ..................... 5", fontsize=12)
    page2.insert_text((72, 210), "Bibliography ........................... 6", fontsize=12)
    
    # Page 3: Chapter 1
    page3 = doc.new_page()
    page3.insert_text((72, 100), "Chapter 1: Introduction", fontsize=16)
    page3.insert_text((72, 140), "This is the introduction chapter. It contains multiple paragraphs", fontsize=12)
    page3.insert_text((72, 160), "with different types of content to test the translation system.", fontsize=12)
    page3.insert_text((72, 200), "• First bullet point", fontsize=12)
    page3.insert_text((72, 220), "• Second bullet point", fontsize=12)
    page3.insert_text((72, 240), "• Third bullet point", fontsize=12)
    
    # Page 4: Chapter 2 with technical content
    page4 = doc.new_page()
    page4.insert_text((72, 100), "Chapter 2: Methodology", fontsize=16)
    page4.insert_text((72, 140), "The research methodology follows these steps:", fontsize=12)
    page4.insert_text((72, 160), "1. Data collection from multiple sources", fontsize=12)
    page4.insert_text((72, 180), "2. Statistical analysis using SPSS", fontsize=12)
    page4.insert_text((72, 200), "3. Validation of results", fontsize=12)
    page4.insert_text((72, 240), "Figure 1: Research Process Diagram", fontsize=10)
    
    # Page 5: Results with mixed content
    page5 = doc.new_page()
    page5.insert_text((72, 100), "Chapter 3: Results", fontsize=16)
    page5.insert_text((72, 140), "The analysis revealed significant findings:", fontsize=12)
    page5.insert_text((72, 160), "- 85% improvement in efficiency", fontsize=12)
    page5.insert_text((72, 180), "- Cost reduction of 40%", fontsize=12)
    page5.insert_text((72, 200), "- User satisfaction increased by 60%", fontsize=12)
    page5.insert_text((72, 240), "Table 1: Statistical Results", fontsize=10)
    page5.insert_text((72, 260), "Mean: 4.5, SD: 1.2, p < 0.05", fontsize=10)
    
    # Page 6: Bibliography
    page6 = doc.new_page()
    page6.insert_text((72, 100), "Bibliography", fontsize=16)
    page6.insert_text((72, 140), "Smith, J. (2023). Research Methods in AI. Tech Press.", fontsize=12)
    page6.insert_text((72, 160), "Jones, M. (2022). Data Analysis Techniques. Science Journal.", fontsize=12)
    page6.insert_text((72, 180), "Brown, K. (2021). Statistical Validation. Math Review.", fontsize=12)
    
    # Save test PDF
    test_pdf_path = "test_document.pdf"
    doc.save(test_pdf_path)
    doc.close()
    
    print(f"✅ Test PDF created: {test_pdf_path}")
    return test_pdf_path

def test_configuration_loading():
    """Test configuration file loading and validation"""
    print("\n🔧 Testing Configuration Loading...")
    
    try:
        from ultimate_pdf_translator import (
            MODEL_NAME, TARGET_LANGUAGE_CONFIG, ENABLE_SMART_GROUPING,
            MAX_GROUP_SIZE_CHARS, EXTRACT_IMAGES, PERFORM_OCR_ON_IMAGES
        )
        
        print(f"✅ Model: {MODEL_NAME}")
        print(f"✅ Target Language: {TARGET_LANGUAGE_CONFIG}")
        print(f"✅ Smart Grouping: {ENABLE_SMART_GROUPING}")
        print(f"✅ Max Group Size: {MAX_GROUP_SIZE_CHARS}")
        print(f"✅ Extract Images: {EXTRACT_IMAGES}")
        print(f"✅ Perform OCR: {PERFORM_OCR_ON_IMAGES}")
        
        return True
    except Exception as e:
        print(f"❌ Configuration loading failed: {e}")
        return False

def test_pdf_extraction(pdf_path):
    """Test PDF content extraction"""
    print("\n📖 Testing PDF Content Extraction...")
    
    try:
        from ultimate_pdf_translator import extract_structured_content_from_pdf
        
        # Test basic PDF opening
        doc = fitz.open(pdf_path)
        print(f"✅ PDF opened successfully: {len(doc)} pages")
        
        # Test text extraction
        total_text = ""
        for page_num in range(len(doc)):
            page_text = doc[page_num].get_text()
            total_text += page_text
        
        print(f"✅ Text extraction: {len(total_text)} characters extracted")
        
        # Test structured content extraction
        structured_content = extract_structured_content_from_pdf(pdf_path, [])
        print(f"✅ Structured extraction: {len(structured_content)} items extracted")
        
        # Analyze extracted content types
        content_types = {}
        for item in structured_content:
            item_type = item.get('type', 'unknown')
            content_types[item_type] = content_types.get(item_type, 0) + 1
        
        print(f"✅ Content types found: {content_types}")
        
        doc.close()
        return structured_content
        
    except Exception as e:
        print(f"❌ PDF extraction failed: {e}")
        import traceback
        traceback.print_exc()
        return None

def test_image_extraction(pdf_path):
    """Test image extraction from PDF"""
    print("\n🖼️ Testing Image Extraction...")
    
    try:
        from ultimate_pdf_translator import extract_images_from_pdf
        
        # Create temporary directory for images
        temp_dir = tempfile.mkdtemp()
        print(f"✅ Temporary directory created: {temp_dir}")
        
        # Extract images
        image_refs = extract_images_from_pdf(pdf_path, temp_dir)
        print(f"✅ Image extraction completed: {len(image_refs)} images found")
        
        # Check if images were actually saved
        saved_images = list(Path(temp_dir).glob("*"))
        print(f"✅ Images saved to disk: {len(saved_images)} files")
        
        # Clean up
        shutil.rmtree(temp_dir)
        print(f"✅ Cleanup completed")
        
        return image_refs
        
    except Exception as e:
        print(f"❌ Image extraction failed: {e}")
        import traceback
        traceback.print_exc()
        return []

def test_smart_grouping(structured_content):
    """Test smart grouping functionality"""
    print("\n📦 Testing Smart Grouping...")
    
    try:
        from ultimate_pdf_translator import SmartGroupingProcessor
        
        grouper = SmartGroupingProcessor(
            max_group_size=12000,
            min_group_items=2,
            preserve_formatting=True
        )
        
        # Test grouping
        groups = grouper.create_smart_groups(structured_content)
        print(f"✅ Smart grouping: {len(structured_content)} items → {len(groups)} groups")
        
        # Analyze grouping efficiency
        original_items = len(structured_content)
        grouped_items = len(groups)
        reduction_percent = ((original_items - grouped_items) / original_items * 100) if original_items > 0 else 0
        
        print(f"✅ Grouping efficiency: {reduction_percent:.1f}% reduction in API calls")
        
        # Test group splitting (simulate translation result)
        test_group = groups[0] if groups else None
        if test_group and len(test_group) > 1 and test_group[0].get('type') == 'grouped_content':
            test_translation = "Translated text part 1.\n\nTranslated text part 2."
            split_results = grouper.split_translated_group(test_translation, test_group[0])
            print(f"✅ Group splitting: {len(split_results)} parts recovered")
        
        return groups
        
    except Exception as e:
        print(f"❌ Smart grouping failed: {e}")
        import traceback
        traceback.print_exc()
        return []

def test_cache_system():
    """Test translation cache functionality"""
    print("\n💾 Testing Cache System...")
    
    try:
        from ultimate_pdf_translator import (
            load_translation_cache, save_translation_cache,
            get_cache_key, TRANSLATION_CACHE
        )
        
        # Test cache loading
        load_translation_cache()
        print(f"✅ Cache loaded: {len(TRANSLATION_CACHE)} entries")
        
        # Test cache key generation
        test_text = "This is a test sentence for caching."
        cache_key = get_cache_key(test_text, "Greek", "gemini-1.5-pro")
        print(f"✅ Cache key generated: {cache_key[:50]}...")
        
        # Test cache saving (add a test entry)
        TRANSLATION_CACHE[cache_key] = "Αυτή είναι μια δοκιμαστική πρόταση για cache."
        save_translation_cache()
        print(f"✅ Cache saved with test entry")
        
        return True
        
    except Exception as e:
        print(f"❌ Cache system failed: {e}")
        import traceback
        traceback.print_exc()
        return False

def test_word_document_creation(structured_content):
    """Test Word document creation"""
    print("\n📝 Testing Word Document Creation...")
    
    try:
        from ultimate_pdf_translator import create_word_document_with_structure
        
        # Create temporary output file
        temp_dir = tempfile.mkdtemp()
        word_output_path = os.path.join(temp_dir, "test_output.docx")
        
        # Simulate translated content
        translated_content = []
        for item in structured_content[:10]:  # Test with first 10 items
            translated_item = item.copy()
            original_text = item.get('text', '')
            if original_text:
                # Simulate translation (just add prefix)
                translated_item['translated_text'] = f"[TRANSLATED] {original_text}"
            translated_content.append(translated_item)
        
        print(f"✅ Prepared {len(translated_content)} items for Word document")
        
        # Create Word document
        create_word_document_with_structure(
            translated_content, 
            word_output_path, 
            None,  # No image folder
            None   # No cover page data
        )
        
        # Verify file was created
        if os.path.exists(word_output_path):
            file_size = os.path.getsize(word_output_path)
            print(f"✅ Word document created: {word_output_path} ({file_size} bytes)")
            
            # Test opening the document
            doc = Document(word_output_path)
            paragraph_count = len(doc.paragraphs)
            print(f"✅ Word document verification: {paragraph_count} paragraphs")
            
            # Clean up
            shutil.rmtree(temp_dir)
            return True
        else:
            print(f"❌ Word document not created")
            return False
        
    except Exception as e:
        print(f"❌ Word document creation failed: {e}")
        import traceback
        traceback.print_exc()
        return False

def test_pdf_conversion():
    """Test Word to PDF conversion"""
    print("\n📄 Testing Word to PDF Conversion...")
    
    try:
        # Create a simple Word document for testing
        temp_dir = tempfile.mkdtemp()
        word_path = os.path.join(temp_dir, "test_conversion.docx")
        pdf_path = os.path.join(temp_dir, "test_conversion.pdf")
        
        # Create test Word document
        doc = Document()
        doc.add_heading('Test Document', 0)
        doc.add_paragraph('This is a test paragraph for PDF conversion.')
        doc.add_paragraph('This is another paragraph with some content.')
        doc.save(word_path)
        
        print(f"✅ Test Word document created: {word_path}")
        
        # Test conversion
        from docx2pdf import convert
        convert(word_path, pdf_path)
        
        if os.path.exists(pdf_path):
            file_size = os.path.getsize(pdf_path)
            print(f"✅ PDF conversion successful: {pdf_path} ({file_size} bytes)")
            
            # Verify PDF content
            pdf_doc = fitz.open(pdf_path)
            page_count = len(pdf_doc)
            text_content = pdf_doc[0].get_text() if page_count > 0 else ""
            pdf_doc.close()
            
            print(f"✅ PDF verification: {page_count} pages, {len(text_content)} characters")
            
            # Clean up
            shutil.rmtree(temp_dir)
            return True
        else:
            print(f"❌ PDF conversion failed - file not created")
            return False
        
    except Exception as e:
        print(f"❌ PDF conversion failed: {e}")
        print("💡 Note: docx2pdf requires Microsoft Word or LibreOffice to be installed")
        import traceback
        traceback.print_exc()
        return False

def test_file_management():
    """Test file and directory management"""
    print("\n📁 Testing File Management...")
    
    try:
        from ultimate_pdf_translator import (
            get_desktop_path, get_specific_output_dir_for_file
        )
        
        # Test desktop path detection
        desktop_path = get_desktop_path()
        print(f"✅ Desktop path detected: {desktop_path}")
        
        # Test output directory creation
        test_pdf_path = "test_document.pdf"
        output_dir = get_specific_output_dir_for_file("/tmp", test_pdf_path)
        print(f"✅ Output directory path: {output_dir}")
        
        # Test directory creation
        if not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)
            print(f"✅ Directory created successfully")
            
            # Clean up
            if os.path.exists(output_dir):
                os.rmdir(output_dir)
                print(f"✅ Directory cleanup completed")
        
        return True
        
    except Exception as e:
        print(f"❌ File management failed: {e}")
        import traceback
        traceback.print_exc()
        return False

def test_optimization_features():
    """Test optimization features"""
    print("\n🚀 Testing Optimization Features...")
    
    try:
        from ultimate_pdf_translator import (
            optimization_manager, QuotaManager, EnhancedErrorRecovery,
            ProgressTracker, validate_configuration, estimate_translation_cost
        )
        
        # Test quota manager
        quota_manager = QuotaManager()
        print(f"✅ Quota manager initialized: {quota_manager.can_make_request()}")
        
        # Test error recovery
        error_recovery = EnhancedErrorRecovery()
        recovery_report = error_recovery.get_recovery_report()
        print(f"✅ Error recovery system ready")
        
        # Test progress tracker
        progress_tracker = ProgressTracker(100)
        progress_tracker.update(completed=25, failed=5)
        print(f"✅ Progress tracking: 25 completed, 5 failed")
        
        # Test configuration validation
        config_valid = validate_configuration()
        print(f"✅ Configuration validation: {'Passed' if config_valid else 'Issues found'}")
        
        # Test optimization manager
        test_content = [
            {'text': 'Test content for optimization', 'type': 'paragraph'},
            {'text': 'Another test paragraph', 'type': 'paragraph'}
        ]
        
        optimized_content, params = optimization_manager.optimize_translation_pipeline(
            test_content, 'Greek'
        )
        print(f"✅ Optimization manager: {len(test_content)} → {len(optimized_content)} items")
        
        return True
        
    except Exception as e:
        print(f"❌ Optimization features failed: {e}")
        import traceback
        traceback.print_exc()
        return False

def main():
    """Run comprehensive capability testing"""
    print("🔍 ULTIMATE PDF TRANSLATOR - COMPREHENSIVE CAPABILITY TEST")
    print("=" * 80)
    
    test_results = {}
    
    # Create test PDF
    test_pdf_path = create_test_pdf()
    
    try:
        # Test each capability
        test_results['configuration'] = test_configuration_loading()
        test_results['pdf_extraction'] = test_pdf_extraction(test_pdf_path)
        test_results['image_extraction'] = test_image_extraction(test_pdf_path)
        test_results['cache_system'] = test_cache_system()
        test_results['file_management'] = test_file_management()
        test_results['optimization_features'] = test_optimization_features()
        
        # Test smart grouping if PDF extraction worked
        if test_results['pdf_extraction']:
            structured_content = test_results['pdf_extraction']
            test_results['smart_grouping'] = test_smart_grouping(structured_content)
            test_results['word_creation'] = test_word_document_creation(structured_content)
        else:
            test_results['smart_grouping'] = False
            test_results['word_creation'] = False
        
        # Test PDF conversion
        test_results['pdf_conversion'] = test_pdf_conversion()
        
    finally:
        # Clean up test PDF
        if os.path.exists(test_pdf_path):
            os.remove(test_pdf_path)
            print(f"✅ Test PDF cleaned up")
    
    # Generate test report
    print("\n" + "=" * 80)
    print("📊 CAPABILITY TEST RESULTS")
    print("=" * 80)
    
    passed_tests = 0
    total_tests = 0
    
    for test_name, result in test_results.items():
        total_tests += 1
        if result:
            passed_tests += 1
            status = "✅ PASS"
        else:
            status = "❌ FAIL"
        
        print(f"{status} {test_name.replace('_', ' ').title()}")
    
    success_rate = (passed_tests / total_tests) * 100 if total_tests > 0 else 0
    
    print(f"\n📈 OVERALL RESULTS: {passed_tests}/{total_tests} tests passed ({success_rate:.1f}%)")
    
    if success_rate >= 80:
        print("🎉 EXCELLENT! Most capabilities are working correctly.")
    elif success_rate >= 60:
        print("⚠️ GOOD! Some issues need attention.")
    else:
        print("🚨 ATTENTION NEEDED! Multiple capabilities require fixes.")
    
    print("\n💡 NEXT STEPS:")
    if not test_results.get('pdf_conversion', False):
        print("• Install Microsoft Word or LibreOffice for PDF conversion")
    if not test_results.get('configuration', False):
        print("• Check configuration file and environment variables")
    if success_rate < 100:
        print("• Review failed tests and fix underlying issues")
    
    return success_rate >= 80

if __name__ == "__main__":
    success = main()
    sys.exit(0 if success else 1)
